#!/usr/bin/env python3
"""
Argos R&D Research Journal — Word Document Generator

Generates a comprehensive, management-ready research documentation
Word document tracking all R&D progress on the Argos visual speech
processing project.

This is a LIVING DOCUMENT — re-run to regenerate with updates.
All content is embedded; no external file dependencies.

Usage:
    python3 generate_research_journal.py

Output:
    english_1k_results/analysis/docx/research_documentation.docx
"""

import os
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

# ── Output ──
OUTPUT_DIR = Path(__file__).parent / "docx"
OUTPUT_FILE = OUTPUT_DIR / "research_documentation.docx"

# ── Colors ──
C_PRIMARY = RGBColor(0x1a, 0x3a, 0x5c)
C_H2 = RGBColor(0x2a, 0x5a, 0x8c)
C_H3 = RGBColor(0x3a, 0x6a, 0x9c)
C_H4 = RGBColor(0x4a, 0x7a, 0xac)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_BLACK = RGBColor(0x00, 0x00, 0x00)
C_DARK = RGBColor(0x33, 0x33, 0x33)
C_GREEN = RGBColor(0x15, 0x57, 0x24)
C_RED = RGBColor(0x72, 0x1c, 0x24)
C_AMBER = RGBColor(0x85, 0x64, 0x04)
C_GRAY = RGBColor(0x66, 0x66, 0x66)

HEADER_BG = "1a3a5c"
ZEBRA_BG = "f0f4f8"
GREEN_BG = "d4edda"
RED_BG = "f8d7da"
AMBER_BG = "fff3cd"

LAST_UPDATED = datetime.now().strftime("%B %d, %Y at %H:%M UTC")


# ═══════════════════════════════════════════════════
# HELPER FUNCTIONS
# ═══════════════════════════════════════════════════

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, bold=False, color=None, size=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = size
    run.font.name = "Calibri"


def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, HEADER_BG)
        set_cell_text(cell, h, bold=True, color=C_WHITE, size=Pt(9))

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            set_cell_text(cell, val, size=Pt(9))
            if r_idx % 2 == 1:
                set_cell_shading(cell, ZEBRA_BG)

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return table


def add_heading(doc, text, level, color=None):
    h = doc.add_heading(text, level=level)
    if color is None:
        color = {1: C_PRIMARY, 2: C_H2, 3: C_H3, 4: C_H4}.get(level, C_PRIMARY)
    for run in h.runs:
        run.font.color.rgb = color
        run.font.name = "Calibri"
    return h


def add_para(doc, text, bold=False, italic=False, size=Pt(11), color=None, space_after=Pt(6)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = size
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = space_after
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = "Calibri"
    return p


def add_bullet_bold_value(doc, label, value, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after = Pt(2)
    run_b = p.add_run(label)
    run_b.bold = True
    run_b.font.size = Pt(10)
    run_b.font.name = "Calibri"
    run_v = p.add_run(value)
    run_v.font.size = Pt(10)
    run_v.font.name = "Calibri"
    return p


def add_header_footer(doc):
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    # Header (not on first page)
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run("Argos R&D Research Journal  |  INTERNAL")
    run.font.size = Pt(8)
    run.font.color.rgb = C_GRAY
    run.font.name = "Calibri"
    run.italic = True

    # Footer with page number (not on first page)
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = C_GRAY
    run.font.name = "Calibri"
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    run._r.append(fld1)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run._r.append(instr)
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld2)


# ═══════════════════════════════════════════════════
# DOCUMENT SECTIONS
# ═══════════════════════════════════════════════════

def create_cover_page(doc):
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ARGOS")
    run.font.size = Pt(42)
    run.font.color.rgb = C_PRIMARY
    run.bold = True
    run.font.name = "Calibri"

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Visual Speech Processing\nR&D Research Journal")
    run2.font.size = Pt(20)
    run2.font.color.rgb = C_H2
    run2.font.name = "Calibri"

    doc.add_paragraph()

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("Development Log, Technical Analysis & Progress Report")
    run3.font.size = Pt(13)
    run3.font.color.rgb = C_H3
    run3.italic = True
    run3.font.name = "Calibri"

    for _ in range(3):
        doc.add_paragraph()

    info_lines = [
        ("Period:", "July 2025 — February 2026"),
        ("Last Updated:", LAST_UPDATED),
        ("Repository:", "github.com/MrYoyoad/Argos"),
        ("Classification:", "Internal R&D Documentation"),
    ]
    for label, value in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rl = p.add_run(label + " ")
        rl.bold = True
        rl.font.size = Pt(11)
        rl.font.color.rgb = C_DARK
        rl.font.name = "Calibri"
        rv = p.add_run(value)
        rv.font.size = Pt(11)
        rv.font.color.rgb = C_DARK
        rv.font.name = "Calibri"

    doc.add_page_break()


def section_1_overview(doc):
    add_heading(doc, "1. Project Overview", 1)

    add_para(doc, (
        "Argos is a production visual speech processing (lip-reading) system that takes raw video "
        "as input and produces text transcriptions by analyzing lip movements. It is built on top of "
        "the VSP-LLM research paper (Yeo et al., arXiv:2402.15151, 2024) which combines a visual "
        "speech encoder with a large language model."
    ))

    add_heading(doc, "Architecture", 2)
    add_para(doc, (
        "The system chains three neural network components: (1) AV-HuBERT, a self-supervised visual "
        "speech encoder that extracts 1024-dimensional features from 88x88 mouth crop videos at 25fps; "
        "(2) a linear bridge layer that projects these features into the 4096-dimensional space of the "
        "language model; and (3) LLaMA-2-7B, a 7-billion parameter language model fine-tuned with QLoRA "
        "(rank 16, 4-bit quantization) to generate text from visual speech features."
    ))
    add_para(doc, (
        "Visual speech units are deduplicated using K-means clustering (200 clusters), reducing "
        "consecutive identical frames by ~50%, before being fed to the language model for text generation."
    ))

    add_heading(doc, "Repository Structure", 2)
    add_styled_table(doc,
        ["Component", "Repository", "Role"],
        [
            ["Argos (main)", "MrYoyoad/Argos", "Pipeline orchestration, UI, deployment, documentation"],
            ["VSP-LLM (fork)", "MrYoyoad/VSP-LLM", "Model architecture, decode, training, clustering"],
            ["auto_avsr (fork)", "MrYoyoad/auto_avsr", "Video preprocessing, mouth cropping, ASR"],
            ["av_hubert", "facebookresearch/av_hubert", "Feature extraction, data preparation"],
        ],
        col_widths=[1.5, 2.0, 3.0]
    )

    add_heading(doc, "Technology Stack", 2)
    add_styled_table(doc,
        ["Layer", "Technology", "Version/Details"],
        [
            ["GPU Runtime", "NVIDIA CUDA", "12.8 (preprocessing), 12.4 (VSP-LLM)"],
            ["Language Model", "LLaMA-2-7B", "4-bit NF4 quantization via BitsAndBytes"],
            ["Visual Encoder", "AV-HuBERT Large", "Pre-trained on LRS3 + VoxCeleb2"],
            ["Fine-tuning", "QLoRA (PEFT)", "Rank 16, targets Q/K/V projections"],
            ["ML Framework", "PyTorch + fairseq", "2.8.0 (preprocessing), 2.5.1 (VSP-LLM)"],
            ["ASR", "OpenAI Whisper", "Medium model, English"],
            ["Face Detection", "MediaPipe", "v0.10.21"],
            ["Web UI", "Flask + JavaScript", "Custom-built, 4,500+ lines"],
            ["Deployment", "Docker + NVIDIA Container Toolkit", "Ubuntu 22.04 base"],
            ["Client OS", "Arch Linux", "With offline NVIDIA driver packages"],
        ],
        col_widths=[1.3, 2.0, 3.2]
    )
    doc.add_page_break()


def section_2_paper_vs_argos(doc):
    add_heading(doc, "2. What the Paper Provides vs What Argos Built", 1)

    add_heading(doc, "Layer 1: Upstream Open-Source Projects", 2)
    add_para(doc, "Argos builds on four upstream open-source projects that provide foundational capabilities:")
    add_styled_table(doc,
        ["Project", "Source", "What It Provides"],
        [
            ["AV-HuBERT", "facebookresearch/av_hubert", "Self-supervised visual speech encoder (315M params)"],
            ["Auto-AVSR", "mpc001/auto_avsr", "Audio-visual preprocessing, mouth cropping, LRS3 format"],
            ["fairseq", "facebookresearch/fairseq", "Sequence-to-sequence training framework (~66K lines)"],
            ["LLaMA-2-7B", "meta-llama/Llama-2-7b-hf", "7B parameter language model backbone"],
        ],
        col_widths=[1.3, 2.3, 3.0]
    )

    add_heading(doc, "Layer 2: The VSP-LLM Paper (Sally-SH/VSP-LLM)", 2)
    add_para(doc, (
        "The paper repository adds the core VSP-LLM model architecture on top of the upstream projects. "
        "It provides: the model definition (vsp_llm.py, ~15KB), K-means clustering pipeline (10 Python files), "
        "QLoRA integration with BitsAndBytes, basic training/decode wrappers (train.sh at 31 lines, decode.sh "
        "at 2 lines), and Hydra configuration files."
    ))
    add_para(doc, (
        "Critically, the paper repo does NOT provide: any video preprocessing pipeline, end-to-end "
        "orchestration, ASR integration, user interface, deployment infrastructure, or reporting beyond "
        "basic WER. The README instructs users to 'follow AV-HuBERT and Auto-AVSR docs separately' for "
        "data preparation — there is no unified pipeline."
    ), bold=False, italic=True)

    add_heading(doc, "Layer 3: Argos — What We Built From Scratch", 2)

    add_heading(doc, "3a. First Client Delivery (Jul–Nov 2024)", 3)
    add_para(doc, (
        "The original VSP-LLM paper code was downloaded (to_download/VSP-LLM.tar.gz, July 2025) and the "
        "work began to turn it into a usable product. The first deliverable was the FOR_CLIENT_OLD package:"
    ))
    add_bullet(doc, "vsp-llm-pipeline.tar.gz — a 90MB Docker image (OCI format, ~2.7GB uncompressed) containing the first containerized pipeline")
    add_bullet(doc, "VSP-LLM-App/ — client-facing application folder with input_videos/ and outputs/ directories")
    add_bullet(doc, "run_pipeline.txt — Docker-based launcher script that maps host volumes and runs the pipeline")
    add_bullet(doc, "vsp-llm.Desktop.txt — Linux desktop icon for one-click execution via gnome-terminal")
    add_bullet(doc, "arch-nvidia-offline/ — offline NVIDIA driver packages (v580.119.02) for the client's Arch Linux system, with pacman installation instructions")
    add_bullet(doc, "installation_order_for_drivers.txt — step-by-step driver installation guide (blacklist nouveau, install DKMS, build kernel modules)")

    add_heading(doc, "3b. Galaxy Export — First Production Pipeline (Nov 2024)", 3)
    add_para(doc, (
        "The galaxy_export/ directory was created as a self-contained deployable package. "
        "This was the foundation for all subsequent container builds:"
    ))
    add_bullet(doc, "run_flat_english_pipeline.sh — 285-line monolithic pipeline built from scratch, integrating all 4 upstream repos into 8 sequential stages")
    add_bullet(doc, "Original bash UI: ui/process_videos.sh (30 lines) — simple CLI prompting users to place videos in a folder and press Enter")
    add_bullet(doc, "Pre-built virtual environments shipped with the export (pre-process-venv ~8.2GB, vsp-llm-yoad-venv ~6.2GB)")
    add_bullet(doc, "Offline requirements and pre-downloaded Whisper model cache for air-gapped deployment")
    add_bullet(doc, "Dockerfile copies this galaxy_export directory into /workspace/ to create the container image")

    add_heading(doc, "3c. Pipeline Evolution (Nov 2024 – Jan 2026)", 3)
    add_para(doc, "The pipeline underwent continuous development, growing from 285 to 803 lines:")

    add_styled_table(doc,
        ["Version", "Lines", "Date", "Key Changes"],
        [
            ["Original (galaxy_export)", "285", "Nov 2024", "8 stages, basic normalization, 4s segments"],
            ["Jan 18 backup (.bak)", "240", "Jan 18, 2026", "Simplified EC2 version (removed Docker paths)"],
            ["updated_script.txt (FOR_CLIENT)", "410", "Dec 2025", "QUIET/NORM controls, TRAIN_KMEANS toggle, Cython check, host export"],
            ["Pre-refactor backup", "612", "Jan 2026", "Container-adapted version with /host/galaxy_export paths"],
            ["Pre-Claude backup", "673", "Feb 1, 2026", "Feature-complete monolithic version"],
            ["Initial git commit", "803–823", "Jan 29, 2026", "Full features: segmentation, overlap, reports, burned videos"],
            ["Current (modular)", "393 + 1,526", "Feb 18, 2026", "11 lib/ modules, 37 tests, environment-aware"],
        ],
        col_widths=[1.8, 0.6, 1.2, 3.0]
    )

    add_heading(doc, "3d. Custom Python Scripts (1,886 lines)", 3)
    add_styled_table(doc,
        ["Script", "Lines", "Purpose"],
        [
            ["fast_segment.py", "197", "Ultra-fast ffmpeg -c copy video segmentation"],
            ["preprocess_with_overlap.py", "370", "Extends upstream to support overlapping segments"],
            ["make_report.py", "809", "Report generation: WER, WWER, NEA (spaCy), HTML/ANSI/CSV/JSON"],
            ["make_burn.py", "510", "Burns subtitles onto original video segments with ffmpeg"],
            ["merge_overlapping_predictions.py", "289", "Merges predictions from overlapping segments with conflict detection"],
            ["transcribe_segments.py", "353", "Post-processing of segment transcriptions"],
            ["calculate_per_video_wer.py", "300", "Per-video WER calculation with overlap deduplication"],
        ],
        col_widths=[2.5, 0.6, 3.5]
    )

    add_heading(doc, "3e. Web UI Built From Scratch (4,564 lines)", 3)
    add_para(doc, "The original 30-line bash UI was replaced with a full-featured web application:")
    add_styled_table(doc,
        ["Component", "File", "Lines", "Description"],
        [
            ["Backend", "server.py", "1,124", "Flask REST API, pipeline execution, progress tracking"],
            ["Pipeline Runner", "pipeline_runner.py", "349", "Subprocess orchestration, log monitoring"],
            ["Progress Tracker", "progress_tracker.py", "206", "Real-time log parsing, stage detection"],
            ["Transcription Mgr", "transcription_manager.py", "280", "CRUD operations, auto/manual metadata"],
            ["Validator", "validator.py", "374", "Input validation, video inspection"],
            ["Frontend", "app.js", "1,921", "Client-side logic, drag-drop, transcription UI"],
            ["Layout", "index.html", "328", "Responsive UI layout"],
            ["Styling", "style.css", "1,191", "Dark theme, responsive design"],
        ],
        col_widths=[1.3, 1.8, 0.5, 3.0]
    )

    add_heading(doc, "Quantitative Comparison", 2)
    add_styled_table(doc,
        ["Aspect", "Paper Repo", "Original Argos", "Current Argos"],
        [
            ["Pipeline orchestration", "0 lines", "285–803 lines", "393 + 1,526 lines (modular)"],
            ["Video preprocessing", "Not included", "Normalization + segmentation", "HDR/10-bit + overlap + fallback"],
            ["ASR integration", "Not included", "Whisper (basic)", "Whisper + transcription reuse"],
            ["User interface", "Not included", "30-line bash → 4,564-line web UI", "Enhanced (9 features)"],
            ["Report generation", "~100 lines", "809 lines (WER/WWER/NEA)", "28.5KB (8x expansion)"],
            ["Burned video output", "Not included", "510 lines (make_burn.py)", "Enhanced with metadata"],
            ["Container deployment", "Not supported", "Dockerfile + galaxy_export", "5 versions, INSTALL.sh, desktop launcher"],
            ["Testing", "No tests", "Manual testing", "37 automated tests"],
            ["Documentation", "Basic README", "Inline comments", "3,000+ lines across 5 docs"],
        ],
        col_widths=[1.5, 1.5, 2.0, 2.0]
    )
    doc.add_page_break()


def section_3_comparison(doc):
    add_heading(doc, "3. Pipeline Parameters: Argos vs Paper", 1)
    add_para(doc, (
        "The following table shows that Argos preserves all core model parameters from the paper. "
        "The only differences are in the production infrastructure built around the model."
    ))

    add_styled_table(doc,
        ["Parameter", "VSP-LLM Paper", "Argos", "Match?"],
        [
            ["Visual encoder", "AV-HuBERT Large", "AV-HuBERT Large", "YES"],
            ["Encoder features", "1024-dim (layer 12)", "1024-dim (layer 12)", "YES"],
            ["LLM backbone", "LLaMA-2-7B", "LLaMA-2-7B", "YES"],
            ["Quantization", "QLoRA 4-bit NF4", "QLoRA 4-bit NF4", "YES"],
            ["LoRA rank", "16", "16", "YES"],
            ["LoRA targets", "Q, V projections", "Q, K, V projections", "K added"],
            ["LoRA dropout", "5%", "5%", "YES"],
            ["K-means clusters", "200", "200", "YES"],
            ["Beam search width", "20", "20", "YES"],
            ["Length penalty", "0.0", "0.0", "YES"],
            ["Learning rate", "5e-4", "5e-4", "YES"],
            ["Training steps", "30K (433h data)", "30K", "YES"],
            ["Optimizer", "Adam (0.9, 0.98)", "Adam (0.9, 0.98)", "YES"],
            ["LR scheduler", "Tri-stage", "Tri-stage", "YES"],
            ["Freeze schedule", "18K steps", "18K steps (finetune)", "YES"],
            ["Training GPUs", "8x RTX 3090", "Config: 8 GPUs", "YES"],
            ["Segmentation", "None (full videos)", "12s segments, 2s overlap", "ADDED"],
            ["Video normalization", "Not discussed", "HDR/10-bit, NVENC/CPU", "ADDED"],
            ["ASR transcription", "Not applicable", "Whisper medium", "ADDED"],
            ["Reporting", "Basic WER", "WER/WWER/NEA/HTML/JSON", "EXTENDED"],
        ],
        col_widths=[1.5, 1.7, 2.0, 1.0]
    )
    doc.add_page_break()


def section_4_dev_log(doc):
    add_heading(doc, "4. Chronological Development Log", 1)

    # ── Pre-Claude ──
    add_heading(doc, "Pre-Claude: Foundation Work (Jul 2025 – Jan 2026)", 2)

    add_heading(doc, "July 2025 — Paper Acquisition & Initial Setup", 3)
    add_bullet(doc, "Downloaded VSP-LLM paper repository from Sally-SH/VSP-LLM (69MB tarball)")
    add_bullet(doc, "Forked to MrYoyoad/VSP-LLM; forked mpc001/auto_avsr to MrYoyoad/auto_avsr")
    add_bullet(doc, "Acquired LLaMA-2-7B model weights from Hugging Face")
    add_bullet(doc, "Acquired AV-HuBERT Large pre-trained checkpoint (large_vox_iter5.pt)")
    add_bullet(doc, "Set up EC2 development environment with Tesla T4 GPU (16GB VRAM)")

    add_heading(doc, "Jul–Nov 2025 — Pipeline Construction", 3)
    add_bullet(doc, "Integrated 4 separate upstream repos into unified pipeline — paper's README provided no guidance on this")
    add_bullet(doc, "Built run_flat_english_pipeline.sh (285 lines) from scratch with 8 processing stages")
    add_bullet(doc, "Set up dual virtual environments: pre-process-venv (PyTorch 2.8+cu128) and vsp-llm-yoad-venv (PyTorch 2.5.1+cu124) — required because preprocessing and VSP-LLM have conflicting package dependencies")
    add_bullet(doc, "Integrated Whisper ASR for automatic transcription (paper assumes pre-existing labels)")
    add_bullet(doc, "Created Whisper wrapper script (asr_to_words_notime.py) for text extraction")
    add_bullet(doc, "Set up K-means clustering with 200 clusters (matching paper)")

    add_heading(doc, "November 2024 — Galaxy Export & First Client Delivery", 3)
    add_bullet(doc, "Created galaxy_export/ as self-contained deployable package (~65GB)")
    add_bullet(doc, "Pre-built virtual environments included for offline deployment")
    add_bullet(doc, "Original UI: simple bash script (ui/process_videos.sh, 30 lines)")
    add_bullet(doc, "Built Docker container (vsp-llm-pipeline.tar.gz, 90MB OCI image)")
    add_bullet(doc, "Prepared offline NVIDIA driver packages for client's Arch Linux system (v580.119.02)")
    add_bullet(doc, "Created VSP-LLM-App/ folder with desktop launcher for client")
    add_bullet(doc, "Delivered FOR_CLIENT_OLD package: Docker image + drivers + installation instructions")

    add_heading(doc, "Dec 2025 – Jan 2026 — Pipeline Growth & Feature Development", 3)
    add_bullet(doc, "Pipeline grew from 285 to 803 lines as features were added")
    add_bullet(doc, "Added video segmentation with configurable overlap (fast_segment.py, 197 lines)")
    add_bullet(doc, "Added overlap-aware preprocessing (preprocess_with_overlap.py, 370 lines)")
    add_bullet(doc, "Added video normalization with NVENC GPU encoding and HDR/10-bit handling")
    add_bullet(doc, "Created make_report.py (809 lines) — extended paper's ~100 lines with WER, WWER, NEA, spaCy, HTML/ANSI output")
    add_bullet(doc, "Created make_burn.py (510 lines) — burns subtitles onto original video segments")
    add_bullet(doc, "Created merge_overlapping_predictions.py (289 lines) — handles overlapping segments")
    add_bullet(doc, "Built complete web UI from scratch (4,564 lines) replacing the 30-line bash UI")
    add_bullet(doc, "Established frame-based segment naming: {video_id}_{seg}_{start_frame}_{end_frame}.mp4")

    # ── Claude R&D Phases ──
    add_heading(doc, "Claude R&D Phase 1: Modular Refactoring (Jan 29, 2026)", 2)
    add_para(doc, "Mission 1: Transform the 823-line monolithic pipeline into a modular architecture.")
    add_bullet(doc, "Initial commit: VSP Pipeline with sync automation (commit 5a0ce3a)")
    add_bullet(doc, "Extracted 11 modules into lib/ directory: common.sh, config.sh, archive.sh, normalization.sh, asr.sh, lrs3_prep.sh, manifests.sh, clustering.sh, decode.sh, outputs.sh, test_all_modules.sh")
    add_bullet(doc, "Result: 52% line reduction (823 to 393 lines), 37 automated tests passing")
    add_bullet(doc, "Environment-aware configuration: auto-detects EC2 vs container paths")
    add_bullet(doc, "Added intelligent transcription matching (Step 0.6) and persistence (Step 1.5)")

    add_heading(doc, "Phase 2: Bug Fixes & Operational Hardening (Jan 29 – Feb 1)", 2)
    add_bullet_bold_value(doc, "CRITICAL — log_info stdout contamination: ", "All logging functions echoed to stdout instead of stderr, contaminating return values. Broke client outputs, archive paths, and derived variables. Fixed: redirect all logging to stderr.")
    add_bullet_bold_value(doc, "Non-segmented video naming: ", "Videos <24s received artificial segment suffixes, breaking transcription matching. Fixed: keep original names.")
    add_bullet_bold_value(doc, "Burned output full-frame fix: ", "Burned videos showed 88x88 mouth crops instead of full-frame originals. Fixed: extract from original video.")
    add_bullet_bold_value(doc, "Segment-first normalization: ", "Restructured to preprocess individual segments instead of full videos. Result: ~90% faster normalization.")
    add_bullet(doc, "Successful end-to-end 12-second segment pipeline run validated")

    add_heading(doc, "Phase 3: Container Sync & Path Corrections (Feb 1–3)", 2)
    add_bullet(doc, "Migrated submodules to MrYoyoad forks (VSP-LLM, auto_avsr)")
    add_bullet(doc, "Systematic path correction: /workspace → /host/galaxy_export across entire codebase")
    add_bullet(doc, "Dynamic transcription paths (uses RAW_DIR instead of hardcoded $HOME)")
    add_bullet(doc, "First deployment package created: vsp_linux_container_FINAL_20260203.tar.gz")
    add_bullet(doc, "Comprehensive deployment documentation: 94KB across 5 files")

    add_heading(doc, "Phase 4: Testing & Documentation (Feb 4)", 2)
    add_bullet(doc, "Comprehensive environment setup documentation")
    add_bullet(doc, "Complete dependency lists for all environments")
    add_bullet(doc, "Installation instruction clarification for container updates")

    add_heading(doc, "Phase 5: Container Hardening (Feb 9–12)", 2)
    add_bullet(doc, "Bugs 11-13 fixed: fairseq max_len configuration, UI networking, host launcher")
    add_bullet(doc, "Added docker-run.sh with correct container start command")
    add_bullet(doc, "Added docker.conf for configurable Docker image names (client vs developer)")
    add_bullet(doc, "Desktop icon launcher tolerates mounted filesystem permissions")

    add_heading(doc, "Phase 6: Deployment Fixes (Feb 15)", 2)
    add_bullet(doc, "Bugs 32-35 fixed: server startup, decode step, terminal detection, docker.conf auto-detect")

    add_heading(doc, "Phase 7: Major Evaluation & Critical Fixes (Feb 17)", 2)
    add_bullet_bold_value(doc, "CRITICAL — NVENC silent corruption (43% video loss): ", "NVENC H.264 encoder silently produced corrupt streams (exit code 0 but undecodable). Combined with bash fd interference that corrupted filenames in loops. 656/1520 videos failed face detection. Fixed: post-encode validation + CPU fallback (libx264). Result: 0% failure rate.")
    add_bullet(doc, "Added lip crop copy to client output folder")
    add_bullet(doc, "Fixed chunked upload for large drag-and-drop files")
    add_bullet(doc, "Split CLAUDE.md into 5 topic-based documentation files")
    add_bullet(doc, "Synced 5 features to standalone container (v1.0.31-35)")
    add_bullet(doc, "Documented 28 pending container sync items with exact code diffs")

    add_heading(doc, "Phase 8: Performance Analysis & Reports (Feb 17–18)", 2)
    add_bullet(doc, "Evaluated 860 segments from ~1,000 YouTube videos (english_1k dataset)")
    add_bullet(doc, "Mean WER: 67.0% — identified 2.6x domain gap vs paper's 25.4% on LRS3")
    add_bullet(doc, "Only 11.4% of segments usable (WER <= 20%); 20.6% catastrophically hallucinated (WER >= 100%)")
    add_bullet(doc, "Created 6 comprehensive analysis reports (executive, hyperparameter, prompt engineering, confidence scoring, beam search, fine-tuning)")
    add_bullet(doc, "All reports converted to PDF and DOCX formats")
    add_bullet(doc, "Tuning experiment framework created with 7 parameter variations")

    add_heading(doc, "Phase 9: Fine-Tuning Infrastructure (Feb 18)", 2)
    add_bullet(doc, "Created run_avspeech_finetune_pipeline.sh — end-to-end fine-tuning pipeline reusing all lib/ modules")
    add_bullet(doc, "Created vsp-llm-avspeech-finetune.yaml — Hydra training config for domain adaptation")
    add_bullet(doc, "Training research notes document covering LoRA rank, length distribution, angle robustness")
    add_bullet(doc, "Dry-run validated: all prerequisites check passed")
    add_bullet(doc, "Expected improvement: 15-25 WER points with full encoder unfreezing")

    doc.add_page_break()


def section_5_containers(doc):
    add_heading(doc, "5. Container Building & Deployment", 1)

    add_heading(doc, "Docker Image Construction", 2)
    add_para(doc, "The container was built from the galaxy_export directory using a 261-line Dockerfile:")
    add_bullet(doc, "Base image: nvidia/cuda:12.8.0-base-ubuntu22.04 (NVIDIA CUDA runtime + Ubuntu)")
    add_bullet(doc, "System packages: ffmpeg, build-essential, Python 3, CUDA development tools")
    add_bullet(doc, "Two isolated virtual environments created with exact package pinning:")
    add_bullet(doc, "pre-process-venv: PyTorch 2.8.0+cu128, Whisper, MediaPipe — 135+ packages, ~8.2GB", level=1)
    add_bullet(doc, "vsp-llm-yoad-venv: PyTorch 2.5.1+cu124, fairseq, transformers, PEFT — 256+ packages, ~6.2GB", level=1)
    add_bullet(doc, "Galaxy export bundle (~65GB) copied into /workspace/: all code + pre-built venvs + model checkpoints")
    add_bullet(doc, "Pre-built offline wheels for air-gapped deployment: spaCy (20), Whisper deps (37), MediaPipe, PyTorch")
    add_bullet(doc, "Final image: ~80-100GB uncompressed; pre-built venvs save 30+ min on startup")

    add_heading(doc, "Container Version History", 2)
    add_styled_table(doc,
        ["Version", "Date", "SHA256 (first 12)", "Changes"],
        [
            ["v1.0.0", "Feb 3, 2026", "5ff5ce59...", "Initial release — modular architecture, core pipeline"],
            ["v1.0.13", "Feb 8, 2026", "e429c21a...", "Bugs 1-13: venv paths, Whisper offline, imports, max_len"],
            ["v1.0.25", "Feb 12, 2026", "8d89ae22...", "Bugs 14-25: CUDA OOM, port access, drag-drop, docker.conf"],
            ["v1.0.31", "Feb 15, 2026", "80e19533...", "Bugs 26-31: upload parser, beam OOM, terminal, server startup"],
            ["v1.0.39", "Feb 17, 2026", "081b1228...", "NVENC corruption fix, bash fd fix, lip crops, docs split"],
        ],
        col_widths=[0.8, 1.2, 1.2, 3.4]
    )

    add_heading(doc, "Bug Fix Summary (37 Bugs)", 2)
    add_styled_table(doc,
        ["Category", "Bug Range", "Key Issues"],
        [
            ["Installation & Setup", "Bugs 1-13", "Venv paths, Whisper offline mode, missing imports, fairseq Cython build, max_len config"],
            ["Deployment & GPU", "Bugs 14-25", "CUDA OOM on 12GB GPU, port access, drag-drop, segment duration, desktop launcher"],
            ["Final Fixes", "Bugs 26-37", "Multipart upload parser (cgi.FieldStorage replaced), beam search OOM, terminal detection, GPU memory management"],
        ],
        col_widths=[1.5, 1.0, 4.0]
    )

    add_heading(doc, "Critical Bug Highlights", 3)
    add_bullet_bold_value(doc, "CUDA OOM (Bug 14): ", "Memory accumulation during decode. Fixed with gc.collect(), torch.cuda.empty_cache(), and dynamic max_new_tokens calculation.")
    add_bullet_bold_value(doc, "NVENC corruption (v1.0.36): ", "GPU encoder silently produced corrupt video (exit code 0). 43% of videos failed face detection. Fixed: post-encode frame extraction test + CPU fallback.")
    add_bullet_bold_value(doc, "Cython auto-build (Bug 29): ", "fairseq Cython extensions not pre-compiled for container's CPU architecture. Fixed: auto-detection and build on first run.")
    add_bullet_bold_value(doc, "Upload parser (Bug 26): ", "Python's cgi.FieldStorage deprecated in 3.11+. Replaced with manual multipart boundary parser.")

    add_heading(doc, "Deployment Infrastructure", 2)
    add_styled_table(doc,
        ["Script", "Lines", "Purpose"],
        [
            ["INSTALL.sh", "459", "Automated installation with timestamped backups, venv validation"],
            ["VERIFY.sh", "~100", "Post-deployment verification (venvs, paths, syntax, packages)"],
            ["vsp-start.sh", "336", "Host launcher: terminal auto-detect, Docker lifecycle, UI startup"],
            ["docker-run.sh", "~30", "Docker run command with GPU, ports, volume mounts"],
            ["build_container.sh", "93", "EC2-to-container sync orchestrator"],
            ["verify_container_sync.sh", "164", "Compares EC2 vs container files, auto-copies"],
            ["check_container_inventory.sh", "136", "In-container validation of directory structure"],
        ],
        col_widths=[2.2, 0.5, 3.8]
    )
    doc.add_page_break()


def section_6_metrics(doc):
    add_heading(doc, "6. Key Metrics & Evaluation Results", 1)

    add_heading(doc, "Evaluation Dataset: english_1k", 2)
    add_para(doc, (
        "860 video segments from ~1,000 YouTube videos across diverse domains (news, vlogs, "
        "interviews, tutorials). Ground truth: Whisper ASR transcriptions. Segments: up to 12 seconds."
    ))

    add_heading(doc, "Performance Distribution", 2)
    add_styled_table(doc,
        ["Quality Tier", "WER Range", "Segments", "% of Total", "Assessment"],
        [
            ["Usable", "0–20%", "98", "11.4%", "Output can be trusted"],
            ["Marginal", "21–40%", "150", "17.4%", "Core meaning preserved"],
            ["Poor", "41–60%", "153", "17.8%", "Meaning distorted"],
            ["Unusable", "61–99%", "282", "32.8%", "Output mostly wrong"],
            ["Hallucinated", "100%+", "177", "20.6%", "Model fabricated text"],
        ],
        col_widths=[1.0, 1.0, 0.8, 0.8, 3.0]
    )

    add_heading(doc, "Aggregate Metrics", 2)
    add_styled_table(doc,
        ["Metric", "Value", "Paper (LRS3)", "Gap"],
        [
            ["Mean WER", "67.0%", "25.4%", "2.6x worse"],
            ["Median WER", "63.8%", "—", "—"],
            ["Corpus WER", "125.5%", "—", "Model outputs more words than exist"],
            ["Named Entity F1", "38.8%", "—", "Poor entity preservation"],
            ["WWER (Weighted)", "61.9%", "—", "Names/numbers penalized 2x"],
            ["Segments WER <= 20%", "11.4%", "~70-80% (estimated)", "~1 in 9 usable"],
            ["Segments WER >= 100%", "20.6%", "~0%", "Catastrophic hallucination"],
        ],
        col_widths=[1.5, 1.0, 1.2, 2.8]
    )

    add_heading(doc, "Domain Gap Analysis", 2)
    add_para(doc, "The 2.6x performance gap is explained by domain mismatch between training data (LRS3: curated TED talks) and evaluation data (YouTube: diverse real-world content):")
    add_styled_table(doc,
        ["Dimension", "LRS3 (Paper)", "YouTube (Real-World)"],
        [
            ["Head pose", "Frontal, stable", "Variable angles, movement"],
            ["Lighting", "Professional, consistent", "Variable, poor conditions"],
            ["Speaker diversity", "TED presenters", "Global creators, all demographics"],
            ["Vocabulary", "Academic, rehearsed", "Casual, slang, brand names"],
            ["Speech style", "Clear, deliberate", "Fast, natural, overlapping"],
            ["Video quality", "Professional cameras", "Phone cameras, webcams"],
        ],
        col_widths=[1.3, 2.5, 2.8]
    )

    add_heading(doc, "Paper's Reported Results (Table 1)", 2)
    add_styled_table(doc,
        ["Configuration", "Data", "WER", "Notes"],
        [
            ["VSP-LLM (frozen encoder)", "433h", "26.7%", "LoRA + bridge only"],
            ["VSP-LLM (fine-tuned)", "433h", "25.4%", "Best result — encoder unfrozen after 18K steps"],
            ["VSP-LLM (low data)", "30h", "29.8%", "Competitive with 433h methods"],
            ["AV-HuBERT (baseline)", "433h", "28.6%", "Prior SOTA"],
        ],
        col_widths=[2.0, 0.8, 0.8, 3.0]
    )
    doc.add_page_break()


def section_7_finetuning(doc):
    add_heading(doc, "7. Fine-Tuning Roadmap", 1)

    add_heading(doc, "Component Impact Analysis", 2)
    add_styled_table(doc,
        ["Component", "Parameters", "Expected WER Improvement", "Risk"],
        [
            ["LoRA adapters only", "~4.2M", "+3-8 points (67% → 59-64%)", "Low — minimal forgetting"],
            ["LoRA + bridge layer", "~8.4M", "+5-10 points (67% → 57-62%)", "Low"],
            ["LoRA + bridge + encoder", "~323M", "+15-25 points (67% → 42-52%)", "Medium — needs careful LR"],
        ],
        col_widths=[1.8, 1.0, 2.5, 1.8]
    )

    add_heading(doc, "Recommended Training Strategy", 2)
    add_para(doc, "Two-phase approach, starting from existing checkpoint_finetune.pt (transfer learning):")
    add_bullet_bold_value(doc, "Phase A (steps 0–5K): ", "Freeze encoder, train LoRA + bridge. Low risk, lets adapters learn new vocabulary/style.")
    add_bullet_bold_value(doc, "Phase B (steps 5K–15K): ", "Unfreeze encoder, train everything. Encoder adapts to diverse visual conditions — this is where the main improvement comes from.")

    add_heading(doc, "Infrastructure Ready", 2)
    add_bullet(doc, "Fine-tuning pipeline script: run_avspeech_finetune_pipeline.sh (validated with dry run)")
    add_bullet(doc, "Training config: vsp-llm-avspeech-finetune.yaml (15K steps, freeze 5K, lr=5e-4)")
    add_bullet(doc, "Multi-GPU support: NUM_GPUS=8 matches paper's setup (~3-5 hour training)")
    add_bullet(doc, "Single-GPU fallback: gradient accumulation (update_freq=8) for Tesla T4")

    add_heading(doc, "Research Recommendations", 2)
    add_bullet_bold_value(doc, "LoRA rank increase: ", "r=16 → r=64 (0.2% → 0.4% of LLM params). Literature suggests r=48-64 is peak for substantial domain shift.")
    add_bullet_bold_value(doc, "Data curation: ", "80% clean + 20% moderate difficulty. Avoid extreme angles where Whisper labels are unreliable.")
    add_bullet_bold_value(doc, "Length strategy: ", "Use all 3-10s clips. RoPE handles length generalization. Do not filter by length.")
    add_bullet_bold_value(doc, "Length penalty tuning: ", "lenpen=0.0 (current) contributes to hallucination. lenpen=0.5-1.0 may reduce corpus WER by 15-30 points.")

    doc.add_page_break()


def section_8_architecture(doc):
    add_heading(doc, "8. Architecture Decisions", 1)

    decisions = [
        ("Segment-level output (no merging)",
         "Early versions merged overlapping segment predictions. Removed because: (a) merge algorithm introduced errors at boundaries, (b) segment-level output is more transparent for evaluation, (c) users can review individual segments."),
        ("12-second segments with 2-second overlap",
         "Paper uses full videos. We segment because: (a) real-world videos are minutes long, (b) model was trained on 3-10s clips, (c) overlap prevents boundary cutoffs. 12s chosen to match model's max_sample_size of 500 frames (20s at 25fps)."),
        ("Modular pipeline (lib/ directory)",
         "823-line monolithic script split into 11 modules. Benefits: each module testable independently, can be reused by fine-tuning pipeline, clear separation of concerns. 37 tests validate all modules."),
        ("Dual virtual environments",
         "Preprocessing (Whisper, MediaPipe) and VSP-LLM (fairseq, transformers) have conflicting PyTorch/CUDA requirements. Two separate venvs avoid dependency hell. ASR venv uses PyTorch 2.8+cu128; VSP-LLM venv uses PyTorch 2.5.1+cu124."),
        ("CPU encoding default (not GPU NVENC)",
         "NVENC silently produced corrupt streams on 43% of videos. CPU encoding (libx264) is 100% reliable. Post-encode validation added as safety net. GPU encoding available as opt-in for trusted hardware."),
        ("Transcription reuse system",
         "Manual transcriptions persist in .transcriptions/ directory across pipeline runs. Whisper automatically skips segments with existing .wrd files. Saves hours of ASR processing on re-runs. Users can manually correct transcriptions and they survive."),
    ]
    for title, desc in decisions:
        add_heading(doc, title, 3)
        add_para(doc, desc)

    doc.add_page_break()


def section_9_executive_summary(doc):
    add_heading(doc, "9. Executive Summary", 1)

    add_para(doc, (
        "This section provides a high-level overview of the Argos project status, achievements, "
        "and next steps for management review."
    ), italic=True, color=C_GRAY)

    add_heading(doc, "Project Status: Production System Deployed, Research Phase Active", 2)
    add_para(doc, (
        "Argos is a working visual speech processing (lip-reading) system based on the VSP-LLM "
        "research paper. The system has been deployed to a client via Docker container on Arch Linux "
        "with NVIDIA GPU support. The production pipeline processes raw videos end-to-end and produces "
        "text transcriptions with quality metrics."
    ))

    add_heading(doc, "Key Accomplishments", 2)

    accomplishments = [
        ("Built production pipeline from research code: ",
         "Integrated 4 separate open-source repositories (VSP-LLM, AV-HuBERT, Auto-AVSR, fairseq) "
         "plus LLaMA-2-7B into a single, end-to-end system. The paper provided only model code with "
         "no pipeline — all orchestration, preprocessing, reporting, and deployment was built from scratch."),
        ("Developed complete web interface: ",
         "Replaced the original 30-line bash script with a 4,564-line web application featuring "
         "drag-and-drop upload, real-time progress tracking, transcription management, and segment review."),
        ("Refactored into modular architecture: ",
         "Transformed 823-line monolithic script into 11 independent modules (52% reduction in main "
         "script). Created 37 automated tests. Enables code reuse across inference and training pipelines."),
        ("Built and deployed 5 production container versions: ",
         "From initial Dockerfile through 5 iterative releases, fixing 37 documented bugs across "
         "installation, deployment, GPU memory, and UI categories. Includes automated installation "
         "and verification scripts."),
        ("Evaluated model on 860 real-world video segments: ",
         "Identified 67% mean WER on YouTube content vs paper's 25.4% on LRS3 benchmark — a 2.6x "
         "domain gap. Only 11.4% of outputs are usable. Root cause: domain mismatch between training "
         "data (curated TED talks) and real-world video."),
        ("Produced 6 comprehensive analysis reports: ",
         "Executive assessment, hyperparameter tuning, prompt engineering, confidence scoring, "
         "beam search aggregation, and fine-tuning analysis — all in PDF and DOCX format."),
        ("Created fine-tuning infrastructure: ",
         "Ready-to-run pipeline script and training configuration for domain adaptation on custom "
         "video datasets. Research indicates 15-25 WER point improvement is achievable."),
        ("Fixed critical NVENC corruption bug: ",
         "Discovered that GPU video encoding silently produced corrupt output on 43% of videos, "
         "causing face detection failures. Implemented post-encode validation and CPU fallback. "
         "Zero failure rate after fix."),
    ]

    for label, value in accomplishments:
        add_bullet_bold_value(doc, label, value)

    add_heading(doc, "Metrics Summary", 2)
    add_styled_table(doc,
        ["Metric", "Value"],
        [
            ["Pipeline architecture", "11 modules, 393-line orchestrator + 1,526 lines of libraries"],
            ["Automated tests", "37 (all passing)"],
            ["Custom code written", "~8,500 lines (pipeline + Python scripts + web UI)"],
            ["Container versions deployed", "5 (Feb 3–17, 2026)"],
            ["Bugs fixed", "37+ across EC2 and container"],
            ["Video segments evaluated", "860 (from ~1,000 YouTube videos)"],
            ["Current mean WER", "67.0% (real-world YouTube)"],
            ["Paper WER (benchmark)", "25.4% (curated LRS3 TED talks)"],
            ["Usable output rate", "11.4% (WER <= 20%)"],
            ["Domain gap factor", "2.6x"],
            ["Expected improvement (fine-tuning)", "15-25 WER points (to 42-52%)"],
            ["Analysis reports produced", "6 (PDF + DOCX)"],
            ["Documentation", "3,000+ lines across 5 documents"],
        ],
        col_widths=[2.5, 4.0]
    )

    add_heading(doc, "Risk Assessment", 2)
    add_styled_table(doc,
        ["Risk", "Severity", "Mitigation"],
        [
            ["Hallucination rate (20.6%)", "High", "Fine-tuning + length penalty tuning expected to reduce"],
            ["Domain gap (2.6x)", "High", "Fine-tuning on AVSpeech-like data is primary mitigation"],
            ["No confidence scoring", "Medium", "Report 4 designs solution; requires code change to enable output_scores"],
            ["GPU memory (T4 16GB)", "Low", "Inference works; training tight but feasible; 8xGPU available"],
            ["Container sync overhead", "Low", "Automated sync tools and documentation in place"],
        ],
        col_widths=[1.8, 0.8, 4.0]
    )

    add_heading(doc, "Next Steps", 2)
    add_bullet_bold_value(doc, "1. Transfer AVSpeech data ", "from external machine to training environment")
    add_bullet_bold_value(doc, "2. Run fine-tuning ", "on 8xGPU AWS instance using prepared pipeline and config (~3-5 hours)")
    add_bullet_bold_value(doc, "3. Evaluate fine-tuned model ", "on english_1k dataset — compare WER to 67% baseline")
    add_bullet_bold_value(doc, "4. Iterate ", "on LoRA rank (16→64), length penalty, and data curation based on results")
    add_bullet_bold_value(doc, "5. Deploy updated model ", "to client container with improved checkpoint")

    doc.add_paragraph()
    add_para(doc,
        f"This document was last updated on {LAST_UPDATED}. "
        "It is a living document that will be updated as the project progresses.",
        italic=True, color=C_GRAY, size=Pt(10)
    )


# ═══════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════

def main():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.27)   # A4
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    add_header_footer(doc)
    create_cover_page(doc)

    section_1_overview(doc)
    section_2_paper_vs_argos(doc)
    section_3_comparison(doc)
    section_4_dev_log(doc)
    section_5_containers(doc)
    section_6_metrics(doc)
    section_7_finetuning(doc)
    section_8_architecture(doc)
    section_9_executive_summary(doc)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_FILE))
    print(f"Generated: {OUTPUT_FILE}")
    print(f"Size: {OUTPUT_FILE.stat().st_size / 1024:.0f} KB")


if __name__ == "__main__":
    main()
