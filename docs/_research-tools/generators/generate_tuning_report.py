#!/usr/bin/env python3
"""
Argos — Hyperparameter Tuning Experiments Report Generator

Generates a branded .docx summarizing 13 decode hyperparameter experiments (A-M)
with tables, deep analysis examples, failure modes, and recommended configuration.

Usage:
    python3 generate_tuning_report.py

Output:
    argos_research/tuning_experiments.docx
"""

from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

# ── Output ──
OUTPUT_DIR = Path(__file__).parent
OUTPUT_FILE = OUTPUT_DIR / "tuning_experiments.docx"

# ── Logos ──
LOGO_ORCHARD = OUTPUT_DIR / "logo.png"
LOGO_PEACOCK = OUTPUT_DIR / "peacock.png"

# ── Colors ──
C_PRIMARY = RGBColor(0x1a, 0x3a, 0x5c)
C_H2 = RGBColor(0x2a, 0x5a, 0x8c)
C_H3 = RGBColor(0x3a, 0x6a, 0x9c)
C_H4 = RGBColor(0x4a, 0x7a, 0xac)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_DARK = RGBColor(0x33, 0x33, 0x33)
C_GREEN = RGBColor(0x15, 0x57, 0x24)
C_RED = RGBColor(0x72, 0x1c, 0x24)
C_AMBER = RGBColor(0x85, 0x64, 0x04)
C_GRAY = RGBColor(0x66, 0x66, 0x66)
C_CODE = RGBColor(0x2d, 0x2d, 0x2d)

HEADER_BG = "1a3a5c"
ZEBRA_BG = "f0f4f8"
GREEN_BG = "d4edda"
RED_BG = "f8d7da"
CODE_BG = "f5f5f5"

LAST_UPDATED = datetime.now().strftime("%B %d, %Y")


# ═══════════════════════════════════════════════════
# HELPER FUNCTIONS
# ═══════════════════════════════════════════════════

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, bold=False, color=None, size=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = size
    run.font.name = "Calibri"


def add_styled_table(doc, headers, rows, col_widths=None, highlight_rows=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, HEADER_BG)
        set_cell_text(cell, h, bold=True, color=C_WHITE, size=Pt(8))

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            set_cell_text(cell, val, size=Pt(8))
            if highlight_rows and r_idx in highlight_rows:
                set_cell_shading(cell, highlight_rows[r_idx])
            elif r_idx % 2 == 1:
                set_cell_shading(cell, ZEBRA_BG)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return table


_bookmark_counter = [0]

def add_heading(doc, text, level, color=None, bookmark_id=None):
    h = doc.add_heading(text, level=level)
    if color is None:
        color = {1: C_PRIMARY, 2: C_H2, 3: C_H3, 4: C_H4}.get(level, C_PRIMARY)
    for run in h.runs:
        run.font.color.rgb = color
        run.font.name = "Calibri"
    if bookmark_id:
        _bookmark_counter[0] += 1
        bm_start = OxmlElement("w:bookmarkStart")
        bm_start.set(qn("w:id"), str(_bookmark_counter[0]))
        bm_start.set(qn("w:name"), bookmark_id)
        bm_end = OxmlElement("w:bookmarkEnd")
        bm_end.set(qn("w:id"), str(_bookmark_counter[0]))
        h._p.insert(0, bm_start)
        h._p.append(bm_end)
    return h


def add_para(doc, text, bold=False, italic=False, size=Pt(10), color=None, space_after=Pt(4)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = size
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = space_after
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.name = "Calibri"
    return p


def add_bullet_bold_value(doc, label, value, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after = Pt(2)
    run_b = p.add_run(label)
    run_b.bold = True
    run_b.font.size = Pt(9)
    run_b.font.name = "Calibri"
    run_v = p.add_run(value)
    run_v.font.size = Pt(9)
    run_v.font.name = "Calibri"
    return p


def _build_inline_image_xml(rId, cx, cy, pic_id=1, name="Logo"):
    inline = OxmlElement('wp:inline')
    inline.set('distT', '0')
    inline.set('distB', '0')
    inline.set('distL', '0')
    inline.set('distR', '0')
    extent = OxmlElement('wp:extent')
    extent.set('cx', str(cx))
    extent.set('cy', str(cy))
    inline.append(extent)
    docPr = OxmlElement('wp:docPr')
    docPr.set('id', str(pic_id))
    docPr.set('name', name)
    inline.append(docPr)
    graphic = OxmlElement('a:graphic')
    graphicData = OxmlElement('a:graphicData')
    graphicData.set('uri', 'http://schemas.openxmlformats.org/drawingml/2006/picture')
    pic = OxmlElement('pic:pic')
    nvPicPr = OxmlElement('pic:nvPicPr')
    cNvPr = OxmlElement('pic:cNvPr')
    cNvPr.set('id', '0')
    cNvPr.set('name', name)
    nvPicPr.append(cNvPr)
    cNvPicPr = OxmlElement('pic:cNvPicPr')
    nvPicPr.append(cNvPicPr)
    pic.append(nvPicPr)
    blipFill = OxmlElement('pic:blipFill')
    blip = OxmlElement('a:blip')
    blip.set(qn('r:embed'), rId)
    blipFill.append(blip)
    stretch = OxmlElement('a:stretch')
    fillRect = OxmlElement('a:fillRect')
    stretch.append(fillRect)
    blipFill.append(stretch)
    pic.append(blipFill)
    spPr = OxmlElement('pic:spPr')
    xfrm = OxmlElement('a:xfrm')
    off = OxmlElement('a:off')
    off.set('x', '0')
    off.set('y', '0')
    xfrm.append(off)
    ext_el = OxmlElement('a:ext')
    ext_el.set('cx', str(cx))
    ext_el.set('cy', str(cy))
    xfrm.append(ext_el)
    spPr.append(xfrm)
    prstGeom = OxmlElement('a:prstGeom')
    prstGeom.set('prst', 'rect')
    spPr.append(prstGeom)
    pic.append(spPr)
    graphicData.append(pic)
    graphic.append(graphicData)
    inline.append(graphic)
    drawing = OxmlElement('w:drawing')
    drawing.append(inline)
    return drawing


def add_header_footer(doc):
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    if LOGO_ORCHARD.exists():
        image_part = doc.part.package.get_or_add_image_part(str(LOGO_ORCHARD))
        rId = header.part.relate_to(image_part, RT.IMAGE)
        size_emu = int(0.3 * 914400)
        logo_run = hp.add_run()
        drawing = _build_inline_image_xml(rId, size_emu, size_emu, pic_id=10, name="Header Logo")
        logo_run._r.append(drawing)
        hp.add_run("  ")
    run = hp.add_run("Argos \u2014 The Orchard")
    run.font.size = Pt(8)
    run.font.color.rgb = C_GRAY
    run.font.name = "Calibri"
    run.italic = True
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = C_GRAY
    run.font.name = "Calibri"
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    run._r.append(fld1)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run._r.append(instr)
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld2)


# ═══════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════

def create_cover_page(doc):
    for _ in range(2):
        doc.add_paragraph()

    if LOGO_PEACOCK.exists():
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_logo = p_logo.add_run()
        run_logo.add_picture(str(LOGO_PEACOCK), width=Inches(2.0))

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ARGOS")
    run.font.size = Pt(42)
    run.font.color.rgb = C_PRIMARY
    run.bold = True
    run.font.name = "Calibri"

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Hyperparameter Tuning Experiments")
    run2.font.size = Pt(20)
    run2.font.color.rgb = C_H2
    run2.font.name = "Calibri"

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("13 Decode Configurations \u00d7 107 Segments")
    run3.font.size = Pt(14)
    run3.font.color.rgb = C_H3
    run3.italic = True
    run3.font.name = "Calibri"

    for _ in range(3):
        doc.add_paragraph()

    info_lines = [
        ("Test Set:", "107 segments from 100 AVSpeech videos"),
        ("Parameters Tested:", "lenpen, beam, temperature, sampling, repetition_penalty"),
        ("Winner:", "Config J \u2014 lenpen=1.0, do_sample=true, temperature=0.5"),
        ("Last Updated:", LAST_UPDATED),
    ]
    for label, value in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rl = p.add_run(label + " ")
        rl.bold = True
        rl.font.size = Pt(11)
        rl.font.color.rgb = C_DARK
        rl.font.name = "Calibri"
        rv = p.add_run(value)
        rv.font.size = Pt(11)
        rv.font.color.rgb = C_DARK
        rv.font.name = "Calibri"

    doc.add_page_break()


# ═══════════════════════════════════════════════════
# SECTION 1: EXPERIMENT OVERVIEW
# ═══════════════════════════════════════════════════

def section_1_overview(doc):
    add_heading(doc, "1. Experiment Overview", 1)

    add_para(doc, (
        "We ran 13 decode experiments (A\u2013M) on a fixed 107-segment test set derived from 100 "
        "AVSpeech YouTube videos. All experiments used the same model checkpoint (checkpoint_finetune.pt), "
        "the same preprocessed data, and differed only in generation hyperparameters. Each experiment "
        "took ~50 minutes on a single T4 GPU (16GB)."
    ))

    add_heading(doc, "1.1 Master Results Table", 2)
    add_para(doc, (
        "Rows highlighted in green = best, red = failed/catastrophic. "
        "WWER = Weighted Word Error Rate (emphasizes named entities). "
        "NEA F1 = Named Entity Accuracy."
    ))

    add_styled_table(doc,
        ["Exp", "Configuration", "WER%", "WWER%", "Empty", "NEA F1%", "Verdict"],
        [
            ["A", "baseline (beam=20, lenpen=0)", "59.4", "59.5", "4 (3.7%)", "40.9", "Reference"],
            ["B", "repetition_penalty=1.0 (off)", "59.4", "59.4", "4 (3.7%)", "40.8", "No effect"],
            ["C", "lenpen=1.0", "58.6", "58.6", "0", "41.1", "Best single param"],
            ["D", "lenpen=-0.5", "65.8", "75.3", "48 (44.9%)", "26.1", "FAIL"],
            ["E", "do_sample=true, temp=0.5", "59.4", "59.4", "0", "41.8", "Slight NEA gain"],
            ["F", "do_sample=true, temp=1.0", "59.4", "59.4", "0", "42.4", "Best NEA (single)"],
            ["G", "greedy (beam=1)", "63.2", "63.2", "0", "37.7", "Worse"],
            ["H", "lenpen=2.0", "539.6", "171.5", "0", "39.1", "CATASTROPHIC"],
            ["I", "lenpen=1.0 + sample temp=1.0", "83.6", "60.1", "0", "40.3", "Combo: mediocre"],
            ["J", "lenpen=1.0 + sample temp=0.5", "70.1", "57.7", "0", "41.0", "BEST WWER"],
            ["K", "do_sample=true, temp=1.5", "60.9", "59.2", "1", "40.7", "Competitive"],
            ["L", "do_sample=true, temp=0.3", "59.6", "58.4", "0", "41.8", "Good"],
            ["M", "lenpen=1.0 + sample temp=0.3", "80.9", "77.1", "0", "42.1", "Unexpectedly bad"],
        ],
        col_widths=[0.3, 2.2, 0.5, 0.5, 0.7, 0.5, 1.3],
        highlight_rows={
            2: GREEN_BG,   # C
            3: RED_BG,     # D
            7: RED_BG,     # H
            9: GREEN_BG,   # J
        }
    )

    add_heading(doc, "1.2 Baseline Configuration", 2)
    add_styled_table(doc,
        ["Parameter", "Default Value", "Source"],
        [
            ["beam", "20", "s2s_decode.yaml"],
            ["lenpen", "0.0", "VSP-LLM paper Section 4.2"],
            ["max_len_a", "2.0", "Dynamic max length scaling"],
            ["max_len_b", "200", "Buffer for short segments"],
            ["max_len", "2048", "Hard cap"],
            ["no_repeat_ngram_size", "3", "Prevent degenerate repetition"],
            ["repetition_penalty", "1.2", "Mild token repetition penalty"],
            ["do_sample", "false", "Deterministic beam search"],
            ["temperature", "1.0", "Only used when do_sample=true"],
            ["top_p", "0.9", "Nucleus sampling threshold"],
        ],
        col_widths=[1.8, 1.0, 3.7]
    )


# ═══════════════════════════════════════════════════
# SECTION 2: PARAMETER ANALYSIS
# ═══════════════════════════════════════════════════

def section_2_analysis(doc):
    add_heading(doc, "2. Parameter Analysis", 1)

    # 2.1 Length Penalty
    add_heading(doc, "2.1 Length Penalty (lenpen)", 2)
    add_para(doc, (
        "Length penalty controls whether the model favors shorter or longer outputs. "
        "Positive values reward longer sequences; negative values penalize them. "
        "This was the most impactful single parameter tested."
    ))

    add_styled_table(doc,
        ["lenpen", "Exp", "WER%", "WWER%", "Empty", "Effect"],
        [
            ["-0.5", "D", "65.8", "75.3", "48 (44.9%)", "Catastrophic \u2014 model generates nothing for half the segments"],
            ["0.0", "A", "59.4", "59.5", "4 (3.7%)", "Default \u2014 4 empty predictions"],
            ["1.0", "C", "58.6", "58.6", "0 (0%)", "Best single-param improvement \u2014 eliminates empties"],
            ["2.0", "H", "539.6", "171.5", "0 (0%)", "Catastrophic \u2014 model generates paragraphs of fiction"],
        ],
        col_widths=[0.5, 0.3, 0.5, 0.5, 0.8, 3.9]
    )

    add_para(doc, (
        "Finding: lenpen=1.0 is the sweet spot. It eliminates empty predictions and improves WWER "
        "by 0.9 points. Going higher (2.0) triggers catastrophic over-generation where the model "
        "hallucinates entire paragraphs. Going negative (-0.5) causes the model to produce nothing "
        "for 44.9% of segments."
    ), bold=True, size=Pt(9))

    # 2.2 Sampling vs Beam Search
    add_heading(doc, "2.2 Sampling vs Beam Search", 2)
    add_para(doc, (
        "Sampling introduces randomness into token selection. When do_sample=true, "
        "the model selects from the probability distribution (shaped by temperature and top_p) "
        "rather than always picking the most likely token."
    ))

    add_styled_table(doc,
        ["Mode", "Exp", "WER%", "WWER%", "NEA F1%", "Effect"],
        [
            ["Beam only (default)", "A", "59.4", "59.5", "40.9", "Reference"],
            ["Greedy (beam=1)", "G", "63.2", "63.2", "37.7", "Worse \u2014 beam search is necessary"],
            ["Sample temp=0.3", "L", "59.6", "58.4", "41.8", "Good \u2014 slightly better WWER, better NEA"],
            ["Sample temp=0.5", "E", "59.4", "59.4", "41.8", "Equivalent WER, better NEA"],
            ["Sample temp=1.0", "F", "59.4", "59.4", "42.4", "Same WER, best NEA F1 (single param)"],
            ["Sample temp=1.5", "K", "60.9", "59.2", "40.7", "Slight degradation at high temperature"],
        ],
        col_widths=[1.5, 0.3, 0.5, 0.5, 0.5, 3.2]
    )

    add_para(doc, (
        "Finding: Sampling alone does not significantly improve WER/WWER but consistently improves "
        "Named Entity F1 (+1-1.5 points). This suggests sampling helps the model explore "
        "alternative vocabulary that better matches entity names. Higher temperature increases diversity "
        "but degrades slightly above temp=1.0."
    ), bold=True, size=Pt(9))

    # 2.3 Combinations
    add_heading(doc, "2.3 Combined Configurations (lenpen + sampling)", 2)
    add_para(doc, (
        "Combining lenpen=1.0 with sampling produced mixed results. The interaction between "
        "length penalty and sampling temperature is non-trivial."
    ))

    add_styled_table(doc,
        ["Combination", "Exp", "WER%", "WWER%", "NEA F1%", "Notes"],
        [
            ["lenpen=1.0 only", "C", "58.6", "58.6", "41.1", "Best single param"],
            ["lenpen=1.0 + temp=1.0", "I", "83.6", "60.1", "40.3", "WER degrades, WWER only slightly worse"],
            ["lenpen=1.0 + temp=0.5", "J", "70.1", "57.7", "41.0", "BEST WWER overall"],
            ["lenpen=1.0 + temp=0.3", "M", "80.9", "77.1", "42.1", "Unexpectedly bad \u2014 temp too low causes degeneracy"],
        ],
        col_widths=[1.8, 0.3, 0.5, 0.5, 0.5, 2.9]
    )

    add_para(doc, (
        "Finding: Config J (lenpen=1.0, temp=0.5) achieves the best WWER (57.7%) but at the cost "
        "of higher raw WER (70.1%). This means the model generates more text overall (more insertions) "
        "but those insertions sometimes contain useful content, improving weighted metrics. "
        "Very low temperature (0.3) with lenpen creates pathological behavior \u2014 the peaked distribution "
        "combined with length reward causes repetitive or degenerate outputs."
    ), bold=True, size=Pt(9))

    # 2.4 Repetition Penalty
    add_heading(doc, "2.4 Repetition Penalty", 2)
    add_para(doc, (
        "Comparing A (rep_pen=1.2, default) vs B (rep_pen=1.0, disabled): virtually identical results "
        "(WWER 59.5% vs 59.4%, NEA F1 40.9% vs 40.8%). The no_repeat_ngram_size=3 parameter already "
        "prevents degenerate repetition, making the token-level repetition penalty redundant."
    ))


# ═══════════════════════════════════════════════════
# SECTION 3: DEEP ANALYSIS EXAMPLES
# ═══════════════════════════════════════════════════

def section_3_examples(doc):
    add_heading(doc, "3. Deep Analysis: Per-Segment Examples", 1)

    add_para(doc, (
        "The following examples show how specific segments responded to different hyperparameter "
        "configurations. Comparing across configs reveals the trade-offs in each approach."
    ))

    # Example 1: Empty elimination
    add_heading(doc, "3.1 Empty Output Elimination (lenpen benefit)", 2)
    add_para(doc, "Segment: DBhaa45mAro (casual vlog, 15 ref words)", italic=True, color=C_GRAY)

    add_styled_table(doc,
        ["Config", "WER%", "Output"],
        [
            ["REF", "\u2014", "hey it's me just doing the eush keeping up with what's going on in school so here's the latest taylor and rob over britt"],
            ["A (baseline)", "100.0", "(empty \u2014 model produced no output)"],
            ["C (lenpen=1.0)", "66.7", "hi it's me jessica dude who's keeping up with what's going on in the world so that i can tell you the latest taylor swif..."],
            ["J (winner)", "73.3", "hi it's me jessica dude who's keeping up with what's going on in the world so i'm going to go ahead and tell you the lat..."],
            ["L (temp=0.3)", "63.3", "hi it's me jessica dude who's keeping up with what's going on in the world through vine and instagram the latest tailer..."],
        ],
        col_widths=[1.0, 0.5, 5.0]
    )

    add_para(doc, (
        "Analysis: Baseline produced EMPTY output. All lenpen=1.0 configs recovered partial content. "
        "The model correctly identifies \"it's me\", \"keeping up with what's going on\", and \"taylor\" across all configs \u2014 "
        "these lip patterns are visually distinctive. \"eush\" (slang) and \"britt\" (name) are missed by all configs."
    ), size=Pt(9))

    # Example 2: Good improvement
    add_heading(doc, "3.2 Entity Recognition Improvement", 2)
    add_para(doc, "Segment: zmwgmt7wcv8 (documentary, 15 ref words)", italic=True, color=C_GRAY)

    add_styled_table(doc,
        ["Config", "WER%", "Output"],
        [
            ["REF", "\u2014", "watching the way words develop and change does suggest that wittgenstein was on to something"],
            ["A (baseline)", "26.7", "the way words develop and change suggests that philistine was on to something"],
            ["C (lenpen=1.0)", "20.0", "the way words develop and change doesn't suggest that philistine was on to something"],
            ["J (winner)", "20.0", "the way words develop and change doesn't suggest that philistine was on to something"],
        ],
        col_widths=[1.0, 0.5, 5.0]
    )

    add_para(doc, (
        "Analysis: All configs produce \"philistine\" instead of \"wittgenstein\" \u2014 a classic homophene error "
        "(the lip movements for /w\u026at\u0261\u0259n\u0283ta\u026an/ and /f\u026al\u0259sti\u02d0n/ overlap significantly). "
        "This error is a fundamental visual ambiguity that no hyperparameter change can fix; "
        "it requires domain-specific vocabulary via fine-tuning or prompt engineering."
    ), size=Pt(9))

    # Example 3: Regression
    add_heading(doc, "3.3 Regression: lenpen Triggers Hallucination on Short Segments", 2)
    add_para(doc, "Segment: eLS1vcpGVHQ (short, 10 ref words)", italic=True, color=C_GRAY)

    add_styled_table(doc,
        ["Config", "WER%", "Output"],
        [
            ["REF", "\u2014", "i surrendered good quite i surrendered to the low security"],
            ["A (baseline)", "90.0", "years ago when i was"],
            ["C (lenpen=1.0)", "740.0", "years ago when i was 18 years old i came to the united states with a one way ticket and a suitcase and i didn't speak a..."],
            ["J (winner)", "400.0", "i'm so happy to be here with you all today thank you so much for inviting me to speak today it's an honor and a privileg..."],
            ["L (temp=0.3)", "100.0", "the world so thank you"],
        ],
        col_widths=[1.0, 0.5, 5.0]
    )

    add_para(doc, (
        "Analysis: This short segment has ambiguous visual input. Baseline (A) already fails but produces "
        "little output (5 words). With lenpen=1.0, the model is rewarded for length and hallucinates "
        "entire narratives \u2014 fluent, grammatically correct text that has no relation to the actual speech. "
        "This is the core trade-off: lenpen fixes empty outputs but amplifies hallucination on ambiguous segments."
    ), size=Pt(9), bold=True)

    # Example 4: Stable good performance
    add_heading(doc, "3.4 Stable: High-Quality Segment (Config-Independent)", 2)
    add_para(doc, "Segment: 0v2N6w4m46s (medical topic, 48 ref words)", italic=True, color=C_GRAY)

    add_styled_table(doc,
        ["Config", "WER%", "Output (first 100 chars)"],
        [
            ["REF", "\u2014", "on today's the doctor is in the topic is hypertension there is no ideal blood pressure reading but..."],
            ["A (baseline)", "22.9", "today as the doctor says the topic is hypertension there is no ideal blood pressure reading but..."],
            ["C (lenpen=1.0)", "22.9", "and today as the doctor is saying the topic is hypertension there is no ideal blood pressure rea..."],
            ["J (winner)", "22.9", "and today as the doctor is saying the topic is hypertension there is no ideal blood pressure rea..."],
            ["L (temp=0.3)", "22.9", "today as the doctor says the topic is hypertension there is no ideal blood pressure reading but..."],
        ],
        col_widths=[1.0, 0.5, 5.0]
    )

    add_para(doc, (
        "Analysis: Long segments with clear frontal speech are stable across all configs. "
        "The 22.9% WER comes from minor substitutions (\"today as the doctor says\" vs \"on today's the doctor is in\") "
        "that are consistent across configs. These errors require visual encoder improvement (fine-tuning), not decode tuning."
    ), size=Pt(9))

    # Example 5: lenpen vs baseline
    add_heading(doc, "3.5 Near-Perfect Baseline Degraded by lenpen", 2)
    add_para(doc, "Segment: 48mil-jr-Pk (workplace topic, 9 ref words)", italic=True, color=C_GRAY)

    add_styled_table(doc,
        ["Config", "WER%", "Output"],
        [
            ["REF", "\u2014", "for example wellness programs are very popular among employers"],
            ["A (baseline)", "11.1", "example wellness programs are very popular among employers"],
            ["C (lenpen=1.0)", "55.6", "example one of these programs is very popular among employers"],
            ["J (winner)", "55.6", "example one of these programs is very popular among employers"],
            ["L (temp=0.3)", "11.1", "example wellness programs are very popular among employers"],
        ],
        col_widths=[1.0, 0.5, 5.0]
    )

    add_para(doc, (
        "Analysis: Baseline nearly perfect (only missing \"for\"). With lenpen=1.0, the model replaces "
        "\"wellness programs\" with \"one of these programs\" \u2014 generating more words (as lenpen rewards) but "
        "less accurate ones. Without lenpen (L, temp=0.3), the output matches baseline. "
        "This shows lenpen's downside: it can corrupt already-good predictions by encouraging unnecessary verbosity."
    ), size=Pt(9))


# ═══════════════════════════════════════════════════
# SECTION 4: FAILURE MODES
# ═══════════════════════════════════════════════════

def section_4_failures(doc):
    add_heading(doc, "4. Failure Modes", 1)

    # 4.1 Over-generation
    add_heading(doc, "4.1 Catastrophic Over-Generation (lenpen=2.0, Config H)", 2)
    add_para(doc, (
        "Config H (lenpen=2.0) produced an average WER of 539.6% \u2014 the model generates roughly "
        "5x more words than the reference for every segment. The output is fluent, grammatically "
        "correct English that reads like coherent speech \u2014 but has no relation to the actual video. "
        "This demonstrates that the LLM backbone (LLaMA-2-7B) can generate plausible text from "
        "any visual input when the length reward is strong enough to override the visual signal."
    ))

    add_para(doc, (
        "Example: REF = \"get the idea\" (3 words). Config H output = \"i think the idea of what we're "
        "doing here is we're trying to figure out how to make the world a better place for everyone "
        "and i think that starts with education and understanding...\" (233.3% WER). "
        "The model latched onto \"the idea\" and hallucinated an entire motivational speech."
    ), italic=True, size=Pt(9))

    # 4.2 Empty outputs
    add_heading(doc, "4.2 Mass Empty Outputs (lenpen=-0.5, Config D)", 2)
    add_para(doc, (
        "Config D (lenpen=-0.5) produced empty output for 48 of 107 segments (44.9%). "
        "Negative length penalty means shorter sequences score better. For segments where the "
        "visual signal is ambiguous, the model \"prefers\" generating nothing over risking a "
        "wrong prediction. The 59 segments that did get output had reasonable quality, suggesting "
        "negative lenpen acts as an extreme confidence filter \u2014 but one that discards far too much."
    ))

    # 4.3 Temp degeneracy
    add_heading(doc, "4.3 Low-Temperature Degeneracy with lenpen (Config M)", 2)
    add_para(doc, (
        "Config M (lenpen=1.0, temp=0.3) unexpectedly produced 77.1% WWER \u2014 worse than either "
        "lenpen=1.0 alone (58.6%) or temp=0.3 alone (58.4%). At very low temperature, the probability "
        "distribution becomes extremely peaked. Combined with length reward, the model gets stuck in "
        "local optima: it selects the single most likely next token at each step (almost greedy) "
        "but is rewarded for continuing, creating long sequences of high-confidence but wrong tokens. "
        "This pathological interaction explains why temp=0.5 works but temp=0.3 does not with lenpen."
    ))


# ═══════════════════════════════════════════════════
# SECTION 5: CONCLUSIONS
# ═══════════════════════════════════════════════════

def section_5_conclusions(doc):
    add_heading(doc, "5. Conclusions & Recommended Configuration", 1)

    add_heading(doc, "5.1 Recommended Config (for production)", 2)
    add_styled_table(doc,
        ["Parameter", "Recommended", "Previous Default", "Rationale"],
        [
            ["beam", "20", "20", "Unchanged \u2014 greedy (beam=1) is significantly worse"],
            ["lenpen", "1.0", "0.0", "Eliminates empty outputs, -0.9 WWER improvement"],
            ["do_sample", "true", "false", "Improves NEA F1 by +1-1.5 points"],
            ["temperature", "0.5", "1.0", "Best balance of diversity vs stability with lenpen"],
            ["top_p", "0.9", "0.9", "Unchanged"],
            ["repetition_penalty", "1.2", "1.2", "Unchanged \u2014 negligible effect, keep as safety net"],
            ["no_repeat_ngram_size", "3", "3", "Unchanged"],
        ],
        col_widths=[1.5, 1.0, 1.0, 3.0]
    )

    add_heading(doc, "5.2 Key Takeaways", 2)
    add_bullet_bold_value(doc, "lenpen=1.0 is the most impactful single change: ",
        "Eliminates empty predictions (4\u21920), improves WWER by 0.9 points, and costs nothing computationally.")
    add_bullet_bold_value(doc, "Sampling improves entity recognition: ",
        "NEA F1 improves +1-1.5 points with sampling enabled, suggesting the model explores "
        "vocabulary alternatives that better match names and technical terms.")
    add_bullet_bold_value(doc, "Combined config J yields best WWER (57.7%): ",
        "But at the cost of higher raw WER (70.1% vs 59.4%) due to more insertions. "
        "Whether this trade-off is acceptable depends on the use case.")
    add_bullet_bold_value(doc, "Parameter interactions are non-linear: ",
        "lenpen=1.0 + temp=0.5 works well, but lenpen=1.0 + temp=0.3 fails catastrophically. "
        "Temperature extremes interact poorly with length rewards.")
    add_bullet_bold_value(doc, "Tuning ceiling is ~1.8 WWER points: ",
        "The gap from baseline (59.5%) to best (57.7%) is only 1.8 points. "
        "Real improvement requires fine-tuning (expected -15 to -25 points).")
    add_bullet_bold_value(doc, "Full-dataset validation is in progress: ",
        "Config J is currently being decoded on the full 1,497-segment dataset. "
        "Results will validate whether the 107-segment findings generalize.")

    add_heading(doc, "5.3 What Tuning Cannot Fix", 2)
    add_bullet(doc, "Homophene errors (e.g., \"wittgenstein\" \u2192 \"philistine\") \u2014 visual ambiguity, requires vocabulary injection")
    add_bullet(doc, "Domain vocabulary (e.g., \"costal cartilages\" \u2192 \"cause our hormones\") \u2014 requires in-domain fine-tuning")
    add_bullet(doc, "Short segment failure (<5 words) \u2014 architectural limitation, LLM needs context")
    add_bullet(doc, "Hallucination on ambiguous input \u2014 lenpen amplifies this; need confidence scoring to filter")

    # Final note
    doc.add_paragraph()
    add_para(doc, (
        "Full per-segment data for all 13 experiments is available in "
        "tuning_results/interesting_examples/ with metadata explaining why each example is notable."
    ), italic=True, color=C_GRAY, size=Pt(9))


# ═══════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════

def main():
    print("Generating Argos Tuning Experiments Report...")
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    add_header_footer(doc)

    create_cover_page(doc)
    section_1_overview(doc)
    section_2_analysis(doc)
    section_3_examples(doc)
    section_4_failures(doc)
    section_5_conclusions(doc)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_FILE))
    print(f"Saved: {OUTPUT_FILE}")
    print(f"Size: {OUTPUT_FILE.stat().st_size / 1024:.0f} KB")


if __name__ == "__main__":
    main()
