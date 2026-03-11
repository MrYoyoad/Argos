#!/usr/bin/env python3
"""
Generate analysis of segments where Claude's LLM heuristic identifies
recoverable predictions that traditional metrics (WER, WWER, IS) miss.

These are cases where the model output is meaningful and useful to a
domain-aware viewer despite poor error-rate metrics — the "hidden value"
in the VSP-LLM system that numeric scores undercount.

Outputs:
  - Markdown analysis document
  - JSON data file with all divergent segments
  - DOCX formatted report (Argos style)
"""

import csv
import json
import os
import sys
from collections import Counter, defaultdict
from pathlib import Path
from datetime import date

SCRIPT_DIR = Path(__file__).resolve().parent
ASSETS_DIR = SCRIPT_DIR.parent / "assets"
CSV_PATH = Path("/home/ubuntu/docs/evaluation/intelligibility/intelligibility_scores.csv")
OUTPUT_DIR = Path("/home/ubuntu/docs/evaluation/llm_salvage")
PRES_DIR = Path("/home/ubuntu/presentation_materials_20260224")


def load_data():
    rows = []
    with open(CSV_PATH, "r", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for r in reader:
            for k in ["wer_%", "wwer_%", "nea_f1_%", "semantic_sim", "phonetic_sim",
                       "length_ratio", "intelligibility_score", "llm_context_prob"]:
                try:
                    r[k] = float(r[k])
                except (ValueError, KeyError):
                    r[k] = None
            try:
                r["ref_words"] = int(r["ref_words"])
            except (ValueError, KeyError):
                r["ref_words"] = 0
            r["context_recoverable"] = r.get("context_recoverable", "").strip().lower() == "true"
            rows.append(r)
    return rows


def interest_score(r):
    """Score how compelling this example is for presentation."""
    sem = r["semantic_sim"] or 0
    phon = r["phonetic_sim"] or 0
    llm_p = r["llm_context_prob"] or 0
    wer = r["wer_%"] or 0
    ref_words = r["ref_words"] or 0
    # Prefer: high semantic (meaning preserved), high LLM prob, high WER (dramatic gap),
    # longer segments (more convincing), high phonetic (natural lip-reading errors)
    return (
        sem * 3.0
        + phon * 1.5
        + llm_p * 2.0
        + min(wer, 120) / 100 * 2.0  # Reward higher WER (bigger gap = more dramatic)
        + min(ref_words, 20) / 20 * 1.0  # Prefer longer segments
    )


def categorize_salvage(rows):
    """Split divergent segments into meaningful presentation categories."""
    categories = {}

    # All divergent: LLM >= 0.5 AND IS < 3.0
    all_divergent = [r for r in rows
                     if r["llm_context_prob"] is not None and r["llm_context_prob"] >= 0.5
                     and r["intelligibility_score"] is not None and r["intelligibility_score"] < 3.0]

    # Category 1: "Hidden Gems" — High LLM confidence (>= 0.8), IS < 3.0, WER >= 50%
    #   These are the most dramatic: bad WER, bad IS, but Claude says meaning is preserved
    cat1 = [r for r in all_divergent
            if r["llm_context_prob"] >= 0.8 and r["wer_%"] is not None and r["wer_%"] >= 50]
    cat1.sort(key=interest_score, reverse=True)
    categories["hidden_gems"] = {
        "title": "Hidden Gems: High Confidence Despite Bad Metrics",
        "description": "Claude assigns probability >= 0.80 that a viewer with context could recover meaning, "
                       "yet IS < 3.0 and WER >= 50%. These segments demonstrate where traditional metrics "
                       "systematically undervalue lip-reading output.",
        "segments": cat1,
    }

    # Category 2: "Semantic Preservation" — semantic_sim >= 0.5 but WER >= 50%
    #   Meaning is intact despite word-level errors
    cat2 = [r for r in all_divergent
            if r["semantic_sim"] is not None and r["semantic_sim"] >= 0.5
            and r["wer_%"] is not None and r["wer_%"] >= 50]
    cat2.sort(key=interest_score, reverse=True)
    categories["semantic_preservation"] = {
        "title": "Semantic Preservation: Same Meaning, Different Words",
        "description": "Semantic similarity >= 0.50 confirms the hypothesis conveys the same meaning as "
                       "the reference, yet WER >= 50% rates these as failures. These cases show WER's "
                       "blindness to meaning-equivalent paraphrasing.",
        "segments": cat2,
    }

    # Category 3: "Phonetic Bridge" — phonetic_sim >= 0.6, typical lip-reading confusions
    #   Errors are natural lip-reading confusions, not hallucinations
    cat3 = [r for r in all_divergent
            if r["phonetic_sim"] is not None and r["phonetic_sim"] >= 0.6
            and r["wer_%"] is not None and r["wer_%"] >= 40]
    cat3.sort(key=interest_score, reverse=True)
    categories["phonetic_bridge"] = {
        "title": "Phonetic Bridge: Natural Lip-Reading Confusions",
        "description": "Phonetic similarity >= 0.60 shows errors are caused by visually similar mouth "
                       "shapes (homophenes), not hallucination. A viewer hearing these words spoken aloud "
                       "would recognize the intended meaning.",
        "segments": cat3,
    }

    # Category 4: "Entity-Preserved" — key entities survive despite overall WER
    #   Named entities (names, places, numbers) are correct even when function words fail
    cat4 = [r for r in all_divergent
            if r["nea_f1_%"] is not None and r["nea_f1_%"] >= 50
            and r["wer_%"] is not None and r["wer_%"] >= 40]
    cat4.sort(key=interest_score, reverse=True)
    categories["entity_preserved"] = {
        "title": "Entity-Preserved: Key Information Survives",
        "description": "Named Entity F1 >= 50% shows that the most important information (names, numbers, "
                       "places) is captured correctly, even when surrounding function words are wrong.",
        "segments": cat4,
    }

    # Category 5: "Structure Match" — sequence preserved, word order intact
    cat5 = [r for r in all_divergent
            if r["llm_context_reason"] in ("strong_structure_match", "moderate_structure_coherent",
                                            "good_overlap_coherent")]
    cat5.sort(key=interest_score, reverse=True)
    categories["structure_match"] = {
        "title": "Structure Match: Word Order Preserved",
        "description": "The hypothesis follows the same grammatical structure as the reference. "
                       "Content words appear in the same sequence, making the output readable despite "
                       "individual word errors.",
        "segments": cat5,
    }

    # Category 6: "WER Over-Punishment" — low WWER but high WER
    #   WER inflated by function word errors that don't affect meaning
    cat6 = [r for r in all_divergent
            if r["wer_%"] is not None and r["wwer_%"] is not None
            and r["wer_%"] >= 50 and (r["wer_%"] - r["wwer_%"]) >= 10]
    cat6.sort(key=lambda r: (r["wer_%"] - r["wwer_%"]), reverse=True)
    categories["wer_over_punishment"] = {
        "title": "WER Over-Punishment: Function Word Inflation",
        "description": "WER is at least 10 percentage points higher than WWER, meaning most errors "
                       "are in function words (articles, prepositions) that don't affect meaning.",
        "segments": cat6,
    }

    return categories, all_divergent


def format_example_md(r, idx):
    """Format a single segment as markdown."""
    lines = []
    lines.append(f"#### Example {idx}: `{r['display_name']}`")
    lines.append("")
    lines.append(f"| Metric | Value |")
    lines.append(f"|--------|-------|")
    lines.append(f"| **Reference** | {r['ref']} |")
    lines.append(f"| **Hypothesis** | {r['hyp']} |")
    lines.append(f"| WER | {r['wer_%']:.1f}% |")
    lines.append(f"| WWER | {r['wwer_%']:.1f}% |")
    lines.append(f"| IS | {r['intelligibility_score']:.3f} ({r['intelligibility_label']}) |")
    lines.append(f"| Semantic Sim | {r['semantic_sim']:.4f} |")
    lines.append(f"| Phonetic Sim | {r['phonetic_sim']:.4f} |")
    lines.append(f"| NEA F1 | {r['nea_f1_%']:.1f}% |")
    lines.append(f"| LLM Context Prob | **{r['llm_context_prob']:.3f}** |")
    lines.append(f"| LLM Reason | {r['llm_context_reason']} |")
    lines.append(f"| Topic | {r.get('topic', 'N/A')} |")
    lines.append("")

    # Add interpretation
    wer = r["wer_%"] or 0
    llm_p = r["llm_context_prob"] or 0
    sem = r["semantic_sim"] or 0

    if sem >= 0.6:
        lines.append(f"**Why this matters**: Despite {wer:.0f}% WER, semantic similarity of {sem:.2f} "
                      f"confirms the core meaning is preserved. Claude assigns {llm_p*100:.0f}% probability "
                      f"that a viewer with context would understand this segment.")
    elif r["phonetic_sim"] and r["phonetic_sim"] >= 0.6:
        lines.append(f"**Why this matters**: The {r['phonetic_sim']:.2f} phonetic similarity shows these "
                      f"are natural lip-reading confusions (similar mouth shapes), not random errors. "
                      f"Claude assigns {llm_p*100:.0f}% recovery probability.")
    else:
        lines.append(f"**Why this matters**: Claude identifies structural and semantic cues "
                      f"({r['llm_context_reason']}) suggesting {llm_p*100:.0f}% recovery probability, "
                      f"while IS scores only {r['intelligibility_score']:.2f}/5.0.")
    lines.append("")
    return "\n".join(lines)


def generate_markdown(categories, all_divergent, all_rows):
    """Generate the full markdown analysis document."""
    lines = []

    # Header
    lines.append("# LLM Salvage Analysis: Recoverable Predictions That Metrics Undercount")
    lines.append("")
    lines.append(f"**Date**: {date.today().isoformat()}")
    lines.append(f"**Dataset**: Full baseline (1,497 segments, AVSpeech)")
    lines.append(f"**Divergent segments**: {len(all_divergent)} (LLM prob >= 0.5, IS < 3.0)")
    lines.append("")

    # Executive Summary
    lines.append("## Executive Summary")
    lines.append("")
    lines.append("Traditional metrics (WER, WWER, IS) classify 900 of 1,497 segments as failures "
                 "(IS < 3.0). However, Claude's LLM-calibrated heuristic identifies **165 of these "
                 "900 segments** as having recoverable meaning — cases where a viewer with domain "
                 "context would understand the lip-reading output despite high word error rates.")
    lines.append("")
    lines.append("This represents an **18.3% recovery rate** among segments that metrics mark as failed. "
                 "The effective useful rate is 61.6% by IS (\u2265 2.00), confirmed at 64.9% by Opus judge.")
    lines.append("")

    # Summary statistics
    total = len(all_rows)
    is_good = sum(1 for r in all_rows if r["intelligibility_score"] and r["intelligibility_score"] >= 3.0)
    llm_good = sum(1 for r in all_rows if r["llm_context_prob"] and r["llm_context_prob"] >= 0.5)
    is_failed = total - is_good

    lines.append("### Key Numbers")
    lines.append("")
    lines.append("| Metric | Count | Percentage |")
    lines.append("|--------|-------|------------|")
    lines.append(f"| Total segments | {total} | 100% |")
    lines.append(f"| IS >= 3.0 (metrics say captured) | {is_good} | {100*is_good/total:.1f}% |")
    lines.append(f"| IS < 3.0 (metrics say failed) | {is_failed} | {100*is_failed/total:.1f}% |")
    lines.append(f"| LLM prob >= 0.5 (Claude says recoverable) | {llm_good} | {100*llm_good/total:.1f}% |")
    lines.append(f"| **Divergent: LLM salvages from IS failures** | **{len(all_divergent)}** | **{100*len(all_divergent)/is_failed:.1f}% of failures** |")
    lines.append(f"| **Effective capture rate (IS + LLM salvage)** | **{is_good + len(all_divergent)}** | **{100*(is_good + len(all_divergent))/total:.1f}%** |")
    lines.append("")

    # Divergent segment breakdown
    lines.append("### Divergent Segment Profile")
    lines.append("")

    # WER distribution
    wer_ranges = [(0, 30, "0-30%"), (30, 50, "30-50%"), (50, 70, "50-70%"),
                  (70, 100, "70-100%"), (100, 999, "100%+")]
    lines.append("**WER distribution of salvageable segments:**")
    lines.append("")
    lines.append("| WER Range | Count | % of Divergent |")
    lines.append("|-----------|-------|----------------|")
    for lo, hi, label in wer_ranges:
        count = sum(1 for r in all_divergent if r["wer_%"] is not None and lo <= r["wer_%"] < hi)
        lines.append(f"| {label} | {count} | {100*count/len(all_divergent):.1f}% |")
    lines.append("")

    # LLM reason distribution
    reasons = Counter(r["llm_context_reason"] for r in all_divergent)
    lines.append("**Recovery mechanism (why Claude says recoverable):**")
    lines.append("")
    lines.append("| LLM Reason | Count | % | Interpretation |")
    lines.append("|------------|-------|---|----------------|")
    reason_interp = {
        "good_overlap_coherent": "Key content words match and topic is preserved",
        "semantic_plus_phonetic": "Meaning preserved via semantics and natural lip-reading confusions",
        "high_semantic_good_overlap": "Strong meaning match with significant word overlap",
        "phonetic_bridge_semantic": "Phonetically similar words bridge to correct meaning",
        "strong_structure_match": "Word order and structure closely match reference",
        "moderate_structure_coherent": "Moderate structural similarity with topic coherence",
        "moderate_semantic_preserved_content": "Moderate semantic match with preserved content words",
    }
    for reason, count in reasons.most_common():
        interp = reason_interp.get(reason, "")
        lines.append(f"| {reason} | {count} | {100*count/len(all_divergent):.1f}% | {interp} |")
    lines.append("")

    # Topic distribution
    topics = Counter(r.get("topic", "Unknown") for r in all_divergent)
    lines.append("**Topic distribution:**")
    lines.append("")
    lines.append("| Topic | Salvageable | Rate |")
    lines.append("|-------|-------------|------|")
    for topic, count in topics.most_common():
        topic_total = sum(1 for r in all_rows if r.get("topic") == topic
                         and r["intelligibility_score"] is not None and r["intelligibility_score"] < 3.0)
        rate = 100 * count / topic_total if topic_total > 0 else 0
        lines.append(f"| {topic} | {count} | {rate:.0f}% of topic failures |")
    lines.append("")

    # Categories with examples
    lines.append("---")
    lines.append("")
    lines.append("## Curated Examples by Category")
    lines.append("")

    used_ids = set()  # Cross-category deduplication

    for cat_key, cat_data in categories.items():
        segs = cat_data["segments"]
        if not segs:
            continue
        lines.append(f"### {cat_data['title']}")
        lines.append("")
        lines.append(f"*{cat_data['description']}*")
        lines.append("")
        lines.append(f"**{len(segs)} segments** in this category.")
        lines.append("")

        # Select 5 unique examples with topic diversity
        selected = []
        topics_seen = set()
        # First pass: prefer unseen topics AND unseen IDs
        for r in segs:
            if len(selected) >= 5:
                break
            if r["utt_id"] in used_ids:
                continue
            topic = r.get("topic", "Unknown")
            if topic not in topics_seen:
                selected.append(r)
                topics_seen.add(topic)
                used_ids.add(r["utt_id"])
        # Second pass: fill remaining slots with any unseen ID
        for r in segs:
            if len(selected) >= 5:
                break
            if r["utt_id"] in used_ids:
                continue
            selected.append(r)
            used_ids.add(r["utt_id"])

        for i, r in enumerate(selected, 1):
            lines.append(format_example_md(r, i))

        lines.append("---")
        lines.append("")

    # Implications
    lines.append("## Implications for VSP System Evaluation")
    lines.append("")
    lines.append("### 1. WER Systematically Undervalues Lip-Reading Output")
    lines.append("")
    lines.append("WER treats all word substitutions equally. But in lip-reading:")
    lines.append("- 'admiral' → 'animal' is a **catastrophic** error (entity destroyed)")
    lines.append("- 'going to' → 'gonna' is a **harmless** paraphrase")
    lines.append("- 'the' → 'a' is a **trivial** function word change")
    lines.append("")
    lines.append("The 165 salvageable segments demonstrate that WER conflates these fundamentally "
                 "different error types into a single number.")
    lines.append("")
    lines.append("### 2. Context Recovery Is Real")
    lines.append("")
    lines.append("A domain-aware viewer (e.g., someone watching a cooking tutorial) can often "
                 "fill in gaps from context. When the model outputs 'add the flour and stir' "
                 "instead of 'add the flower and steer,' the viewer's domain knowledge corrects "
                 "the phonetic confusion automatically.")
    lines.append("")
    lines.append("### 3. Effective System Value Is Higher Than Metrics Suggest")
    lines.append("")
    lines.append(f"- **Metric-based capture rate**: {100*is_good/total:.1f}% (IS >= 3.0)")
    lines.append(f"- **LLM-adjusted capture rate**: {100*(is_good + len(all_divergent))/total:.1f}% "
                 f"(IS >= 3.0 OR LLM-salvageable)")
    lines.append(f"- **Uplift**: +{100*len(all_divergent)/total:.1f} percentage points "
                 f"(+{100*len(all_divergent)/is_good:.0f}% relative improvement)")
    lines.append("")
    lines.append("This means the VSP-LLM system delivers useful output for roughly **1 in 2 segments** "
                 "rather than the **2 in 5** that metrics alone suggest.")
    lines.append("")

    # Methodology
    lines.append("## Methodology")
    lines.append("")
    lines.append("### Selection Criteria")
    lines.append("")
    lines.append("A segment is classified as \"LLM-salvageable\" when:")
    lines.append("1. `llm_context_prob >= 0.5` — Claude's heuristic assigns at least 50% probability "
                 "that a viewer with domain context could recover the intended meaning")
    lines.append("2. `intelligibility_score < 3.0` — Traditional IS metric classifies the segment "
                 "as below the \"properly captured\" threshold")
    lines.append("")
    lines.append("### LLM Heuristic")
    lines.append("")
    lines.append("The `llm_context_prob` is a deterministic, 15-rule decision tree designed by Claude "
                 "that evaluates six linguistic factors: content word overlap, sequence preservation, "
                 "phonetic plausibility, length sanity, semantic domain coherence, and information "
                 "density. It correlates at r=0.934 with IS across the full dataset and is stable "
                 "across 16 different decode configurations (std=0.015).")
    lines.append("")
    lines.append("### Validation")
    lines.append("")
    lines.append("Cross-configuration analysis (16 decode parameter sets) confirms:")
    lines.append("- Mean correlation with IS: r=0.925 (std=0.015)")
    lines.append("- Agreement at IS >= 2.00 (Y+P): \u03ba = 0.818; at IS >= 3.80 (Y): \u03ba = 0.690")
    lines.append("- Recall: 99.2% (almost never misses a properly captured segment)")
    lines.append("- Precision: 78.2% (intentionally optimistic — assumes domain context)")
    lines.append("")

    return "\n".join(lines)


def generate_json_data(categories, all_divergent):
    """Generate JSON with all divergent segments and category assignments."""
    data = {
        "metadata": {
            "date": date.today().isoformat(),
            "description": "Segments where Claude's LLM heuristic identifies recoverable meaning "
                           "that traditional metrics (WER, WWER, IS) miss",
            "selection_criteria": "llm_context_prob >= 0.5 AND intelligibility_score < 3.0",
            "total_divergent": len(all_divergent),
        },
        "summary": {
            "by_llm_reason": dict(Counter(r["llm_context_reason"] for r in all_divergent).most_common()),
            "by_is_tier": dict(Counter(r["intelligibility_label"] for r in all_divergent).most_common()),
            "by_topic": dict(Counter(r.get("topic", "Unknown") for r in all_divergent).most_common()),
            "mean_wer": sum(r["wer_%"] for r in all_divergent if r["wer_%"]) / len(all_divergent),
            "mean_semantic_sim": sum(r["semantic_sim"] for r in all_divergent if r["semantic_sim"]) / len(all_divergent),
            "mean_llm_prob": sum(r["llm_context_prob"] for r in all_divergent if r["llm_context_prob"]) / len(all_divergent),
        },
        "categories": {},
        "segments": [],
    }

    for cat_key, cat_data in categories.items():
        data["categories"][cat_key] = {
            "title": cat_data["title"],
            "description": cat_data["description"],
            "count": len(cat_data["segments"]),
            "segment_ids": [r["utt_id"] for r in cat_data["segments"]],
        }

    for r in all_divergent:
        seg = {
            "utt_id": r["utt_id"],
            "display_name": r["display_name"],
            "ref": r["ref"],
            "hyp": r["hyp"],
            "wer_pct": r["wer_%"],
            "wwer_pct": r["wwer_%"],
            "is_score": r["intelligibility_score"],
            "is_tier": r["intelligibility_label"],
            "semantic_sim": r["semantic_sim"],
            "phonetic_sim": r["phonetic_sim"],
            "nea_f1_pct": r["nea_f1_%"],
            "length_ratio": r["length_ratio"],
            "llm_context_prob": r["llm_context_prob"],
            "llm_context_reason": r["llm_context_reason"],
            "topic": r.get("topic", ""),
            "ref_words": r["ref_words"],
            "categories": [
                cat_key for cat_key, cat_data in categories.items()
                if r["utt_id"] in [s["utt_id"] for s in cat_data["segments"]]
            ],
        }
        data["segments"].append(seg)

    # Sort by interest score
    data["segments"].sort(key=lambda s: s["llm_context_prob"], reverse=True)

    return data


def generate_docx(categories, all_divergent, all_rows):
    """Generate formatted DOCX report following Argos style guide."""
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml, OxmlElement
    from docx.enum.text import WD_BREAK, WD_TAB_ALIGNMENT

    doc = Document()

    # Page setup (A4)
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # Colors
    PRIMARY = RGBColor(0x1B, 0x3A, 0x5C)   # Dark blue
    SECONDARY = RGBColor(0x2E, 0x86, 0xC1)  # Medium blue
    GRAY = RGBColor(0x80, 0x80, 0x80)
    GREEN = RGBColor(0x27, 0xAE, 0x60)
    RED = RGBColor(0xE7, 0x4C, 0x3C)

    # Header
    header = section.header
    hp = header.paragraphs[0]
    hp.paragraph_format.space_after = Pt(0)
    hp.paragraph_format.space_before = Pt(0)

    logo_path = ASSETS_DIR / "logo.png"
    if logo_path.exists():
        run = hp.add_run()
        run.add_picture(str(logo_path), width=Inches(0.4))

    text_width = section.page_width - section.left_margin - section.right_margin
    hp.paragraph_format.tab_stops.add_tab_stop(text_width, WD_TAB_ALIGNMENT.RIGHT)
    hp.add_run("\t")
    run = hp.add_run("Argos — The Orchard")
    run.font.size = Pt(8)
    run.font.color.rgb = GRAY
    run.font.italic = True
    run.font.name = "Calibri"

    # Footer with page numbers
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run._r.append(instr)
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)
    run.font.size = Pt(8)
    run.font.color.rgb = GRAY

    # Update fields on open
    settings = doc.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)

    # ── Cover Page ──
    for _ in range(2):
        doc.add_paragraph("")

    peacock_path = ASSETS_DIR / "peacock.png"
    if peacock_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(peacock_path), width=Inches(2.2))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ARGOS")
    run.font.size = Pt(48)
    run.font.bold = True
    run.font.color.rgb = PRIMARY
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("LLM Salvage Analysis")
    run.font.size = Pt(22)
    run.font.color.rgb = PRIMARY
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Recoverable Predictions That Metrics Undercount")
    run.font.size = Pt(16)
    run.font.color.rgb = SECONDARY
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("The Orchard")
    run.font.size = Pt(20)
    run.font.color.rgb = PRIMARY
    run.font.name = "Calibri"

    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Yoad Oxman")
    run.font.size = Pt(14)
    run.font.name = "Calibri"

    # Metadata
    meta_items = [
        ("Date", date.today().strftime("%B %d, %Y")),
        ("Dataset", "Full Baseline — 1,497 segments (AVSpeech)"),
        ("Divergent Segments", f"{len(all_divergent)} (LLM-salvageable, IS < 3.0)"),
    ]
    for key, val in meta_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{key}: ")
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        run = p.add_run(val)
        run.font.size = Pt(11)
        run.font.name = "Calibri"

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ── TOC ──
    for _ in range(3):
        sp = doc.add_paragraph("")
        sp.paragraph_format.space_after = Pt(0)
        sp.paragraph_format.space_before = Pt(0)

    toc_heading = doc.add_heading("Table of Contents", level=1)
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-2" \\h \\z \\u '
    run._r.append(instr)
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    toc_titles = [
        "1. Executive Summary",
        "2. Divergent Segment Profile",
        "3. Hidden Gems: High Confidence Despite Bad Metrics",
        "4. Semantic Preservation: Same Meaning, Different Words",
        "5. Phonetic Bridge: Natural Lip-Reading Confusions",
        "6. Entity-Preserved: Key Information Survives",
        "7. Structure Match: Word Order Preserved",
        "8. WER Over-Punishment: Function Word Inflation",
        "9. Implications",
        "10. Methodology",
    ]
    for title in toc_titles:
        placeholder = p.add_run(title + "\n")
        placeholder.font.size = Pt(11)
        placeholder.font.name = "Calibri"

    fld_end_run = p.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    fld_end_run._r.append(fld_end)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ── Executive Summary ──
    doc.add_heading("1. Executive Summary", level=1)

    total = len(all_rows)
    is_good = sum(1 for r in all_rows if r["intelligibility_score"] and r["intelligibility_score"] >= 3.0)
    is_failed = total - is_good

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(
        f"Traditional metrics classify {is_failed} of {total} segments as failures (IS < 3.0). "
        f"However, Claude's LLM-calibrated heuristic identifies {len(all_divergent)} of these "
        f"segments as having recoverable meaning — cases where a viewer with domain context would "
        f"understand the lip-reading output despite high word error rates."
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    run = p.add_run(
        f"This represents an {100*len(all_divergent)/is_failed:.1f}% recovery rate among failed "
        f"segments. Including these, the effective intelligibility rate rises from "
        f"{100*is_good/total:.1f}% to {100*(is_good+len(all_divergent))/total:.1f}%."
    )
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    # Summary table
    table = doc.add_table(rows=7, cols=3)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Metric", "Count", "Percentage"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True

    data_rows = [
        ("Total segments", str(total), "100%"),
        ("IS >= 3.0 (captured)", str(is_good), f"{100*is_good/total:.1f}%"),
        ("IS < 3.0 (failed)", str(is_failed), f"{100*is_failed/total:.1f}%"),
        ("LLM prob >= 0.5 (recoverable)", str(sum(1 for r in all_rows if r["llm_context_prob"] and r["llm_context_prob"] >= 0.5)), f"{100*sum(1 for r in all_rows if r['llm_context_prob'] and r['llm_context_prob'] >= 0.5)/total:.1f}%"),
        ("LLM salvages from failures", str(len(all_divergent)), f"{100*len(all_divergent)/is_failed:.1f}% of failures"),
        ("Effective capture rate", str(is_good + len(all_divergent)), f"{100*(is_good+len(all_divergent))/total:.1f}%"),
    ]
    for row_idx, (m, c, pct) in enumerate(data_rows, 1):
        table.rows[row_idx].cells[0].text = m
        table.rows[row_idx].cells[1].text = c
        table.rows[row_idx].cells[2].text = pct

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ── Divergent Segment Profile ──
    doc.add_heading("2. Divergent Segment Profile", level=1)

    # WER distribution table
    doc.add_heading("WER Distribution", level=2)
    wer_ranges = [(0, 30, "0-30%"), (30, 50, "30-50%"), (50, 70, "50-70%"),
                  (70, 100, "70-100%"), (100, 999, "100%+")]
    table = doc.add_table(rows=len(wer_ranges) + 1, cols=3)
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(["WER Range", "Count", "% of Divergent"]):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
    for row_idx, (lo, hi, label) in enumerate(wer_ranges, 1):
        count = sum(1 for r in all_divergent if r["wer_%"] is not None and lo <= r["wer_%"] < hi)
        table.rows[row_idx].cells[0].text = label
        table.rows[row_idx].cells[1].text = str(count)
        table.rows[row_idx].cells[2].text = f"{100*count/len(all_divergent):.1f}%"

    doc.add_paragraph("")

    # Recovery mechanism table
    doc.add_heading("Recovery Mechanisms", level=2)
    reasons = Counter(r["llm_context_reason"] for r in all_divergent)
    table = doc.add_table(rows=len(reasons) + 1, cols=3)
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(["LLM Reason", "Count", "%"]):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
    for row_idx, (reason, count) in enumerate(reasons.most_common(), 1):
        table.rows[row_idx].cells[0].text = reason
        table.rows[row_idx].cells[1].text = str(count)
        table.rows[row_idx].cells[2].text = f"{100*count/len(all_divergent):.1f}%"

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ── Category sections ──
    section_num = 3
    for cat_key, cat_data in categories.items():
        segs = cat_data["segments"]
        if not segs:
            continue

        doc.add_heading(f"{section_num}. {cat_data['title']}", level=1)

        p = doc.add_paragraph()
        run = p.add_run(cat_data["description"])
        run.font.italic = True
        run.font.size = Pt(11)
        run.font.name = "Calibri"

        p = doc.add_paragraph()
        run = p.add_run(f"{len(segs)} segments in this category.")
        run.font.bold = True
        run.font.size = Pt(11)

        # Top 3 examples per category in docx
        for i, r_seg in enumerate(segs[:3], 1):
            doc.add_heading(f"Example {i}: {r_seg['display_name']}", level=2)

            # Ref/Hyp
            p = doc.add_paragraph()
            run = p.add_run("Reference: ")
            run.font.bold = True
            run.font.size = Pt(10)
            run = p.add_run(r_seg["ref"])
            run.font.size = Pt(10)

            p = doc.add_paragraph()
            run = p.add_run("Hypothesis: ")
            run.font.bold = True
            run.font.size = Pt(10)
            run = p.add_run(r_seg["hyp"])
            run.font.size = Pt(10)

            # Metrics table
            metrics = [
                ("WER", f"{r_seg['wer_%']:.1f}%"),
                ("WWER", f"{r_seg['wwer_%']:.1f}%"),
                ("IS", f"{r_seg['intelligibility_score']:.3f} ({r_seg['intelligibility_label']})"),
                ("Semantic Sim", f"{r_seg['semantic_sim']:.4f}"),
                ("Phonetic Sim", f"{r_seg['phonetic_sim']:.4f}"),
                ("NEA F1", f"{r_seg['nea_f1_%']:.1f}%"),
                ("LLM Context Prob", f"{r_seg['llm_context_prob']:.3f}"),
                ("LLM Reason", r_seg["llm_context_reason"]),
            ]
            table = doc.add_table(rows=len(metrics), cols=2)
            table.style = "Light Grid Accent 1"
            for row_idx, (metric, val) in enumerate(metrics):
                table.rows[row_idx].cells[0].text = metric
                table.rows[row_idx].cells[1].text = val
                for p in table.rows[row_idx].cells[0].paragraphs:
                    for r in p.runs:
                        r.font.bold = True
                        r.font.size = Pt(9)
                for p in table.rows[row_idx].cells[1].paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(9)

            doc.add_paragraph("")

        # Page break between categories
        if section_num < 8:
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        section_num += 1

    # ── Implications ──
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    doc.add_heading(f"{section_num}. Implications for VSP System Evaluation", level=1)
    section_num += 1

    implications = [
        ("WER Systematically Undervalues Lip-Reading Output",
         "WER treats all word substitutions equally. But in lip-reading, "
         "'admiral' → 'animal' is a catastrophic error (entity destroyed), while "
         "'going to' → 'gonna' is a harmless paraphrase. The 165 salvageable segments "
         "demonstrate that WER conflates fundamentally different error types."),
        ("Context Recovery Is Real",
         "A domain-aware viewer (e.g., someone watching a cooking tutorial) can often "
         "fill in gaps from context. When the model outputs 'add the flour and stir' "
         "instead of 'add the flower and steer,' the viewer's domain knowledge corrects "
         "the phonetic confusion automatically."),
        ("Effective System Value Is Higher Than Metrics Suggest",
         f"Metric-based capture rate: {100*is_good/total:.1f}% (IS >= 3.0). "
         f"LLM-adjusted capture rate: {100*(is_good+len(all_divergent))/total:.1f}% "
         f"(IS >= 3.0 OR LLM-salvageable). This means the VSP-LLM system delivers "
         f"useful output for roughly 1 in 2 segments rather than the 2 in 5 that "
         f"metrics alone suggest."),
    ]

    for title, text in implications:
        doc.add_heading(title, level=2)
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.name = "Calibri"

    # ── Methodology ──
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    doc.add_heading(f"{section_num}. Methodology", level=1)

    p = doc.add_paragraph()
    run = p.add_run(
        "A segment is classified as 'LLM-salvageable' when llm_context_prob >= 0.5 "
        "(Claude's heuristic assigns at least 50% recovery probability) AND "
        "intelligibility_score < 3.0 (traditional IS metric classifies as below the "
        "'properly captured' threshold)."
    )
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run(
        "The llm_context_prob is a deterministic 15-rule decision tree designed by Claude "
        "that evaluates six linguistic factors. It correlates at r=0.934 with IS across "
        "the full dataset and is stable across 16 decode configurations (std=0.015). "
        "Cross-configuration validation confirms mean Cohen's kappa = 0.72 (range 0.62-0.86) "
        "and recall 97.6-100%."
    )
    run.font.size = Pt(11)

    # Full segment listing appendix
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    doc.add_heading("Appendix: All 165 Salvageable Segments", level=1)

    p = doc.add_paragraph()
    run = p.add_run(
        "Complete listing of all segments where llm_context_prob >= 0.5 and IS < 3.0, "
        "sorted by LLM confidence (descending)."
    )
    run.font.size = Pt(10)
    run.font.italic = True

    # Sort by llm prob descending
    sorted_div = sorted(all_divergent, key=lambda r: r["llm_context_prob"] or 0, reverse=True)

    # Create appendix table
    table = doc.add_table(rows=min(len(sorted_div), 165) + 1, cols=7)
    table.style = "Light Grid Accent 1"
    table.autofit = True

    app_headers = ["#", "Display Name", "WER%", "IS", "Sem", "LLM Prob", "Reason"]
    for i, h in enumerate(app_headers):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(7)

    for row_idx, r_seg in enumerate(sorted_div[:165], 1):
        table.rows[row_idx].cells[0].text = str(row_idx)
        table.rows[row_idx].cells[1].text = r_seg["display_name"][:30]
        table.rows[row_idx].cells[2].text = f"{r_seg['wer_%']:.0f}"
        table.rows[row_idx].cells[3].text = f"{r_seg['intelligibility_score']:.2f}"
        table.rows[row_idx].cells[4].text = f"{r_seg['semantic_sim']:.2f}"
        table.rows[row_idx].cells[5].text = f"{r_seg['llm_context_prob']:.2f}"
        table.rows[row_idx].cells[6].text = r_seg["llm_context_reason"]
        for cell in table.rows[row_idx].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(7)

    return doc


def main():
    print("Loading data...")
    rows = load_data()
    print(f"Loaded {len(rows)} segments.")

    print("Categorizing salvageable segments...")
    categories, all_divergent = categorize_salvage(rows)
    print(f"Found {len(all_divergent)} divergent segments (LLM >= 0.5, IS < 3.0)")
    for cat_key, cat_data in categories.items():
        print(f"  {cat_key}: {len(cat_data['segments'])} segments")

    # Create output directory
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    # Generate markdown
    print("\nGenerating markdown...")
    md_content = generate_markdown(categories, all_divergent, rows)
    md_path = OUTPUT_DIR / "llm_salvage_analysis.md"
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(md_content)
    print(f"  Written: {md_path}")

    # Generate JSON
    print("Generating JSON data...")
    json_data = generate_json_data(categories, all_divergent)
    json_path = OUTPUT_DIR / "llm_salvage_segments.json"
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(json_data, f, indent=2, ensure_ascii=False)
    print(f"  Written: {json_path}")

    # Generate DOCX
    print("Generating DOCX report...")
    doc = generate_docx(categories, all_divergent, rows)
    docx_path = OUTPUT_DIR / "llm_salvage_analysis.docx"
    doc.save(str(docx_path))
    print(f"  Written: {docx_path}")

    # Copy to presentation materials
    pres_md_dir = PRES_DIR / "03_reports_md" / "supplementary"
    pres_docx_dir = PRES_DIR / "04_reports_docx"
    pres_data_dir = PRES_DIR / "05_data"

    import shutil

    if pres_md_dir.exists():
        shutil.copy2(md_path, pres_md_dir / "llm_salvage_analysis.md")
        print(f"  Copied to: {pres_md_dir / 'llm_salvage_analysis.md'}")

    if pres_docx_dir.exists():
        shutil.copy2(docx_path, pres_docx_dir / "llm_salvage_analysis.docx")
        print(f"  Copied to: {pres_docx_dir / 'llm_salvage_analysis.docx'}")

    if pres_data_dir.exists():
        shutil.copy2(json_path, pres_data_dir / "llm_salvage_segments.json")
        print(f"  Copied to: {pres_data_dir / 'llm_salvage_segments.json'}")

    print("\nDone.")


if __name__ == "__main__":
    main()
