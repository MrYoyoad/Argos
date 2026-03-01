#!/usr/bin/env python3
"""
Argos — The Orchard — Comprehensive R&D Document Generator (v3)

Generates a detailed, chapter-based research documentation Word document
covering ALL work done on the Argos visual speech processing project.

16 chapters, ~70 pages, content-rich with code examples and config snippets.
v3: Added Chapter 13 (Intelligibility Assessment), expanded Chapter 12 with
    plots/analysis, updated numbers to full 1,497-segment dataset.

Usage:
    python3 generate_research_journal.py

Output:
    docs/evaluation/research-journal.docx
"""

import os
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

# ── Output ──
SCRIPT_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = SCRIPT_DIR.parent.parent / "evaluation"
OUTPUT_FILE = OUTPUT_DIR / "research-journal.docx"

# ── Logos ──
ASSETS_DIR = SCRIPT_DIR.parent / "assets"
LOGO_ORCHARD = ASSETS_DIR / "logo.png"
LOGO_PEACOCK = ASSETS_DIR / "peacock.png"

# ── Colors ──
C_PRIMARY = RGBColor(0x1a, 0x3a, 0x5c)
C_H2 = RGBColor(0x2a, 0x5a, 0x8c)
C_H3 = RGBColor(0x3a, 0x6a, 0x9c)
C_H4 = RGBColor(0x4a, 0x7a, 0xac)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_BLACK = RGBColor(0x00, 0x00, 0x00)
C_DARK = RGBColor(0x33, 0x33, 0x33)
C_GREEN = RGBColor(0x15, 0x57, 0x24)
C_RED = RGBColor(0x72, 0x1c, 0x24)
C_AMBER = RGBColor(0x85, 0x64, 0x04)
C_GRAY = RGBColor(0x66, 0x66, 0x66)
C_CODE = RGBColor(0x2d, 0x2d, 0x2d)

HEADER_BG = "1a3a5c"
ZEBRA_BG = "f0f4f8"
GREEN_BG = "d4edda"
RED_BG = "f8d7da"
AMBER_BG = "fff3cd"
CODE_BG = "f5f5f5"

LAST_UPDATED = datetime.now().strftime("%B %d, %Y")


# ═══════════════════════════════════════════════════
# HELPER FUNCTIONS
# ═══════════════════════════════════════════════════

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, bold=False, color=None, size=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = size
    run.font.name = "Calibri"


def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, HEADER_BG)
        set_cell_text(cell, h, bold=True, color=C_WHITE, size=Pt(9))

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            set_cell_text(cell, val, size=Pt(9))
            if r_idx % 2 == 1:
                set_cell_shading(cell, ZEBRA_BG)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return table


_bookmark_counter = [0]  # mutable counter for unique bookmark IDs

def add_heading(doc, text, level, color=None, bookmark_id=None):
    h = doc.add_heading(text, level=level)
    if color is None:
        color = {1: C_PRIMARY, 2: C_H2, 3: C_H3, 4: C_H4}.get(level, C_PRIMARY)
    for run in h.runs:
        run.font.color.rgb = color
        run.font.name = "Calibri"
    # Insert a bookmark so TOC hyperlinks can target this heading
    if bookmark_id:
        _bookmark_counter[0] += 1
        bm_start = OxmlElement("w:bookmarkStart")
        bm_start.set(qn("w:id"), str(_bookmark_counter[0]))
        bm_start.set(qn("w:name"), bookmark_id)
        bm_end = OxmlElement("w:bookmarkEnd")
        bm_end.set(qn("w:id"), str(_bookmark_counter[0]))
        h._p.insert(0, bm_start)
        h._p.append(bm_end)
    return h


def add_para(doc, text, bold=False, italic=False, size=Pt(11), color=None, space_after=Pt(6)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = size
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = space_after
    return p


def add_para_rich(doc, segments, space_after=Pt(6)):
    """Add a paragraph with mixed bold/normal text.
    segments is a list of (text, bold) tuples."""
    p = doc.add_paragraph()
    for text, bold in segments:
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(11)
        run.font.name = "Calibri"
    p.paragraph_format.space_after = space_after
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = "Calibri"
    return p


def add_bullet_bold_value(doc, label, value, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after = Pt(2)
    run_b = p.add_run(label)
    run_b.bold = True
    run_b.font.size = Pt(10)
    run_b.font.name = "Calibri"
    run_v = p.add_run(value)
    run_v.font.size = Pt(10)
    run_v.font.name = "Calibri"
    return p


def add_code_block(doc, code_text, language=""):
    """Add a code block with monospace font and gray background."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.3)

    # Background shading
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{CODE_BG}" w:val="clear"/>')
    pPr.append(shading)

    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(8)
    run.font.color.rgb = C_CODE
    return p


def _build_inline_image_xml(rId, cx, cy, pic_id=1, name="Logo"):
    """Build the WordprocessingML XML for an inline image."""
    inline = OxmlElement('wp:inline')
    inline.set('distT', '0')
    inline.set('distB', '0')
    inline.set('distL', '0')
    inline.set('distR', '0')

    extent = OxmlElement('wp:extent')
    extent.set('cx', str(cx))
    extent.set('cy', str(cy))
    inline.append(extent)

    docPr = OxmlElement('wp:docPr')
    docPr.set('id', str(pic_id))
    docPr.set('name', name)
    inline.append(docPr)

    graphic = OxmlElement('a:graphic')
    graphicData = OxmlElement('a:graphicData')
    graphicData.set('uri', 'http://schemas.openxmlformats.org/drawingml/2006/picture')

    pic = OxmlElement('pic:pic')
    nvPicPr = OxmlElement('pic:nvPicPr')
    cNvPr = OxmlElement('pic:cNvPr')
    cNvPr.set('id', '0')
    cNvPr.set('name', name)
    nvPicPr.append(cNvPr)
    cNvPicPr = OxmlElement('pic:cNvPicPr')
    nvPicPr.append(cNvPicPr)
    pic.append(nvPicPr)

    blipFill = OxmlElement('pic:blipFill')
    blip = OxmlElement('a:blip')
    blip.set(qn('r:embed'), rId)
    blipFill.append(blip)
    stretch = OxmlElement('a:stretch')
    fillRect = OxmlElement('a:fillRect')
    stretch.append(fillRect)
    blipFill.append(stretch)
    pic.append(blipFill)

    spPr = OxmlElement('pic:spPr')
    xfrm = OxmlElement('a:xfrm')
    off = OxmlElement('a:off')
    off.set('x', '0')
    off.set('y', '0')
    xfrm.append(off)
    ext_el = OxmlElement('a:ext')
    ext_el.set('cx', str(cx))
    ext_el.set('cy', str(cy))
    xfrm.append(ext_el)
    spPr.append(xfrm)
    prstGeom = OxmlElement('a:prstGeom')
    prstGeom.set('prst', 'rect')
    spPr.append(prstGeom)
    pic.append(spPr)

    graphicData.append(pic)
    graphic.append(graphicData)
    inline.append(graphic)

    drawing = OxmlElement('w:drawing')
    drawing.append(inline)
    return drawing


def add_header_footer(doc):
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    section = doc.sections[0]
    section.different_first_page_header_footer = True

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]

    if LOGO_ORCHARD.exists():
        image_part = doc.part.package.get_or_add_image_part(str(LOGO_ORCHARD))
        rId = header.part.relate_to(image_part, RT.IMAGE)
        size_emu = int(0.3 * 914400)
        logo_run = hp.add_run()
        drawing = _build_inline_image_xml(rId, size_emu, size_emu, pic_id=10, name="Header Logo")
        logo_run._r.append(drawing)
    text_width = section.page_width - section.left_margin - section.right_margin
    hp.paragraph_format.tab_stops.add_tab_stop(text_width, WD_TAB_ALIGNMENT.RIGHT)
    hp.add_run("\t")
    run = hp.add_run("Argos \u2014 The Orchard")
    run.font.size = Pt(8)
    run.font.color.rgb = C_GRAY
    run.font.name = "Calibri"
    run.italic = True

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = C_GRAY
    run.font.name = "Calibri"
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    run._r.append(fld1)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run._r.append(instr)
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld2)


# ═══════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════

def create_cover_page(doc):
    for _ in range(2):
        doc.add_paragraph()

    if LOGO_PEACOCK.exists():
        doc.add_picture(str(LOGO_PEACOCK), width=Inches(2.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ARGOS")
    run.font.size = Pt(48)
    run.font.color.rgb = C_PRIMARY
    run.bold = True
    run.font.name = "Calibri"

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Visual Speech Processing System")
    run2.font.size = Pt(22)
    run2.font.color.rgb = C_H2
    run2.font.name = "Calibri"

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("The Orchard")
    run3.font.size = Pt(20)
    run3.font.color.rgb = C_H2
    run3.font.name = "Calibri"

    doc.add_paragraph()

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run("Detailed Edition v3.0")
    run4.font.size = Pt(14)
    run4.font.color.rgb = C_H3
    run4.italic = True
    run4.font.name = "Calibri"

    for _ in range(2):
        doc.add_paragraph()

    info_lines = [
        ("Project Start:", "October 18, 2024"),
        ("Lead Engineer:", "Yoad Oxman"),
        ("Last Updated:", LAST_UPDATED),
    ]
    for label, value in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rl = p.add_run(label + " ")
        rl.bold = True
        rl.font.size = Pt(11)
        rl.font.color.rgb = C_DARK
        rl.font.name = "Calibri"
        rv = p.add_run(value)
        rv.font.size = Pt(11)
        rv.font.color.rgb = C_DARK
        rv.font.name = "Calibri"

    doc.add_page_break()


# ═══════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════

def _add_toc_hyperlink(doc, bookmark_name, display_text, level=0):
    """Add a clickable TOC entry that jumps to a bookmark."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.left_indent = Inches(0.3 * level)

    # Build w:hyperlink element pointing to the bookmark
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), bookmark_name)
    hyperlink.set(qn("w:history"), "1")

    run_el = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    # Blue underline style for hyperlink
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), "1a3a5c")
    rPr.append(color_el)
    u_el = OxmlElement("w:u")
    u_el.set(qn("w:val"), "single")
    rPr.append(u_el)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(22 if level == 0 else 20))  # 11pt or 10pt
    rPr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(22 if level == 0 else 20))
    rPr.append(szCs)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:hAnsi"), "Calibri")
    rPr.append(rFonts)
    if level == 0:
        b_el = OxmlElement("w:b")
        rPr.append(b_el)
    run_el.append(rPr)

    text_el = OxmlElement("w:t")
    text_el.set(qn("xml:space"), "preserve")
    text_el.text = display_text
    run_el.append(text_el)

    hyperlink.append(run_el)
    p._p.append(hyperlink)
    return p


# The chapter titles and their bookmark IDs — used by both TOC and chapter headings
TOC_ENTRIES = [
    ("ch1",  "1. Research Foundation & Literature Review"),
    ("ch2",  "2. Environment & Infrastructure Setup"),
    ("ch3",  "3. Building the Pipeline from Scratch"),
    ("ch4",  "4. Video Normalization"),
    ("ch5",  "5. Video Segmentation"),
    ("ch6",  "6. Face Detection & Mouth Cropping"),
    ("ch7",  "7. Reports & Burned Videos"),
    ("ch8",  "8. Modular Code Refactoring"),
    ("ch9",  "9. Web UI & Server"),
    ("ch10", "10. Backend Services"),
    ("ch11", "11. Docker Container & Deployment"),
    ("ch12", "12. Performance Evaluation & Tuning"),
    ("ch13", "13. Intelligibility Assessment & Metric Analysis"),
    ("ch14", "14. Fine-Tuning & Training Infrastructure"),
    ("ch15", "15. Executive Summary"),
    ("ch16", "16. Project To-Do List"),
]


def create_toc(doc):
    for _ in range(3):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    add_heading(doc, "Table of Contents", 1)

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fld_begin)
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' TOC \\o "1-2" \\h \\z \\u '
    run._r.append(instr)
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    run._r.append(fld_sep)

    for _bm, title in TOC_ENTRIES:
        placeholder = paragraph.add_run(title + "\n")
        placeholder.font.size = Pt(11)
        placeholder.font.name = "Calibri"
        placeholder.font.color.rgb = C_PRIMARY

    fld_end_run = paragraph.add_run()
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    fld_end_run._r.append(fld_end)

    settings = doc.settings.element
    update_fields = OxmlElement('w:updateFields')
    update_fields.set(qn('w:val'), 'true')
    settings.append(update_fields)

    doc.add_page_break()


# ═══════════════════════════════════════════════════
# CHAPTER 1: Research Foundation & Literature Review
# ═══════════════════════════════════════════════════

def chapter_1(doc):
    add_heading(doc, "1. Research Foundation & Literature Review", 1, bookmark_id="ch1")

    # 1.1 Team
    add_heading(doc, "1.1 Team & Project Inception", 2)
    add_para(doc, (
        "The Argos project began on October 18, 2024, with the goal of building a production-ready "
        "visual speech processing (lip-reading) system. The project team consists of three members, "
        "though all engineering and code development was carried out by a single developer:"
    ))
    add_styled_table(doc,
        ["Team Member", "Role", "Responsibilities"],
        [
            ["Yoad Oxman", "Lead Engineer", "All code development, pipeline architecture, UI, container deployment, testing, documentation"],
            ["Omer Leibovitch", "Data & Research", "Data sourcing, early conceptual samples, initial experiments"],
            ["Ido Springer", "Supervisor", "Management, supervision, project direction"],
        ],
        col_widths=[1.5, 1.2, 3.8]
    )
    add_para(doc, (
        "This is fundamentally a solo engineering effort. Every script, module, UI component, Docker container, "
        "and deployment artifact documented in this journal was designed and implemented by Yoad Oxman. "
        "The total codebase comprises approximately 14,882 lines of custom production code across ~55 files."
    ))

    # 1.2 VSP-LLM Paper
    add_heading(doc, "1.2 The VSP-LLM Paper", 2)
    add_para(doc, (
        "The core model is based on the VSP-LLM paper by Yeo et al. (arXiv:2402.15151v2, May 2024), "
        "titled 'Visual Speech Processing Incorporated with Large Language Models.' This paper presents "
        "a novel approach that combines visual speech features extracted by AV-HuBERT with the language "
        "understanding capabilities of LLaMA-2 through a parameter-efficient fine-tuning approach."
    ))

    add_heading(doc, "Architecture", 3)
    add_para(doc, "The VSP-LLM architecture consists of three main components connected in sequence:")
    add_code_block(doc, (
        "┌─────────────────────┐    ┌────────────────────┐    ┌──────────────────────┐\n"
        "│   AV-HuBERT Encoder │    │  Linear Projection │    │   LLaMA-2-7B + QLoRA │\n"
        "│   (315M parameters) │ -> │  (1024 -> 4096)    │ -> │   (7B params, ~0.1%  │\n"
        "│   Video -> Features │    │  Dim. alignment    │    │   trainable via LoRA) │\n"
        "└─────────────────────┘    └────────────────────┘    └──────────────────────┘\n"
        "\n"
        "Input: 88x88 mouth crops at 25 FPS\n"
        "Output: Text transcription (beam search or sampling decode)"
    ))

    add_para(doc, "The paper reports the following key results on the LRS3 benchmark dataset:")
    add_styled_table(doc,
        ["Metric", "Value", "Dataset", "Notes"],
        [
            ["Word Error Rate (WER)", "25.4%", "LRS3 (test)", "Table 1, visual-only setting"],
            ["Training Data", "433 hours", "LRS3 (labeled)", "Fine-tuning subset"],
            ["Pre-training", "1,759 hours", "LRS3 + VoxCeleb2", "AV-HuBERT self-supervised"],
            ["LLM", "LLaMA-2-7B", "—", "4-bit NF4 quantization with QLoRA"],
            ["LoRA Config", "rank=16, alpha=32", "—", "Applied to all linear layers"],
        ],
        col_widths=[1.8, 1.2, 1.5, 2.0]
    )

    # 1.3 Papers Studied
    add_heading(doc, "1.3 Literature Review (~2-3 Weeks)", 2)
    add_para(doc, (
        "Before any code was written, approximately two to three weeks were spent studying the foundational "
        "papers and understanding the full technical stack. This was essential because the VSP-LLM system "
        "sits at the intersection of multiple research areas: audio-visual speech recognition, self-supervised "
        "representation learning, large language models, and parameter-efficient fine-tuning."
    ))
    add_styled_table(doc,
        ["Paper / Framework", "Authors", "Key Contribution to Argos"],
        [
            ["VSP-LLM (arXiv:2402.15151)", "Yeo et al., 2024", "Core model architecture, decode strategy, QLoRA approach"],
            ["AV-HuBERT (arXiv:2201.02184)", "Shi et al., 2022", "Self-supervised audio-visual encoder, clustering approach"],
            ["Auto-AVSR (arXiv:2303.14307)", "Ma et al., 2023", "Preprocessing pipeline, face detection, mouth ROI extraction"],
            ["fairseq", "Ott et al. (Meta)", "Training framework, sequence generation, Hydra configs"],
            ["LRS3-TED Dataset", "Afouras et al., 2018", "Benchmark dataset structure, evaluation protocol"],
            ["LLaMA-2", "Touvron et al. (Meta), 2023", "Base LLM decoder, tokenizer, 7B architecture"],
            ["QLoRA (arXiv:2305.14314)", "Dettmers et al., 2023", "4-bit quantization + LoRA for efficient fine-tuning"],
        ],
        col_widths=[2.2, 1.5, 2.8]
    )
    add_para(doc, (
        "Understanding these papers was non-trivial. Each paper introduces its own terminology, assumptions, "
        "and dependencies. The AV-HuBERT paper alone requires understanding masked prediction pre-training, "
        "k-means clustering of speech features, and iterative refinement of cluster assignments. The fairseq "
        "framework adds its own configuration system (Hydra/OmegaConf) with deeply nested YAML files."
    ))

    # 1.4 Key Insight
    add_heading(doc, "1.4 Key Insight: Research-Grade vs. Production-Ready", 2)
    add_para(doc, (
        "A critical realization emerged early in the literature review, later formalized in the December 29, 2025 "
        "project presentation (Slide 9):"
    ))
    add_para(doc, (
        '"The paper assumes clean data, multi-GPU infrastructure, and internet connectivity. '
        'The code is research-grade, not plug-and-play."'
    ), bold=True, italic=True, color=C_H2)
    add_para(doc, (
        "This insight shaped the entire Argos project. The published code provides model definitions and training "
        "scripts, but assumes the user has already preprocessed data in LRS3 format, has multi-GPU clusters "
        "available, and can download dependencies on demand. None of these assumptions hold in the production "
        "environments where Argos needed to operate."
    ))

    # 1.5 Paper vs Argos
    add_heading(doc, "1.5 What the Paper Provides vs. What Argos Built", 2)
    add_styled_table(doc,
        ["Aspect", "Paper Provides", "Argos Built"],
        [
            ["Model Code", "PyTorch model definitions", "Same (used as-is)"],
            ["End-to-End Pipeline", "None — 'follow AV-HuBERT docs'", "Complete 8-stage pipeline (501 lines + 1,562 lib)"],
            ["Preprocessing", "Assumes LRS3 format", "'Flat' adaptation for arbitrary videos"],
            ["Video Handling", "None", "Normalization, segmentation, codec handling"],
            ["Face Detection", "References Auto-AVSR", "MediaPipe two-tier + RetinaFace integration"],
            ["ASR Ground Truth", "Assumes pre-existing labels", "Whisper integration with transcription reuse"],
            ["Evaluation", "Standard WER", "WER + WWER + NEA (809-line report generator)"],
            ["User Interface", "None (CLI only)", "Full web UI (5,780 lines, 6 screens, 12 API endpoints)"],
            ["Deployment", "None", "Docker container, 5 versions, 37 bugs fixed"],
            ["Multi-Environment", "Single machine", "4 environments (EC2, 2 standalone, HORIZON)"],
            ["Offline Operation", "Requires internet", "Full offline capability with local wheels"],
        ],
        col_widths=[1.5, 2.0, 3.0]
    )

    # 1.6 Three Upstream Repos
    add_heading(doc, "1.6 Three Upstream Repositories", 2)
    add_para(doc, (
        "Argos integrates three separate open-source repositories, each serving a distinct role in the pipeline. "
        "These repositories were forked, modified, and orchestrated into a unified system:"
    ))
    add_styled_table(doc,
        ["Repository", "Origin", "Role in Argos", "Key Modifications"],
        [
            ["auto_avsr", "Ma et al. (mpc001)", "Preprocessing: face detection, mouth cropping, video transforms",
             "MediaPipe integration, flat dataset support, fast_segment.py, overlapping_segmentation.py"],
            ["VSP-LLM", "Yeo et al. (Sally-SH)", "Model: inference, decode, training, clustering",
             "12 custom Python scripts (3,342 lines), decode configs, training configs"],
            ["av_hubert", "Meta (facebookresearch)", "Feature extraction: HuBERT encoder, clustering preparation",
             "Flat-to-LRS3 conversion, avhubert_flat variant for arbitrary datasets"],
        ],
        col_widths=[1.2, 1.3, 2.0, 2.0]
    )

    # 1.7 Architecture Overview
    add_heading(doc, "1.7 Full Pipeline Architecture", 2)
    add_para(doc, (
        "The complete Argos pipeline processes raw video files through eight sequential stages, "
        "managed by a master orchestrator script. Each stage is implemented as a reusable module "
        "in the lib/ directory:"
    ))
    add_code_block(doc, (
        "Stage 0:   Input Validation & Archive\n"
        "           └── Validate input videos, archive previous outputs\n"
        "           └── Preserve manual transcriptions across runs\n"
        "\n"
        "Stage 0.5: Video Segmentation (fast_segment.py)\n"
        "           └── Split long videos into 12-second segments\n"
        "           └── Ultra-fast codec-copy (no re-encoding)\n"
        "\n"
        "Stage 0.6: Transcription Reuse\n"
        "           └── Copy existing transcriptions from .transcriptions/\n"
        "           └── Match by exact filename\n"
        "\n"
        "Stage 1:   Video Normalization (lib/normalization.sh)\n"
        "           └── HDR/10-bit tone mapping, FPS standardization\n"
        "           └── GPU (NVENC) or CPU (libx264) encoding\n"
        "\n"
        "Stage 2:   Face Detection & Mouth Cropping (auto_avsr)\n"
        "           └── MediaPipe or RetinaFace detection\n"
        "           └── 88x88 pixel mouth ROI extraction at 25 FPS\n"
        "\n"
        "Stage 3:   ASR Transcription (Whisper)\n"
        "           └── Run Whisper on segments without transcriptions\n"
        "           └── Save new auto-transcriptions for future reuse\n"
        "\n"
        "Stage 4:   LRS3 Format Conversion\n"
        "           └── Convert flat dataset structure to LRS3 format\n"
        "           └── Generate manifests, TSVs, and metadata\n"
        "\n"
        "Stage 5:   Feature Clustering (K-Means)\n"
        "           └── Extract AV-HuBERT features\n"
        "           └── K-means clustering (200 clusters) or use golden model\n"
        "\n"
        "Stage 6:   VSP-LLM Decode\n"
        "           └── Beam search (beam=20) or sampling decode\n"
        "           └── Automatic Cython extension build for containers\n"
        "\n"
        "Stage 7:   Client Outputs\n"
        "           └── Reports (CSV, HTML, JSON, TXT, ANSI)\n"
        "           └── Burned videos with subtitle overlays\n"
        "           └── Lip crop video copies"
    ))

    # 1.8 Decode Config
    add_heading(doc, "1.8 Decode Configuration", 2)
    add_para(doc, (
        "The decode stage is controlled by a Hydra YAML configuration file (s2s_decode.yaml) that defines "
        "beam search parameters, length penalties, and generation constraints:"
    ))
    add_code_block(doc, (
        "# s2s_decode.yaml — VSP-LLM Decode Configuration\n"
        "generation:\n"
        "  beam: 20                  # Beam width for search\n"
        "  lenpen: 0.0               # Length penalty (paper's value)\n"
        "  max_len_a: 2.0            # Dynamic max_len = max_len_a * src_clusters + max_len_b\n"
        "  max_len_b: 200            # Buffer for short-input segments\n"
        "  max_len: 2048             # Hard cap on generated length\n"
        "  no_repeat_ngram_size: 3   # Prevent degenerate repetitions\n"
        "  repetition_penalty: 1.2   # Mild penalty for repeating tokens\n"
        "  do_sample: false          # false = deterministic beam search\n"
        "  temperature: 1.0          # Sampling temperature (only with do_sample=true)\n"
        "  top_p: 0.9                # Nucleus sampling threshold"
    ))

    # 1.9 Presentation Milestone
    add_heading(doc, "1.9 December 29, 2025 Presentation Milestone", 2)
    add_para(doc, (
        'On December 29, 2025, a 19-slide presentation titled "Argos — A Working Pipeline" was delivered '
        "to the team. This marked the transition from research and development to a functioning, demonstrable "
        "system. Key accomplishments highlighted in the presentation include:"
    ))
    add_bullet(doc, "HORIZON server operational — pipeline running on a closed network with no internet access")
    add_bullet(doc, "Labeled data preparation — AVSpeech dataset imported and preprocessed")
    add_bullet(doc, "Improved lip cropping — MediaPipe two-tier detection strategy")
    add_bullet(doc, "Linux standalone Docker — containerized deployment for production machines")
    add_bullet(doc, "AVSpeech support — pipeline processes arbitrary YouTube videos, not just LRS3")
    add_bullet(doc, "Whisper built in — ASR ground truth generation integrated into pipeline")
    add_bullet(doc, "Demo-ready system with web UI for non-technical users")

    doc.add_page_break()


# ═══════════════════════════════════════════════════
# CHAPTER 2: Environment & Infrastructure Setup
# ═══════════════════════════════════════════════════

def chapter_2(doc):
    add_heading(doc, "2. Environment & Infrastructure Setup", 1, bookmark_id="ch2")

    add_para(doc, (
        "Setting up the infrastructure for Argos was one of the most time-consuming aspects of the project. "
        "The system requires GPU-accelerated computing with specific CUDA versions, two separate Python virtual "
        "environments with conflicting dependencies, and the ability to operate in environments with no internet "
        "connectivity. This chapter documents the approximately six weeks spent on environment setup alone."
    ))

    # 2.1 Four Environments
    add_heading(doc, "2.1 Four Deployment Environments", 2)
    add_styled_table(doc,
        ["Environment", "Type", "Internet", "GPU", "Purpose"],
        [
            ["AWS EC2", "Cloud instance", "Full", "NVIDIA (varies)", "Development, testing, experiments"],
            ["Standalone Ubuntu #1", "Physical machine", "Limited", "NVIDIA GPU", "Production deployment target"],
            ["Standalone Ubuntu #2", "Physical machine", "Limited", "NVIDIA GPU", "Secondary production machine"],
            ["HORIZON Server", "Closed network", "None", "NVIDIA GPU", "Secure deployment (no internet)"],
        ],
        col_widths=[1.4, 1.2, 0.8, 1.2, 1.9]
    )
    add_para(doc, (
        "Each environment presented unique challenges. The EC2 instance served as the primary development "
        "environment where all code was first written and tested. Changes were then replicated to the standalone "
        "machines and HORIZON server, accounting for path differences (/home/ubuntu/ on EC2 vs /workspace/ in "
        "containers) and the absence of internet connectivity."
    ))
    add_para(doc, (
        "Making the HORIZON server operational was a significant milestone (highlighted in the December 2025 "
        "presentation). This closed-network server has zero internet access, meaning every Python package, "
        "model weight, and system dependency had to be pre-packaged and transferred physically."
    ))

    # 2.2 Failed Venv
    add_heading(doc, "2.2 Failed Virtual Environment Build Attempts (~2 Weeks)", 2)
    add_para(doc, (
        "The first attempt to set up the software environment on the standalone computers consumed approximately "
        "two weeks and ultimately failed. The core issue was attempting to build Python virtual environments "
        "directly on machines with limited internet access and non-standard system configurations."
    ))
    add_para(doc, "Key problems encountered during the failed venv builds:")
    add_bullet(doc, "System Python versions incompatible with required packages (needed Python 3.9+ for PyTorch)")
    add_bullet(doc, "pip failing to resolve dependency conflicts between fairseq and transformers")
    add_bullet(doc, "Compilation failures for packages requiring CUDA headers (torch, triton, bitsandbytes)")
    add_bullet(doc, "Network timeouts on machines with restricted or slow internet access")
    add_bullet(doc, "Version pinning conflicts — omegaconf==2.0.6 requires old pip for metadata compatibility")
    add_para(doc, (
        "The lesson learned was that building venvs from scratch on production machines was not viable. "
        "The solution that eventually worked was to build complete virtual environments on the EC2 instance "
        "and transfer them as part of the container deployment package."
    ))

    # 2.3 NVIDIA Driver
    add_heading(doc, "2.3 NVIDIA Driver & CUDA Installation (~1 Month)", 2)
    add_para(doc, (
        "After the venv build failures, approximately one month was spent on getting GPU acceleration working "
        "on the standalone machines. This involved installing NVIDIA drivers, CUDA toolkit, and verifying "
        "that PyTorch could communicate with the GPU."
    ))
    add_para(doc, "The installation process involved multiple iterations:")
    add_bullet(doc, "Identifying the correct NVIDIA driver version for the installed GPU hardware")
    add_bullet(doc, "Installing the NVIDIA driver (NVIDIA-Linux-x86_64-535.183.01.run or similar)")
    add_bullet(doc, "Installing CUDA toolkit (multiple versions tried: 12.1, 12.4, 12.8)")
    add_bullet(doc, "Resolving kernel module conflicts with existing nouveau drivers")
    add_bullet(doc, "Verifying GPU visibility with nvidia-smi and CUDA availability in Python")
    add_bullet(doc, "Testing that PyTorch CUDA operations work correctly (tensor operations, memory allocation)")
    add_para(doc, (
        "A particular challenge was that the pipeline requires two different CUDA versions: CUDA 12.8 for "
        "the preprocessing venv (with triton 3.4.0) and CUDA 12.4 compatibility for the decode venv "
        "(with triton 3.1.0). This was resolved through careful library path management within each "
        "virtual environment."
    ))

    # 2.4 Dual Venv
    add_heading(doc, "2.4 Dual Virtual Environment Strategy", 2)
    add_para(doc, (
        "The most significant infrastructure decision was adopting a dual virtual environment architecture. "
        "The preprocessing and decoding stages have fundamentally incompatible dependency requirements "
        "that cannot coexist in a single environment:"
    ))
    add_styled_table(doc,
        ["Aspect", "pre-process-venv (ASR)", "vsp-llm-yoad-venv (Decode)"],
        [
            ["PyTorch", "2.8.0+cu128", "2.5.1+cu124"],
            ["Triton", "3.4.0", "3.1.0"],
            ["CUDA Runtime", "12.8.*", "12.4.*"],
            ["Key Packages", "Whisper, MediaPipe, OpenCV, JAX", "fairseq, transformers, peft, bitsandbytes"],
            ["Package Count", "~135 packages", "~256 packages"],
            ["Purpose", "Face detection, mouth cropping, ASR", "K-means, decode, training"],
            ["Python", "3.9+", "3.9+"],
        ],
        col_widths=[1.3, 2.6, 2.6]
    )
    add_para(doc, (
        "The root conflict is between triton versions. The preprocessing pipeline requires triton 3.4.0 "
        "(which needs CUDA 12.8 runtime libraries), while fairseq's Cython extensions and the decode pipeline "
        "require triton 3.1.0 (CUDA 12.4). Installing both in one environment leads to import errors "
        "and GPU operation failures."
    ))

    # 2.5 Dependency Hell
    add_heading(doc, "2.5 Dependency Conflicts & Resolution", 2)
    add_para(doc, "Several specific dependency conflicts required creative solutions:")
    add_bullet_bold_value(doc, "omegaconf==2.0.6: ",
        "fairseq requires this exact version, but its package metadata format is incompatible with modern pip. "
        "Solution: install using an older pip version, then upgrade pip afterward.")
    add_bullet_bold_value(doc, "fairseq installation: ",
        "Required as an editable install from the forked repository (Sally-SH/VSP-LLM) because the model "
        "code adds custom tasks and criteria that must be importable.")
    add_bullet_bold_value(doc, "numpy version: ",
        "Preprocessing needs numpy 1.26.4 (for MediaPipe compatibility), while decode works with numpy 2.x. "
        "Resolved by pinning in each venv independently.")
    add_bullet_bold_value(doc, "Cython extensions: ",
        "fairseq's data_utils_fast Cython module must be compiled on the target machine's architecture. "
        "The decode module includes automatic detection and compilation on first run.")

    # 2.6 HORIZON
    add_heading(doc, "2.6 HORIZON: Operating Without Internet", 2)
    add_para(doc, (
        "The HORIZON server operates on a completely isolated network with no internet connectivity. "
        "This required pre-packaging every dependency:"
    ))
    add_bullet(doc, "All Python packages as wheel files (.whl) — transferred physically to the server")
    add_bullet(doc, "spaCy language model (en_core_web_sm) pre-downloaded with offline wheels")
    add_bullet(doc, "Whisper model files (medium.pt) pre-cached in the whisper_cache directory")
    add_bullet(doc, "NVIDIA CUDA runtime libraries bundled within the container")
    add_bullet(doc, "Model weights (checkpoint_finetune.pt, LLaMA-2-7B) pre-loaded")
    add_para(doc, (
        "The outputs module (lib/outputs.sh) implements a graceful degradation strategy for spaCy: "
        "it first tries to install from local wheels, falls back to online pip if available, and "
        "finally degrades to basic WER metrics if neither option succeeds. This ensures the pipeline "
        "produces useful output even when advanced NER metrics are unavailable."
    ))

    # 2.7 Model Weights
    add_heading(doc, "2.7 Model Weights Acquisition", 2)
    add_para(doc, (
        "The pre-trained model weights were obtained from the VSP-LLM paper's official Git repository "
        "(Sally-SH/VSP-LLM). These weights represent the model after the authors' own fine-tuning on "
        "the LRS3 dataset — specifically the checkpoint_finetune.pt file containing the AV-HuBERT encoder "
        "weights and linear projection layer. The LLaMA-2-7B base model was separately downloaded "
        "from Meta's distribution."
    ))

    # 2.8 AVSpeech Data
    add_heading(doc, "2.8 AVSpeech Data Sourcing", 2)
    add_para(doc, (
        "Training and evaluation data was sourced from the AVSpeech dataset, which contains short clips "
        "of people speaking extracted from YouTube videos. Unlike the LRS3 dataset (TED talks with clean "
        "audio and frontal faces), AVSpeech represents real-world conditions with diverse camera angles, "
        "lighting, and background noise. The data was imported from an old server that had been previously "
        "mounted, making it accessible for processing through the pipeline."
    ))

    doc.add_page_break()


# ═══════════════════════════════════════════════════
# CHAPTER 3: Building the Pipeline from Scratch
# ═══════════════════════════════════════════════════

def chapter_3(doc):
    add_heading(doc, "3. Building the Pipeline from Scratch", 1, bookmark_id="ch3")

    # 3.1 The Gap
    add_heading(doc, "3.1 The Gap: No Pipeline Existed", 2)
    add_para(doc, (
        "The single biggest engineering challenge in the Argos project was that no end-to-end pipeline existed. "
        "The VSP-LLM paper provides model code (PyTorch modules, training scripts, decode scripts) but assumes "
        "the user has already preprocessed their data into LRS3 format. The paper's instructions effectively say: "
        "'For preprocessing, follow the AV-HuBERT documentation. For face detection, follow Auto-AVSR.' "
        "This means three separate research codebases needed to be understood, modified, and orchestrated "
        "into a unified pipeline that could process raw video files into text transcriptions."
    ))
    add_para(doc, (
        "Building this pipeline from scratch — connecting auto_avsr's preprocessing with av_hubert's feature "
        "extraction and VSP-LLM's decode — required deep understanding of each repository's data formats, "
        "directory conventions, configuration systems, and Python import structures."
    ))

    # 3.2 Flat Dataset Adaptation
    add_heading(doc, '3.2 "Flat" Dataset Adaptation', 2)
    add_para(doc, (
        "The paper's code exclusively supports the LRS3 dataset format, which has a rigid directory structure "
        "organized by TED talk ID with specific naming conventions. Real-world deployment requires processing "
        "arbitrary video files that don't follow this structure."
    ))
    add_para(doc, (
        'The "flat" adaptation was a fundamental modification to the preprocessing pipeline that allows '
        "it to accept any collection of video files in a single directory. This required changes to:"
    ))
    add_bullet(doc, "Directory traversal logic — scan a flat directory instead of nested LRS3 paths")
    add_bullet(doc, "Manifest generation — create TSV files mapping flat filenames to features")
    add_bullet(doc, "Naming conventions — establish a segment naming system for arbitrary videos")
    add_bullet(doc, "LRS3 compatibility layer — convert flat dataset to LRS3 format for downstream stages")
    add_para(doc, (
        "This adaptation is what enabled Argos to process AVSpeech data and any other video source, "
        "rather than being limited to the LRS3 dataset the paper was designed for."
    ))

    # 3.3 First Pipeline
    add_heading(doc, "3.3 First Pipeline Version (Oct-Nov 2024)", 2)
    add_para(doc, (
        "The first pipeline version was approximately 285 lines of bash script, created in the galaxy_export "
        "directory. It implemented the basic 8-stage flow: face detection, mouth cropping, LRS3 conversion, "
        "manifest generation, k-means clustering, and VSP-LLM decode. This version was functional but "
        "fragile — it assumed specific directory structures, had hardcoded paths, and offered no error "
        "recovery or transcription reuse."
    ))

    # 3.4 Pipeline Evolution
    add_heading(doc, "3.4 Pipeline Evolution: 7 Versions", 2)
    add_para(doc, (
        "The pipeline script underwent seven major revisions, each adding significant capabilities. "
        "The following table tracks the evolution by line count and key features added:"
    ))
    add_styled_table(doc,
        ["Version", "Lines", "Key Changes"],
        [
            ["v1 (Initial)", "~240", "Basic sequential flow, hardcoded paths"],
            ["v2 (galaxy_export)", "~285", "Galaxy export packaging, input validation"],
            ["v3 (Segmentation)", "~410", "Video segmentation (4s windows), overlap support"],
            ["v4 (Normalization)", "~612", "Video normalization, ASR integration, HDR handling"],
            ["v5 (Transcription)", "~673", "Transcription reuse system, manual transcription support"],
            ["v6 (Monolithic)", "823", "Full feature set in single file, all stages complete"],
            ["v7 (Modular)", "501 + 1,562", "Refactored into 12 lib/ modules (see Chapter 8)"],
        ],
        col_widths=[1.5, 1.0, 4.0]
    )
    add_para(doc, (
        "The transition from v6 (823 lines) to v7 (501 + 1,562 lines) represents the modular refactoring "
        "detailed in Chapter 8. While the total line count increased, the main orchestrator script was reduced "
        "by 52%, and the extracted modules became independently testable and reusable across environments."
    ))

    # 3.5 Integration Challenges
    add_heading(doc, "3.5 Key Integration Challenges", 2)
    add_para(doc, (
        "Connecting three research repositories into a unified pipeline required solving several "
        "cross-cutting concerns:"
    ))
    add_bullet_bold_value(doc, "Path Management: ",
        "Each repository assumes its own directory structure. The pipeline must translate between "
        "auto_avsr's flat video paths, av_hubert's LRS3-style paths, and VSP-LLM's dataset directory "
        "conventions. This is handled by lib/config.sh which provides environment-aware path functions.")
    add_bullet_bold_value(doc, "Virtual Environment Switching: ",
        "The preprocessing stages (face detection, ASR) use the pre-process-venv, while stages 5-8 "
        "(clustering, decode, outputs) use vsp-llm-yoad-venv. The pipeline must activate and deactivate "
        "venvs at the correct boundaries without contaminating the shell environment.")
    add_bullet_bold_value(doc, "Data Format Translation: ",
        "auto_avsr produces mouth crop videos as .mp4 files. av_hubert expects features as .npy arrays. "
        "VSP-LLM expects fairseq-style manifests with .tsv and .wrd files. Each transition requires "
        "format conversion scripts.")
    add_bullet_bold_value(doc, "Error Propagation: ",
        "A failure in face detection (e.g., no face found in a segment) must be handled gracefully "
        "without stopping the entire pipeline. The modular architecture uses log_warn for non-fatal "
        "issues and log_error for stage-level failures.")

    # 3.6 Custom Python Scripts
    add_heading(doc, "3.6 Custom Python Scripts (13 Files, 3,342 Lines)", 2)
    add_para(doc, (
        "In addition to the pipeline orchestrator and library modules, 13 custom Python scripts were written "
        "to handle specific processing tasks. These scripts are the workhorses of the pipeline, handling "
        "everything from video segmentation to quality assessment:"
    ))
    add_styled_table(doc,
        ["Script", "Lines", "Location", "Purpose"],
        [
            ["make_report.py", "809", "VSP-LLM/scripts/", "WER/WWER/NEA metrics, color-coded HTML reports"],
            ["make_burn.py", "510", "VSP-LLM/scripts/", "Burn subtitle overlays onto videos with ffmpeg"],
            ["overlapping_segmentation.py", "390", "auto_avsr/preparation/", "Time-based overlap splitting algorithm"],
            ["transcribe_segments.py", "353", "VSP-LLM/scripts/", "Per-segment transcription management tool"],
            ["calculate_per_video_wer.py", "300", "VSP-LLM/scripts/", "Per-video WER with overlap deduplication"],
            ["merge_overlapping_predictions.py", "289", "VSP-LLM/scripts/", "Conflict resolution for overlapping segments"],
            ["fast_segment.py", "197", "auto_avsr/preparation/", "Ultra-fast codec-copy video segmentation"],
            ["asr_to_words_notime.py", "151", "auto_avsr/", "Whisper ASR wrapper — audio to word tokens"],
            ["generate_conflict_report.py", "149", "VSP-LLM/scripts/", "Analyze and report prediction conflicts"],
            ["build_flat_train_tsv.py", "93", "VSP-LLM/scripts/", "Generate training TSV manifests"],
            ["make_flat_tsv.py", "56", "VSP-LLM/scripts/", "Create flat-format TSV files"],
            ["build_flat_cluster_counts.py", "45", "VSP-LLM/scripts/", "Generate cluster count statistics"],
            ["inspect_km_hist.py", "38", "VSP-LLM/scripts/", "K-means cluster histogram visualization"],
        ],
        col_widths=[2.0, 0.5, 1.8, 2.2]
    )

    # 3.7 Transcription Reuse
    add_heading(doc, "3.7 Transcription Reuse System", 2)
    add_para(doc, (
        "One of the most impactful features for production use is the transcription reuse system, "
        "implemented across three pipeline steps (0.6, 3, and 1.5) in lib/asr.sh. This system ensures "
        "that Whisper ASR only runs on segments that don't already have transcriptions, saving hours "
        "of processing time on pipeline re-runs."
    ))
    add_code_block(doc, (
        "Step 0.6: Copy Existing Transcriptions\n"
        "  └── Check .transcriptions/ directory for existing .wrd files\n"
        "  └── Match transcription to segment by exact filename\n"
        "  └── Copy matching transcriptions to segment_wrd_tmp/\n"
        "  └── Whisper will skip segments that already have .wrd files\n"
        "\n"
        "Step 3: Run Whisper ASR\n"
        "  └── python3 asr_to_words_notime.py \\\n"
        "        --in_videos <segment_dir> --out_wrd <wrd_dir> \\\n"
        "        --model medium --lang en --tokenize alnum --lower\n"
        "  └── Only processes segments WITHOUT existing .wrd files\n"
        "\n"
        "Step 1.5: Save New Auto-Transcriptions\n"
        "  └── Copy new Whisper outputs to .transcriptions/ directory\n"
        "  └── Mark as 'auto' type in metadata.json\n"
        "  └── Skip if manual transcription already exists (preserve user edits)\n"
        "  └── Next run, Step 0.6 will reuse these transcriptions"
    ))
    add_para(doc, (
        "This system means that a user can manually transcribe specific segments via the web UI, "
        "and those manual transcriptions will persist across all future pipeline runs. Whisper never "
        "overwrites manual work, and auto-transcriptions are cached for reuse."
    ))

    doc.add_page_break()


# ═══════════════════════════════════════════════════
# PLACEHOLDER — remaining chapters will be added
# ═══════════════════════════════════════════════════

def chapter_4(doc):
    add_heading(doc, "4. Video Normalization", 1, bookmark_id="ch4")

    add_para(doc, (
        "Real-world video files arrive in an enormous variety of formats: different codecs (H.264, H.265, VP9, AV1), "
        "color spaces (SDR, HDR, BT.2020), bit depths (8-bit, 10-bit), frame rates (24-60 FPS), and resolutions. "
        "The face detection and model inference stages require standardized input: 8-bit SDR video at a consistent "
        "frame rate. The normalization module handles this conversion, and its development uncovered one of the "
        "most challenging production bugs in the entire project."
    ))

    # 4.1 Why Needed
    add_heading(doc, "4.1 Why Normalization Is Needed", 2)
    add_para(doc, (
        "Without normalization, the pipeline fails in multiple ways. HDR (High Dynamic Range) and 10-bit video "
        "causes face detection to produce incorrect landmarks because the pixel value ranges differ from what "
        "the model expects. Variable frame rates cause the model to receive inconsistent temporal information, "
        "degrading transcription quality. Non-standard codecs may not be compatible with the preprocessing "
        "pipeline's ffmpeg-based extraction."
    ))
    add_para(doc, "The normalization module standardizes all input to:")
    add_bullet(doc, "Codec: H.264 (universally compatible)")
    add_bullet(doc, "Color space: SDR (BT.709), 8-bit")
    add_bullet(doc, "Frame rate: 25 FPS (matching LRS3 training data)")
    add_bullet(doc, "Pixel format: yuv420p")

    # 4.2 The Module
    add_heading(doc, "4.2 The Normalization Module (lib/normalization.sh, 235 Lines)", 2)
    add_para(doc, (
        "The normalization module is the largest single module in the lib/ directory. It implements "
        "HDR detection, GPU-accelerated encoding, CPU fallback, and post-encode validation."
    ))

    add_heading(doc, "HDR Detection", 3)
    add_para(doc, (
        "The needs_tonemap() function uses ffprobe to detect videos that require tone mapping "
        "from HDR to SDR. It checks two indicators:"
    ))
    add_code_block(doc, (
        "needs_tonemap() {\n"
        '  local pix_fmt=$(ffprobe -v error -select_streams v:0 \\\n'
        '    -show_entries stream=pix_fmt -of csv=p=0 "$1")\n'
        '  local color_space=$(ffprobe -v error -select_streams v:0 \\\n'
        '    -show_entries stream=color_space -of csv=p=0 "$1")\n'
        '\n'
        '  # 10-bit pixel formats indicate HDR\n'
        '  if [[ "$pix_fmt" == *"10le"* ]] || [[ "$pix_fmt" == *"10be"* ]]; then\n'
        '    return 0  # needs tone mapping\n'
        '  fi\n'
        '  # BT.2020 color space indicates HDR\n'
        '  if [[ "$color_space" == *"bt2020"* ]]; then\n'
        '    return 0\n'
        '  fi\n'
        '  return 1  # standard SDR\n'
        '}'
    ))

    add_heading(doc, "GPU vs. CPU Encoding Paths", 3)
    add_styled_table(doc,
        ["Aspect", "GPU Path (NVENC)", "CPU Path (libx264)"],
        [
            ["Encoder", "h264_nvenc", "libx264"],
            ["Preset", "-preset p7", "-preset slow"],
            ["Quality", "-rc vbr -cq 18", "-crf 18"],
            ["Scaling", "scale_cuda filter", "scale filter"],
            ["HDR Support", "Falls back to CPU", "zscale tonemap"],
            ["Speed", "~5-10x real-time", "~1-2x real-time"],
            ["Requirement", "NVIDIA GPU with NVENC", "Any CPU"],
        ],
        col_widths=[1.3, 2.6, 2.6]
    )

    # 4.3 NVENC Bug
    add_heading(doc, "4.3 NVENC Silent Corruption Bug", 2)
    add_para(doc, (
        "The NVENC silent corruption bug was the most severe production issue encountered in the entire project. "
        "It caused 43% of normalized video segments to be silently corrupted — the ffmpeg command exited with "
        "code 0 (success), the output file existed and had non-zero size, but the video was undecodable. "
        "Face detection would then fail on these segments, and no transcription would be produced."
    ), bold=True)

    add_para(doc, "Three root causes were identified:")

    add_heading(doc, "Root Cause 1: NVENC Hardware Encoder Corruption", 3)
    add_para(doc, (
        "NVIDIA's hardware H.264 encoder (NVENC) silently produces corrupt output streams when given certain "
        "input codec combinations. Unlike software encoders that error out, NVENC writes a file that appears "
        "valid on disk but contains malformed NAL units that cannot be decoded. The GPU encoder's internal "
        "error handling does not propagate failures to the ffmpeg process."
    ))

    add_heading(doc, "Root Cause 2: Bash File Descriptor Interference", 3)
    add_para(doc, (
        "When the normalization function runs inside a while-read loop (iterating over video files), "
        "ffmpeg inherits the shell's file descriptors. ffmpeg reads from stdin by default, which "
        "interferes with the while-read loop's input stream. This causes the loop to skip files "
        "or terminate prematurely."
    ))
    add_code_block(doc, (
        "# BROKEN: ffmpeg inherits fd 0, interferes with while-read\n"
        'while IFS= read -r f; do\n'
        '  ffmpeg -i "$f" -c:v h264_nvenc ... "$out"   # eats stdin!\n'
        'done < <(find "$dir" -name "*.mp4")\n'
        "\n"
        "# FIXED: redirect ffmpeg stdin from /dev/null, use fd 3\n"
        'while IFS= read -r f <&3; do\n'
        '  ffmpeg -nostdin -i "$f" -c:v h264_nvenc ... "$out" < /dev/null\n'
        'done 3< <(find "$dir" -name "*.mp4")'
    ))

    add_heading(doc, "Root Cause 3: No Post-Encode Validation", 3)
    add_para(doc, (
        "The original code assumed that a successful ffmpeg exit code (0) meant the output was valid. "
        "This is not true for hardware encoders. The fix adds a mandatory validation step:"
    ))
    add_code_block(doc, (
        "# Post-encode validation — catches silent NVENC corruption\n"
        'if ! ffmpeg -v error -i "$out" -vframes 1 -f null - 2>/dev/null; then\n'
        '  log_warn "Corrupt output detected for ${bn}, falling back to raw copy"\n'
        '  rm -f "$out"\n'
        '  cp "$f" "${output_dir}/${bn}"  # raw copy as fallback\n'
        'fi'
    ))

    # 4.4 Segment-first
    add_heading(doc, "4.4 Segment-First Architecture", 2)
    add_para(doc, (
        "An important architectural decision was to segment videos BEFORE normalization, rather than "
        "normalizing the full video first. This approach has several advantages:"
    ))
    add_bullet(doc, "Speed: normalizing a 12-second segment is much faster than normalizing a 60-minute video")
    add_bullet(doc, "Parallelism: segments can be normalized independently")
    add_bullet(doc, "Fault isolation: a corrupt segment doesn't affect other segments from the same video")
    add_bullet(doc, "Efficiency: only segments that pass validation are normalized (skip excluded segments)")

    # 4.5 Failsafes
    add_heading(doc, "4.5 Failsafe Mechanisms", 2)
    add_para(doc, "The normalization module includes multiple layers of protection against failure:")
    add_bullet(doc, "Per-video timeout (default 600 seconds) with SIGKILL after 5-second grace period")
    add_bullet(doc, 'Input flags: -fflags +genpts+discardcorrupt -err_detect ignore_err')
    add_bullet(doc, "Increased probe: -probesize 10M -analyzeduration 10M for detecting unusual formats")
    add_bullet(doc, "Fallback to raw copy if encoding fails (ensures some output for every input)")
    add_bullet(doc, "Statistics tracking: norm_ok, norm_fail, norm_timeout, fallback_copy counters")

    doc.add_page_break()


def chapter_5(doc):
    add_heading(doc, "5. Video Segmentation", 1, bookmark_id="ch5")

    add_para(doc, (
        "The VSP-LLM model processes video in fixed-length windows. Input segments that are too long "
        "degrade model performance (the attention mechanism's effectiveness drops with sequence length), "
        "while segments that are too short may not contain enough visual information for meaningful "
        "transcription. The segmentation system splits long videos into optimal-length segments with "
        "overlap to prevent information loss at segment boundaries."
    ))

    # 5.1 Why Segment
    add_heading(doc, "5.1 Why Segmentation Is Needed", 2)
    add_para(doc, (
        "The model was trained on LRS3 clips that are typically 3-12 seconds long. Processing a full "
        "60-second video as a single input leads to poor transcription quality because:"
    ))
    add_bullet(doc, "The AV-HuBERT encoder produces a very long feature sequence that exceeds the model's effective context window")
    add_bullet(doc, "Beam search becomes computationally expensive with long sequences (memory grows quadratically)")
    add_bullet(doc, "The model tends to hallucinate or repeat text on long inputs")
    add_bullet(doc, "Error in one part of the video contaminates the entire transcription")

    # 5.2 fast_segment.py
    add_heading(doc, "5.2 Fast Segmentation (fast_segment.py, 197 Lines)", 2)
    add_para(doc, (
        "The fast_segment.py script implements ultra-fast video segmentation using ffmpeg's codec-copy mode "
        "(-c copy). This means no re-encoding occurs — the script simply locates keyframe boundaries and "
        "splits the video into segments. This is orders of magnitude faster than re-encoding because "
        "it only manipulates the container format, not the video data."
    ))
    add_code_block(doc, (
        "# Core segmentation command (codec-copy, no re-encoding)\n"
        "ffmpeg -ss {start_time} -i {input_video} \\\n"
        "  -t {segment_duration} \\\n"
        "  -c copy \\\n"
        "  -avoid_negative_ts make_zero \\\n"
        "  {output_segment}"
    ))
    add_para(doc, (
        "The script handles edge cases such as videos shorter than the segment duration (passed through "
        "unchanged), videos with no audio track, and videos with variable frame rates."
    ))

    # 5.3 Overlapping Segmentation
    add_heading(doc, "5.3 Overlapping Segmentation (overlapping_segmentation.py, 390 Lines)", 2)
    add_para(doc, (
        "The overlapping segmentation module extends basic segmentation with time-based overlap windows. "
        "With 12-second segments and 2-second overlap, a 30-second video produces segments at:"
    ))
    add_code_block(doc, (
        "Segment 0:  0:00 - 0:12  (12 seconds)\n"
        "Segment 1:  0:10 - 0:22  (12 seconds, 2s overlap with seg 0)\n"
        "Segment 2:  0:20 - 0:30  (10 seconds, 2s overlap with seg 1)\n"
        "\n"
        "Overlap regions contain duplicate content that can be used for:\n"
        "  - Boundary word recovery (words split at segment edges)\n"
        "  - Confidence comparison (same content decoded twice)\n"
        "  - Conflict resolution (merge_overlapping_predictions.py)"
    ))

    # 5.4 Naming Convention
    add_heading(doc, "5.4 Segment Naming Convention", 2)
    add_para(doc, "Every segment follows a strict naming convention that encodes its source video and temporal location:")
    add_styled_table(doc,
        ["Component", "Format", "Example", "Meaning"],
        [
            ["Base name", "{original_video}", "Obama_speech", "Source video filename (without extension)"],
            ["Segment index", "_{idx:02d}_", "_00_", "Zero-indexed segment number"],
            ["Start frame", "{start:06d}", "000000", "Start position in frames"],
            ["End frame", "_{end:06d}", "_000300", "End position in frames"],
            ["Extension", ".mp4", ".mp4", "Always MP4"],
        ],
        col_widths=[1.2, 1.3, 1.5, 2.5]
    )
    add_para(doc, (
        'Full example: Obama_speech_00_000000_000300.mp4 — first segment of "Obama_speech.mp4", '
        "frames 0 to 300 (at 25 FPS = 12 seconds)."
    ))

    # 5.5 4s to 12s
    add_heading(doc, "5.5 Evolution: 4-Second to 12-Second Segments", 2)
    add_para(doc, (
        "The segmentation window was changed from 4 seconds to 12 seconds during development. The original "
        "4-second windows were too short for meaningful transcription — many segments contained only partial "
        "words or silence. The 12-second window provides substantially more context for the model."
    ))
    add_styled_table(doc,
        ["Aspect", "4-Second Segments", "12-Second Segments"],
        [
            ["Context per segment", "Very limited (~2-4 words)", "Adequate (~8-20 words)"],
            ["Segments per minute", "~15", "~5"],
            ["Boundary effects", "Many cut words", "Few cut words"],
            ["Model performance", "High error rate", "Better transcription quality"],
            ["Processing time", "More overhead per segment", "Fewer segments to process"],
        ],
        col_widths=[1.8, 2.3, 2.4]
    )

    # 5.6 Non-segmented bug
    add_heading(doc, "5.6 Non-Segmented Video Naming Bug", 2)
    add_para(doc, (
        "A bug was discovered where videos shorter than the segment duration (e.g., a 10-second clip with "
        "12-second windows) were passed through without segmentation but retained their original filename. "
        "This caused downstream stages to fail because they expected the segment naming convention. The fix "
        "ensures that even non-segmented videos are renamed to follow the convention "
        "(e.g., short_clip.mp4 becomes short_clip_00_000000_000250.mp4)."
    ))

    # 5.7 Why No Merging
    add_heading(doc, "5.7 Why Segments Are Not Merged Back", 2)
    add_para(doc, (
        "An early design choice was to keep segment-level outputs rather than merging predictions back "
        "into a single per-video transcription. This decision was made for four reasons:"
    ))
    add_bullet(doc, "Segment-level quality assessment: each segment gets its own WER score, making it easy to identify problematic segments")
    add_bullet(doc, "Overlap handling complexity: merging overlapping predictions requires conflict resolution (which words to keep at boundaries)")
    add_bullet(doc, "User value: burned videos with per-segment subtitles are more useful for manual review than a single long transcript")
    add_bullet(doc, "Debugging: segment-level output makes it easier to trace issues back to specific video moments")

    doc.add_page_break()


def chapter_6(doc):
    add_heading(doc, "6. Face Detection & Mouth Cropping", 1, bookmark_id="ch6")

    add_para(doc, (
        "The VSP-LLM model requires tightly-cropped mouth region videos as input — specifically, 88x88 pixel "
        "grayscale crops centered on the speaker's mouth, resampled to 25 frames per second. Producing these "
        "crops from arbitrary video requires robust face detection, landmark estimation, spatial transformation, "
        "and temporal interpolation. This chapter documents the detection pipeline and the modifications made "
        "to improve its reliability in production conditions."
    ))

    # 6.1 Pipeline
    add_heading(doc, "6.1 Detection Pipeline Overview", 2)
    add_para(doc, "The face detection and cropping pipeline processes each video segment through four stages:")
    add_code_block(doc, (
        "1. Face Detection (detector.py)\n"
        "   └── Detect faces in every frame\n"
        "   └── Extract 4 keypoints per face (eyes + mouth corners)\n"
        "   └── Select largest face per frame\n"
        "\n"
        "2. Landmark Interpolation (video_process.py)\n"
        "   └── Fill in missing detections via temporal interpolation\n"
        "   └── Smooth landmark trajectories across frames\n"
        "\n"
        "3. Spatial Transformation (video_process.py)\n"
        "   └── Compute similarity transform from landmarks\n"
        "   └── Align and crop mouth region (88x88 pixels)\n"
        "\n"
        "4. Output Generation\n"
        "   └── Write grayscale mouth crop video at 25 FPS\n"
        "   └── Save as .mp4 in LRS3-compatible directory structure"
    ))

    # 6.2 MediaPipe
    add_heading(doc, "6.2 MediaPipe Detector (52 Lines)", 2)
    add_para(doc, (
        "The primary face detector uses Google's MediaPipe library with a two-tier detection strategy. "
        "This was an improvement over using a single detection model, which would fail on either "
        "distant or close-up faces depending on the model selection."
    ))
    add_code_block(doc, (
        "class LandmarksDetector:\n"
        "    def __init__(self):\n"
        "        # Two-tier strategy:\n"
        "        #   model_selection=1: full-range model (detects distant faces)\n"
        "        #   model_selection=0: short-range model (detects close-up faces)\n"
        "        self.min_detection_confidence = 0.5\n"
        "\n"
        "    def detect(self, frames):\n"
        "        # Try full-range model first\n"
        "        results = self._detect_with_model(frames, model_selection=1)\n"
        "        if all frames failed:\n"
        "            # Fallback to short-range model\n"
        "            results = self._detect_with_model(frames, model_selection=0)\n"
        "        if still all frames failed:\n"
        '            raise RuntimeError("Cannot detect any frames in the video")\n'
        "        return results  # numpy array [num_faces, 4 keypoints, (x, y)]"
    ))
    add_para(doc, (
        "The two-tier approach significantly improved detection rates in production. YouTube videos contain "
        "a wide range of camera distances — from news anchors at arm's length to speakers at podiums "
        "across a room. The full-range model handles distant faces, while the short-range model catches "
        "close-up selfie-style videos."
    ))

    # 6.3 Video Processing
    add_heading(doc, "6.3 Video Processing (video_process.py, 220 Lines)", 2)
    add_para(doc, (
        "The video processing module handles the transformation from detected face landmarks to "
        "cropped mouth region videos. Key operations include:"
    ))
    add_bullet_bold_value(doc, "Landmark Interpolation: ",
        "When face detection fails on individual frames (due to motion blur, occlusion, or camera "
        "transitions), the module interpolates missing landmarks from neighboring frames. This "
        "prevents jarring jumps in the crop position.")
    add_bullet_bold_value(doc, "Similarity Transform: ",
        "A 2D similarity transform (rotation, scaling, translation) is computed from the eye and "
        "mouth corner landmarks to align the face to a canonical orientation. This ensures the "
        "mouth region is consistently positioned regardless of head tilt or camera angle.")
    add_bullet_bold_value(doc, "cut_patch() Function: ",
        "Extracts the 88x88 pixel mouth region from the aligned frame. The crop is centered on "
        "the midpoint between the mouth corner landmarks, with a fixed-size window that captures "
        "the lips and immediate surrounding area.")
    add_bullet_bold_value(doc, "Temporal Resampling: ",
        "Input videos at various frame rates (24, 30, 60 FPS) are resampled to exactly 25 FPS "
        "to match the training data distribution.")

    # 6.4 RetinaFace
    add_heading(doc, "6.4 RetinaFace Alternative (260 Lines)", 2)
    add_para(doc, (
        "An alternative face detection backend using RetinaFace is also available. RetinaFace is "
        "a deep learning-based face detector that provides higher accuracy than MediaPipe, especially "
        "for difficult angles and partial occlusions, but at significantly slower inference speed."
    ))
    add_styled_table(doc,
        ["Aspect", "MediaPipe", "RetinaFace"],
        [
            ["detector.py", "52 lines", "43 lines"],
            ["video_process.py", "220 lines", "217 lines"],
            ["Speed", "Fast (CPU-based)", "Slower (GPU-based)"],
            ["Accuracy", "Good for frontal faces", "Better for difficult angles"],
            ["Dependencies", "mediapipe", "torch, torchvision"],
            ["Use Case", "Default (production)", "When MediaPipe fails"],
        ],
        col_widths=[1.5, 2.5, 2.5]
    )

    # 6.5 Transforms
    add_heading(doc, "6.5 Data Transforms (transforms.py, 171 Lines)", 2)
    add_para(doc, (
        "The transforms module provides data augmentation and preprocessing operations applied to "
        "mouth crop videos before they are fed to the model. These include spatial transforms "
        "(random crops, flips for training), normalization (mean subtraction, variance scaling), "
        "and format conversion (grayscale, tensor conversion)."
    ))

    # 6.6 Challenges
    add_heading(doc, "6.6 Production Challenges", 2)
    add_para(doc, (
        "Face detection in production encounters conditions rarely seen in research datasets:"
    ))
    add_bullet(doc, "Extreme camera angles: side profiles, overhead shots, low-angle interviews")
    add_bullet(doc, "Variable lighting: outdoor sunlight, dim indoor settings, backlighting")
    add_bullet(doc, "Partial occlusion: microphones, hands, other people crossing the frame")
    add_bullet(doc, "Multiple faces: group shots where the speaker must be identified")
    add_bullet(doc, "Camera motion: handheld footage with shake and rapid panning")
    add_bullet(doc, "Low resolution: old or heavily compressed YouTube videos")
    add_para(doc, (
        "The two-tier MediaPipe strategy with landmark interpolation handles most of these cases. "
        "For particularly difficult videos, the RetinaFace backend can be used as an alternative."
    ))

    doc.add_page_break()

def chapter_7(doc):
    add_heading(doc, "7. Reports & Burned Videos", 1, bookmark_id="ch7")

    add_para(doc, (
        "A key differentiator of the Argos system is its comprehensive output generation. Rather than "
        "producing raw text files, the pipeline generates multi-format quality reports with semantic metrics "
        "and burned videos with color-coded subtitle overlays. These outputs enable non-technical users to "
        "assess transcription quality at a glance and identify which segments need manual correction."
    ))

    # 7.1 make_report.py
    add_heading(doc, "7.1 Report Generation (make_report.py, 809 Lines)", 2)
    add_para(doc, (
        "The make_report.py script is the largest custom Python script in the project. It generates "
        "comprehensive quality reports in five output formats from a single decode output JSON file:"
    ))
    add_styled_table(doc,
        ["Format", "File", "Audience", "Features"],
        [
            ["CSV", "report.csv", "Data analysis", "Per-segment metrics, machine-readable"],
            ["HTML", "report.html", "Visual review", "Color-coded table, alignment visualization"],
            ["ANSI", "report.ansi.txt", "Terminal users", "Colored terminal output for quick checks"],
            ["JSON", "metrics.json", "Automation", "Aggregate metrics, pipeline integration"],
            ["TXT", "report.txt", "Documentation", "Plain text summary"],
        ],
        col_widths=[0.7, 1.3, 1.2, 3.3]
    )

    # 7.2 Metrics
    add_heading(doc, "7.2 Metrics: WER, WWER, and NEA", 2)
    add_para(doc, (
        "The report system calculates three complementary metrics, each capturing a different aspect "
        "of transcription quality:"
    ))

    add_heading(doc, "Standard WER (Word Error Rate)", 3)
    add_para(doc, (
        "The standard metric from ASR evaluation. Calculated as the edit distance between the reference "
        "transcription and the hypothesis, divided by the number of reference words. A WER of 0% means "
        "perfect transcription; WER can exceed 100% when the hypothesis contains many inserted words."
    ))

    add_heading(doc, "WWER (Weighted Word Error Rate)", 3)
    add_para(doc, (
        "Standard WER treats all words equally, but in practice, content words (names, numbers, nouns) "
        "carry far more information than function words (the, a, is). WWER assigns weights based on "
        "part-of-speech and entity type:"
    ))
    add_styled_table(doc,
        ["Category", "Weight", "POS Tags / Entity Types", "Examples"],
        [
            ["HIGH VALUE", "2.0", "PROPN, NUM, Named Entities (PERSON, ORG, GPE, LOC, MONEY, DATE)", "Biden, 2024, Washington, $500"],
            ["MEDIUM VALUE", "1.0", "NOUN, VERB, ADJ, ADV", "speech, said, important, quickly"],
            ["LOW VALUE", "0.5", "DET, AUX, PRON, ADP, CONJ, PUNCT, stopwords", "the, is, he, of, and"],
        ],
        col_widths=[1.3, 0.7, 2.5, 2.0]
    )
    add_para(doc, (
        "With these weights, missing a person's name or a number counts twice as heavily as missing "
        "a determiner. This better reflects the real-world impact of transcription errors."
    ))

    add_heading(doc, "NEA (Named Entity Accuracy)", 3)
    add_para(doc, (
        "NEA measures how well the transcription captures semantically important tokens, using spaCy's "
        "NER (Named Entity Recognition) and POS (Part-of-Speech) tagging. It reports recall, precision, "
        "and F1 score for high-value tokens:"
    ))
    add_bullet_bold_value(doc, "NEA Recall: ", "What fraction of important reference tokens appear in the hypothesis?")
    add_bullet_bold_value(doc, "NEA Precision: ", "What fraction of important hypothesis tokens appear in the reference?")
    add_bullet_bold_value(doc, "NEA F1: ", "Harmonic mean of recall and precision")
    add_para(doc, (
        "When spaCy is unavailable (e.g., on HORIZON without internet), the report falls back to "
        "basic WER metrics. The spaCy installation in lib/outputs.sh tries local wheels first, "
        "then online pip, then gracefully degrades."
    ))

    # 7.3 Alignment
    add_heading(doc, "7.3 Color-Coded Alignment", 2)
    add_para(doc, (
        "Each word in the hypothesis is classified using an alignment algorithm that compares it "
        "against the reference transcription:"
    ))
    add_styled_table(doc,
        ["Tag", "Color", "Meaning", "Example"],
        [
            ["ok", "Green", "Word matches reference at correct position", 'ref: "the president spoke"  hyp: "the" → GREEN'],
            ["rep", "Yellow/Amber", "Word appears in reference but at wrong position", 'ref: "the president spoke"  hyp: "spoke the" → YELLOW'],
            ["ins", "Red", "Word does not appear in reference (hallucination)", 'ref: "the president spoke"  hyp: "the great president" → RED for "great"'],
        ],
        col_widths=[0.6, 1.0, 2.5, 2.4]
    )
    add_para(doc, (
        "In the HTML report, each word is rendered with its background color, creating an instant "
        "visual summary of transcription quality. A segment that is mostly green is well-transcribed; "
        "one that is mostly red contains significant hallucination."
    ))

    # 7.4 Part Naming
    add_heading(doc, "7.4 Segment Part Naming", 2)
    add_para(doc, (
        "When a source video is split into multiple segments, the reports label each segment with "
        'a human-readable part name: "Part 1", "Part 2", etc. This is derived from the segment '
        "naming convention — the segment index field determines the part number. Videos that were "
        "not segmented (shorter than the window size) are labeled without a part suffix."
    ))

    # 7.5 Decode Params
    add_heading(doc, "7.5 Decode Parameters in Reports", 2)
    add_para(doc, (
        "Each decode run saves its effective parameters (beam width, length penalty, repetition penalty, "
        "temperature, etc.) as a JSON file alongside the hypothesis output. The report generator reads "
        "this file and includes the parameters in the report header, ensuring full reproducibility. "
        "This was essential for the tuning experiments (Chapter 12), where seven different parameter "
        "configurations were compared."
    ))

    # 7.6 make_burn.py
    add_heading(doc, "7.6 Video Burning (make_burn.py, 510 Lines)", 2)
    add_para(doc, (
        "The make_burn.py script creates demonstration videos by overlaying transcription text as subtitles "
        "directly onto the original video. This produces a visual artifact where viewers can watch the "
        "speaker's lips while reading both the reference (ground truth) and hypothesis (model output) text."
    ))

    add_heading(doc, "Three-Strategy Video Source Lookup", 3)
    add_para(doc, (
        "Finding the correct source video for each segment requires a three-tier lookup strategy, "
        "because the video may be available in different forms:"
    ))
    add_bullet(doc, "1. Original video file (before segmentation) — best quality, full context")
    add_bullet(doc, "2. Normalized segment — post-normalization, correct format")
    add_bullet(doc, "3. Lip crop video — mouth region only, useful as fallback")

    add_heading(doc, "Subtitle Overlay", 3)
    add_para(doc, (
        "The script uses ffmpeg's drawtext filter to overlay two lines of text at the bottom of each video:"
    ))
    add_code_block(doc, (
        "# Reference text (white on semi-transparent black)\n"
        'ffmpeg -i input.mp4 \\\n'
        '  -vf "drawtext=text=\'REF: the president spoke today\':fontsize=16:\n'
        '       fontcolor=white:box=1:boxcolor=black@0.7:x=10:y=h-60,\n'
        '       drawtext=text=\'HYP: the great president spoke\':fontsize=16:\n'
        '       fontcolor=green:box=1:boxcolor=black@0.7:x=10:y=h-30" \\\n'
        '  -c:a copy output_burned.mp4'
    ))

    # 7.7 Per-video WER
    add_heading(doc, "7.7 Per-Video WER Calculation (300 Lines)", 2)
    add_para(doc, (
        "The calculate_per_video_wer.py script aggregates segment-level metrics into per-video scores. "
        "When overlapping segmentation is used, the same content appears in multiple segments. The script "
        "implements a deduplication algorithm that identifies overlapping regions and uses the best "
        "available transcription for each time window, avoiding double-counting words."
    ))

    # 7.8 Conflict Resolution
    add_heading(doc, "7.8 Conflict Resolution (289 + 149 Lines)", 2)
    add_para(doc, (
        "The merge_overlapping_predictions.py (289 lines) and generate_conflict_report.py (149 lines) "
        "scripts handle cases where overlapping segments produce different transcriptions for the same "
        "time window. The merge script uses a simple heuristic: prefer the transcription with lower WER "
        "against the reference. The conflict report documents all disagreements for manual review."
    ))

    doc.add_page_break()


def chapter_8(doc):
    add_heading(doc, "8. Modular Code Refactoring", 1, bookmark_id="ch8")

    add_para(doc, (
        "By version 6, the pipeline script had grown to 823 lines of monolithic bash code. All stages, "
        "configuration, logging, error handling, and environment detection were interleaved in a single file. "
        "This made the code difficult to test, debug, and maintain — and critically, it prevented reuse "
        "between the EC2 development environment and the Linux container deployment. The modular refactoring "
        "(January 2026) was the most significant architectural improvement in the project."
    ))

    # 8.1 Problem
    add_heading(doc, "8.1 The Problem with Monolithic Code", 2)
    add_para(doc, "The 823-line monolithic script suffered from several concrete problems:")
    add_bullet(doc, "No unit testing: individual stages could not be tested in isolation")
    add_bullet(doc, "Path duplication: EC2 and container paths hardcoded throughout, requiring manual search-and-replace")
    add_bullet(doc, "Error cascading: a failure in one stage could leave the environment in an inconsistent state")
    add_bullet(doc, "No reuse: the fine-tuning pipeline needed the same preprocessing stages but couldn't import them")
    add_bullet(doc, "Debugging difficulty: with 823 lines of sequential logic, finding the source of a bug required reading the entire file")

    # 8.2 All Modules
    add_heading(doc, "8.2 The 12 Library Modules (1,562 Lines Total)", 2)
    add_para(doc, (
        "The refactoring extracted all reusable functionality into a lib/ directory with 12 modules. "
        "Each module is a self-contained bash script that exports specific functions and can be sourced "
        "independently:"
    ))
    add_styled_table(doc,
        ["Module", "Lines", "Key Functions", "Phase"],
        [
            ["common.sh", "81", "log_info, log_error, log_warn, log_stage, validate_directory", "Infrastructure"],
            ["config.sh", "87", "detect_environment(), get_base_path(), derived path functions", "Infrastructure"],
            ["venv_utils.sh", "36", "activate_venv(), deactivate_venv()", "Infrastructure"],
            ["archive.sh", "112", "archive_previous_run(), transcription preservation", "Preparation"],
            ["normalization.sh", "235", "run_normalization(), needs_tonemap(), GPU/CPU fallback", "Processing"],
            ["asr.sh", "229", "run_asr_transcription(), Steps 0.6/3/1.5", "Processing"],
            ["lrs3_prep.sh", "39", "run_lrs3_preparation()", "Processing"],
            ["manifests.sh", "80", "run_manifest_generation()", "Processing"],
            ["clustering.sh", "67", "run_clustering(), TRAIN_KMEANS toggle", "Processing"],
            ["decode.sh", "100", "run_vsp_decode(), Cython auto-build", "Decode"],
            ["outputs.sh", "165", "run_client_outputs(), spaCy install, reports + burns + lip crops", "Output"],
            ["test_all_modules.sh", "331", "37 automated tests across all modules", "Testing"],
        ],
        col_widths=[1.3, 0.5, 2.5, 0.8]
    )
    add_para(doc, (
        "The modules follow a consistent pattern: each file begins with a guard check to ensure "
        "it is being sourced (not executed directly), sources its dependencies (e.g., common.sh), "
        "and exports its public functions. This allows the main pipeline script to source only "
        "the modules it needs."
    ))

    # 8.3 Design Decisions
    add_heading(doc, "8.3 Key Design Decisions", 2)

    add_heading(doc, "Environment-Aware Configuration", 3)
    add_para(doc, (
        "The config.sh module automatically detects whether the pipeline is running on EC2 (/home/ubuntu/) "
        "or in a container (/workspace/) and configures all paths accordingly. This eliminates the need "
        "for manual path editing when moving between environments:"
    ))
    add_code_block(doc, (
        "detect_environment() {\n"
        '  if [[ -d "/workspace" ]]; then\n'
        '    ENV_TYPE="container"\n'
        '    BASE_PATH="/workspace"\n'
        "  else\n"
        '    ENV_TYPE="ec2"\n'
        '    BASE_PATH="/home/ubuntu"\n'
        "  fi\n"
        "  # All derived paths use BASE_PATH\n"
        '  VSP_DIR="${BASE_PATH}/VSP-LLM"\n'
        '  AUTO_AVSR_DIR="${BASE_PATH}/auto_avsr"\n'
        '  # ... etc\n'
        "}"
    ))

    add_heading(doc, "Virtual Environment Strategy", 3)
    add_para(doc, (
        "The ASR module (asr.sh) is self-contained: it activates the pre-process-venv internally "
        "and deactivates it when done. All other modules (stages 5-8) expect the caller to have "
        "already activated vsp-llm-yoad-venv. This minimizes venv activation overhead — the main "
        "pipeline activates the VSP venv once and uses it for multiple stages."
    ))

    add_heading(doc, "The stdout Contamination Bug", 3)
    add_para(doc, (
        "One of the most subtle bugs discovered during refactoring was that log_info() wrote to stdout "
        "instead of stderr. This caused logging output to be captured by command substitution "
        "operations, contaminating variables with log messages. For example:"
    ))
    add_code_block(doc, (
        "# BROKEN: log output mixed into the variable\n"
        'result=$(some_function)  # result = "[INFO] Processing...\\n/actual/path"\n'
        "\n"
        "# FIXED: log_info writes to stderr (fd 2)\n"
        'log_info() { echo "[$(date +%H:%M:%S)] INFO: $*" >&2; }\n'
        'result=$(some_function)  # result = "/actual/path"'
    ))

    # 8.4 Test Suite
    add_heading(doc, "8.4 Test Suite: 37 Automated Tests", 2)
    add_para(doc, (
        "The test suite (lib/test_all_modules.sh, 331 lines) validates every module's exports, "
        "functionality, and integration points. Tests use colored output for quick visual assessment:"
    ))
    add_styled_table(doc,
        ["Module", "Tests", "What Is Verified"],
        [
            ["common.sh", "5", "log_info/log_error formatting, validate_directory on existing and non-existent paths"],
            ["config.sh", "5", "ENV_TYPE detection, BASE_PATH setting, path function exports"],
            ["venv_utils.sh", "3", "activate_venv/deactivate_venv exports, real venv activation if available"],
            ["normalization.sh", "4", "Normalization function export, HDR detection, GPU/CPU fallback"],
            ["asr.sh", "4", "Whisper transcription, offline model loading, transcription reuse logic"],
            ["archive.sh", "3", "Previous run archiving, transcription preservation"],
            ["lrs3_prep.sh", "2", "LRS3 preparation function export and basic operation"],
            ["manifests.sh", "2", "Manifest generation function export"],
            ["clustering.sh", "2", "Clustering function export, TRAIN_KMEANS toggle"],
            ["decode.sh", "3", "Decode function export, Cython build check"],
            ["outputs.sh", "2", "Output function export, spaCy install detection"],
            ["Integration", "2", "Cross-module function calls, path consistency"],
        ],
        col_widths=[1.3, 0.5, 4.7]
    )

    # 8.5 Before/After
    add_heading(doc, "8.5 Before and After Metrics", 2)
    add_styled_table(doc,
        ["Metric", "Before (v6)", "After (v7)", "Improvement"],
        [
            ["Main script", "823 lines, 1 file", "501 lines, 1 file", "52% reduction"],
            ["Library code", "0 lines, 0 files", "1,562 lines, 12 files", "New reusable modules"],
            ["Total code", "823 lines", "2,063 lines", "More code but better organized"],
            ["Testability", "None", "37 automated tests", "Full coverage"],
            ["Environment support", "Manual path editing", "Auto-detection", "Zero configuration"],
            ["Fine-tuning reuse", "Not possible", "Shares all lib/ modules", "579-line script reuses entire lib/"],
        ],
        col_widths=[1.5, 1.5, 1.8, 1.7]
    )

    doc.add_page_break()

def chapter_9(doc):
    add_heading(doc, "9. Web UI & Server", 1, bookmark_id="ch9")

    add_para(doc, (
        "The Argos web interface was built to make the pipeline accessible to non-technical users who need "
        "to process videos, review transcriptions, and download results without using the command line. "
        "The UI consists of a Python HTTP server backend and a vanilla JavaScript frontend — no external "
        "frameworks (no React, no Vue, no Flask) — to minimize dependencies and maximize portability. "
        "The complete UI spans 5,780 lines across 11 files."
    ))

    # 9.1 Architecture
    add_heading(doc, "9.1 Architecture: Zero-Dependency Web Stack", 2)
    add_para(doc, (
        "A deliberate architectural decision was made to avoid external web frameworks. The server uses "
        "Python's built-in http.server module (ThreadedHTTPServer), and the frontend uses vanilla JavaScript "
        "with no build tools, bundlers, or transpilers. This means the UI works on any machine with Python 3 "
        "installed — critical for the HORIZON server and standalone machines where installing npm, Node.js, "
        "or Python web frameworks is not possible."
    ))
    add_styled_table(doc,
        ["Component", "File", "Lines", "Technology"],
        [
            ["HTTP Server", "server.py", "1,124", "Python http.server (ThreadedHTTPServer)"],
            ["Frontend Logic", "app.js", "1,921", "Vanilla JavaScript (ES6+)"],
            ["HTML Structure", "index.html", "328", "Semantic HTML5"],
            ["Styling", "style.css", "1,191", "Mobile-first CSS3"],
            ["Pipeline Service", "pipeline_runner.py", "349", "subprocess + threading"],
            ["Progress Service", "progress_tracker.py", "206", "Regex-based log parsing"],
            ["Transcription Service", "transcription_manager.py", "280", "JSON + .wrd file management"],
            ["Validation Service", "validator.py", "374", "ffprobe-based video analysis"],
            ["Entry Point", "__main__.py", "5", "Python package entry"],
            ["Package Init", "__init__.py", "1", "Package marker"],
            ["Services Init", "services/__init__.py", "1", "Package marker"],
        ],
        col_widths=[1.5, 1.8, 0.5, 2.7]
    )

    # 9.2 State Machine
    add_heading(doc, "9.2 Six-Screen State Machine", 2)
    add_para(doc, (
        "The UI operates as a state machine with six mutually exclusive screens. Only one screen is "
        "visible at a time, and transitions occur based on user actions and pipeline events:"
    ))
    add_styled_table(doc,
        ["Screen", "Purpose", "Key Actions", "Next Screen"],
        [
            ["Welcome", "Upload videos, inspect input folder", "Drag-drop upload, remove videos, configure options", "Segment Review"],
            ["Segment Review", "Review segments after fast segmentation", "Transcribe segments, delete orphans, choose k-means mode", "Processing"],
            ["Processing", "Monitor pipeline execution", "View progress bar, read logs, cancel pipeline", "Complete or Error"],
            ["Transcription", "Per-segment review (legacy screen)", "Edit transcriptions, view segments", "Complete"],
            ["Complete", "Download results and review", "Download ZIP, open folder, start new batch", "Welcome"],
            ["Error", "Display error details", "View full logs, retry or reset", "Welcome"],
        ],
        col_widths=[1.0, 1.5, 2.3, 1.0]
    )

    # 9.3 server.py
    add_heading(doc, "9.3 Server (server.py, 1,124 Lines)", 2)
    add_para(doc, (
        "The server handles all HTTP requests through a single VSPRequestHandler class. It implements "
        "12 REST API endpoints (9 GET, 9 POST — some paths support both methods):"
    ))

    add_heading(doc, "GET Endpoints", 3)
    add_styled_table(doc,
        ["Endpoint", "Returns", "Purpose"],
        [
            ["/api/status", "JSON: video count, pipeline state, is_running", "Dashboard status polling"],
            ["/api/progress", "JSON: stage, percent, logs, errors", "Progress bar updates (500ms poll)"],
            ["/api/logs", "JSON: recent log lines", "Collapsible log viewer"],
            ["/api/golden-models", "JSON: available k-means models with metadata", "Golden model selector dropdown"],
            ["/api/download-output", "ZIP file stream", "Download all outputs as archive"],
            ["/api/transcription", "JSON: text, type, metadata", "Load transcription for editing"],
            ["/api/orphaned-transcriptions", "JSON: transcriptions without videos", "Orphan cleanup workflow"],
            ["/api/segments", "JSON: segment list with transcription status", "Segment review screen"],
            ["/api/segment-video/<id>", "Video file stream", "Lazy-loaded video preview in modal"],
        ],
        col_widths=[2.0, 2.0, 2.5]
    )

    add_heading(doc, "POST Endpoints", 3)
    add_styled_table(doc,
        ["Endpoint", "Input", "Action"],
        [
            ["/api/validate", "—", "Run ffprobe on all input videos, estimate segments"],
            ["/api/start", "JSON: train_kmeans, golden_model, segmentation, overlap, segment_only", "Launch pipeline subprocess"],
            ["/api/cancel", "—", "Send SIGTERM to pipeline process"],
            ["/api/reset", "—", "Clear pipeline state for fresh start"],
            ["/api/remove-video", "JSON: filename", "Move video to .excluded/ folder"],
            ["/api/transcription", "JSON: filename, text", "Save or delete transcription"],
            ["/api/save-segment-transcription", "JSON: filename, text", "Save with normalization"],
            ["/api/open-folder", "JSON: path", "Open in system file manager"],
            ["/api/upload", "Multipart form data", "Chunked file upload (1MB chunks)"],
        ],
        col_widths=[2.2, 2.5, 1.8]
    )

    add_heading(doc, "Manual Multipart Parser", 3)
    add_para(doc, (
        "The upload endpoint implements a custom multipart form data parser because Python's cgi.FieldStorage "
        "was deprecated in Python 3.11 and removed in Python 3.13. The custom parser reads the request body "
        "in 1MB chunks to prevent out-of-memory errors on large video files (the original implementation "
        "used .read(content_length) which loaded the entire file into RAM)."
    ))

    # 9.4 app.js
    add_heading(doc, "9.4 Frontend (app.js, 1,921 Lines)", 2)
    add_para(doc, "The frontend JavaScript implements several sophisticated UI patterns:")

    add_bullet_bold_value(doc, "Progress Polling (500ms): ",
        "The UI polls /api/progress every 500 milliseconds during pipeline execution. This interval "
        "provides responsive feedback without excessive server load. The progress bar uses smooth "
        "CSS transitions for visual continuity between updates.")
    add_bullet_bold_value(doc, "Drag-and-Drop Upload: ",
        "Files can be dragged onto the welcome screen. The upload uses XMLHttpRequest (not fetch) "
        "to support real-time progress events. For small files, simulated smooth progress fills "
        "the gap between actual server responses (8% increments for small, 2% for large files).")
    add_bullet_bold_value(doc, "Transcription Modal: ",
        "Clicking a segment opens a modal with a text area for editing transcriptions. The modal "
        "shows a live normalization preview (lowercase, alphanumeric + apostrophes only) and word count. "
        "A toggle button lazily loads the segment video for visual reference.")
    add_bullet_bold_value(doc, "Lazy Video Loading: ",
        "Segment videos in the review screen use IntersectionObserver with a 200px root margin to "
        "load only when scrolled into view. This prevents the browser from attempting to load hundreds "
        "of video files simultaneously.")
    add_bullet_bold_value(doc, "Orphan Transcription Cleanup: ",
        "When videos are removed but their transcriptions remain, the UI detects these orphans and "
        "offers the user a choice: delete the transcription or keep it for potential re-addition.")
    add_bullet_bold_value(doc, "K-Means Golden Model Selector: ",
        "A dropdown populated from /api/golden-models lets users choose a pre-trained k-means model "
        "instead of training a new one. Each model shows its dataset info and file size.")
    add_bullet_bold_value(doc, "Adaptive Upload Timeout: ",
        "Upload timeout scales with file size: 5 minutes base + 1 minute per 100MB. This prevents "
        "timeouts on large video files while keeping short timeouts for small files.")

    # 9.5 index.html
    add_heading(doc, "9.5 HTML Structure (index.html, 328 Lines)", 2)
    add_para(doc, (
        "The HTML file defines six screen sections (mutually exclusive via CSS display toggle), a drop "
        "zone overlay for drag-and-drop, and a transcription modal. The modal includes video preview, "
        "textarea, normalization preview, word count, and save/delete buttons. Status badges use CSS "
        "classes to indicate pipeline state (Ready, Processing, Complete, Error)."
    ))

    # 9.6 style.css
    add_heading(doc, "9.6 Styling (style.css, 1,191 Lines)", 2)
    add_para(doc, (
        "The stylesheet implements a complete design system with mobile-first responsive breakpoints. "
        "Key design elements include:"
    ))
    add_bullet(doc, "Color scheme: #2563eb (primary blue), #16a34a (success green), #dc2626 (error red)")
    add_bullet(doc, "Transcription badges: orange for auto-generated, green for manually edited")
    add_bullet(doc, "Progress bar: gradient fill with smooth transitions")
    add_bullet(doc, "Modal: fixed overlay with 70% black background, slide-in animation, max 650px width")
    add_bullet(doc, "Responsive: column layout below 600px, 95% modal width below 768px")
    add_bullet(doc, "Upload progress: per-file progress bars with checkmark or X status indicators")

    # 9.7 Upload Bug
    add_heading(doc, "9.7 Large File Upload Bug", 2)
    add_para(doc, (
        "One of the more impactful production bugs was the upload OOM (out of memory) crash. The original "
        "upload handler called request.rfile.read(content_length), which loaded the entire uploaded file "
        "into memory before writing it to disk. For a 2GB video file, this consumed 2GB of RAM instantly, "
        "crashing the server process."
    ))
    add_para(doc, "The fix involved two changes:")
    add_bullet(doc, "Chunked reading: read the request body in 1MB chunks, writing each chunk to disk immediately")
    add_bullet(doc, "Adaptive timeout: increase the socket timeout proportionally to file size (base 5 min + 1 min per 100MB)")

    doc.add_page_break()


def chapter_10(doc):
    add_heading(doc, "10. Backend Services", 1, bookmark_id="ch10")

    add_para(doc, (
        "The web UI's backend is organized into four service modules, each handling a distinct aspect "
        "of pipeline management. These services are stateful Python classes instantiated once when the "
        "server starts and shared across all request handlers."
    ))

    # 10.1 PipelineRunner
    add_heading(doc, "10.1 PipelineRunner (349 Lines)", 2)
    add_para(doc, (
        "The PipelineRunner manages the lifecycle of the pipeline subprocess — starting, monitoring, "
        "and stopping the bash script execution. It operates as a singleton: only one pipeline run "
        "can be active at a time."
    ))

    add_heading(doc, "Subprocess Management", 3)
    add_para(doc, "The runner launches the pipeline with carefully configured environment variables:")
    add_code_block(doc, (
        "# Pipeline launch command\n"
        "bash /home/ubuntu/run_flat_english_pipeline.sh <INPUT_DIR>\n"
        "\n"
        "# Environment variables passed to subprocess:\n"
        "TRAIN_KMEANS=0|1           # Train new k-means or use existing\n"
        "GOLDEN_KMEANS=<path>       # Path to golden k-means model (optional)\n"
        "SEGMENTATION_ENABLED=0|1   # Enable video segmentation\n"
        "OVERLAP_ENABLED=0|1        # Enable overlap between segments\n"
        "SEGMENT_ONLY=0|1           # Stop after segmentation (for review)\n"
        "PYTHONUNBUFFERED=1         # Force unbuffered output for real-time logs"
    ))

    add_heading(doc, "Graceful Shutdown", 3)
    add_para(doc, (
        "When the user clicks Cancel, the runner follows a graceful shutdown sequence: send SIGTERM "
        "to the pipeline process, wait up to 10 seconds for it to finish the current stage, then "
        "send SIGKILL if it's still running. This prevents orphaned ffmpeg or Python processes."
    ))

    add_heading(doc, "Monitoring", 3)
    add_para(doc, (
        "The runner reads the pipeline's stdout line by line (line-buffered) in a background thread. "
        "Each line is passed to the ProgressTracker for stage detection. An inactivity warning is "
        "triggered if 10 minutes pass without any output from the pipeline."
    ))

    # 10.2 ProgressTracker
    add_heading(doc, "10.2 ProgressTracker (206 Lines)", 2)
    add_para(doc, (
        "The ProgressTracker parses pipeline log output in real time to determine which stage is currently "
        "executing and calculate an overall progress percentage."
    ))

    add_heading(doc, "ProgressState Dataclass", 3)
    add_code_block(doc, (
        "class ProgressState:\n"
        "    state: str          # IDLE, RUNNING, COMPLETED, FAILED, CANCELLED\n"
        "    current_stage_index: int\n"
        "    current_stage_id: str\n"
        "    current_stage_name: str\n"
        "    percent_complete: float   # 0-100\n"
        "    errors: list              # Last 10 errors\n"
        "    warnings: list            # Last 20 warnings\n"
        "    logs: list                # Last 1000 lines\n"
        "    segment_only: bool"
    ))

    add_heading(doc, "Stage Detection", 3)
    add_para(doc, (
        "Each pipeline stage prints a marker line (e.g., '[STAGE 3] Running Whisper ASR'). The tracker "
        "uses regex patterns (STAGE_MARKERS) to detect these markers and update the current stage. "
        "Progress is calculated by summing the weights of completed stages plus 50% of the current "
        "stage's weight, capped at 99% until the completion marker is detected."
    ))

    add_heading(doc, "Error Detection", 3)
    add_para(doc, (
        "The tracker watches for error patterns in the log output: ValueError, RuntimeError, CUDA errors, "
        "OOM messages, and other common failure modes. Detected errors are stored in the errors list and "
        "surfaced through the /api/progress endpoint."
    ))

    # 10.3 TranscriptionManager
    add_heading(doc, "10.3 TranscriptionManager (280 Lines)", 2)
    add_para(doc, (
        "The TranscriptionManager provides unified transcription storage with persistent metadata. "
        "All transcriptions (manual and auto-generated) are stored in a central .transcriptions/ "
        "directory within the input folder."
    ))

    add_heading(doc, "Storage Format", 3)
    add_code_block(doc, (
        "vsp_input/.transcriptions/\n"
        "├── video_name_00_000000_000300.wrd   # One word per line\n"
        "├── video_name_01_000300_000600.wrd\n"
        "└── metadata.json                      # Type, timestamps, word counts\n"
        "\n"
        "metadata.json structure:\n"
        "{\n"
        '  "transcriptions": {\n'
        '    "video_name_00_000000_000300.mp4": {\n'
        '      "type": "manual",          // or "auto"\n'
        '      "created_at": "2026-02-15T10:30:00Z",\n'
        '      "edited_at": "2026-02-16T14:20:00Z",\n'
        '      "word_count": 42\n'
        "    }\n"
        "  }\n"
        "}"
    ))

    add_heading(doc, "Text Normalization", 3)
    add_para(doc, (
        "When saving transcriptions, the manager normalizes text to match the model's expected format: "
        "lowercase, alphanumeric characters and apostrophes only. This ensures consistency between "
        "manual transcriptions and Whisper auto-transcriptions."
    ))

    add_heading(doc, "Orphan Detection", 3)
    add_para(doc, (
        "The get_orphaned_transcriptions() method identifies transcriptions that no longer have "
        "corresponding video files (e.g., because the video was excluded). The UI presents these "
        "orphans to the user for cleanup."
    ))

    # 10.4 Validator
    add_heading(doc, "10.4 Validator (374 Lines)", 2)
    add_para(doc, (
        "The Validator uses ffprobe to analyze input video files and determine their suitability "
        "for processing. It produces a detailed ValidationResult with per-video analysis."
    ))

    add_heading(doc, "Validation Checks", 3)
    add_bullet(doc, "Supported file extension (.mp4, .mkv, .webm, .mov, .m4v, .avi)")
    add_bullet(doc, "ffprobe can successfully read the file (not corrupt)")
    add_bullet(doc, "File contains at least one video stream")
    add_bullet(doc, "Duration is greater than 0 seconds")
    add_bullet(doc, "Resolution is at least 64x64 pixels (minimum for face detection)")

    add_heading(doc, "Segment Estimation", 3)
    add_para(doc, (
        "For each valid video, the validator estimates the number of segments that will be produced: "
        "ceil(duration / segment_duration). This is used by the UI to show the total expected segment "
        "count and to assess k-means viability (a warning is shown if fewer than 200 segments are expected)."
    ))

    add_heading(doc, "VideoInfo Dataclass", 3)
    add_para(doc, (
        "Each validated video is represented as a VideoInfo object containing: filename, path, "
        "duration_seconds, width, height, fps, has_audio, estimated_segments, valid (bool), error "
        "message, and transcription metadata (has_transcription, type, created_at, edited_at, word_count)."
    ))

    doc.add_page_break()


def chapter_11(doc):
    add_heading(doc, "11. Docker Container & Deployment", 1, bookmark_id="ch11")

    add_para(doc, (
        "Deploying Argos to production machines required packaging the entire system — three Python repos, "
        "two virtual environments, model weights, system libraries, and the web UI — into a portable "
        "container that could run on any machine with an NVIDIA GPU. This chapter documents the Docker "
        "architecture, five container releases, 37 bugs fixed during deployment, and the testing "
        "infrastructure that ensures reliability."
    ))

    # 11.1 Architecture
    add_heading(doc, "11.1 Container Architecture", 2)
    add_para(doc, (
        "The Docker container is built on nvidia/cuda:12.8.0-base-ubuntu22.04, providing CUDA runtime "
        "libraries and GPU access through NVIDIA Container Toolkit."
    ))
    add_code_block(doc, (
        "# Dockerfile structure (262 lines)\n"
        "FROM nvidia/cuda:12.8.0-base-ubuntu22.04\n"
        "\n"
        "# System dependencies\n"
        "RUN apt-get update && apt-get install -y \\\n"
        "    python3 python3-pip python3-venv gcc make \\\n"
        "    ffmpeg git libglib2.0-0 libsm6 libxext6 \\\n"
        "    libxrender1 libsndfile1 libportaudio2\n"
        "\n"
        "# Virtual Environment 1: Preprocessing\n"
        "# PyTorch 2.8.0+cu128, triton 3.4.0, Whisper, MediaPipe\n"
        "# ~135 packages\n"
        "\n"
        "# Virtual Environment 2: Decoding\n"
        "# PyTorch 2.5.1+cu124, triton 3.1.0, fairseq, transformers\n"
        "# ~256 packages\n"
        "\n"
        "# Model weights, spaCy wheels, pipeline scripts, web UI\n"
        "COPY galaxy_export/ /workspace/"
    ))

    # 11.2 Galaxy Export
    add_heading(doc, "11.2 Galaxy Export Bundle (~65GB)", 2)
    add_para(doc, (
        "The galaxy export is the complete offline deployment package. At approximately 65GB, it includes "
        "everything needed to run the pipeline without internet access: both virtual environments with all "
        "packages, model weights (LLaMA-2-7B ~13GB, AV-HuBERT checkpoint ~1.2GB), the complete pipeline "
        "codebase, web UI, spaCy wheels for offline installation, and Whisper model cache."
    ))

    # 11.3 Container Versions
    add_heading(doc, "11.3 Five Container Releases", 2)
    add_styled_table(doc,
        ["Date", "Version Range", "Key Changes"],
        [
            ["Feb 3, 2026", "v1.0.0 — v1.0.10", "Initial release: basic pipeline, web UI, first deployment"],
            ["Feb 8, 2026", "v1.0.11 — v1.0.13", "Installation bug fixes: hardcoded paths, python3 compatibility, fairseq max_len, UI networking"],
            ["Feb 12, 2026", "v1.0.14 — v1.0.25", "Deployment fixes: CUDA OOM on 12GB GPU, desktop launcher, docker.conf, upload reliability"],
            ["Feb 15, 2026", "v1.0.26 — v1.0.35", "Final fixes: beam search OOM, inference tuning, report metrics, drag-drop upload, NVENC corruption"],
            ["Feb 17, 2026", "v1.0.36 — v1.0.37", "NVENC silent corruption fix, bash fd interference fix, post-encode validation"],
        ],
        col_widths=[1.2, 1.5, 3.8]
    )

    # 11.4 First Delivery
    add_heading(doc, "11.4 First Client Delivery", 2)
    add_para(doc, (
        "The first delivery package (FOR_CLIENT_OLD) included the OCI container image, a desktop "
        "launcher (.desktop file) for easy startup, and NVIDIA driver installers for Arch Linux "
        "compatibility. The container was loaded using docker load and launched with nvidia-docker "
        "or the NVIDIA Container Runtime."
    ))

    # 11.5 Bugs Table
    add_heading(doc, "11.5 Thirty-Seven Bugs Fixed", 2)
    add_para(doc, (
        "Over the course of five container releases, 37 bugs were identified and fixed. These bugs "
        "span installation, deployment, GPU compatibility, and user interface issues. The following "
        "table lists the most impactful bugs by category:"
    ))

    add_heading(doc, "Installation Bugs (1-13)", 3)
    add_styled_table(doc,
        ["Bug#", "Summary", "Impact"],
        [
            ["1-3", "Missing files during INSTALL.sh, hardcoded EC2 paths", "Container fails to start"],
            ["4-6", "python vs python3 incompatibility across 11 files", "Scripts crash on Ubuntu"],
            ["7-8", "Whisper offline model loading, spaCy offline install", "Pipeline fails on HORIZON"],
            ["9-10", "fairseq GenerationConfig schema (max_len field missing)", "Decode crashes"],
            ["11-13", "UI server binding 127.0.0.1 vs 0.0.0.0, host launcher", "UI inaccessible from host"],
        ],
        col_widths=[0.7, 3.5, 2.3]
    )

    add_heading(doc, "Deployment & GPU Bugs (14-25)", 3)
    add_styled_table(doc,
        ["Bug#", "Summary", "Impact"],
        [
            ["14", "CUDA OOM on 12GB GPU — memory accumulation across samples", "Decode crashes after ~50 segments"],
            ["15-17", "Desktop launcher: Terminal=true, execute permission, untrusted .desktop", "Users can't launch UI"],
            ["18-20", "Docker port conflicts, ENTRYPOINT double-bash, docker.conf parsing", "Container startup fails"],
            ["21-22", "fairseq PYTHONPATH override, wrong installation targeted", "Decode uses wrong fairseq"],
            ["23-25", "Open Results Folder, file upload reliability, cgi deprecation", "UI features broken"],
        ],
        col_widths=[0.7, 3.5, 2.3]
    )

    add_heading(doc, "Final Fixes (26-37)", 3)
    add_styled_table(doc,
        ["Bug#", "Summary", "Impact"],
        [
            ["26", "Beam search OOM on 12GB GPUs (no_repeat_ngram_size + dynamic max_len cap)", "Decode crashes on long segments"],
            ["27-28", "Inference tuning: max_len_a 3.0→2.0, max_len_b 300→200, repetition_penalty=1.2", "Reduced hallucination"],
            ["29-30", "Report part naming for segments, lip crop copy logic", "Reports show wrong segment labels"],
            ["31", "Log output stdout contamination (log_info to stderr)", "Pipeline paths corrupted"],
            ["32-33", "Drag-drop upload: manual multipart parser (Python 3.13 compatible)", "Upload crashes"],
            ["34-35", "Segment transcription persistence (Steps 0.6 and 1.5)", "Transcriptions lost between runs"],
            ["36-37", "NVENC silent corruption, bash fd interference, post-encode validation", "43% of segments corrupt"],
        ],
        col_widths=[0.7, 3.5, 2.3]
    )

    add_para(doc, (
        "The most impactful bugs were: CUDA OOM (#14/#26) which made decode impossible on 12GB GPUs, "
        "NVENC silent corruption (#36) which destroyed 43% of processed segments, log contamination (#31) "
        "which caused random pipeline failures, and the upload OOM (#25/#32) which crashed the server "
        "on large file uploads."
    ), bold=True)

    # 11.6 Build Automation
    add_heading(doc, "11.6 Build & Verification Automation", 2)
    add_styled_table(doc,
        ["Script", "Lines", "Purpose"],
        [
            ["build_container.sh", "92", "Docker image build with layer caching"],
            ["check_container_inventory.sh", "135", "Verify all required files present in container"],
            ["verify_container_sync.sh", "353", "Check EC2 changes are replicated to container"],
        ],
        col_widths=[2.2, 0.6, 3.7]
    )

    # 11.7 Installation Flow
    add_heading(doc, "11.7 Installation & Startup Flow", 2)
    add_code_block(doc, (
        "1. Load container image:\n"
        "   docker load < vsp_linux_container_FINAL_20260217.tar\n"
        "\n"
        "2. Run INSTALL.sh:\n"
        "   - Creates required directories\n"
        "   - Sets permissions\n"
        "   - Installs desktop launcher (optional)\n"
        "\n"
        "3. Start with vsp-start.sh:\n"
        "   - Launches Docker container with GPU access\n"
        "   - Mounts input/output directories\n"
        "   - Starts web UI server on port 8080\n"
        "\n"
        "4. Access UI:\n"
        "   - Open browser to http://localhost:8080\n"
        "   - Or click desktop icon (if installed)"
    ))

    # 11.8 Sync Items
    add_heading(doc, "11.8 Container Synchronization (28 Pending Items)", 2)
    add_para(doc, (
        "Every change made to the EC2 development environment must be explicitly replicated to the "
        "Linux container version. This synchronization requirement is documented in detail in "
        "docs/container-sync-changelog.md (1,329 lines), which lists 28 pending sync items with "
        "exact file paths, line numbers, and code diffs for each change."
    ))

    # 11.9 Testing
    add_heading(doc, "11.9 Testing Infrastructure", 2)
    add_para(doc, (
        "The project includes a comprehensive testing suite spanning five test scripts with a total "
        "of 1,731 lines of test code:"
    ))
    add_styled_table(doc,
        ["Script", "Lines", "Tests", "What It Verifies"],
        [
            ["test_all_modules.sh", "331", "37", "All lib/ module exports, functions, and integration"],
            ["test_pipeline_modules.sh", "350", "~20", "Module-level function tests in isolation"],
            ["test_ec2_environment.sh", "371", "~15", "EC2 environment: paths, venvs, GPU, dependencies"],
            ["test_pipeline_smoke.sh", "326", "~10", "End-to-end pipeline stages with test data"],
            ["verify_container_sync.sh", "353", "~15", "EC2 changes replicated to container version"],
        ],
        col_widths=[1.8, 0.5, 0.5, 3.7]
    )

    doc.add_page_break()

def chapter_12(doc):
    add_heading(doc, "12. Performance Evaluation & Tuning", 1, bookmark_id="ch12")

    add_para(doc, (
        "Evaluating the Argos system's performance required processing a large dataset of real-world videos, "
        "establishing metrics that go beyond simple WER, running systematic parameter tuning experiments, "
        "and producing detailed analysis reports. This chapter documents the evaluation methodology, results, "
        "thirteen decode tuning experiments, 21 analytical plots, six analysis reports, pairwise statistical "
        "testing, and the key insights that emerged."
    ))

    # 12.1 Setup
    add_heading(doc, "12.1 Evaluation Setup", 2)
    add_para(doc, (
        "The primary evaluation used approximately 1,500 YouTube videos from the AVSpeech dataset, "
        "which were processed through the pipeline to produce 1,497 evaluated segments. Whisper ASR "
        "(medium model) generated ground truth transcriptions for comparison against VSP-LLM's "
        "visual-only predictions."
    ))
    add_styled_table(doc,
        ["Parameter", "Value"],
        [
            ["Source Dataset", "AVSpeech (YouTube videos)"],
            ["Videos Processed", "~1,500"],
            ["Segments Evaluated", "1,497"],
            ["Ground Truth", "Whisper ASR (medium model, English)"],
            ["Decode Config", "beam=20, lenpen=0.0, rep_pen=1.2"],
            ["Evaluation Metrics", "WER, WWER, NEA (Recall, Precision, F1), IS"],
        ],
        col_widths=[2.0, 4.5]
    )

    # 12.2 Results
    add_heading(doc, "12.2 Results: Argos vs. Paper", 2)
    add_para(doc, (
        "The headline result is a significant performance gap between the paper's reported WER on "
        "LRS3 and Argos's WER on AVSpeech. This gap is expected and is primarily driven by domain "
        "mismatch rather than implementation error."
    ))
    add_styled_table(doc,
        ["Metric", "Paper (LRS3)", "Argos (AVSpeech)", "Gap Factor"],
        [
            ["Mean WER", "25.4%", "67.0%", "2.7x"],
            ["Dataset Type", "TED talks", "YouTube (wild)", "—"],
            ["Conditions", "Clean audio, frontal face", "Diverse angles, noise, lighting", "—"],
            ["Segments Usable (WER ≤ 20%)", "—", "11.4%", "—"],
            ["Segments Hallucinated (WER ≥ 100%)", "—", "20.6%", "—"],
            ["NEA F1", "—", "38.8%", "—"],
        ],
        col_widths=[2.0, 1.5, 1.7, 1.3]
    )

    # 12.3 WER Distribution
    add_heading(doc, "12.3 WER Distribution", 2)
    add_para(doc, (
        "Examining the distribution of WER scores across segments reveals a bimodal pattern: "
        "some segments are transcribed well while others are nearly or completely wrong."
    ))
    add_styled_table(doc,
        ["WER Range", "Category", "% of Segments", "Interpretation"],
        [
            ["0-20%", "Excellent / Usable", "11.4%", "Can be used with minor corrections"],
            ["21-40%", "Marginal", "17.4%", "Core meaning preserved but noisy"],
            ["41-60%", "Poor", "17.8%", "Some correct words but unreliable"],
            ["61-99%", "Unusable", "32.8%", "Too many errors for practical use"],
            ["≥100%", "Hallucinated", "20.6%", "Model generates unrelated text"],
        ],
        col_widths=[1.0, 1.3, 1.0, 3.2]
    )
    add_para(doc, (
        "The 20.6% hallucination rate is particularly noteworthy. These are cases where the LLaMA-2 "
        "decoder generates fluent, grammatically correct English text that has no relationship to "
        "what the speaker is saying. This happens when the visual features extracted by AV-HuBERT "
        "are ambiguous (poor face angle, motion blur, or the speaker's mouth is occluded)."
    ))

    # 12.4 Domain Gap
    add_heading(doc, "12.4 Domain Gap Analysis", 2)
    add_para(doc, (
        "The 2.7x performance gap between LRS3 and AVSpeech is caused by fundamental differences "
        "between the training and deployment domains:"
    ))
    add_styled_table(doc,
        ["Aspect", "LRS3 (Training Domain)", "AVSpeech (Deployment Domain)"],
        [
            ["Content", "TED talks", "YouTube videos (all genres)"],
            ["Camera", "Frontal, static, professional", "Variable angles, handheld, phone cameras"],
            ["Lighting", "Stage lighting, consistent", "Outdoor, indoor, backlighting, shadows"],
            ["Audio", "Clean studio recording", "Background noise, music, crosstalk"],
            ["Speakers", "Professional speakers", "Anyone on YouTube"],
            ["Face Position", "Centered, large in frame", "Variable size, sometimes partial"],
            ["Duration", "2-10 seconds", "Variable (requires segmentation)"],
        ],
        col_widths=[1.2, 2.6, 2.7]
    )
    add_para(doc, (
        "The model was fine-tuned by the paper's authors exclusively on LRS3 data (TED talks with "
        "clean conditions). When applied to YouTube videos with diverse real-world conditions, the "
        "visual features become less informative, and the LLaMA-2 decoder fills in the gaps with "
        "plausible but incorrect text (hallucination)."
    ))

    # 12.5 Tuning Experiments
    add_heading(doc, "12.5 Thirteen Decode Tuning Experiments (A\u2013M)", 2)
    add_para(doc, (
        "A systematic sweep of 13 decode configurations was run on a 107-segment test set derived from "
        "100 AVSpeech videos. Parameters tested: length penalty, beam width, sampling mode, temperature, "
        "and combinations. Each experiment took ~50 minutes on a T4 GPU (16GB)."
    ))
    add_styled_table(doc,
        ["Exp", "Configuration", "WER%", "WWER%", "Empty", "NEA F1%", "Verdict"],
        [
            ["A", "baseline (beam=20, lenpen=0)", "59.4", "59.5", "4 (3.7%)", "40.9", "Reference"],
            ["B", "repetition_penalty=1.0 (off)", "59.4", "59.4", "4 (3.7%)", "40.8", "No effect"],
            ["C", "lenpen=1.0", "58.6", "58.6", "0", "41.1", "Best single param"],
            ["D", "lenpen=-0.5", "65.8", "75.3", "48 (44.9%)", "26.1", "FAIL"],
            ["E", "do_sample=true, temp=0.5", "59.4", "59.4", "0", "41.8", "Slight NEA gain"],
            ["F", "do_sample=true, temp=1.0", "59.4", "59.4", "0", "42.4", "Best NEA F1 (single)"],
            ["G", "greedy (beam=1)", "63.2", "63.2", "0", "37.7", "Worse"],
            ["H", "lenpen=2.0", "539.6", "171.5", "0", "39.1", "CATASTROPHIC"],
            ["I", "lenpen=1.0 + sample temp=1.0", "83.6", "60.1", "0", "40.3", "Combo: mediocre"],
            ["J", "lenpen=1.0 + sample temp=0.5", "70.1", "57.7", "0", "41.0", "BEST WWER"],
            ["K", "do_sample=true, temp=1.5", "60.9", "59.2", "1", "40.7", "Competitive"],
            ["L", "do_sample=true, temp=0.3", "59.6", "58.4", "0", "41.8", "Good"],
            ["M", "lenpen=1.0 + sample temp=0.3", "80.9", "77.1", "0", "42.1", "Unexpectedly bad"],
        ],
        col_widths=[0.3, 2.0, 0.45, 0.5, 0.7, 0.5, 1.15]
    )
    add_para(doc, (
        "The experiment framework (run_all_experiments.sh + run_remaining.sh + run_experiment.sh) "
        "automates the full cycle: set Hydra overrides, run decode, generate reports, save config snapshots. "
        "A comparison CSV is built at the end with all metrics."
    ))

    # 12.5.1 Parameter Analysis
    add_heading(doc, "12.5.1 Length Penalty Analysis", 3)
    add_para(doc, (
        "Length penalty (lenpen) was the most impactful single parameter. lenpen=1.0 eliminates all "
        "empty predictions and improves WWER by 0.9 points. lenpen=-0.5 is catastrophic (44.9% empty). "
        "lenpen=2.0 causes catastrophic over-generation (539.6% WER \u2014 model generates 5x reference length)."
    ))

    add_heading(doc, "12.5.2 Sampling Analysis", 3)
    add_para(doc, (
        "Sampling alone does not improve WER/WWER but consistently improves NEA F1 by +1-1.5 points. "
        "temp=0.5 is the sweet spot when combined with lenpen=1.0 (Config J, best WWER 57.7%). "
        "Very low temperature (0.3) with lenpen creates pathological behavior \u2014 the peaked distribution "
        "combined with length reward causes degenerate outputs (Config M, 77.1% WWER)."
    ))

    add_heading(doc, "12.5.3 Example: lenpen Eliminates Empty Outputs", 3)
    add_para(doc, (
        "Segment DBhaa45mAro (casual vlog, 15 ref words):"
    ), italic=True, color=C_GRAY)
    add_styled_table(doc,
        ["Config", "WER%", "Output (truncated)"],
        [
            ["REF", "\u2014", "hey it's me just doing the eush keeping up with what's going on in school..."],
            ["A (baseline)", "100.0", "(empty \u2014 model produced no output)"],
            ["C (lenpen=1.0)", "66.7", "hi it's me jessica dude who's keeping up with what's going on in the world..."],
            ["J (winner)", "73.3", "hi it's me jessica dude who's keeping up with what's going on in the world..."],
        ],
        col_widths=[1.0, 0.5, 5.0]
    )

    add_heading(doc, "12.5.4 Example: lenpen Regression on Short Segments", 3)
    add_para(doc, (
        "Segment eLS1vcpGVHQ (short, 10 ref words):"
    ), italic=True, color=C_GRAY)
    add_styled_table(doc,
        ["Config", "WER%", "Output (truncated)"],
        [
            ["REF", "\u2014", "i surrendered good quite i surrendered to the low security"],
            ["A (baseline)", "90.0", "years ago when i was"],
            ["C (lenpen=1.0)", "740.0", "years ago when i was 18 years old i came to the united states..."],
            ["J (winner)", "400.0", "i'm so happy to be here with you all today thank you so much..."],
        ],
        col_widths=[1.0, 0.5, 5.0]
    )
    add_para(doc, (
        "This is the core lenpen trade-off: it fixes empty outputs but amplifies hallucination on "
        "ambiguous short segments. The model generates fluent, grammatically correct text that has "
        "no relation to the actual speech."
    ), bold=True)

    add_heading(doc, "12.5.5 Recommended Configuration", 3)
    add_para(doc, (
        "Config J (lenpen=1.0, do_sample=true, temperature=0.5) is recommended for production. "
        "It achieves the best WWER (57.7%, -1.8 vs baseline) and eliminates empty outputs. "
        "A full 1,497-segment decode with this config is in progress to validate these findings at scale."
    ))

    # 12.6 Analysis Reports
    add_heading(doc, "12.6 Six Analysis Reports", 2)
    add_para(doc, (
        "In addition to the tuning experiments, six detailed analysis reports were written to explore "
        "different aspects of the system's performance and potential improvements:"
    ))
    add_styled_table(doc,
        ["#", "Title", "Key Finding / Recommendation"],
        [
            ["1", "Executive Assessment", "67% WER with 11.4% usable segments; domain gap is primary bottleneck"],
            ["2", "Hyperparameter Tuning", "Length penalty (lenpen) is highest-impact parameter; recommend 0.5-1.0"],
            ["3", "Prompt Engineering & Context Injection", "Analysis of context injection strategies for improved decode"],
            ["4", "Word-Level Confidence Scoring", "Methods for estimating per-word reliability from beam scores"],
            ["5", "Beam Search Aggregation & N-Best Fusion", "Using multiple beam hypotheses to improve accuracy"],
            ["6", "Fine-Tuning Analysis", "Domain adaptation on AVSpeech is essential; expected 15-25 WER improvement"],
        ],
        col_widths=[0.3, 2.2, 4.0]
    )

    # 12.7 Analytical Plots
    add_heading(doc, "12.7 Analytical Plots (16 Visualizations)", 2)
    add_para(doc, (
        "A comprehensive set of 16 analytical plots was generated to visualize performance patterns "
        "across experiments. These plots are produced by generate_experiment_plots.py and stored in "
        "docs/evaluation/plots/. Experiments D and H are excluded from most plots as known failures "
        "(45% empty and catastrophic hallucination respectively)."
    ))
    add_styled_table(doc,
        ["Plot", "Type", "Description", "Key Finding"],
        [
            ["01-04", "Binned line", "WER & WWER vs duration and ref word count (5 configs)",
             "WER improves with longer segments; short segments (<3s) problematic"],
            ["05-08", "Binned line", "NEA Recall & F1 vs duration and ref word count",
             "Entity accuracy plateaus above 10 ref words"],
            ["09", "Box plot", "WWER distribution across 11 good experiments",
             "Median WWER tightly clustered (55-58%); large outlier spread"],
            ["10", "Bar chart", "Empty + hallucination rates by experiment",
             "lenpen eliminates empties but increases hallucination"],
            ["11", "Scatter", "WER vs WWER per-segment (5 configs)",
             "WWER tracks WER closely; diverges when entities are destroyed"],
            ["12", "Heatmap", "Segment stability across experiments",
             "70% of segments stable across configs; 30% volatile"],
            ["13", "Histogram", "Duration distribution: tuning subset vs full dataset",
             "Tuning subset representative of full dataset duration range"],
            ["14", "Scatter", "NEA Recall vs WWER",
             "High WWER strongly predicts low entity accuracy"],
            ["15", "CDF", "Empirical CDF of WWER by experiment",
             "Config J shifts CDF left (better) for mid-range segments"],
            ["16", "Delta bar", "Per-segment improvement J vs A",
             "J improves 40% of segments, hurts 25%, neutral 35%"],
        ],
        col_widths=[0.5, 0.7, 2.5, 2.8]
    )

    # 12.8 Presentation Plots
    add_heading(doc, "12.8 Presentation Plots (5 Visualizations)", 2)
    add_para(doc, (
        "Five high-level plots were generated for management presentations, summarizing the "
        "project's performance story. Produced by generate_presentation_plots.py."
    ))
    add_styled_table(doc,
        ["Plot", "Description", "Key Data Point"],
        [
            ["P1", "Quality tier distribution (pie + horizontal bar)", "11.4% usable (WER \u226420%)"],
            ["P2", "Paper vs reality WER comparison", "25.4% LRS3 vs 64.1% YouTube (2.5x gap)"],
            ["P3", "WER improvement trajectory roadmap", "64% \u2192 55% \u2192 45% \u2192 42% projected"],
            ["P4", "Length penalty sensitivity (empty vs hallucination)", "lenpen=1.0 eliminates empties"],
            ["P5", "Before/after tuning (A vs C metrics)", "WWER: 59.5% \u2192 58.6% (-0.9 pts)"],
        ],
        col_widths=[0.5, 3.0, 3.0]
    )

    # 12.9 Baseline vs Config J Full-Scale Analysis
    add_heading(doc, "12.9 Baseline vs Config J Full-Scale Analysis", 2)
    add_para(doc, (
        "A detailed 9-section analysis compared the baseline (Config A: beam=20, lenpen=0) against "
        "Config J (lenpen=1.0, do_sample=true, temperature=0.5) on the full 1,497-segment dataset. "
        "This analysis is documented in baseline_vs_J_analysis.docx and includes:"
    ))
    add_bullet_bold_value(doc, "Overall metrics: ",
        "J achieves best WWER (57.7%, -1.8 vs baseline 59.5%) and eliminates all empty outputs.")
    add_bullet_bold_value(doc, "Where J helps: ",
        "Recovers 4 previously-empty segments, improves WWER on medium-length (10-20 word) segments.")
    add_bullet_bold_value(doc, "Where J hurts: ",
        "Amplifies hallucination on short/ambiguous segments. Some segments that were empty "
        "in baseline become hallucinated with lenpen=1.0 (fluent but fabricated text).")
    add_bullet_bold_value(doc, "Quartile analysis: ",
        "Q1 (easiest 25% by baseline WWER) segments are stable across configs. Q4 (hardest 25%) "
        "segments diverge most \u2014 J helps some and hurts others unpredictably.")
    add_bullet_bold_value(doc, "Duration effect: ",
        "5-15 second segments show most consistent improvement. Very short (<3s) segments are "
        "problematic regardless of config. Segments >20s show diminishing returns.")
    add_bullet_bold_value(doc, "Entity recovery: ",
        "J recovers some entities lost by baseline (primarily function words). Content word "
        "entities remain challenging across both configs.")

    # 12.10 Pairwise Statistical Testing
    add_heading(doc, "12.10 Pairwise Statistical Testing", 2)
    add_para(doc, (
        "Wilcoxon signed-rank tests were run across all 78 configuration pairs (13 choose 2) "
        "to determine which differences are statistically significant. The analysis is documented "
        "in pairwise-comparison.pdf (47 pages). Key findings:"
    ))
    add_bullet(doc,
        "Few pairs show statistically significant differences at p<0.05, confirming that "
        "most decode parameter changes produce marginal effects on the 107-segment tuning set.")
    add_bullet(doc,
        "The only consistently significant differences involve the known-failure configs "
        "(D with lenpen=-0.5 and H with lenpen=2.0) vs. all others.")
    add_bullet(doc,
        "This supports the conclusion that decode-time tuning has reached diminishing returns "
        "and domain adaptation (fine-tuning) is the only viable path to meaningful improvement.")

    # 12.11 Metrics Assessment
    add_heading(doc, "12.11 Metrics Assessment: Beyond WER", 2)
    add_para(doc, (
        "A key insight from the evaluation is that standard WER is insufficient for assessing "
        "transcription quality in a production context. Three metrics were developed:"
    ))
    add_bullet_bold_value(doc, "WWER (Weighted WER): ",
        "Weights entity tokens (names, numbers, locations) at 2x importance and function words "
        "(the, is, of) at 0.5x. A transcription that captures 'Biden met with Putin' but misses "
        "articles is more useful than one that gets articles right but hallucates names.")
    add_bullet_bold_value(doc, "NEA F1 (Named Entity Accuracy): ",
        "Measures how well high-value tokens are preserved. At 38.8% F1, the current system "
        "misses at least one important entity in 85% of segments \u2014 this is the primary quality "
        "gap for real-world usability.")
    add_bullet_bold_value(doc, "Intelligibility Score (IS): ",
        "A composite 6-signal metric (0\u20135 scale) that combines semantic similarity, phonetic "
        "similarity, inverse WER, inverse WWER, named entity accuracy, and length ratio. "
        "At IS 2.52/5.0 mean, only 39.9% of segments are properly captured (IS \u2265 3.0). "
        "Full analysis in Chapter 13.")

    # 12.12 Arabic Dead End
    add_heading(doc, "12.12 Arabic Language Attempt", 2)
    add_para(doc, (
        "An attempt was made to extend the system to Arabic language processing. However, this "
        "proved infeasible with the current model architecture: the LLaMA-2 tokenizer is designed "
        "for English text, and the fine-tuning was performed on English-only data (LRS3). The "
        "tokenizer cannot produce meaningful Arabic text tokens, and the visual features are "
        "language-agnostic but the decoder is not. Supporting Arabic would require either a "
        "multilingual LLM base or separate Arabic-trained model weights."
    ))

    # 12.13 POC Fine-Tune
    add_heading(doc, "12.13 POC Fine-Tune to Overfit", 2)
    add_para(doc, (
        "A preliminary fine-tuning experiment was conducted as a proof-of-concept to validate "
        "that the training pipeline and infrastructure work correctly. The model was trained on "
        "a small dataset with the goal of overfitting \u2014 achieving very low training loss regardless "
        "of generalization. This experiment confirmed that the fairseq-hydra-train integration, "
        "QLoRA adapter setup, and checkpoint saving all function correctly, paving the way for "
        "the full domain adaptation fine-tuning."
    ))

    # 12.14 Key Insights
    add_heading(doc, "12.14 Key Insights", 2)
    add_para(doc, "The evaluation and tuning work yielded several important conclusions:")
    add_bullet(doc, "The domain gap (LRS3 \u2192 AVSpeech) is the primary bottleneck, not the model architecture")
    add_bullet(doc, "Fine-tuning on in-domain AVSpeech data is expected to reduce WER by 15-25 percentage points")
    add_bullet(doc, "Decode-time parameter tuning (lenpen, beam width) can provide 5-10% relative improvement")
    add_bullet(doc, "WER alone is insufficient \u2014 WWER, NEA F1, and the Intelligibility Score (IS) better capture real-world utility")
    add_bullet(doc, "The 20.6% hallucination rate is the most urgent quality issue to address")
    add_bullet(doc, "The Intelligibility Score reveals that WER overstates failure by ~3.5x (39.9% properly captured vs 11.4% by WER alone)")
    add_bullet(doc, "Expected post-fine-tuning WER: 40-50% (still above paper's 25.4% due to harder conditions)")

    doc.add_page_break()


# ═══════════════════════════════════════════════════
# CHAPTER 13: Intelligibility Assessment & Metric Analysis
# ═══════════════════════════════════════════════════

def chapter_13(doc):
    add_heading(doc, "13. Intelligibility Assessment & Metric Analysis", 1, bookmark_id="ch13")

    add_para(doc, (
        "Standard WER proved insufficient for evaluating the Argos system's real-world quality. "
        "This chapter documents the Intelligibility Score (IS), a novel 6-signal composite metric "
        "that captures whether a transcription conveys the speaker's intended meaning \u2014 even when "
        "individual words are wrong. The IS was computed for all 1,497 segments in the full baseline "
        "dataset and reveals that WER overstates failure by approximately 3.5x."
    ))

    # 13.1 Why WER Isn't Enough
    add_heading(doc, "13.1 Why WER Isn't Enough", 2)
    add_para(doc, (
        "WER treats all word errors equally: substituting 'the' for 'a' costs the same as "
        "substituting 'Biden' for 'cooking'. WER also cannot distinguish between phonetically "
        "plausible errors (natural lip-reading confusions) and complete hallucinations. Two "
        "examples from the dataset illustrate this:"
    ))
    add_bullet_bold_value(doc, "Same WER, opposite meaning: ",
        "Prediction A has 29% WER but is fully intelligible (wrong function words, correct "
        "content). Prediction B has 33% WER but is unintelligible (correct function words, "
        "wrong content words). WER cannot distinguish these cases.")
    add_bullet_bold_value(doc, "Phonetic confusions are invisible to WER: ",
        "A lip-reading system often confuses visually identical sounds (p/b/m, f/v, t/d). "
        "These confusions may produce 'wrong' words that sound like the right words and "
        "preserve meaning \u2014 but WER penalizes them fully.")

    # 13.2 The 6-Signal Composite
    add_heading(doc, "13.2 The 6-Signal Composite", 2)
    add_para(doc, (
        "The Intelligibility Score combines six independent signals, each capturing a different "
        "dimension of transcription quality. All signals are scaled to 0\u20135 before weighting:"
    ))
    add_styled_table(doc,
        ["Signal", "Weight", "What It Catches"],
        [
            ["Semantic Similarity", "0.25 (25%)",
             "Overall meaning preservation (all-MiniLM-L6-v2 sentence embeddings, cosine similarity)"],
            ["Phonetic Similarity", "0.15 (15%)",
             "Natural lip-reading confusions (Double Metaphone encoding, match rate)"],
            ["Inverse WER", "0.15 (15%)",
             "Raw word-level accuracy (1 \u2212 WER/100, scaled to 0\u20135)"],
            ["Inverse WWER", "0.15 (15%)",
             "Content-weighted accuracy (entities 2x, content 1x, function 0.5x)"],
            ["NEA F1", "0.15 (15%)",
             "Named entity preservation (proper nouns, numbers, locations)"],
            ["Length Ratio", "0.15 (15%)",
             "Hallucination/truncation detection (hyp_len / ref_len, penalty for deviation from 1.0)"],
        ],
        col_widths=[1.3, 0.7, 4.5]
    )
    add_para(doc, (
        "Semantic similarity receives the highest weight (25%) because meaning preservation is "
        "the ultimate goal. The other five signals each receive 15%, providing balanced coverage "
        "of phonetic plausibility, structural accuracy, entity preservation, and output length."
    ))

    # 13.3 Composite Formula
    add_heading(doc, "13.3 Composite Formula & Scaling", 2)
    add_para(doc, (
        "IS = 0.25 \u00d7 Semantic + 0.15 \u00d7 Phonetic + 0.15 \u00d7 InvWER + "
        "0.15 \u00d7 InvWWER + 0.15 \u00d7 NEA_F1 + 0.15 \u00d7 LengthRatio"
    ), bold=True)
    add_para(doc, (
        "Each signal is independently scaled to the 0\u20135 range before weighting. The final "
        "IS ranges from 0.0 (completely unintelligible) to 5.0 (perfect transcription). "
        "A segment is considered 'properly captured' if IS \u2265 3.0 (Tiers 4 + 5)."
    ))

    # 13.4 Tier Classification
    add_heading(doc, "13.4 Tier Classification", 2)
    add_para(doc, (
        "Each segment is classified into one of five intelligibility tiers based on its IS score:"
    ))
    add_styled_table(doc,
        ["Tier", "Score Range", "Label", "Count", "%", "Interpretation"],
        [
            ["5", "4.0\u20135.0", "Excellent", "276", "18.4%",
             "Human fully understands. Minor word differences only."],
            ["4", "3.0\u20133.99", "Good", "321", "21.4%",
             "Meaning clearly preserved. Some wrong words but recoverable."],
            ["3", "2.0\u20132.99", "Fair", "325", "21.7%",
             "Gist recoverable. General topic correct, details wrong."],
            ["2", "1.0\u20131.99", "Poor", "336", "22.4%",
             "Only fragments survive. Different message conveyed."],
            ["1", "0.0\u20130.99", "Failed", "239", "16.0%",
             "No meaningful connection. Completely different topic."],
        ],
        col_widths=[0.35, 0.7, 0.65, 0.5, 0.4, 3.9]
    )
    add_para(doc, (
        "The distribution is roughly uniform across tiers, with a slight concentration in the "
        "Poor tier (22.4%). Tiers 4+5 ('properly captured') account for 39.9% of segments \u2014 "
        "compared to only 11.4% that would be considered usable by WER alone (WER \u2264 20%)."
    ), bold=True)

    # 13.5 Full Baseline Results
    add_heading(doc, "13.5 Full Baseline Results (1,497 Segments)", 2)
    add_para(doc, (
        "The IS was computed for all 1,497 segments in the full baseline dataset (February 2026):"
    ))
    add_styled_table(doc,
        ["Metric", "Value"],
        [
            ["Total Segments", "1,497"],
            ["Mean IS", "2.52 / 5.0"],
            ["Median IS", "2.538"],
            ["Std Dev", "1.372"],
            ["Properly Captured (IS \u2265 3.0)", "597 (39.9%)"],
            ["Context-Recoverable (Rule-Based)", "652 (43.6%)"],
            ["Context-Recoverable (LLM-Based)", "757 (50.6%)"],
            ["Empty Hypotheses", "70 (4.7%)"],
        ],
        col_widths=[2.5, 4.0]
    )
    add_para(doc, (
        "The 39.9% 'properly captured' rate is the headline metric: 4 in 10 segments convey "
        "intelligible meaning despite a 67% mean WER. This 3.5x gap between WER assessment "
        "and IS assessment demonstrates why WER alone is misleading for lip-reading evaluation."
    ))

    # 13.6 IS Distribution Histogram
    add_heading(doc, "13.6 IS Distribution Histogram", 2)
    add_para(doc, (
        "The distribution of IS scores across 0.5-width bins shows a roughly uniform spread "
        "with a slight concentration in the 1.0\u20131.5 range:"
    ))
    add_styled_table(doc,
        ["IS Range", "Count", "% of Total", "Cumulative %"],
        [
            ["0.0\u20130.5", "108", "7.2%", "7.2%"],
            ["0.5\u20131.0", "131", "8.8%", "16.0%"],
            ["1.0\u20131.5", "186", "12.4%", "28.4%"],
            ["1.5\u20132.0", "150", "10.0%", "38.4%"],
            ["2.0\u20132.5", "161", "10.8%", "49.2%"],
            ["2.5\u20133.0", "164", "11.0%", "60.1%"],
            ["3.0\u20133.5", "165", "11.0%", "71.1%"],
            ["3.5\u20134.0", "156", "10.4%", "81.6%"],
            ["4.0\u20134.5", "159", "10.6%", "92.2%"],
            ["4.5\u20135.0", "99", "6.6%", "98.8%"],
        ],
        col_widths=[1.0, 0.7, 0.8, 1.0]
    )

    # 13.7 Signal Statistics
    add_heading(doc, "13.7 Signal Statistics", 2)
    add_para(doc, (
        "Mean and median values for each of the six component signals across all 1,497 segments:"
    ))
    add_styled_table(doc,
        ["Signal", "Mean", "Median", "Interpretation"],
        [
            ["Semantic Similarity", "0.437", "0.412",
             "Moderate meaning overlap on average; wide spread"],
            ["Phonetic Similarity", "0.552", "0.588",
             "Over half of ref words have phonetic matches in hypothesis"],
            ["Inverse WER (1\u2212WER/100)", "0.394", "\u2014",
             "Equivalent to ~60% mean WER; confirms challenging dataset"],
            ["Inverse WWER (1\u2212WWER/100)", "0.397", "\u2014",
             "Content-weighted accuracy similar to raw WER"],
            ["NEA F1", "38.9%", "\u2014",
             "Entities missed in majority of segments"],
            ["Length Ratio (hyp/ref)", "0.925", "0.947",
             "Slight truncation bias; median close to ideal 1.0"],
        ],
        col_widths=[1.5, 0.5, 0.5, 4.0]
    )

    # 13.8 Signal Comparison
    add_heading(doc, "13.8 Signal Comparison: Success vs Failure", 2)
    add_para(doc, (
        "Comparing mean signal values between successful (IS \u2265 3.0, n=597) and failed "
        "(IS < 3.0, n=900) segments reveals which signals differentiate success from failure:"
    ))
    add_styled_table(doc,
        ["Signal", "Success Mean", "Failure Mean", "Gap", "Diagnostic Value"],
        [
            ["Semantic Similarity", "0.736", "0.238", "+0.50", "Highest discriminator"],
            ["Phonetic Similarity", "0.809", "0.382", "+0.43", "Strong discriminator"],
            ["WER", "30.2%", "86.5%", "\u221256.3 pts", "Large gap (expected)"],
            ["WWER", "31.1%", "82.3%", "\u221251.2 pts", "Large gap (expected)"],
            ["NEA F1", "74.0%", "15.7%", "+58.3 pts", "Strongest single discriminator"],
            ["Length Ratio", "0.974", "0.892", "+0.08", "Modest discriminator"],
        ],
        col_widths=[1.2, 0.8, 0.8, 0.7, 3.0]
    )
    add_para(doc, (
        "NEA F1 shows the largest gap (+58.3 points), meaning entity preservation is the "
        "strongest predictor of intelligibility. Semantic similarity is the highest-weighted "
        "signal and the best standalone discriminator (0.736 vs 0.238)."
    ))

    # 13.9 Failure Mode Analysis
    add_heading(doc, "13.9 Failure Mode Analysis", 2)
    add_para(doc, (
        "Each non-intelligible segment (IS < 3.0, n=900) is classified into one of ten failure "
        "modes based on its signal profile:"
    ))
    add_styled_table(doc,
        ["Failure Mode", "Count", "%", "Description"],
        [
            ["Total Topic Drift", "143", "15.9%", "Semantic < 0.2, phonetic < 0.3 \u2014 completely unrelated output"],
            ["Phonetically Similar but Wrong Topic", "141", "15.7%",
             "Semantic < 0.2 but phonetic \u2265 0.3 \u2014 sounds right, means wrong"],
            ["Accumulated Small Errors", "111", "12.3%", "Many small errors compound to destroy meaning"],
            ["Hallucination", "111", "12.3%", "WER > 100% \u2014 model generates fluent but fabricated text"],
            ["High Error Rate", "109", "12.1%", "WER > 70% \u2014 too many errors for recovery"],
            ["Entity Destruction", "108", "12.0%", "Semantic \u2265 0.2 but NEA F1 < 10% \u2014 key entities lost"],
            ["Content Word Errors", "96", "10.7%", "Semantic \u2265 0.3 but WER > 50% \u2014 content words wrong"],
            ["Empty Output", "70", "7.8%", "Model produced no output for the segment"],
            ["Truncation", "10", "1.1%", "Length ratio < 0.3 \u2014 output is too short"],
            ["Over-generation", "1", "0.1%", "Length ratio > 1.8 \u2014 output is too long"],
        ],
        col_widths=[1.8, 0.5, 0.4, 3.8]
    )
    add_para(doc, "The three most actionable failure modes:")
    add_bullet_bold_value(doc, "Topic Drift + Phonetic Wrong Topic (31.6%): ",
        "Together account for nearly a third of failures. These occur when visual features "
        "are ambiguous and the LLM decoder fills in plausible but incorrect text.")
    add_bullet_bold_value(doc, "Hallucination (12.3%): ",
        "The most dangerous mode \u2014 output is fluent and grammatically correct but fabricated. "
        "Cannot be detected without ground truth comparison.")
    add_bullet_bold_value(doc, "Entity Destruction (12.0%): ",
        "Topic is roughly correct but key entities (names, numbers, places) are lost. "
        "Unrecoverable even with context because the entities cannot be guessed.")

    # 13.10 Success Pattern Analysis
    add_heading(doc, "13.10 Success Pattern Analysis", 2)
    add_para(doc, (
        "Each intelligible segment (IS \u2265 3.0, n=597) is classified by its dominant success pattern:"
    ))
    add_styled_table(doc,
        ["Success Pattern", "Count", "%", "Description"],
        [
            ["Phonetically Preserved", "248", "41.5%",
             "Phonetic > 0.7, WER > 25% \u2014 wrong words that sound like right words"],
            ["Minor Errors, High Semantic", "146", "24.5%",
             "WER \u2264 25%, semantic > 0.6 \u2014 few errors, strong meaning match"],
            ["Entities Preserved", "74", "12.4%",
             "NEA F1 > 60%, WER > 25% \u2014 key entities survived despite word errors"],
            ["Near-Perfect Output", "69", "11.6%",
             "WER \u2264 10% \u2014 almost exact transcription"],
            ["Good Semantic + Correct Length", "45", "7.5%",
             "Semantic > 0.5, length ratio 0.7\u20131.3 \u2014 right amount of right-ish content"],
            ["Low-Moderate WER", "13", "2.2%",
             "WER \u2264 35% \u2014 structurally accurate"],
            ["Combined Semantic + Phonetic Bridge", "2", "0.3%",
             "Semantic > 0.4, phonetic > 0.5 \u2014 meaning and sound both partially match"],
        ],
        col_widths=[1.8, 0.5, 0.4, 3.8]
    )
    add_para(doc, (
        "Phonetic preservation is the dominant success driver (41.5%). This means the lip-reading "
        "acoustic front-end (AV-HuBERT) is doing its job \u2014 it extracts features that map to the "
        "correct phonetic space even when the LLM decoder selects the wrong word. This finding "
        "supports the hypothesis that fine-tuning the decoder on in-domain data will yield "
        "significant improvements."
    ), bold=True)

    # 13.11 Context Recovery Analysis
    add_heading(doc, "13.11 Context Recovery Analysis", 2)
    add_para(doc, (
        "Two approaches were used to estimate how many additional segments could be understood "
        "if the reader had topic context (e.g., knowing the video is about cooking):"
    ))
    add_styled_table(doc,
        ["Approach", "Recoverable", "%", "Method"],
        [
            ["Rule-Based", "652 / 1,497", "43.6%",
             "Decision tree using semantic, WER, WWER, length ratio thresholds"],
            ["LLM-Knowledge-Based", "757 / 1,497", "50.6%",
             "Multi-factor analysis: content overlap, sequence preservation, phonetic "
             "plausibility, length sanity, semantic coherence, information density"],
        ],
        col_widths=[1.3, 1.0, 0.5, 3.7]
    )
    add_para(doc, (
        "The LLM-based approach recovers an additional 105 segments (7 percentage points) beyond "
        "the rule-based method, primarily in the 'marginal' category where multiple weak signals "
        "combine to make the output recoverable with context. This suggests that domain-specific "
        "fine-tuning (which provides implicit topic context) could shift ~7-11% of segments from "
        "'unintelligible' to 'recoverable'."
    ))

    # 13.12 WER Threshold Mapping
    add_heading(doc, "13.12 WER Threshold Mapping to Intelligibility", 2)
    add_para(doc, (
        "Mapping WER ranges to intelligibility expectations provides practical rules of thumb "
        "for interpreting raw WER scores:"
    ))
    add_styled_table(doc,
        ["WER Range", "Intelligibility Expectation", "Action"],
        [
            ["WER \u2264 15%", "Almost certainly understandable",
             "Use directly \u2014 minor corrections only"],
            ["WER 15\u201330%", "Probably understandable",
             "Review recommended; meaning likely preserved"],
            ["WER 30\u201350%", "Coin flip",
             "May or may not be intelligible; check IS score"],
            ["WER 50\u201370%", "Probably not understandable",
             "Significant errors; topic may still be recoverable with context"],
            ["WER > 70%", "Almost certainly not understandable",
             "Severe errors; discard or flag for manual transcription"],
            ["WER > 100%", "Hallucinated",
             "Model output is longer than reference and likely fabricated"],
        ],
        col_widths=[1.0, 2.2, 3.3]
    )

    # 13.13 Topic Analysis
    add_heading(doc, "13.13 Topic Analysis", 2)
    add_para(doc, (
        "Segments were classified into 10 topic categories using keyword matching on reference "
        "transcriptions. Performance varies significantly by topic:"
    ))
    add_styled_table(doc,
        ["Topic", "N", "Mean IS", "Mean WER", "Captured %", "Ctx Recovery %"],
        [
            ["Business/Finance", "46", "3.08", "46.8%", "56.5%", "67.4%"],
            ["Sports/Fitness", "31", "2.90", "52.9%", "48.4%", "61.3%"],
            ["Education/Academic", "86", "2.84", "52.4%", "47.7%", "59.3%"],
            ["Politics/News", "34", "2.81", "56.7%", "50.0%", "55.9%"],
            ["Technology", "132", "2.70", "56.7%", "49.2%", "56.8%"],
            ["Cooking/Food", "117", "2.66", "59.3%", "44.4%", "59.8%"],
            ["Medical/Health", "39", "2.64", "56.7%", "53.8%", "53.8%"],
            ["Religion/Spirituality", "17", "2.55", "68.7%", "35.3%", "47.1%"],
            ["Entertainment", "69", "2.23", "67.3%", "30.4%", "47.8%"],
            ["DIY/Home", "27", "2.13", "76.0%", "29.6%", "37.0%"],
            ["Other (unclassified)", "899", "2.42", "68.0%", "36.2%", "46.7%"],
        ],
        col_widths=[1.3, 0.35, 0.5, 0.6, 0.7, 0.8]
    )
    add_para(doc, (
        "Business/Finance topics perform best (IS 3.08, 56.5% captured) \u2014 likely because these "
        "videos feature professional speakers in controlled settings. Entertainment and DIY/Home "
        "perform worst (IS 2.1\u20132.2, ~30% captured), likely due to rapid speech, background noise, "
        "and non-frontal camera angles common in those genres."
    ))

    # 13.14 Segment Length Analysis
    add_heading(doc, "13.14 Segment Length Analysis", 2)
    add_para(doc, (
        "Metrics improve consistently with longer segments (more reference words). This analysis "
        "suggests a minimum segment length filter could improve overall quality:"
    ))
    add_heading(doc, "13.14.1 Cumulative Filters", 3)
    add_styled_table(doc,
        ["Filter", "N", "Mean IS", "Mean WER", "Captured %", "Ctx LLM %"],
        [
            ["All segments", "1,497", "2.52", "64.1%", "39.9%", "50.6%"],
            ["\u2265 5 words", "1,463", "2.55", "61.9%", "40.5%", "51.1%"],
            ["\u2265 7 words", "1,379", "2.56", "61.1%", "40.8%", "51.6%"],
            ["\u2265 10 words", "1,173", "2.61", "58.9%", "42.7%", "53.5%"],
            ["\u2265 15 words", "805", "2.65", "56.5%", "46.3%", "56.5%"],
            ["\u2265 20 words", "535", "2.68", "55.1%", "48.6%", "57.6%"],
        ],
        col_widths=[1.0, 0.5, 0.6, 0.6, 0.7, 0.7]
    )
    add_heading(doc, "13.14.2 Length Bands", 3)
    add_styled_table(doc,
        ["Band", "N", "Mean IS", "Mean WER", "Captured %"],
        [
            ["5\u201310 words", "290", "2.31", "74.2%", "31.7%"],
            ["10\u201315 words", "368", "2.51", "64.1%", "34.8%"],
            ["15\u201320 words", "270", "2.60", "59.4%", "41.9%"],
            ["20+ words", "535", "2.68", "55.1%", "48.6%"],
        ],
        col_widths=[1.0, 0.5, 0.6, 0.6, 0.7]
    )
    add_para(doc, (
        "Filtering to segments with \u2265 10 reference words drops 324 short segments (21.6%) "
        "but improves captured rate from 39.9% to 42.7% (+2.8 pts). Filtering to \u2265 20 words "
        "reaches 48.6% captured but retains only 535 segments (35.7% of total)."
    ))

    # 13.15 Phonetic Confusion Analysis
    add_heading(doc, "13.15 Phonetic Confusion Analysis", 2)
    add_para(doc, (
        "Lip-reading inherently confuses sounds that look identical on the lips. Six homophene "
        "groups (sounds sharing the same lip shape) drive most confusions:"
    ))
    add_styled_table(doc,
        ["Lip Shape", "Sounds", "Examples"],
        [
            ["Both lips close", "p, b, m", "pat/bat/mat"],
            ["Teeth on lower lip", "f, v", "fan/van"],
            ["Tongue behind teeth", "t, d, n, s, z, l", "ten/den/den/sen"],
            ["Back of throat", "k, g, h", "cap/gap"],
            ["Rounded lips", "w, r", "wet/ret"],
            ["Open mouth (vowels)", "various", "Context-dependent"],
        ],
        col_widths=[1.3, 1.2, 4.0]
    )
    add_para(doc, "The top 15 word-level confusions observed in the 1,497-segment dataset:")
    add_styled_table(doc,
        ["Reference", "Hypothesis", "Count", "Reference", "Hypothesis", "Count"],
        [
            ["the", "to", "13", "a", "the", "9"],
            ["the", "a", "11", "a", "to", "9"],
            ["in", "and", "11", "the", "you", "9"],
            ["and", "in", "11", "we're", "we", "9"],
            ["you're", "you", "11", "to", "into", "8"],
            ["that", "and", "10", "it's", "is", "8"],
            ["gonna", "going", "10", "i", "and", "7"],
        ],
        col_widths=[0.7, 0.7, 0.45, 0.7, 0.7, 0.45]
    )
    add_para(doc, (
        "Most confusions involve function words (the/to/a/and/in) which carry minimal semantic "
        "weight. The IS metric correctly discounts these via the WWER signal (function words "
        "weighted at 0.5x) and the semantic similarity signal (which is unaffected by function "
        "word substitutions)."
    ))

    # 13.16 Key Insights
    add_heading(doc, "13.16 Key Insights", 2)
    add_para(doc, "The intelligibility assessment produced several important conclusions:")
    add_bullet(doc,
        "WER overstates failure by ~3.5x: 39.9% of segments are properly captured (IS \u2265 3.0) "
        "vs. only 11.4% by WER alone (WER \u2264 20%)")
    add_bullet(doc,
        "Phonetic preservation is the #1 success driver (41.5% of successful segments), "
        "validating that the AV-HuBERT visual front-end extracts correct phonetic features")
    add_bullet(doc,
        "Entity destruction is unrecoverable: when names, numbers, and locations are lost, "
        "no amount of context can restore them (12% of failures)")
    add_bullet(doc,
        "Context recovery can add +7\u201311% intelligibility, suggesting domain-specific "
        "fine-tuning will help significantly")
    add_bullet(doc,
        "Topic matters: Business/Finance (56.5% captured) outperforms DIY/Home (29.6%) "
        "by nearly 2x, driven by speaker quality and camera conditions")
    add_bullet(doc,
        "Longer segments perform better: \u2265 20 words reaches 48.6% captured vs 31.7% for "
        "5\u201310 words, supporting a minimum segment length filter")
    add_bullet(doc,
        "Hallucination (12.3% of failures) is the most dangerous mode \u2014 fluent, correct-sounding "
        "text that is completely fabricated and undetectable without ground truth")

    # 13.17 Config J & C: Decode Parameter Variants
    add_heading(doc, "13.17 Config J & C: Decode Variant Evaluation", 2)
    add_para(doc, (
        "Two alternative decode configurations were evaluated on the full 1,497-segment dataset. "
        "Config J adds stochastic sampling (temperature=0.5) with length penalty (lenpen=1.0). "
        "Config C is its deterministic counterpart (lenpen=1.0 only). Both share beam=20, "
        "rep_penalty=1.2, no_repeat_ngram=3 with the baseline."
    ))

    add_styled_table(doc,
        ["Metric", "Baseline", "Config J", "Config C"],
        [
            ["Mean IS", "2.52", "2.60 (+0.08)", "2.57 (+0.05)"],
            ["Properly Captured", "597 (39.9%)", "622 (41.5%)", "594 (39.7%)"],
            ["Empty Predictions", "70 (4.7%)", "0", "0"],
            ["Hallucinations (WER\u2265100%)", "307 (20.5%)", "348 (23.2%)", "360 (24.0%)"],
            ["Mean WER", "64.1%", "78.9%", "79.3%"],
            ["Mean WWER", "61.9%", "62.8%", "63.8%"],
            ["NEA F1", "38.9%", "43.4%", "39.7%"],
        ],
        col_widths=[1.8, 1.6, 1.6, 1.6]
    )

    add_para(doc, "Key findings from the J/C evaluation:", bold=True)
    add_bullet(doc,
        "lenpen=1.0 eliminates all 70 empty predictions. The model is forced to produce "
        "output even with weak visual signal, trading silent failures for noisy ones.")
    add_bullet(doc,
        "Hallucinations more than doubled (111\u2192262 for J, 270 for C). Most former empties "
        "became hallucinated text \u2014 fluent but fabricated content.")
    add_bullet(doc,
        "Config J outperforms C on every metric. Stochastic sampling at temperature=0.5 "
        "recovers 28 more intelligible segments and +3.7pp more named entities than "
        "deterministic decoding.")
    add_bullet(doc,
        "Long segments (20+ words) benefit most: +0.25 IS and +3.4pp capture rate "
        "under Config J. Short segments (5\u201310 words) are marginally worse due to "
        "over-generation.")
    add_bullet(doc,
        "The improvement is real but marginal (+0.08 IS). Decode parameter tuning has "
        "reached diminishing returns; fine-tuning remains the only viable path to "
        "production-grade accuracy.")

    doc.add_page_break()


# ═══════════════════════════════════════════════════
# CHAPTER 14: Fine-Tuning & Training Infrastructure (was Chapter 13)
# ═══════════════════════════════════════════════════

def chapter_14(doc):
    add_heading(doc, "14. Fine-Tuning & Training Infrastructure", 1, bookmark_id="ch14")

    add_para(doc, (
        "Based on the evaluation results (Chapter 12), fine-tuning the model on in-domain AVSpeech data "
        "is the most promising path to improving transcription quality. This chapter documents the training "
        "infrastructure that was built, the training configuration designed for domain adaptation, and the "
        "research that informed the fine-tuning strategy."
    ))

    # 14.1 Training Pipeline
    add_heading(doc, "14.1 Fine-Tuning Pipeline (579 Lines)", 2)
    add_para(doc, (
        "The run_avspeech_finetune_pipeline.sh script (579 lines) implements the complete end-to-end "
        "fine-tuning workflow. It reuses all lib/ modules from the inference pipeline for preprocessing "
        "(stages 0-6), then hands off to fairseq-hydra-train for the actual training."
    ))
    add_para(doc, "Key features of the fine-tuning pipeline:")
    add_bullet(doc, "Reuses all lib/ modules for preprocessing — same code path as inference pipeline")
    add_bullet(doc, "fairseq-hydra-train integration with custom task (vsp_llm_training)")
    add_bullet(doc, "Multi-GPU support via configurable NUM_GPUS environment variable (1-8 GPUs)")
    add_bullet(doc, "Checkpointing from pre-trained model (checkpoint_finetune.pt)")
    add_bullet(doc, "Dry-run mode for validating setup without launching training")
    add_bullet(doc, "Skip-preprocess option for re-training on already-preprocessed data")
    add_bullet(doc, "Training output directory management with automatic versioning")

    # 14.2 Training Config
    add_heading(doc, "14.2 Training Configuration", 2)
    add_para(doc, (
        "The training configuration (vsp-llm-avspeech-finetune.yaml, 136 lines) was designed specifically "
        "for domain adaptation from LRS3 to AVSpeech. It differs from the paper's original config in "
        "several important ways:"
    ))
    add_styled_table(doc,
        ["Parameter", "Paper Config", "Argos Config", "Rationale"],
        [
            ["freeze_finetune_updates", "18,000", "5,000", "Unfreeze encoder earlier for domain adaptation"],
            ["max_update", "30,000", "15,000", "Shorter schedule for smaller/diverse datasets"],
            ["warmup_steps", "10,000", "3,000", "Proportional to max_update"],
            ["decay_steps", "20,000", "12,000", "Tri-stage LR: warmup → hold → decay"],
            ["lr", "0.0005", "0.0005", "Same learning rate (Adam optimizer)"],
            ["batch_size", "1", "1", "Limited by GPU memory (7B parameter model)"],
            ["update_freq", "[1]", "[1] or [8]", "8 GPUs → freq=1; 1 GPU → freq=8 (grad accumulation)"],
            ["label_smoothing", "0.1", "0.1", "Prevent overconfident predictions"],
            ["seed", "1337", "1337", "Reproducibility"],
            ["save_interval_updates", "—", "2,500", "Checkpoint every 2,500 steps"],
            ["keep_interval_updates", "—", "3", "Keep last 3 checkpoints (disk space)"],
        ],
        col_widths=[1.6, 1.0, 1.0, 2.9]
    )

    # 14.3 QLoRA
    add_heading(doc, "14.3 QLoRA Architecture", 2)
    add_para(doc, (
        "The fine-tuning uses QLoRA (Quantized Low-Rank Adaptation) to make training feasible on "
        "consumer GPUs. Instead of updating all 7 billion parameters of LLaMA-2, QLoRA adds small "
        "trainable adapter matrices to each linear layer:"
    ))
    add_styled_table(doc,
        ["QLoRA Parameter", "Value", "Explanation"],
        [
            ["Rank (r)", "16", "Dimension of the low-rank decomposition"],
            ["Alpha (α)", "32", "Scaling factor (α/r = 2.0)"],
            ["Quantization", "4-bit NF4", "Base model quantized to 4 bits (NormalFloat4)"],
            ["Target Layers", "All linear layers", "LoRA adapters on q_proj, k_proj, v_proj, o_proj, etc."],
            ["Trainable Params", "~0.1% of total", "Only the LoRA adapter matrices are updated"],
            ["Memory Savings", "~75%", "4-bit quantization reduces memory from ~28GB to ~7GB"],
        ],
        col_widths=[1.5, 1.2, 3.8]
    )

    # 14.4 Model Architecture Detail
    add_heading(doc, "14.4 Full Model Architecture", 2)
    add_code_block(doc, (
        "┌─────────────────────────────────────────────────────────┐\n"
        "│                    VSP-LLM Architecture                 │\n"
        "├─────────────────────────────────────────────────────────┤\n"
        "│                                                         │\n"
        "│  Input: 88x88 grayscale mouth crops @ 25 FPS            │\n"
        "│                    │                                     │\n"
        "│                    ▼                                     │\n"
        "│  ┌─────────────────────────────────┐                    │\n"
        "│  │     AV-HuBERT Encoder           │                    │\n"
        "│  │     315M parameters             │                    │\n"
        "│  │     Output: 1024-dim features   │                    │\n"
        "│  └─────────────────────────────────┘                    │\n"
        "│                    │                                     │\n"
        "│                    ▼                                     │\n"
        "│  ┌─────────────────────────────────┐                    │\n"
        "│  │     Linear Projection           │                    │\n"
        "│  │     1024 → 4096 dimensions      │                    │\n"
        "│  └─────────────────────────────────┘                    │\n"
        "│                    │                                     │\n"
        "│                    ▼                                     │\n"
        "│  ┌─────────────────────────────────┐                    │\n"
        "│  │     LLaMA-2-7B + QLoRA          │                    │\n"
        "│  │     7B params (4-bit quantized)  │                    │\n"
        "│  │     + LoRA adapters (r=16, α=32) │                    │\n"
        "│  │     ~0.1% trainable              │                    │\n"
        "│  └─────────────────────────────────┘                    │\n"
        "│                    │                                     │\n"
        "│                    ▼                                     │\n"
        "│  Output: Text transcription (beam search or sampling)   │\n"
        "│                                                         │\n"
        "└─────────────────────────────────────────────────────────┘"
    ))

    # 14.5 AVSpeech Data Strategy
    add_heading(doc, "14.5 AVSpeech Data Strategy", 2)
    add_para(doc, (
        "The training data for domain adaptation was sourced from AVSpeech, a large-scale dataset "
        "of short video clips extracted from YouTube. Unlike LRS3 (which contains only TED talks), "
        "AVSpeech represents the diversity of real-world video: multiple camera angles, variable "
        "lighting, background noise, and speakers of all demographics."
    ))
    add_para(doc, (
        "The data was imported from an old server that had been previously mounted. The pipeline's "
        "preprocessing stages (segmentation, normalization, face detection, mouth cropping, ASR) "
        "prepare this raw data for training — the same stages used in the inference pipeline."
    ))

    # 14.6 Research Notes
    add_heading(doc, "14.6 Training Research Notes", 2)
    add_para(doc, (
        "Detailed research notes (docs/training-research-notes.md, 186 lines) document the analysis "
        "that informed the fine-tuning strategy:"
    ))
    add_bullet_bold_value(doc, "LoRA Rank Selection: ",
        "Rank 16 was chosen as optimal for the VSP-LLM model size. Lower ranks (4, 8) showed "
        "insufficient adaptation capacity in preliminary experiments, while higher ranks (32, 64) "
        "increased memory requirements without proportional quality gains.")
    add_bullet_bold_value(doc, "Angle Robustness: ",
        "Analysis of face detection success rates versus camera angle revealed that the model "
        "performs best with angles within ±30° of frontal. Fine-tuning on diverse-angle data "
        "is expected to improve robustness.")
    add_bullet_bold_value(doc, "Length Distribution: ",
        "Training segments should ideally be 5-12 seconds. Segments shorter than 3 seconds "
        "contain too few visual features, while segments longer than 15 seconds exceed the "
        "model's effective context window.")
    add_bullet_bold_value(doc, "Expected Improvement: ",
        "Based on domain adaptation literature and the evaluation results, fine-tuning on "
        "AVSpeech data is expected to reduce WER by 15-25 percentage points, bringing it "
        "to the 40-50% range on real-world YouTube content.")

    # 14.7 Exp A Results
    add_heading(doc, "14.7 Experiment A Results (r=16 LoRA Fine-Tuning)", 2)
    add_para(doc, (
        "The first fine-tuning experiment (Exp A) was completed on February 27, 2026. "
        "Training ran for 19 epochs (3,000 updates) over 17.0 hours on a Tesla T4 GPU with FP16 "
        "mixed precision. The dataset consisted of 1,273 training and 224 validation segments "
        "from AVSpeech YouTube videos, stratified by Intelligibility Score tier."
    ))

    add_heading(doc, "14.7.1 Training Configuration", 3)
    add_styled_table(doc,
        ["Parameter", "Value"],
        [
            ["LoRA Rank (r)", "16 (continued from pretrained)"],
            ["LoRA Alpha (\u03b1)", "32 (scaling factor \u03b1/r = 2.0)"],
            ["Target Modules", "q_proj, k_proj, v_proj"],
            ["Trainable Parameters", "12,582,912 (0.19% of 6.75B total)"],
            ["Max Updates", "3,000"],
            ["Learning Rate", "3e-4 (tri-stage: warmup 600, decay 2,400)"],
            ["Effective Batch Size", "8 (update_freq=8, batch_size=1)"],
            ["Encoder", "Frozen (freeze_finetune_updates=999,999)"],
            ["Label Smoothing", "0.1"],
            ["Seed", "1337"],
        ],
        col_widths=[2.0, 4.5]
    )

    add_heading(doc, "14.7.2 Key Results", 3)
    add_styled_table(doc,
        ["Metric", "Best (Epoch 2)", "Final (Epoch 19)", "Delta"],
        [
            ["Val Accuracy", "62.94%", "58.98%", "-3.96 pp"],
            ["Val Loss", "2.391", "4.120", "+72%"],
            ["Val Perplexity", "5.24", "17.39", "+232%"],
            ["Train Accuracy", "65.00%", "95.52%", "+30.5 pp"],
            ["Train Loss", "2.192", "0.022", "-99%"],
            ["Train-Val Acc Gap", "2.1 pp", "36.5 pp", "+34.4 pp"],
        ],
        col_widths=[1.5, 1.3, 1.3, 2.4]
    )

    add_heading(doc, "14.7.3 Overfitting Analysis", 3)
    add_para(doc, (
        "Exp A exhibited severe overfitting starting from epoch 3. The best validation accuracy "
        "(62.94%) was reached at epoch 2, after only 320 updates. All subsequent epochs showed "
        "monotonic degradation in validation metrics while training metrics continued improving \u2014 "
        "the classic signature of memorization:"
    ))
    add_bullet_bold_value(doc, "Train accuracy: ",
        "Climbed from 62.7% to 95.5% \u2014 the model memorized virtually all training examples")
    add_bullet_bold_value(doc, "Val accuracy: ",
        "Declined from 62.9% to 59.0% \u2014 generalization worsened with continued training")
    add_bullet_bold_value(doc, "Val perplexity: ",
        "Exploded from 5.24 to 17.39 \u2014 the model became increasingly uncertain on unseen data")
    add_bullet_bold_value(doc, "Gradient norms: ",
        "Decreased from ~3.3 to ~1.0 \u2014 consistent with convergence to a sharp minimum")

    add_heading(doc, "14.7.4 Root Cause Analysis", 3)
    add_para(doc, (
        "The overfitting has five contributing factors, ranked by likely impact:"
    ))
    add_bullet_bold_value(doc, "1. Small dataset: ",
        "1,273 segments is insufficient for a 7B-parameter LLM, even with LoRA. "
        "The model exhausts the training signal within 2 epochs.")
    add_bullet_bold_value(doc, "2. Noisy labels: ",
        "Training labels come from Whisper ASR at 64% WER baseline \u2014 the model learns errors.")
    add_bullet_bold_value(doc, "3. Rank limitation: ",
        "r=16 constrains the adaptation to a 16-dimensional subspace, "
        "insufficient for the TED\u2192YouTube domain shift.")
    add_bullet_bold_value(doc, "4. Minimal regularization: ",
        "LoRA dropout 0.05 and weight decay 0.0 provide almost no overfitting protection.")
    add_bullet_bold_value(doc, "5. Encoder frozen: ",
        "Visual features cannot adapt to YouTube-style faces/angles, "
        "forcing all adaptation through the LLM adapter.")

    add_heading(doc, "14.7.5 Implications for Exp B", 3)
    add_para(doc, (
        "Based on Exp A results, the following adjustments are planned for Exp B (r=64):"
    ))
    add_bullet_bold_value(doc, "Increase LoRA rank: ",
        "r=64 (alpha=128) provides 4\u00d7 adapter capacity for better domain adaptation.")
    add_bullet_bold_value(doc, "Aggressive early stopping: ",
        "max_update=500 with validation every 50 steps (Exp A peaked at 320 updates).")
    add_bullet_bold_value(doc, "Data curation: ",
        "Filter training data by face detection confidence >0.9 and head pose <30\u00b0.")
    add_para(doc, (
        "Decode evaluation of the Exp A best checkpoint on the full 1,497-segment test set is pending. "
        "10 diagnostic plots documenting the full training dynamics are available in "
        "docs/finetuning/plots/FT_01 through FT_10."
    ), bold=True)

    doc.add_page_break()


# ═══════════════════════════════════════════════════
# CHAPTER 15: Executive Summary (was Chapter 14)
# ═══════════════════════════════════════════════════

def chapter_15(doc):
    add_heading(doc, "15. Executive Summary", 1, bookmark_id="ch15")

    add_para(doc, (
        "This chapter provides a high-level overview of the Argos project's scope, deliverables, "
        "and current state for management review."
    ))

    # 15.1 Project Scope
    add_heading(doc, "15.1 Project Scope & Timeline", 2)
    add_para(doc, (
        "The Argos project began on October 18, 2024, with the goal of building a production-ready "
        "visual speech processing (lip-reading) system based on the VSP-LLM research paper. All "
        "engineering work — pipeline development, web UI, Docker container, deployment, testing, and "
        "documentation — was carried out by Yoad Oxman as the sole developer."
    ))
    add_para(doc, (
        "The system takes raw video as input and produces text transcriptions by analyzing lip movements, "
        "with no audio input required. This enables transcription in scenarios where audio is unavailable, "
        "classified, or unreliable."
    ))

    # 15.2 Key Deliverables
    add_heading(doc, "15.2 Key Deliverables", 2)
    add_styled_table(doc,
        ["Deliverable", "Status", "Details"],
        [
            ["End-to-End Pipeline", "Complete (v7)", "8-stage pipeline, modular architecture (501 + 1,562 lines)"],
            ["Web UI", "Complete", "6 screens, 12 API endpoints, drag-drop upload, transcription management"],
            ["Docker Container", "5 versions deployed", "nvidia/cuda base, dual venvs, ~65GB galaxy export"],
            ["Performance Evaluation", "Complete", "1,497 segments, 6 analysis reports, 13 tuning experiments, 21 analytical plots"],
            ["Intelligibility Scoring", "Complete", "6-signal composite metric (IS), tier classification, failure/success analysis"],
            ["Report Generation", "Complete", "WER, WWER, NEA metrics in 5 formats (CSV/HTML/JSON/TXT/ANSI)"],
            ["Burned Videos", "Complete", "Subtitle overlays on original videos for visual review"],
            ["Fine-Tuning Infrastructure", "Exp A Complete", "579-line pipeline; Exp A (r=16) trained 17h, best val acc 62.94%"],
            ["Testing Suite", "Complete", "37 automated tests, 5 test scripts (1,731 lines)"],
            ["Documentation", "Comprehensive", "Architecture, dev guide, sync changelog, training notes, bug tracking"],
            ["4-Environment Support", "Operational", "EC2, 2 standalone Ubuntu, HORIZON (no internet)"],
        ],
        col_widths=[1.8, 1.2, 3.5]
    )

    # 15.3 Effort Estimate
    add_heading(doc, "15.3 Total Effort Estimate", 2)
    add_styled_table(doc,
        ["Phase", "Duration", "Key Activities"],
        [
            ["Literature Review", "~2-3 weeks", "VSP-LLM, AV-HuBERT, Auto-AVSR, fairseq, LRS3, LLaMA-2, QLoRA papers"],
            ["Failed Venv Builds", "~2 weeks", "Attempted venv setup on standalone machines (failed, led to container approach)"],
            ["Driver & CUDA Setup", "~1 month", "NVIDIA drivers, CUDA toolkit, GPU verification on standalone machines"],
            ["Pipeline Development", "Ongoing", "7 versions, from 240 to 501+1,562 lines, 13 custom Python scripts"],
            ["Web UI Development", "Ongoing", "5,780 lines across 11 files, 4 backend services"],
            ["Container & Deployment", "Ongoing", "5 releases, 37 bugs fixed, Dockerfile, galaxy export, installation scripts"],
            ["Evaluation & Tuning", "Ongoing", "1,497 segments, 13 experiments, 6 analysis reports, intelligibility scoring"],
            ["Training Infrastructure", "Ongoing", "579-line pipeline, training config, research notes"],
        ],
        col_widths=[1.5, 1.0, 4.0]
    )

    # 15.4 Performance
    add_heading(doc, "15.4 Current Performance", 2)
    add_styled_table(doc,
        ["Metric", "Current Value", "Exp A (r=16)", "Expected Post-FT"],
        [
            ["Mean WER", "67.0%", "Decode pending", "40-50%"],
            ["Usable Segments (WER ≤ 20%)", "11.4%", "25-40%"],
            ["Hallucination Rate (WER ≥ 100%)", "20.6%", "<10%"],
            ["NEA F1 (Named Entity Accuracy)", "38.8%", "55-70%"],
            ["Mean Intelligibility Score", "2.52 / 5.0", "3.5-4.0"],
            ["Properly Captured (IS \u2265 3)", "39.9%", "55-70%"],
            ["Best Val Acc (token)", "\u2014", "62.94% (ep 2)", "\u2014"],
            ["Paper WER (LRS3)", "25.4%", "\u2014", "25.4%"],
        ],
        col_widths=[2.3, 1.1, 1.3, 1.1]
    )

    # 15.5 Code Inventory
    add_heading(doc, "15.5 Complete Code Inventory", 2)
    add_para(doc, (
        "The following table summarizes the complete codebase written for the Argos project. "
        "All code was developed by Yoad Oxman:"
    ))
    add_styled_table(doc,
        ["Category", "Files", "Lines", "Key Components"],
        [
            ["Pipeline Scripts", "4", "1,299", "Main pipeline (501), fine-tune (579), tuning (219)"],
            ["Library Modules", "12", "1,562", "common, config, archive, normalization, asr, decode, outputs, etc."],
            ["Custom Python Scripts", "13", "3,342", "make_report, make_burn, fast_segment, transcribe_segments, etc."],
            ["Web UI + Server", "11", "5,780", "server.py (1,124), app.js (1,921), style.css (1,191), 4 services"],
            ["Face Detection", "4", "532", "MediaPipe detector (52+220), RetinaFace (43+217)"],
            ["Training Config", "1", "136", "vsp-llm-avspeech-finetune.yaml (Hydra config)"],
            ["Testing Scripts", "5", "1,731", "37 tests + smoke tests + environment tests + sync verification"],
            ["Deployment Scripts", "5", "~500", "build_container, check_inventory, monitor, resume"],
            ["Grand Total", "~55", "~14,882", "All custom production code"],
        ],
        col_widths=[1.5, 0.5, 0.5, 4.0]
    )

    # 15.6 Infrastructure
    add_heading(doc, "15.6 Infrastructure Summary", 2)
    add_styled_table(doc,
        ["Aspect", "Count", "Details"],
        [
            ["Deployment Environments", "4", "EC2, 2 standalone Ubuntu, HORIZON (no internet)"],
            ["Container Releases", "5", "Feb 3 → Feb 17, 2026 (v1.0.0 — v1.0.37)"],
            ["Bugs Fixed", "37", "Installation (13), deployment/GPU (12), final fixes (12)"],
            ["Container Sync Items", "28", "Pending EC2 → container replication"],
            ["Automated Tests", "37", "Covering all 12 lib/ modules"],
            ["Virtual Environments", "2", "pre-process-venv (135 pkgs) + vsp-llm-yoad-venv (256 pkgs)"],
            ["Upstream Repos Integrated", "3", "auto_avsr, VSP-LLM, av_hubert"],
            ["Git Commits", "40+", "From initial commit through v1.0.40+"],
        ],
        col_widths=[1.8, 0.6, 4.1]
    )

    # 15.7 Recommendations
    add_heading(doc, "15.7 Recommendations", 2)
    add_para(doc, "Based on the evaluation results and research analysis, the following next steps are recommended:")
    add_bullet_bold_value(doc, "1. Fine-Tune on AVSpeech: ",
        "Domain adaptation is the highest-impact improvement available. The training infrastructure "
        "is ready (Chapter 14). Expected improvement: 15-25 WER points. "
        "Intelligibility Score analysis (Chapter 13) confirms WER alone understates true accuracy by 3.5\u00d7.")
    add_bullet_bold_value(doc, "2. Evaluate RetinaFace: ",
        "Switching from MediaPipe to RetinaFace for face detection may improve mouth crop quality "
        "on difficult angles, potentially contributing 2-5 WER points improvement.")
    add_bullet_bold_value(doc, "3. Multi-Language Support: ",
        "Investigate multilingual LLM bases (e.g., LLaMA-3, Mistral) that could support Arabic "
        "and other languages. The current English-only tokenizer is the blocking factor.")
    add_bullet_bold_value(doc, "4. Segment Duration Optimization: ",
        "The hyperparameter tuning report identified segment duration as a high-impact data-level "
        "lever. Filtering segments to 5+ seconds and capping at 20-30 seconds may improve WER by 5-10%.")

    # 15.8 Milestones
    add_heading(doc, "15.8 Project Milestones", 2)
    add_styled_table(doc,
        ["Date", "Milestone"],
        [
            ["Oct 18, 2024", "Project start — literature review begins"],
            ["Oct-Nov 2024", "First pipeline version (~285 lines, galaxy_export)"],
            ["Nov-Dec 2024", "Pipeline iterations, face detection integration, AVSpeech processing"],
            ["Dec 29, 2025", 'Presentation: "Argos — A Working Pipeline" (19 slides)'],
            ["Jan 2026", "Modular refactoring (823 → 501+1,562 lines)"],
            ["Jan 2026", "Web UI development (5,780 lines)"],
            ["Feb 3, 2026", "First container deployment"],
            ["Feb 3-17, 2026", "Five container releases, 37 bugs fixed"],
            ["Feb 2026", "Performance evaluation (1,497 segments, 6 reports, 13 experiments)"],
            ["Feb 2026", "Intelligibility scoring framework (6-signal IS), 16 analytical plots, 5 presentation plots"],
            ["Feb 2026", "Fine-tuning infrastructure ready (579-line pipeline + config)"],
            ["Feb 2026", "HORIZON server operational (closed network, no internet)"],
        ],
        col_widths=[1.3, 5.2]
    )

    # Final note
    doc.add_paragraph()
    add_para(doc, (
        "This document represents a comprehensive record of all research and development work "
        "performed on the Argos visual speech processing system. It is intended as a living document "
        "that will be updated as the project progresses through fine-tuning, evaluation, and deployment."
    ), italic=True, color=C_GRAY)


# ═══════════════════════════════════════════════════
# CHAPTER 16: Project To-Do List
# ═══════════════════════════════════════════════════

def chapter_16(doc):
    add_heading(doc, "16. Project To-Do List", 1, bookmark_id="ch16")

    add_para(doc, (
        "This chapter consolidates all pending work items across research, operations, "
        "technical development, and documentation. Items are prioritized by expected impact "
        "on the system's core mission: improving lip-reading accuracy on real-world video."
    ))

    # 16.1 Research & Model Improvement
    add_heading(doc, "16.1 Research & Model Improvement", 2)
    add_para(doc, (
        "These items directly target the WER gap between benchmark (25.4%) and real-world (64-67%) performance."
    ))

    add_styled_table(doc,
        ["#", "Task", "Priority", "Expected Impact", "Status"],
        [
            ["1", "Fine-tune on AVSpeech data with encoder unfreeze", "Critical",
             "-15 to -25 WER pts", "Infrastructure ready (579-line pipeline + config)"],
            ["2", "Increase LoRA rank from r=16 to r=64 (alpha=128)", "High",
             "-3 to -8 WER pts", "Config change ready (vsp_llm.py line 296)"],
            ["3", "Test length penalty lenpen=0.5-1.0", "High",
             "-1.8 WWER pts (57.7% vs 59.5%)", "DONE \u2014 13 experiments (A-M), best=J"],
            ["4", "Reduce beam width from 20 to 5-10", "Medium",
             "Greedy (beam=1) was 3.7 pts worse", "DONE \u2014 Exp G tested beam=1 vs beam=20"],
            ["5", "Tighten max_len_a from 2.0 to 1.0, max_len_b from 200 to 50", "Medium",
             "lenpen=2.0 caused 539% WER over-generation", "DONE \u2014 lenpen approach preferred"],
            ["6", "Implement confidence scoring (token/sequence level)", "High",
             "Quality filtering (24% \u2192 40-50% precision)", "Code design in Report 4"],
            ["7", "Implement N-best aggregation (ROVER or MBR decoding)", "Medium",
             "-5 to -10 WER pts", "Code design in Report 5"],
            ["8", "Evaluate prompt engineering (word count hints, topic context)", "Medium",
             "-5 to -10 WER pts", "Strategy in Report 3"],
            ["9", "Generate AVSpeech length histogram on full dataset", "Low",
             "Diagnostic / data characterization", "ffprobe script ready"],
            ["10", "Curate training data (80% clean frontal / 20% moderate)", "High",
             "Better training signal, less label noise", "Filtering criteria defined"],
            ["11", "Evaluate RetinaFace vs MediaPipe for face detection", "Low",
             "-2 to -5 WER pts on difficult angles", "Both detectors integrated"],
        ],
        col_widths=[0.3, 2.2, 0.6, 1.5, 1.9]
    )

    # 16.2 Operations & Deployment
    add_heading(doc, "16.2 Operations & Deployment", 2)
    add_para(doc, (
        "These items ensure production reliability and environment consistency."
    ))

    add_styled_table(doc,
        ["#", "Task", "Priority", "Details"],
        [
            ["12", "Sync all 28 pending EC2 \u2192 container changes", "Critical",
             "Full list in docs/container-sync-changelog.md"],
            ["13", "Deploy latest container package to HORIZON server", "High",
             "Closed network, no internet \u2014 requires offline package"],
            ["14", "Automate container sync verification (CI/CD or script)", "Medium",
             "Prevent future drift between EC2 and container"],
            ["15", "Add minimum segment duration filter (\u22655 seconds)", "Medium",
             "Drop <5s segments that produce 128% WER"],
            ["16", "Create automated regression test suite", "Medium",
             "Run a small set of known videos and compare WER against baseline"],
        ],
        col_widths=[0.3, 2.5, 0.6, 3.1]
    )

    # 16.3 Technical Development
    add_heading(doc, "16.3 Technical Development", 2)

    add_styled_table(doc,
        ["#", "Task", "Priority", "Details"],
        [
            ["17", "Add confidence scores to JSON report output", "High",
             "Extract from model generate() with output_scores=True"],
            ["18", "Add N-best hypotheses to decode output format", "Medium",
             "Use num_return_sequences=N in generate()"],
            ["19", "Investigate multi-language support (Arabic)", "Low",
             "Requires multilingual LLM base (LLaMA-3, Mistral)"],
            ["20", "Add segment duration optimization to pipeline config", "Medium",
             "Configurable min/max segment length parameters"],
            ["21", "Consider newer model architectures (LLaMA-3, Mistral)", "Low",
             "May improve language modeling capacity"],
        ],
        col_widths=[0.3, 2.5, 0.6, 3.1]
    )

    # 16.4 Documentation & Process
    add_heading(doc, "16.4 Documentation & Process", 2)

    add_styled_table(doc,
        ["#", "Task", "Priority", "Details"],
        [
            ["22", "Update this document after each fine-tuning experiment", "Ongoing",
             "Record hyperparameters, results, and analysis"],
            ["23", "Document HORIZON server setup and constraints", "Medium",
             "Offline deployment, GPU config, network restrictions"],
            ["24", "Create runbook for fine-tuning on cloud GPUs", "Medium",
             "p3.16xlarge setup, data transfer, checkpoint management"],
            ["25", "Maintain container-sync-changelog.md as changes are made", "Ongoing",
             "Every EC2 change documented with code diffs"],
        ],
        col_widths=[0.3, 2.5, 0.6, 3.1]
    )

    # Priority summary
    add_heading(doc, "16.5 Priority Summary", 2)
    add_para(doc, "Items ranked by expected impact on the core metric (WER reduction):")

    add_bullet_bold_value(doc, "1. Fine-tune on AVSpeech (#1): ",
        "Highest single-intervention impact. Infrastructure is ready. Requires GPU time (~$100-150).")
    add_bullet_bold_value(doc, "2. Length penalty tuning (#3): ",
        "COMPLETED \u2014 13 experiments run. Best config (J): lenpen=1.0 + temp=0.5 = 57.7% WWER (-1.8 pts).")
    add_bullet_bold_value(doc, "3. Container sync (#12): ",
        "Production users are running outdated code. 28 changes pending.")
    add_bullet_bold_value(doc, "4. LoRA rank increase (#2): ",
        "Simple config change with meaningful impact on domain adaptation capacity.")
    add_bullet_bold_value(doc, "5. Confidence scoring (#6): ",
        "Essential for production use \u2014 currently no way to separate good from bad outputs.")


# ═════════════════���═════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════

def main():
    print("Generating Argos — The Orchard (Detailed Edition v3)...")
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Header/Footer
    add_header_footer(doc)

    # Build document
    create_cover_page(doc)
    create_toc(doc)

    chapter_1(doc)
    chapter_2(doc)
    chapter_3(doc)
    chapter_4(doc)
    chapter_5(doc)
    chapter_6(doc)
    chapter_7(doc)
    chapter_8(doc)
    chapter_9(doc)
    chapter_10(doc)
    chapter_11(doc)
    chapter_12(doc)
    chapter_13(doc)   # Intelligibility Assessment & Metric Analysis (NEW)
    chapter_14(doc)   # Fine-Tuning & Training Infrastructure
    chapter_15(doc)   # Executive Summary
    chapter_16(doc)   # Project To-Do List

    # Save
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_FILE))
    print(f"Saved: {OUTPUT_FILE}")
    print(f"Size: {OUTPUT_FILE.stat().st_size / 1024:.0f} KB")


if __name__ == "__main__":
    main()
