#!/usr/bin/env python3
"""
Argos — The Orchard — Q2 2026 Quarterly Summary Generator

Generates a branded .docx summarizing the Argos project's second-quarter
work (April 1 – June 15, 2026) in 4 parts:
  Part 1. Research & Evaluation
  Part 2. Client Engagement & Operations
  Part 3. Technical Development
  Part 4. Insights, Lessons & Next Quarter

Style mirrors generate_summary.py (the living project summary). This is the
standalone quarter-by-quarter artifact; the prior quarter snapshot is
docs/evaluation/project-summary.docx (Last Updated March 11, 2026).

Usage:
    python3 generate_q2_summary.py

Output:
    docs/evaluation/Q2-2026-summary.docx
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

# ── Output ──
SCRIPT_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = SCRIPT_DIR.parent.parent / "evaluation"
OUTPUT_FILE = OUTPUT_DIR / "Q2-2026-summary.docx"

# ── Logos ──
ASSETS_DIR = SCRIPT_DIR.parent / "assets"
LOGO_ORCHARD = ASSETS_DIR / "logo.png"
LOGO_PEACOCK = ASSETS_DIR / "peacock.png"

# ── Colors ──
C_PRIMARY = RGBColor(0x1a, 0x3a, 0x5c)
C_H2 = RGBColor(0x2a, 0x5a, 0x8c)
C_H3 = RGBColor(0x3a, 0x6a, 0x9c)
C_H4 = RGBColor(0x4a, 0x7a, 0xac)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_BLACK = RGBColor(0x00, 0x00, 0x00)
C_DARK = RGBColor(0x33, 0x33, 0x33)
C_GREEN = RGBColor(0x15, 0x57, 0x24)
C_RED = RGBColor(0x72, 0x1c, 0x24)
C_AMBER = RGBColor(0x85, 0x64, 0x04)
C_GRAY = RGBColor(0x66, 0x66, 0x66)
C_CODE = RGBColor(0x2d, 0x2d, 0x2d)

HEADER_BG = "1a3a5c"
ZEBRA_BG = "f0f4f8"
GREEN_BG = "d4edda"
RED_BG = "f8d7da"
AMBER_BG = "fff3cd"
CODE_BG = "f5f5f5"

# Fixed dates for this quarterly snapshot (Date.now() avoided for reproducibility).
PERIOD = "April 1 – June 15, 2026"
LAST_UPDATED = "June 15, 2026"


# ═══════════════════════════════════════════════════
# HELPER FUNCTIONS (shared with generate_summary.py)
# ═══════════════════════════════════════════════════

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, bold=False, color=None, size=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = size
    run.font.name = "Calibri"


def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, HEADER_BG)
        set_cell_text(cell, h, bold=True, color=C_WHITE, size=Pt(9))

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            set_cell_text(cell, val, size=Pt(9))
            if r_idx % 2 == 1:
                set_cell_shading(cell, ZEBRA_BG)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return table


_bookmark_counter = [0]

def add_heading(doc, text, level, color=None, bookmark_id=None):
    h = doc.add_heading(text, level=level)
    if color is None:
        color = {1: C_PRIMARY, 2: C_H2, 3: C_H3, 4: C_H4}.get(level, C_PRIMARY)
    for run in h.runs:
        run.font.color.rgb = color
        run.font.name = "Calibri"
    if bookmark_id:
        _bookmark_counter[0] += 1
        bm_start = OxmlElement("w:bookmarkStart")
        bm_start.set(qn("w:id"), str(_bookmark_counter[0]))
        bm_start.set(qn("w:name"), bookmark_id)
        bm_end = OxmlElement("w:bookmarkEnd")
        bm_end.set(qn("w:id"), str(_bookmark_counter[0]))
        h._p.insert(0, bm_start)
        h._p.append(bm_end)
    return h


def add_para(doc, text, bold=False, italic=False, size=Pt(11), color=None, space_after=Pt(6)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = size
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = space_after
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = "Calibri"
    return p


def add_bullet_bold_value(doc, label, value, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after = Pt(2)
    run_b = p.add_run(label)
    run_b.bold = True
    run_b.font.size = Pt(10)
    run_b.font.name = "Calibri"
    run_v = p.add_run(value)
    run_v.font.size = Pt(10)
    run_v.font.name = "Calibri"
    return p


def _build_inline_image_xml(rId, cx, cy, pic_id=1, name="Logo"):
    inline = OxmlElement('wp:inline')
    inline.set('distT', '0')
    inline.set('distB', '0')
    inline.set('distL', '0')
    inline.set('distR', '0')

    extent = OxmlElement('wp:extent')
    extent.set('cx', str(cx))
    extent.set('cy', str(cy))
    inline.append(extent)

    docPr = OxmlElement('wp:docPr')
    docPr.set('id', str(pic_id))
    docPr.set('name', name)
    inline.append(docPr)

    graphic = OxmlElement('a:graphic')
    graphicData = OxmlElement('a:graphicData')
    graphicData.set('uri', 'http://schemas.openxmlformats.org/drawingml/2006/picture')

    pic = OxmlElement('pic:pic')
    nvPicPr = OxmlElement('pic:nvPicPr')
    cNvPr = OxmlElement('pic:cNvPr')
    cNvPr.set('id', '0')
    cNvPr.set('name', name)
    nvPicPr.append(cNvPr)
    cNvPicPr = OxmlElement('pic:cNvPicPr')
    nvPicPr.append(cNvPicPr)
    pic.append(nvPicPr)

    blipFill = OxmlElement('pic:blipFill')
    blip = OxmlElement('a:blip')
    blip.set(qn('r:embed'), rId)
    blipFill.append(blip)
    stretch = OxmlElement('a:stretch')
    fillRect = OxmlElement('a:fillRect')
    stretch.append(fillRect)
    blipFill.append(stretch)
    pic.append(blipFill)

    spPr = OxmlElement('pic:spPr')
    xfrm = OxmlElement('a:xfrm')
    off = OxmlElement('a:off')
    off.set('x', '0')
    off.set('y', '0')
    xfrm.append(off)
    ext_el = OxmlElement('a:ext')
    ext_el.set('cx', str(cx))
    ext_el.set('cy', str(cy))
    xfrm.append(ext_el)
    spPr.append(xfrm)
    prstGeom = OxmlElement('a:prstGeom')
    prstGeom.set('prst', 'rect')
    spPr.append(prstGeom)
    pic.append(spPr)

    graphicData.append(pic)
    graphic.append(graphicData)
    inline.append(graphic)

    drawing = OxmlElement('w:drawing')
    drawing.append(inline)
    return drawing


def add_header_footer(doc):
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    section = doc.sections[0]
    section.different_first_page_header_footer = True

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]

    if LOGO_ORCHARD.exists():
        image_part = doc.part.package.get_or_add_image_part(str(LOGO_ORCHARD))
        rId = header.part.relate_to(image_part, RT.IMAGE)
        size_emu = int(0.3 * 914400)
        logo_run = hp.add_run()
        drawing = _build_inline_image_xml(rId, size_emu, size_emu, pic_id=10, name="Header Logo")
        logo_run._r.append(drawing)
    text_width = section.page_width - section.left_margin - section.right_margin
    hp.paragraph_format.tab_stops.add_tab_stop(text_width, WD_TAB_ALIGNMENT.RIGHT)
    hp.add_run("\t")
    run = hp.add_run("Argos — The Orchard  ·  Q2 2026 Quarterly Summary")
    run.font.size = Pt(8)
    run.font.color.rgb = C_GRAY
    run.font.name = "Calibri"
    run.italic = True

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = C_GRAY
    run.font.name = "Calibri"
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    run._r.append(fld1)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run._r.append(instr)
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld2)


# ═══════════════════════════════════════════════════
# TOC
# ═══════════════════════════════════════════════════

TOC_ENTRIES = [
    ("part1", "Part 1. Research & Evaluation"),
    ("part2", "Part 2. Client Engagement & Operations"),
    ("part3", "Part 3. Technical Development"),
    ("part4", "Part 4. Insights, Lessons & Next Quarter"),
]


def create_cover_page(doc):
    for _ in range(2):
        doc.add_paragraph()

    if LOGO_PEACOCK.exists():
        doc.add_picture(str(LOGO_PEACOCK), width=Inches(2.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ARGOS")
    run.font.size = Pt(48)
    run.font.color.rgb = C_PRIMARY
    run.bold = True
    run.font.name = "Calibri"

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Visual Speech Processing System")
    run2.font.size = Pt(22)
    run2.font.color.rgb = C_H2
    run2.font.name = "Calibri"

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("The Orchard")
    run3.font.size = Pt(20)
    run3.font.color.rgb = C_H2
    run3.font.name = "Calibri"

    doc.add_paragraph()

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run("Quarterly Summary — Q2 2026 (April–June)")
    run4.font.size = Pt(16)
    run4.font.color.rgb = C_H3
    run4.italic = True
    run4.font.name = "Calibri"

    for _ in range(2):
        doc.add_paragraph()

    info_lines = [
        ("Period Covered:", PERIOD),
        ("Prior Snapshot:", "Project Summary, March 11, 2026 (Q1)"),
        ("Project Start:", "October 18, 2024"),
        ("Lead Engineer:", "Yoad Oxman"),
        ("Last Updated:", LAST_UPDATED),
    ]
    for label, value in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rl = p.add_run(label + " ")
        rl.bold = True
        rl.font.size = Pt(11)
        rl.font.color.rgb = C_DARK
        rl.font.name = "Calibri"
        rv = p.add_run(value)
        rv.font.size = Pt(11)
        rv.font.color.rgb = C_DARK
        rv.font.name = "Calibri"

    doc.add_page_break()


def create_toc(doc):
    for _ in range(3):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    add_heading(doc, "Table of Contents", 1)

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fld_begin)
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' TOC \\o "1-2" \\h \\z \\u '
    run._r.append(instr)
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    run._r.append(fld_sep)

    for _bm, title in TOC_ENTRIES:
        placeholder = paragraph.add_run(title + "\n")
        placeholder.font.size = Pt(11)
        placeholder.font.name = "Calibri"
        placeholder.font.color.rgb = C_PRIMARY

    fld_end_run = paragraph.add_run()
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    fld_end_run._r.append(fld_end)

    settings = doc.settings.element
    update_fields = OxmlElement('w:updateFields')
    update_fields.set(qn('w:val'), 'true')
    settings.append(update_fields)

    doc.add_page_break()


# ═══════════════════════════════════════════════════
# QUARTER AT A GLANCE
# ═══════════════════════════════════════════════════

def quarter_at_a_glance(doc):
    add_heading(doc, "Quarter at a Glance", 1, bookmark_id="glance")
    add_para(doc, (
        "Q2 2026 was the quarter Argos moved from research artifacts to a deployed product in a "
        "client's hands. Two threads ran in parallel. On the research side, the quarter's evaluation "
        "work shipped to production: N-best aggregation (Mission 6) and a per-word confidence stack "
        "that tells an operator which words to trust. On the delivery side, the system was packaged as "
        "a single self-contained Docker image and stood up on an air-gapped Windows 11 + RTX 5090 "
        "client machine, supported by a 2-hour stakeholder meeting and a redesigned presentation suite."
    ))

    add_styled_table(doc,
        ["Area", "Q2 Headline", "Status"],
        [
            ["Decode default", "MBR aggregation now the production default output (was top-1)",
             "Shipped May 2"],
            ["N-best aggregation (M6)", "5 offline methods; hyp_mbr selected; judge Y+P 68.4% → 71.1%",
             "Shipped"],
            ["Confidence scoring (M4)", "Per-word Trust/Salvage/Strip bands live in every run",
             "Shipped Apr 30"],
            ["Client deployment", "Single Docker image on air-gapped Win 11 + RTX 5090",
             "Delivered"],
            ["Stakeholder meeting", "2-hour, ~12-person session; full presentation suite rebuilt",
             "Delivered"],
            ["Llama 3.1 migration", "Code + smoke test pass; blocked on LRS3 data + GPU quota",
             "Prep complete"],
            ["Human-IS study (Path B)", "Pre-study estimates: model ≈ unaided deaf reader",
             "Estimates only"],
        ],
        col_widths=[1.7, 3.5, 1.3]
    )
    doc.add_page_break()


# ═══════════════════════════════════════════════════
# PART 1: RESEARCH & EVALUATION
# ═══════════════════════════════════════════════════

def part_1_research(doc):
    add_heading(doc, "Part 1. Research & Evaluation", 1, bookmark_id="part1")
    add_para(doc, (
        "Q2's research effort focused on a single question carried over from Q1: the model produces "
        "useful output ~62% of the time, but which output can a user trust? Two deliverables answered "
        "it — N-best aggregation (a better single answer) and per-word confidence (a per-word trust "
        "signal). Both shipped to production. All numbers below are computed on the full 1,497-segment "
        "evaluation set unless noted."
    ))

    # 1.1 MBR default
    add_heading(doc, "1.1 Production Decode Upgrade: MBR Now the Default", 2)
    add_para(doc, (
        "Since May 2, 2026 the production pipeline emits the Minimum Bayes Risk (MBR) aggregated "
        "hypothesis as its displayed output instead of the raw top-1 beam result. The change is gated "
        "by VSP_NBEST=1 (now default) and wired via make_report.py --display-method (default hyp_mbr "
        "when aggregated.json exists; override with VSP_DISPLAY_METHOD=top1). MBR improves every "
        "headline metric except a negligible hallucination tick, and the gain is statistically "
        "significant on the LLM judge."
    ))
    add_styled_table(doc,
        ["Metric (1,497 seg)", "Top-1 (old)", "MBR (new default)", "Δ"],
        [
            ["Mean WER", "64.1%", "63.8%", "−0.22 pp"],
            ["Mean IS", "2.532", "2.547", "+0.015"],
            ["Median IS", "2.559", "2.600", "+0.041"],
            ["NIV-Y — clearly conveyed (IS ≥ 3.80)", "359 (24.0%)", "358 (23.9%)", "−1 seg"],
            ["NIV-Y+P — any useful (IS ≥ 2.00)", "923 (61.7%)", "927 (61.9%)", "+4 seg"],
            ["LLM Judge Y+P (Opus 4.7)", "68.4%", "71.1%", "+2.67 pp"],
            ["Hallucination (WER ≥ 100%)", "20.5%", "20.7%", "+0.20 pp"],
        ],
        col_widths=[2.9, 1.3, 1.5, 0.9]
    )
    add_para(doc, (
        "The judge improvement is the meaningful one: on a blind paired re-evaluation of all 1,497 "
        "pairs (Opus 4.7, dual-confidence prompt, 5,988 verdicts), MBR wins the ‘any useful output’ "
        "operating point with 74 method-only gains vs 34 baseline-only losses — McNemar p = 0.0002. "
        "The strict ‘clearly conveyed’ (Y) verdict is unchanged (p = 0.29). MBR also lifts ~11 "
        "segments from IS tier 3 into tier 4. Past videos can adopt the new default without re-decoding "
        "by re-running stage 8 on existing VSP_NBEST=1 artifacts."
    ), bold=True)

    # 1.2 N-best
    add_heading(doc, "1.2 N-Best Aggregation — Mission 6 Shipped", 2)
    add_para(doc, (
        "Mission 6 shipped five offline aggregation methods (pure-Python/CPU, no re-decode) that "
        "combine the beam's N hypotheses into one improved answer. They were evaluated three ways on "
        "the full set — deterministic WER, the IS metric, and the Opus LLM judge — which disagree, "
        "and that disagreement drove the final choice."
    ))
    add_styled_table(doc,
        ["Method", "What it does", "Best at"],
        [
            ["hyp_mbr", "Min expected-WER hypothesis under score-weighted voting", "Judge Y+P (selected default)"],
            ["hyp_vote_conf", "ROVER vote weighted by score × per-word confidence", "WER (−1.56 pp) & lowest halluc."],
            ["hyp_vote_score", "ROVER vote weighted by sequence score", "— (n.s. on judge)"],
            ["hyp_safe", "Top-1, swap only very-low-conf words with beam consensus", "Conservative / neutral"],
            ["hyp_xseg_merge", "Cross-segment overlap fusion", "No-op on this dataset"],
        ],
        col_widths=[1.5, 3.3, 1.7]
    )
    add_para(doc, (
        "The three lenses split: vote_conf wins WER (62.5%, −1.56 pp) and has the lowest hallucination "
        "(19.0%); IS cannot distinguish the methods (all four within 0.015, paired p ≈ 1.0); the judge "
        "favors aggregation broadly, with MBR strongest (Y+P +40 wins, p = 0.0002) and vote_conf second "
        "(+31, p = 0.0026). hyp_mbr was selected as the shipped default for three reasons beyond the "
        "judge edge: (1) highest intra-rater stability at 86.7%, matching the gold-standard judge run; "
        "(2) it emits a calibrated per-word posterior compatible with the confidence-band UI, where the "
        "voting methods only emit a narrow [0.4, 0.8] agreement score; (3) a hybrid sentence-confidence "
        "gate was tested and rejected (+36 vs +37 rescues — one segment, not worth the threshold)."
    ))
    add_para(doc, (
        "Judge methodology note: the first LLM-judge run (v1) injected only each method's own "
        "confidence into the prompt and was contaminated — 27% of byte-identical hypotheses got "
        "different verdicts, and it falsely reported vote_conf as a significant loss. The v3 prompt adds "
        "a shared baseline-confidence anchor identical across all four methods; identical-text drift "
        "dropped to 23% and became balanced. v1 was archived; all rankings above use v3."
    ), italic=True)

    # 1.3 Confidence
    add_heading(doc, "1.3 Per-Word Confidence: the Trust / Salvage / Strip Stack", 2)
    add_para(doc, (
        "The quarter's flagship evaluation deliverable. Since April 30, every pipeline run emits a "
        "per-word confidence signal, and the report colors each word so an operator can read selectively "
        "rather than trust or discard a whole segment. The core finding from Q1 — that no flat "
        "confidence threshold works — was resolved with a two-layer design: a segment-level tier "
        "decides whether to color at all, and a per-word band rule decides the color."
    ))
    add_para(doc, (
        "Outer layer (segment tier, keyed on the segment's mean confidence):"
    ))
    add_styled_table(doc,
        ["Tier", "Segment mean conf", "Reliability of green", "UI behavior"],
        [
            ["Trust", "≥ 0.82", "green ≥ 85% correct", "Full per-word coloring"],
            ["Salvage", "0.65 – 0.82", "green ~65–85%", "Coloring + ‘uncertain’ banner"],
            ["Strip", "< 0.65", "green < 50%", "Coloring removed; plain grey text"],
        ],
        col_widths=[1.1, 1.5, 1.7, 2.2]
    )
    add_para(doc, (
        "Distribution on the 1,497-segment baseline: Trust 22.7%, Salvage 35.7%, Strip 36.9% "
        "(plus 70 empty-hypothesis segments). The thresholds (T_trust 0.89 ≥ 90% reliable, "
        "T_safe 0.82 ≥ 85%, T_salvage 0.74 ≥ 75%, strip boundary 0.65) are Llama-2-7b-specific and "
        "must be refit if the LLM backbone changes."
    ))

    add_heading(doc, "1.3.1 Agreement-Aware Bands", 3)
    add_para(doc, (
        "The inner per-word rule was upgraded on May 2 from a single confidence threshold to a joint "
        "rule combining top-1 confidence with beam agreement (how many of the N hypotheses chose the "
        "same word):"
    ))
    add_bullet_bold_value(doc, "Green: ", "top-1 conf ≥ 0.95 AND beam agreement ≥ 0.80  (≈ 85–94% correct)")
    add_bullet_bold_value(doc, "Yellow: ", "top-1 conf ≥ 0.65 AND beam agreement ≥ 0.50  (≈ 40–80% correct)")
    add_bullet_bold_value(doc, "Numbers / currency: ", "clamped at yellow regardless of confidence (digits are visually near-random)")
    add_para(doc, (
        "Beam agreement is roughly twice as informative as top-1 confidence in the high-confidence "
        "region: at fixed agreement, raising confidence 0.65→0.95 adds +37 pp reliability; at fixed "
        "confidence, raising agreement 0→0.95 adds +54 pp. A word the model is 95%-confident about but "
        "the beam disagrees on is only ~40% reliable — exactly the ‘gonna→going’, ‘billion→million’ "
        "failure cell that single-confidence coloring missed."
    ))

    add_heading(doc, "1.3.2 Band Reliability Is Conditional on Segment Quality", 3)
    add_para(doc, (
        "Measured on 23,261 aligned words / 1,427 segments, the reliability of a green word swings by "
        "more than 70 points depending on the segment it sits in — which is why the outer tier exists:"
    ))
    add_styled_table(doc,
        ["Segment mean conf", "P(correct | green)"],
        [
            ["≥ 0.85 (very high)", "92.8%"],
            ["0.75 – 0.85 (high)", "83.8%"],
            ["0.65 – 0.75 (mid)", "69.6%"],
            ["0.55 – 0.65 (mid-low)", "41.3%"],
            ["0.40 – 0.55 (low)", "21.8%"],
            ["< 0.40 (very low)", "18.2%"],
            ["Overall", "80.6%"],
        ],
        col_widths=[2.6, 2.0]
    )

    add_heading(doc, "1.3.3 Band Reliability by Outcome, and Green Leakage", 3)
    add_para(doc, (
        "Within useful content (the ~62% of segments rated Y or P), the per-word flag genuinely "
        "separates right from wrong: P(correct | green / yellow / red) = 87.2% / 48.9% / 24.7% — a "
        "62.5-point spread. The flag does its heaviest lifting on partial (P) segments (80% / 41% / 20%) "
        "and correctly stops trusting green on failed (N) segments (37% / 17% / 7%), confirming the "
        "Strip policy."
    ))
    add_para(doc, (
        "Residual risk — ‘green leakage’: 2,192 of 23,261 words (9.4%) are wrong yet colored green; "
        "605 of those sit in below-0.65 segments the Strip tier removes coloring from. The dangerous "
        "leaks are confident numeric/entity swaps — ‘1 billion→1 million’ (conf 0.965, off by 1000×), "
        "‘2011→2000’ (0.894), ‘1024→24’ (0.958) — which is why numbers are now capped at yellow."
    ))

    # 1.4 Human IS
    add_heading(doc, "1.4 Human Intelligibility Baseline (Path B Estimates)", 2)
    add_para(doc, (
        "To answer ‘how good is 2.5/5, really?’ we estimated what humans would score on the same wild "
        "1,497-segment task using the IS formula (these are pre-study estimates, validated to reproduce "
        "the model’s own measured per-bin IS within ±0.05; a human pilot is not yet run):"
    ))
    add_styled_table(doc,
        ["Reader", "Estimated IS (mid)", "vs Model 2.55"],
        [
            ["Lay hearing, no context", "0.92", "Model wins by +1.6"],
            ["Deaf adult, no context", "2.74", "Rough tie (−0.2)"],
            ["Expert / forensic lip-reader, no context", "3.03", "Model loses by −0.5"],
            ["Lay + topic + model-assisted review", "3.83", "Model loses by −1.3"],
            ["Model alone (measured, MBR)", "2.55", "—"],
        ],
        col_widths=[3.2, 1.5, 1.6]
    )
    add_para(doc, (
        "The deployment story falls out of the last two rows: the model alone roughly matches an "
        "unaided deaf lip-reader and loses to an unaided expert by ~0.5 IS, but model + a context-aware "
        "human reviewer scores ~1.3 IS higher than either — which is exactly the workflow the confidence "
        "report enables. A separate isolation test showed the IS metric covertly rewards the model for "
        "guessing rather than skipping words (the length-ratio term), worth ~+0.4 IS for a lay reader."
    ))

    # 1.5 Llama 3.1
    add_heading(doc, "1.5 Llama 3.1 Migration — Prepared, Blocked on Data", 2)
    add_para(doc, (
        "Groundwork was laid in Q2 to swap the frozen Llama-2-7B backbone for Llama-3.1-8B. The target "
        "was set to the BASE model (not Instruct) — matching the original paper and the 2025 VSR "
        "literature, since visual features are injected via embeddings, not chat instructions. The "
        "integration is code-complete and validated end-to-end but deliberately not yet on the "
        "production decode path."
    ))
    add_bullet_bold_value(doc, "Load test: ", "PASS — 4-bit load 5.70 GB VRAM, LoRA wrap 9.4M trainable params (0.117%)")
    add_bullet_bold_value(doc, "Fairseq smoke test: ", "FULL PASS — 2-update real training, loss 11.5→8.7, full data→encoder→projector→LLM chain validated")
    add_bullet_bold_value(doc, "Both models downloaded: ", "Base (15 GB) + Instruct backup (30 GB); Meta license granted May 11")
    add_bullet_bold_value(doc, "Blocked: ", "LRS3 dataset delisted from every public mirror; AWS p4d.24xlarge quota pending")
    add_bullet_bold_value(doc, "Expected gain: ", "modest — ~1–2 pp WER; the frozen visual encoder, not the LLM, is the real bottleneck")
    doc.add_page_break()


# ═══════════════════════════════════════════════════
# PART 2: CLIENT ENGAGEMENT & OPERATIONS
# ═══════════════════════════════════════════════════

def part_2_client(doc):
    add_heading(doc, "Part 2. Client Engagement & Operations", 1, bookmark_id="part2")

    # 2.1 The meeting
    add_heading(doc, "2.1 The Stakeholder Meeting & Use-Case Reframe", 2)
    add_para(doc, (
        "The quarter centered on a 2-hour, ~12-person stakeholder meeting with a mixed audience: the "
        "existing client team and a second internal team, two technical reps from a peer/partner "
        "company, and two representatives of a prospective new client. The meeting drove a substantial "
        "narrative reframe of how Argos is positioned."
    ))
    add_para(doc, (
        "The system was repositioned away from media-transcription/accessibility framing toward its "
        "actual use case: recovering speech from footage where audio does not exist or cannot be "
        "obtained (investigations, archive review, footage from cameras without microphones), indoor "
        "and outdoor, from 20-second to hour-long clips. The comparative anchor that resonated: there "
        "is no incumbent vendor — the alternatives are manual expert lip-reading (45–52% word "
        "accuracy) or nothing, and Argos + a human reviewer reaches 55–70% with near-zero "
        "hallucination risk and audits in minutes per video."
    ))
    add_styled_table(doc,
        ["Client priority", "Rank", "Addressed this quarter?"],
        [
            ["Trust / truthfulness of output", "1", "Yes — confidence stack + MBR"],
            ["Multi-speaker support", "2", "Roadmap (Q3)"],
            ["Multi-language (Arabic)", "3", "Roadmap (future)"],
            ["Ease of use", "4", "Yes — UI bundle, user guide, deploy kit"],
        ],
        col_widths=[2.8, 0.8, 2.6]
    )
    add_para(doc, (
        "Explicitly NOT client concerns: inference cost, latency, per-minute pricing, language "
        "detection. The dominant request across deck reviews was to strip research jargon from "
        "client-visible material (no LoRA / PCA / κ / NIV / ‘MBR’ on slides — ‘MBR’ became ‘20 "
        "alternative readings, aggregated’) and to lead with graceful-degradation examples rather than "
        "worst-case failures."
    ))

    # 2.2 Decks
    add_heading(doc, "2.2 Presentation Suite Rebuilt", 2)
    add_para(doc, (
        "Three deck families were produced and iterated heavily over Apr 30 – May 11, every round "
        "driven by 17–43 specific written feedback items per review:"
    ))
    add_styled_table(doc,
        ["Deck", "Audience", "Scope"],
        [
            ["Client deck (Rounds 0–10)", "Existing + prospective client", "48 → ~66 slides; surveillance-reframe, trust section, case studies"],
            ["Academic deck (‘For Orchard’)", "Research peers / internal", "~88 slides, 5-section narrative, all numbers rebased on MBR"],
            ["v13 deck", "Mixed / archive", "99-slide rebuild via dual-mode generator"],
        ],
        col_widths=[1.9, 1.9, 2.5]
    )
    add_para(doc, (
        "All decks were rebased onto MBR-default numbers and audited slide-by-slide via PNG renders "
        "(the project rule: never proof a PPTX by PDF conversion). The full round-by-round history is in "
        "presentation_materials_20260224/DECK_CHANGELOG.md and docs/paper/presentation-remarks-log.md."
    ))

    # 2.3 Competitive
    add_heading(doc, "2.3 Competitive Landscape", 2)
    add_para(doc, (
        "Competitor questions were identified as a gap and answered with a one-page Q&A cheat-sheet and "
        "a fuller reference. Named comparators: Symphonic Labs (‘Read Their Lips’), LipReadPro, "
        "big-tech speech APIs, and 121 Captions (human forensic lip-readers). Key talking points: "
        "LipReadPro’s ‘95% accuracy’ claim is implausible against the published LRS3 state of the art "
        "(~27% WER) and is web-upload, not on-premise — a non-starter for air-gapped clients."
    ))

    # 2.4 Deployment
    add_heading(doc, "2.4 Client Deployment: Single Image, Air-Gapped", 2)
    add_para(doc, (
        "Argos was deployed to an air-gapped Windows 11 + RTX 5090 (Blackwell) client running Docker "
        "Desktop, fully offline after install. The deployment model was deliberately changed from Q1’s "
        "layered patch overlays to a single, clean, self-contained Docker image (one tag, one tarball, "
        "full rebuild on any source change) — layered patches caused operator confusion and were "
        "deprecated."
    ))
    add_bullet_bold_value(doc, "Image: ", "vsp-llm-pipeline (client-build-001→003), ~58 GB after ~8 GB bloat cleanup")
    add_bullet_bold_value(doc, "Two entry points: ", "vsp-start.sh (UI mode, browser at :8080) and vsp-pipeline.sh (CLI mode, folder picker)")
    add_bullet_bold_value(doc, "Windows parity: ", "PowerShell launchers + pre/post/diagnostics .ps1 self-test scripts")
    add_bullet_bold_value(doc, "Offline hardening: ", "HF_HUB_OFFLINE / TRANSFORMERS_OFFLINE set; Whisper + spaCy models pre-baked into the image")
    add_bullet_bold_value(doc, "Update mechanism: ", "apply_update.sh <tarball> <tag> for whole-image swaps")

    # 2.5 Bugs & lessons
    add_heading(doc, "2.5 Deployment Bugs & Lessons", 2)
    add_para(doc, (
        "17 distinct bugs were found and fixed across three build cycles plus one in-field follow-up "
        "(a real client decode crash on May 12, latent since the Mission 4/6 merge). The recurring "
        "theme: a clean exit code does not mean features worked. Validation now checks actual artifacts "
        "(does aggregated.json exist, does it contain hyp_mbr, does report.csv carry the NIV column)."
    ))
    add_bullet(doc, "Syncing lib/ alone is insufficient — auto_avsr/ and the local fairseq fork must move together; two fairseq installs live side-by-side and are reconciled by a decode.sh monkey-patch.")
    add_bullet(doc, "Read-only input mounts forced a VSP_TRANSCRIPTIONS_DIR env var; the server bound to localhost-in-container needed VSP_UI_HOST/PORT overrides.")
    add_bullet(doc, "spaCy has no cp39 wheel — it must be pre-baked (~21 min build cost), or features silently degrade.")
    add_bullet(doc, "Client terminology rarely matches button labels (‘restart’, ‘loading’, ‘frozen’, ‘.mtk’). Two diagnostic questions resolve most calls: what is on screen, and what was clicked immediately before.")
    add_bullet(doc, "Always stamp documents with their version date — the client mistook a February roadmap deck for current state.")
    doc.add_page_break()


# ═══════════════════════════════════════════════════
# PART 3: TECHNICAL DEVELOPMENT
# ═══════════════════════════════════════════════════

def part_3_technical(doc):
    add_heading(doc, "Part 3. Technical Development", 1, bookmark_id="part3")
    add_para(doc, (
        "Beyond the research wiring (MBR default, confidence sidecars, agreement bands), Q2 shipped a "
        "bundle of client-facing usability features driven directly by deployment feedback."
    ))

    # 3.1 Report plumbing
    add_heading(doc, "3.1 Confidence & MBR in the Report", 2)
    add_para(doc, (
        "report.csv gained sentence_confidence, min_word_conf, and n_low_conf_words columns plus a tier "
        "label; report.html renders a per-segment Confidence line in a distinct blue/orange/purple "
        "palette (kept separate from the green/yellow/red accuracy palette to avoid confusion). The "
        "decode path now writes per-word score sidecars by default (VSP_OUTPUT_SCORES=1), with zero new "
        "Python dependencies. The client-facing dark view was renamed argos_demo → "
        "confidence_breakdown."
    ))

    # 3.2 UX bundle
    add_heading(doc, "3.2 Client UX Bundle", 2)
    add_para(doc, (
        "A cluster of UI fixes shipped from real operator pain points: a fix for the transcription "
        "restart loop, host-path visibility (VSP_HOST_INPUT_DIR shows the operator a real Windows "
        "path), archive handling, and drag-drop hardening with a file-picker fallback for an Edge "
        "drag-drop regression."
    ))

    # 3.3 Audio injection
    add_heading(doc, "3.3 Audio-Aligned Transcription Injection", 2)
    add_para(doc, (
        "A new ‘Inject from audio’ workflow lets an operator feed a clean transcript from separate "
        "audio into a silent or bad-audio video, using a two-offset alignment scheme (--audio-start and "
        "--video-start). Injected transcripts are cached (type: audio-injected) and auto-reused on "
        "subsequent runs, with a distinct purple badge in the report. Available as both a CLI script "
        "and a UI form."
    ))

    # 3.4 Watch with CC
    add_heading(doc, "3.4 Watch-with-CC", 2)
    add_para(doc, (
        "A whole-video closed-caption overlay was added so operators can review a full clip with the "
        "decoded text burned in as captions, in addition to the existing per-segment players."
    ))

    # 3.5 Formats
    add_heading(doc, "3.5 Expanded Video Format Support", 2)
    add_para(doc, (
        "Accepted input formats expanded from 6 to 11 containers, adding .mts / .m2ts / .ts (AVCHD "
        "camcorders — a real client gap), .wmv, and .flv. config.py::SUPPORTED_EXTENSIONS is the single "
        "source of truth with four mirror sites; the change was test-driven on generated fixtures, "
        "since the failure mode is silent rejection rather than a decode error."
    ))

    # 3.6 Burns
    add_heading(doc, "3.6 Burned-Video Improvements", 2)
    add_para(doc, (
        "Burned subtitle videos now use the MBR hypothesis matching the HTML report and per-word "
        "confidence coloring, with fixes for a dark-top-half rendering bug and broken-source fallback. "
        "A per-video CLAHE brightness enhancement was added to improve lip-reading on dark footage."
    ))

    # 3.7 Tooling
    add_heading(doc, "3.7 Deployment Kit & Tooling", 2)
    add_para(doc, (
        "The May-2026 deployment kit bundles the launchers, Windows PowerShell counterparts, cp310 "
        "wheels, air-gapped install scripts, and the container update mechanism. A cross-session "
        "/handover skill was also added to pass working context between Claude Code sessions and "
        "between the two working accounts on this project."
    ))
    doc.add_page_break()


# ═══════════════════════════════════════════════════
# PART 4: INSIGHTS, LESSONS & NEXT QUARTER
# ═══════════════════════════════════════════════════

def part_4_insights(doc):
    add_heading(doc, "Part 4. Insights, Lessons & Next Quarter", 1, bookmark_id="part4")

    # 4.1 Research insights
    add_heading(doc, "4.1 Key Research Insights", 2)
    add_bullet_bold_value(doc, "Confidence is conditional, not absolute: ",
        "the same green word is 92.8% reliable in a clean segment and 18% in a noisy one. A flat threshold cannot work; the segment tier must gate the per-word color.")
    add_bullet_bold_value(doc, "Beam agreement beats raw confidence: ",
        "at high confidence, whether the beam agrees is ~2× more predictive of correctness than the confidence value itself — the single most useful confidence signal found this quarter.")
    add_bullet_bold_value(doc, "Aggregation helps where IS cannot see it: ",
        "MBR and voting are statistically better to the LLM judge but invisible to the IS metric — they rescue marginal segments below the IS tier boundary. Pick the metric to the decision.")
    add_bullet_bold_value(doc, "The metric rewards guessing: ",
        "IS’s length-ratio term covertly rewards the model for filling in uncertain words rather than skipping them — a known bias to watch when comparing to humans, who skip.")
    add_bullet_bold_value(doc, "The model is reviewer-grade, not expert-grade: ",
        "alone it matches an unaided deaf lip-reader and trails an expert by ~0.5 IS, but paired with a human reviewer it exceeds both by ~1.3 IS. The product is the pairing.")

    # 4.2 Operational lessons
    add_heading(doc, "4.2 Operational & Client Lessons", 2)
    add_bullet_bold_value(doc, "One clean image beats layered patches: ",
        "a single tag/tarball with full rebuild removed the operator confusion that layered overlays caused; the patch-overlay update flow was deprecated.")
    add_bullet_bold_value(doc, "Exit code 0 is not validation: ",
        "features degraded silently while the pipeline exited cleanly. Validate on real artifacts, not return codes.")
    add_bullet_bold_value(doc, "Test input-boundary changes on real fixtures: ",
        "new file types / host paths / env vars fail as silent rejection, which syntax checks never catch — they need an end-to-end test with a real file.")
    add_bullet_bold_value(doc, "Translate client language: ",
        "a non-engineer operator’s words rarely map to UI labels; ask what is on screen and what was clicked before guessing.")
    add_bullet_bold_value(doc, "Strip jargon for client audiences: ",
        "every deck review asked to remove research terminology; client trust comes from plain language and graceful-degradation examples, not metrics.")

    # 4.3 Timeline
    add_heading(doc, "4.3 Q2 Development Timeline", 2)
    add_styled_table(doc,
        ["Date", "Milestone"],
        [
            ["Apr 30", "Per-word confidence live in every pipeline run; client deck born (Round 0)"],
            ["May 1", "N-best aggregation (5 methods) shipped; band-reliability-by-segment analysis"],
            ["May 2", "MBR becomes default output; agreement-aware bands; v3 judge confirms +2.7 pp Y+P"],
            ["May 3", "Human-IS Path B estimates; competitive cheat-sheet; client deck Rounds 7–9"],
            ["May 6–8", "Academic / ‘For Orchard’ deck rebuilt on MBR numbers; full PNG slide audit"],
            ["May 6–13", "Single-image deployment to air-gapped Win 11 + RTX 5090; 17 bugs fixed"],
            ["May 11–12", "Llama 3.1 prep — load + fairseq smoke tests pass; in-field crash fix (Bug 17)"],
            ["May 25", "Client UX bundle, user guide, audio-injection, 11-format support committed"],
        ],
        col_widths=[1.2, 5.3]
    )

    # 4.4 Next quarter
    add_heading(doc, "4.4 Next Quarter (Q3 2026) To-Do", 2)
    add_heading(doc, "Research & Model", 3)
    add_bullet(doc, "Acquire LRS3 (out-of-band) and launch Llama-3.1-8B base fine-tuning; re-baseline vs paper and current WER/IS.")
    add_bullet(doc, "Run the Path A human-IS pilot (n≈50, deaf-adult-with-context stratum first) to confirm the estimated baselines.")
    add_bullet(doc, "Confidence calibration (Mission 4.1) once B3 sidecar data is complete; quality-signal additivity test (Mission 4.2).")
    add_heading(doc, "Client-Driven Features", 3)
    add_bullet(doc, "Multi-speaker support (client priority #2).")
    add_bullet(doc, "Arabic language path (client priority #3).")
    add_bullet(doc, "Save trained k-means weights to the golden list (Mission 16).")
    add_heading(doc, "Operations", 3)
    add_bullet(doc, "Repaint demo-MP4 burn overlays + UI screenshots to the blue/orange/purple palette.")
    add_bullet(doc, "Rebuild the vsp_docker/ image to fold in the per-commit container-overlay changes deferred this quarter.")

    doc.add_paragraph()
    add_para(doc, (
        "This summary covers Argos work for Q2 2026 (April 1 – June 15). It is the standalone "
        "quarterly companion to the living Project Summary (last full snapshot March 11, 2026). For "
        "detailed analysis see the per-topic reports under docs/ (confidence, beam-search, evaluation, "
        "finetuning) and the deployment lessons under docs/guides/."
    ), italic=True, color=C_GRAY)


# ═══════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════

def main():
    doc = Document()

    # Base style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    create_cover_page(doc)
    create_toc(doc)
    add_header_footer(doc)

    quarter_at_a_glance(doc)
    part_1_research(doc)
    part_2_client(doc)
    part_3_technical(doc)
    part_4_insights(doc)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_FILE))
    print(f"✓ Saved: {OUTPUT_FILE}")
    print(f"  Paragraphs: {len(doc.paragraphs)}  Tables: {len(doc.tables)}")


if __name__ == "__main__":
    main()
