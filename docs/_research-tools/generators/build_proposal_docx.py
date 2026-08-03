"""Build the two new-hire mission proposal docx files from their markdown sources.

Outputs:
- docs/backlog/new_hire_project_video_quality.docx
- docs/backlog/new_hire_project_multi_speaker.docx

Follows STYLE_GUIDE.md §1: Calibri, A4, 2.5/2.0 cm margins, peacock cover,
"Argos -- The Orchard" header, page numbers in footer, author Yoad Oxman.
"""

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

SCRIPT_DIR = Path(__file__).resolve().parent
ASSETS_DIR = SCRIPT_DIR.parent / "assets"
REPO_ROOT = SCRIPT_DIR.parent.parent.parent
OUT_DIR = REPO_ROOT / "docs" / "backlog"

PRIMARY = RGBColor(0x14, 0x4B, 0x55)
BODY = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x66, 0x66, 0x66)
ACCENT = RGBColor(0x00, 0x7A, 0x8A)


def _styled_run(p, text, *, bold=False, italic=False, size=11, color=BODY, font="Calibri"):
    run = p.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    return run


def _add_inline_runs(p, text, *, size=11, color=BODY):
    """Render markdown inline emphasis: **bold**, *italic*, `code`."""
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            _styled_run(p, part[2:-2], bold=True, size=size, color=color)
        elif part.startswith("*") and part.endswith("*"):
            _styled_run(p, part[1:-1], italic=True, size=size, color=color)
        elif part.startswith("`") and part.endswith("`"):
            _styled_run(p, part[1:-1], size=size, color=ACCENT, font="Consolas")
        else:
            _styled_run(p, part, size=size, color=color)


def _heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    sizes = {1: 18, 2: 14, 3: 12}
    _styled_run(p, text, bold=True, size=sizes.get(level, 12), color=PRIMARY)
    style_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}
    if level in style_map:
        try:
            p.style = doc.styles[style_map[level]]
        except KeyError:
            pass
    return p


def _para(doc, text, *, size=11, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if italic:
        _styled_run(p, text, italic=True, size=size)
    else:
        _add_inline_runs(p, text, size=size)
    return p


def _bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(2)
    _add_inline_runs(p, text)
    return p


def _hrule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_break(WD_BREAK.PAGE)


def _set_row_cant_split(row):
    trPr = row._tr.get_or_add_trPr()
    cantSplit = OxmlElement("w:cantSplit")
    trPr.append(cantSplit)


def _code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    _styled_run(p, text, size=10, color=BODY, font="Consolas")


def _setup_page(doc):
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    return section


def _setup_header(doc, section):
    hp = section.header.paragraphs[0]
    text_width = section.page_width - section.left_margin - section.right_margin
    hp.paragraph_format.tab_stops.add_tab_stop(text_width, WD_TAB_ALIGNMENT.RIGHT)
    logo_path = ASSETS_DIR / "logo.png"
    if logo_path.exists():
        run = hp.add_run()
        try:
            run.add_picture(str(logo_path), height=Cm(0.7))
        except Exception:
            pass
    hp.add_run("\t")
    right = hp.add_run("Argos — The Orchard")
    right.font.name = "Calibri"
    right.font.size = Pt(8)
    right.italic = True
    right.font.color.rgb = MUTED


def _setup_footer(doc, section):
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run._r.append(instr)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED


def _cover(doc, *, title, subtitle, metadata_lines):
    for _ in range(2):
        doc.add_paragraph()
    peacock = ASSETS_DIR / "peacock.png"
    if peacock.exists():
        doc.add_picture(str(peacock), width=Cm(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(18)
    _styled_run(title_p, "ARGOS", bold=True, size=48, color=PRIMARY)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _styled_run(subtitle_p, title, bold=True, size=22, color=BODY)

    if subtitle:
        sub2 = doc.add_paragraph()
        sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _styled_run(sub2, subtitle, size=14, color=MUTED)

    orch = doc.add_paragraph()
    orch.alignment = WD_ALIGN_PARAGRAPH.CENTER
    orch.paragraph_format.space_before = Pt(12)
    _styled_run(orch, "The Orchard", size=20, color=BODY)

    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.paragraph_format.space_before = Pt(36)
    _styled_run(author, "Yoad Oxman", size=14, color=BODY)

    for line in metadata_lines:
        m = doc.add_paragraph()
        m.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _styled_run(m, line, size=11, color=MUTED)

    pb = doc.add_paragraph()
    pb.add_run().add_break(WD_BREAK.PAGE)


def _render_md(doc, md_text):
    """Minimal markdown renderer: H1/H2/H3, paragraphs, bullets, code blocks, tables, hrules.

    Skips the metadata block at top (lines starting '**Author:**', etc., between the H1 and the first hrule).
    """
    lines = md_text.splitlines()
    i = 0
    n = len(lines)
    # Skip the front matter — H1 title line and metadata until first --- hrule
    in_front_matter = True
    while i < n:
        line = lines[i]
        stripped = line.rstrip()

        if in_front_matter:
            if stripped == "---":
                in_front_matter = False
                i += 1
                continue
            i += 1
            continue

        if not stripped.strip():
            i += 1
            continue

        if stripped == "---":
            # Treat the body/schedule separator as a hard page break so the
            # schedule table always starts on its own page.
            _page_break(doc)
            i += 1
            continue

        if stripped.startswith("### "):
            _heading(doc, stripped[4:].strip(), level=3)
            i += 1
            continue
        if stripped.startswith("## "):
            _heading(doc, stripped[3:].strip(), level=2)
            i += 1
            continue
        if stripped.startswith("# "):
            _heading(doc, stripped[2:].strip(), level=1)
            i += 1
            continue

        if stripped.startswith("```"):
            # code block
            i += 1
            code_lines = []
            while i < n and not lines[i].rstrip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            _code_block(doc, "\n".join(code_lines))
            if i < n:
                i += 1
            continue

        if stripped.startswith("- "):
            _bullet(doc, stripped[2:].strip())
            i += 1
            continue

        if stripped.startswith("| ") and i + 1 < n and re.match(r"^\|[-:\s|]+\|$", lines[i + 1].rstrip()):
            # table
            header = [c.strip() for c in stripped.strip("|").split("|")]
            i += 2
            body = []
            while i < n and lines[i].rstrip().startswith("| "):
                row = [c.strip() for c in lines[i].rstrip().strip("|").split("|")]
                body.append(row)
                i += 1
            _render_table(doc, header, body)
            continue

        # default: paragraph (may span multiple lines until blank)
        para_lines = [stripped]
        i += 1
        while i < n and lines[i].strip() and not _is_block_start(lines[i].rstrip()):
            para_lines.append(lines[i].rstrip())
            i += 1
        _para(doc, " ".join(para_lines))


def _is_block_start(line):
    if not line.strip():
        return True
    if line.startswith("#"):
        return True
    if line.startswith("- "):
        return True
    if line.startswith("|"):
        return True
    if line.startswith("```"):
        return True
    if line.strip() == "---":
        return True
    return False


def _render_table(doc, header, body):
    cols = len(header)
    table = doc.add_table(rows=1 + len(body), cols=cols)
    table.autofit = True
    # Header
    hcells = table.rows[0].cells
    for c_idx, txt in enumerate(header):
        p = hcells[c_idx].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        hcells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.TOP
        _styled_run(p, txt, bold=True, size=10, color=PRIMARY)
    _set_row_cant_split(table.rows[0])
    for r_idx, row in enumerate(body, start=1):
        rcells = table.rows[r_idx].cells
        for c_idx, txt in enumerate(row):
            if c_idx >= cols:
                break
            p = rcells[c_idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            rcells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.TOP
            _add_inline_runs(p, txt, size=10)
        _set_row_cant_split(table.rows[r_idx])
    doc.add_paragraph()


def build_doc(md_path, out_path, *, cover_title, cover_subtitle, cover_meta):
    md_text = md_path.read_text()
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    section = _setup_page(doc)
    _setup_header(doc, section)
    _setup_footer(doc, section)
    _cover(doc, title=cover_title, subtitle=cover_subtitle, metadata_lines=cover_meta)
    _render_md(doc, md_text)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"Wrote {out_path}")


def main():
    common_meta = [
        "End-of-training-phase project",
        "5 days",
    ]
    build_doc(
        OUT_DIR / "training_project_video_quality.md",
        OUT_DIR / "training_project_video_quality.docx",
        cover_title="Pre-Decode Video Quality Score",
        cover_subtitle="A small end-of-training-phase project",
        cover_meta=common_meta,
    )
    build_doc(
        OUT_DIR / "training_project_multi_speaker.md",
        OUT_DIR / "training_project_multi_speaker.docx",
        cover_title="Multi-Speaker Lip-Reading",
        cover_subtitle="A small end-of-training-phase project",
        cover_meta=common_meta,
    )
    build_doc(
        OUT_DIR / "training_project_additivity_test.md",
        OUT_DIR / "training_project_additivity_test.docx",
        cover_title="Reusable Additivity Test for Quality Signals",
        cover_subtitle="A small end-of-training-phase project",
        cover_meta=common_meta,
    )


if __name__ == "__main__":
    main()
