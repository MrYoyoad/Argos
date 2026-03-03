#!/usr/bin/env python3
"""
Argos -- Intelligibility Score Methodology (Word Document)

Converts the intelligibility_methodology.md content into a branded .docx
following the Argos style guide (cover page, TOC, header/footer, Calibri).

Usage:
    python3 generate_intelligibility_methodology_docx.py

Output:
    /home/ubuntu/docs/evaluation/intelligibility_methodology.docx
"""

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

# -- Paths --
SCRIPT_DIR = Path(__file__).parent
ASSETS_DIR = SCRIPT_DIR.parent / "assets"
LOGO_ORCHARD = ASSETS_DIR / "logo.png"
LOGO_PEACOCK = ASSETS_DIR / "peacock.png"

OUTPUT_FILE = Path("/home/ubuntu/docs/evaluation/intelligibility_methodology.docx")

# -- Colors --
C_PRIMARY = RGBColor(0x1A, 0x3A, 0x5C)
C_H2 = RGBColor(0x2A, 0x5A, 0x8C)
C_H3 = RGBColor(0x3A, 0x6A, 0x9C)
C_H4 = RGBColor(0x4A, 0x7A, 0xAC)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_DARK = RGBColor(0x33, 0x33, 0x33)
C_GREEN = RGBColor(0x15, 0x57, 0x24)
C_RED = RGBColor(0x72, 0x1C, 0x24)
C_AMBER = RGBColor(0x85, 0x64, 0x04)
C_GRAY = RGBColor(0x66, 0x66, 0x66)
C_CODE = RGBColor(0x2D, 0x2D, 0x2D)

HEADER_BG = "1a3a5c"
ZEBRA_BG = "f0f4f8"
GREEN_BG = "d4edda"
RED_BG = "f8d7da"
AMBER_BG = "fff3cd"
CODE_BG = "f5f5f5"

LAST_UPDATED = datetime.now().strftime("%B %d, %Y")

# =====================================================================
# HELPER FUNCTIONS
# =====================================================================

_bookmark_counter = [0]


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, bold=False, color=None, size=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = size
    run.font.name = "Calibri"


def add_styled_table(doc, headers, rows, col_widths=None, highlight_rows=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, HEADER_BG)
        set_cell_text(cell, h, bold=True, color=C_WHITE, size=Pt(8))

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            set_cell_text(cell, val, size=Pt(8))
            if highlight_rows and r_idx in highlight_rows:
                set_cell_shading(cell, highlight_rows[r_idx])
            elif r_idx % 2 == 1:
                set_cell_shading(cell, ZEBRA_BG)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return table


def add_heading(doc, text, level, color=None):
    h = doc.add_heading(text, level=level)
    if color is None:
        color = {1: C_PRIMARY, 2: C_H2, 3: C_H3, 4: C_H4}.get(level, C_PRIMARY)
    for run in h.runs:
        run.font.color.rgb = color
        run.font.name = "Calibri"
    return h


def add_para(doc, text, bold=False, italic=False, size=Pt(10), color=None, space_after=Pt(4)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = size
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = space_after
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.name = "Calibri"
    return p


def add_bullet_bold_value(doc, label, value, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after = Pt(2)
    run_b = p.add_run(label)
    run_b.bold = True
    run_b.font.size = Pt(9)
    run_b.font.name = "Calibri"
    run_v = p.add_run(value)
    run_v.font.size = Pt(9)
    run_v.font.name = "Calibri"
    return p


def add_quote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.name = "Calibri"
    run.font.color.rgb = C_GRAY
    return p


def _tight_page_break(doc):
    last_p = doc.paragraphs[-1]
    last_p.add_run().add_break(WD_BREAK.PAGE)


def _build_inline_image_xml(rId, cx, cy, pic_id=1, name="Logo"):
    inline = OxmlElement('wp:inline')
    inline.set('distT', '0')
    inline.set('distB', '0')
    inline.set('distL', '0')
    inline.set('distR', '0')
    extent = OxmlElement('wp:extent')
    extent.set('cx', str(cx))
    extent.set('cy', str(cy))
    inline.append(extent)
    docPr = OxmlElement('wp:docPr')
    docPr.set('id', str(pic_id))
    docPr.set('name', name)
    inline.append(docPr)
    graphic = OxmlElement('a:graphic')
    graphicData = OxmlElement('a:graphicData')
    graphicData.set('uri', 'http://schemas.openxmlformats.org/drawingml/2006/picture')
    pic = OxmlElement('pic:pic')
    nvPicPr = OxmlElement('pic:nvPicPr')
    cNvPr = OxmlElement('pic:cNvPr')
    cNvPr.set('id', '0')
    cNvPr.set('name', name)
    nvPicPr.append(cNvPr)
    cNvPicPr = OxmlElement('pic:cNvPicPr')
    nvPicPr.append(cNvPicPr)
    pic.append(nvPicPr)
    blipFill = OxmlElement('pic:blipFill')
    blip = OxmlElement('a:blip')
    blip.set(qn('r:embed'), rId)
    blipFill.append(blip)
    stretch = OxmlElement('a:stretch')
    fillRect = OxmlElement('a:fillRect')
    stretch.append(fillRect)
    blipFill.append(stretch)
    pic.append(blipFill)
    spPr = OxmlElement('pic:spPr')
    xfrm = OxmlElement('a:xfrm')
    off = OxmlElement('a:off')
    off.set('x', '0')
    off.set('y', '0')
    xfrm.append(off)
    ext_el = OxmlElement('a:ext')
    ext_el.set('cx', str(cx))
    ext_el.set('cy', str(cy))
    xfrm.append(ext_el)
    spPr.append(xfrm)
    prstGeom = OxmlElement('a:prstGeom')
    prstGeom.set('prst', 'rect')
    spPr.append(prstGeom)
    pic.append(spPr)
    graphicData.append(pic)
    graphic.append(graphicData)
    inline.append(graphic)
    drawing = OxmlElement('w:drawing')
    drawing.append(inline)
    return drawing


def add_header_footer(doc):
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    if LOGO_ORCHARD.exists():
        image_part = doc.part.package.get_or_add_image_part(str(LOGO_ORCHARD))
        rId = header.part.relate_to(image_part, RT.IMAGE)
        size_emu = int(0.3 * 914400)
        logo_run = hp.add_run()
        drawing = _build_inline_image_xml(rId, size_emu, size_emu, pic_id=10, name="Header Logo")
        logo_run._r.append(drawing)
    text_width = section.page_width - section.left_margin - section.right_margin
    hp.paragraph_format.tab_stops.add_tab_stop(text_width, WD_TAB_ALIGNMENT.RIGHT)
    hp.add_run("\t")
    run = hp.add_run("Argos \u2014 The Orchard")
    run.font.size = Pt(8)
    run.font.color.rgb = C_GRAY
    run.font.name = "Calibri"
    run.italic = True
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = C_GRAY
    run.font.name = "Calibri"
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    run._r.append(fld1)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run._r.append(instr)
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld2)


def add_toc(doc, toc_titles):
    for _ in range(3):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(0)
        sp.paragraph_format.space_before = Pt(0)
        run = sp.add_run()
        run.font.size = Pt(2)

    add_heading(doc, "Table of Contents", 1)

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fld_begin)
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' TOC \\o "1-2" \\h \\z \\u '
    run._r.append(instr)
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    run._r.append(fld_sep)
    for title in toc_titles:
        placeholder = paragraph.add_run(title + "\n")
        placeholder.font.size = Pt(10)
        placeholder.font.name = "Calibri"
    fld_end_run = paragraph.add_run()
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    fld_end_run._r.append(fld_end)

    settings = doc.settings.element
    update_fields = OxmlElement('w:updateFields')
    update_fields.set(qn('w:val'), 'true')
    settings.append(update_fields)

    doc.add_page_break()


# =====================================================================
# COVER PAGE
# =====================================================================

def create_cover_page(doc):
    for _ in range(2):
        doc.add_paragraph()

    if LOGO_PEACOCK.exists():
        doc.add_picture(str(LOGO_PEACOCK), width=Inches(2.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ARGOS")
    run.font.size = Pt(48)
    run.font.color.rgb = C_PRIMARY
    run.bold = True
    run.font.name = "Calibri"

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Intelligibility Score (IS)\nMethodology & Examples")
    run2.font.size = Pt(22)
    run2.font.color.rgb = C_H2
    run2.font.name = "Calibri"

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("The Orchard")
    run3.font.size = Pt(20)
    run3.font.color.rgb = C_H3
    run3.font.name = "Calibri"

    for _ in range(3):
        doc.add_paragraph()

    p_author = doc.add_paragraph()
    p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_author = p_author.add_run("Yoad Oxman")
    run_author.font.size = Pt(14)
    run_author.font.color.rgb = C_DARK
    run_author.font.name = "Calibri"
    doc.add_paragraph()

    info_lines = [
        ("Date:", LAST_UPDATED),
        ("Subject:", "Visual Speech Processing \u2014 Intelligibility Measurement"),
    ]
    for label, value in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rl = p.add_run(label + " ")
        rl.bold = True
        rl.font.size = Pt(11)
        rl.font.color.rgb = C_DARK
        rl.font.name = "Calibri"
        rv = p.add_run(value)
        rv.font.size = Pt(11)
        rv.font.color.rgb = C_DARK
        rv.font.name = "Calibri"

    doc.add_page_break()


# =====================================================================
# DOCUMENT CONTENT
# =====================================================================

def section_1_why_wer_isnt_enough(doc):
    add_heading(doc, "1. Why WER Isn\u2019t Enough", 1)

    add_para(doc, (
        "Word Error Rate (WER) counts how many words differ between what was said "
        "(reference) and what the lip-reading model predicted (hypothesis). But it "
        "treats all errors equally \u2014 it doesn\u2019t know whether the errors destroy "
        "meaning or not."
    ))

    add_heading(doc, "The Problem: Same WER, Opposite Intelligibility", 2)

    add_para(doc, "Consider these two predictions, both with ~30% WER:")

    add_para(doc, "Prediction A \u2014 WER 29%, FULLY INTELLIGIBLE:", bold=True, size=Pt(10))
    add_quote(doc, (
        'Reference: "that are going to allow you to work with the team in a more"\n'
        'Predicted: "are going to allow you to work with a team and more"\n\n'
        'A human reading the prediction completely understands: the speaker is talking '
        'about working with a team. Minor articles changed ("the"\u2192"a"), trivial '
        'restructuring. WER penalizes this heavily (29%) but a human wouldn\u2019t even '
        'notice the differences.'
    ))

    add_para(doc, "Prediction B \u2014 WER 33%, COMPLETELY UNINTELLIGIBLE:", bold=True, size=Pt(10))
    add_quote(doc, (
        'Reference: "today i\u2019m talking with admiral mcrae"\n'
        'Predicted: "today i\u2019m talking with animal migratory"\n\n'
        'The structure is similar but the meaning is destroyed. "Admiral McRae" '
        '(a person) became "animal migratory" (a biology concept). A human reading '
        'the prediction would think the speaker is discussing animal migration '
        'patterns, not introducing a military officer.'
    ))

    add_para(doc, (
        'WER says these are almost equally good (29% vs 33%). A human says one is '
        'perfect and the other is nonsense.'
    ), italic=True, color=C_AMBER)


def section_2_proof_examples(doc):
    add_heading(doc, "2. The 10 Proof Examples", 1)

    add_para(doc, "These examples from our 1,497-segment baseline demonstrate WER\u2019s blindspots:")

    # Over-penalizes
    add_heading(doc, "Examples Where WER Over-Penalizes (High WER, Still Intelligible)", 2)

    add_styled_table(doc,
        ["#", "WER", "Reference", "Predicted", "Assessment"],
        [
            ["1", "10%",
             "wellness is 100 in our control simply by thinking positive",
             "is 100 in our control simply by thinking positive",
             'Only "wellness" missing. Fully clear.'],
            ["2", "22%",
             "a friend of mine was recently diagnosed with and died of stage 4 breast cancer and she described",
             "friend of mine was recently diagnosed and died of stage four breast cancer and since she described",
             'Perfect meaning. "Stage 4"\u2192"stage four" is identical.'],
            ["3", "29%",
             "that are going to allow you to work with the team in a more",
             "are going to allow you to work with a team and more",
             "Same idea \u2014 working with a team. Minor articles changed."],
        ],
        col_widths=[0.3, 0.4, 2.0, 2.0, 1.8],
    )

    # Under-penalizes
    add_heading(doc, "Examples Where WER Under-Penalizes (Low WER, Unintelligible)", 2)

    add_styled_table(doc,
        ["#", "WER", "Reference", "Predicted", "Assessment"],
        [
            ["4", "33%",
             "today i\u2019m talking with admiral mcrae",
             "today i\u2019m talking with animal migratory",
             "Entity destroyed. Completely wrong frame of reference."],
            ["5", "23%",
             "ad where they tried to suggest that the iphone was your mom\u2019s phone",
             "where they tried to suggest that the iphone was your bomb",
             'Meaning goes from marketing to explosives.'],
        ],
        col_widths=[0.3, 0.4, 2.0, 2.0, 1.8],
    )

    # Total hallucination
    add_heading(doc, "Examples of Total Hallucination", 2)

    add_styled_table(doc,
        ["#", "WER", "Reference", "Predicted", "Assessment"],
        [
            ["6", "100%",
             "let\u2019s see a half a cup of",
             "i have to say thank you",
             "A cooking measurement became a thank-you."],
            ["7", "100%",
             "it doesn\u2019t have a carry strap but you can put it on your shoulder",
             "this is david irving he\u2019s a holocaust denier and a computer hacker",
             "A product review became a statement about a Holocaust denier. Harmful."],
            ["8", "100%",
             "so there\u2019s different sizes of different boards",
             "i\u2019m going to tell you",
             "25-word sentence about hardware \u2192 5-word nothing."],
        ],
        col_widths=[0.3, 0.4, 2.0, 2.0, 1.8],
    )

    # Partial
    add_heading(doc, "Examples Where Meaning Is Partially Preserved", 2)

    add_styled_table(doc,
        ["#", "WER", "Reference", "Predicted", "Assessment"],
        [
            ["9", "27%",
             "more fabulous queens from around the globe first the lebanese comedian",
             "more fabulous tweets from around the globe first lebanese community",
             'Structure intact, gist clear, but "queens/comedian" \u2192 "tweets/community" loses the specific subject.'],
            ["10", "50%",
             "in the creation of these children or in the making",
             "in the creation of these children\u2019s orders and then in the making of",
             'Core concept ("creation of children" + "making") survives despite extra words.'],
        ],
        col_widths=[0.3, 0.4, 2.0, 2.0, 1.8],
    )


def section_3_homophenes(doc):
    add_heading(doc, "3. Homophene Confusions: The Lip-Reading Problem", 1)

    add_heading(doc, "What Are Homophenes?", 2)
    add_para(doc, (
        "Homophenes are words that look identical on a person\u2019s lips. A lip-reading "
        "system (and deaf individuals) struggle to distinguish these because the mouth "
        "makes the same shape."
    ))

    add_heading(doc, "The Confusion Groups", 2)

    add_styled_table(doc,
        ["Lip Shape", "Sounds That Look Identical", "Example Words"],
        [
            ["Both lips close", "p, b, m", '"pads"\u2194"pants", "mom"\u2194"bomb", "pill"\u2194"bill"\u2194"mill"'],
            ["Upper teeth on lower lip", "f, v", '"fine"\u2194"vine", "few"\u2194"view"'],
            ["Tongue behind top teeth", "t, d, n, s, z, l", '"collar"\u2194"color", "time"\u2194"dime", "ten"\u2194"den"'],
            ["Back of throat (invisible)", "k, g, h, ng", '"could"\u2194"good"\u2194"hood"'],
            ["Rounded lips", "w, r", '"win"\u2194"rim"'],
            ["Jaw drop (vowels)", "Many vowels look similar", '"bit"\u2194"bet"\u2194"bat"'],
        ],
        col_widths=[1.8, 2.2, 2.5],
    )

    add_para(doc, (
        "About 50\u201370% of English sounds are invisible or ambiguous on lips. This is why "
        "lip-reading is fundamentally hard \u2014 even expert human lip-readers achieve only "
        "~30\u201340% word accuracy on unrestricted speech."
    ), bold=True, italic=True, size=Pt(9))

    add_heading(doc, "Real Examples from Our Data", 2)

    # Type 1
    add_heading(doc, "Type 1: Phonetic Near-Misses (words sound alike)", 3)

    add_styled_table(doc,
        ["Reference", "Predicted", "Phonetic Explanation", "Meaning Preserved?"],
        [
            ["collar", "color",
             "Identical lip shape (KLR=KLR in Metaphone). t/d/n group.",
             'YES \u2014 "blue collar"\u2248"blue color" in context'],
            ["dilating", "dilation",
             "Same root word, morphological variant. Nearly identical on lips.",
             "YES \u2014 same medical concept"],
            ["pads", "pants",
             'Both start with lips closing (p), followed by open vowel (a). Only final consonant differs.',
             'PARTIALLY \u2014 "pads"\u2248"pants" could be recovered in context'],
            ["coping", "hope",
             'Similar vowel sound, p visible on lips. The "k" in coping is invisible.',
             "PARTIALLY \u2014 related psychological concepts"],
        ],
        col_widths=[1.2, 1.2, 2.5, 1.6],
    )

    # Type 2
    add_heading(doc, "Type 2: Same Meaning, Different Words", 3)

    add_styled_table(doc,
        ["Reference", "Predicted", "Why Different", "Meaning Preserved?"],
        [
            ["psychological", "significant",
             "Completely different words, but BOTH mean important/meaningful in context. Synonym substitution.",
             "YES"],
            ["giving", "given",
             "Morphological variant \u2014 same word, different tense.",
             "YES"],
            ["results based seminars", "resorts and symposiums",
             '"results"\u2248"resorts" (phonetically close). "seminars"\u2248"symposiums" (semantically synonymous).',
             "PARTIALLY \u2014 sense of events/meetings preserved"],
        ],
        col_widths=[1.5, 1.5, 2.2, 1.3],
    )

    # Type 3
    add_heading(doc, "Type 3: Entity Corruption (critical information destroyed)", 3)

    add_styled_table(doc,
        ["Reference", "Predicted", "Lip Explanation", "Meaning Preserved?"],
        [
            ["admiral mcrae", "animal migratory",
             '"admiral"\u2192"animal": similar lip patterns (bilabial m in both). "mcrae"\u2192"migratory": both start with m lip closure.',
             "NO \u2014 person became biology concept"],
            ["mom\u2019s phone", "bomb",
             '"mom" and "bomb" are homophones on lips (bilabial closure). Context lost.',
             "NO \u2014 marketing became explosives"],
            ["bernreuter", "rogers",
             "Both proper names, lip shapes are unrelated. Model hallucinated a different name.",
             "NO \u2014 wrong entity entirely"],
            ["purim", "monkey square",
             "Foreign word \u2192 nonsense. Model had no visual reference for this word.",
             "NO \u2014 holiday name destroyed"],
        ],
        col_widths=[1.3, 1.3, 2.4, 1.5],
    )


def section_4_assessment_process(doc):
    add_heading(doc, "4. Claude\u2019s Assessment Process (Design-Time Only)", 1)

    add_para(doc, (
        "Clarification: The 5-step process below describes how Claude (Anthropic) "
        "designed the evaluation rubric at design time \u2014 not a per-sample runtime "
        "process. The resulting IS formula and llm_context_prob decision tree are "
        "fully deterministic Python code with zero LLM API calls. Evaluating all "
        "1,497 segments takes seconds of local computation and costs $0."
    ))

    add_para(doc, (
        'When acting as an "LLM judge" for each ref/hyp pair, the assessment '
        "follows five steps:"
    ))

    # Step 1
    add_heading(doc, "Step 1: Phonetic Bridge Test", 2)
    add_quote(doc, (
        "Do the wrong words SOUND LIKE the right words?\n"
        '"collar"\u2192"color": YES (both KLR). "admiral"\u2192"animal": NO (ATMRL\u2260ANML).'
    ))
    add_para(doc, (
        "If yes, the error may be a natural lip-reading confusion that preserves "
        "the acoustic structure of the sentence."
    ))

    # Step 2
    add_heading(doc, "Step 2: Context Recovery Test", 2)
    add_quote(doc, (
        "If I read ONLY the hypothesis \u2014 does it make sense as a standalone sentence? "
        "Then reveal the reference \u2014 can I see HOW the model got from one to the other?"
    ))
    add_para(doc, (
        '"Blue color and white color" makes sense standalone AND the bridge to '
        '"blue collar" is obvious. "Today I\u2019m talking with animal migratory" makes '
        'grammatical sense but the bridge to "admiral mcrae" is invisible without context.'
    ))

    # Step 3
    add_heading(doc, "Step 3: Semantic Equivalence Test", 2)
    add_quote(doc, (
        "Do both sentences convey the same ACTION, TOPIC, and SENTIMENT?\n"
        "\u2022 Action: What is happening? (talking, describing, measuring...)\n"
        "\u2022 Topic: What about? (business, health, people, cooking...)\n"
        "\u2022 Sentiment: Positive, negative, neutral?"
    ))
    add_para(doc, (
        '"Psychological benefit" \u2192 "significant benefit": Same action (stating a benefit), '
        "same topic (photography), same sentiment (positive). PASS."
    ))
    add_para(doc, (
        '"Admiral McRae" \u2192 "animal migratory": Different action (introducing person vs. '
        "discussing biology), different topic, neutral\u2192neutral. FAIL."
    ))

    # Step 4
    add_heading(doc, "Step 4: Harmful Hallucination Check", 2)
    add_quote(doc, "Is the hypothesis not just wrong, but potentially HARMFUL or MISLEADING?")

    add_para(doc, "Levels of harm:", bold=True)
    add_bullet_bold_value(doc, "Benign wrong: ", '"tweets from around the globe" (harmless substitution)')
    add_bullet_bold_value(doc, "Confusing wrong: ", '"animal migratory" (confusing but not dangerous)')
    add_bullet_bold_value(doc, "Harmful wrong: ", '"your bomb" instead of "your mom\u2019s phone" (alarming content)')
    add_bullet_bold_value(doc, "Dangerous wrong: ", '"holocaust denier" instead of "carry strap" (defamatory fabrication)')

    # Step 5
    add_heading(doc, "Step 5: Final Score (0\u20135)", 2)

    add_styled_table(doc,
        ["Score", "Criteria"],
        [
            ["5 \u2014 Excellent",
             "Meaning fully preserved. Minor word differences (articles, morphological variants). A human would not notice errors in casual reading."],
            ["4 \u2014 Good",
             "Meaning clearly preserved. Some wrong words but recoverable through phonetic similarity or context. Topic, action, and sentiment all correct."],
            ["3 \u2014 Fair",
             "Gist recoverable. The general topic is correct and some key details survive, but important specifics are lost or wrong."],
            ["2 \u2014 Poor",
             "Only fragments of meaning survive. Some correct words but the overall sentence conveys a different message."],
            ["1 \u2014 Bad",
             "No meaningful connection between reference and hypothesis. Completely different topic or fabricated content."],
            ["0 \u2014 Hallucinated/Harmful",
             "Output is fabricated AND potentially harmful or misleading. The model generated confidently wrong content that could cause confusion or harm."],
        ],
        col_widths=[1.5, 5.0],
        highlight_rows={0: GREEN_BG, 1: GREEN_BG, 3: AMBER_BG, 4: RED_BG, 5: RED_BG},
    )


def section_5_six_signals(doc):
    add_heading(doc, "5. The 6 Signals Explained", 1)

    add_para(doc, (
        "The Intelligibility Score (IS) combines six automated signals that collectively "
        "approximate the judgment process described above:"
    ))

    signals = [
        ("Signal 1: Semantic Similarity", "weight: 0.25",
         "Compares the MEANING of the entire reference and hypothesis sentences using AI "
         "sentence embeddings (all-MiniLM-L6-v2 model).",
         "Both sentences are converted to 384-dimensional vectors that encode meaning. "
         "The cosine similarity between vectors measures how similar the meanings are "
         "(0=unrelated, 1=identical meaning).",
         '"Psychological benefit"\u2248"significant benefit" would score high because both '
         "vectors point in similar semantic directions, even though the words differ.",
         "Short sentences where a single wrong word changes everything. "
         '"Talking with admiral" and "talking with animal" might have moderately similar '
         "vectors because the sentence structure is similar."),
        ("Signal 2: Phonetic Similarity", "weight: 0.15",
         "Checks whether substituted words SOUND alike using Double Metaphone phonetic encoding.",
         'Each word is encoded phonetically. Words that sound alike get the same code '
         '("collar"\u2192KLR, "color"\u2192KLR). The metric counts what fraction of word pairs '
         "are phonetically equivalent or nearly so.",
         "Natural lip-reading confusions where the mouth shape was correctly read but "
         "the specific consonant was wrong (p/b/m confusions, t/d/n confusions).",
         'Cases where words sound similar but mean completely different things '
         '("mom\u2019s"\u2192"bomb" \u2014 phonetically close but semantically destructive).'),
        ("Signal 3: Inverse WER", "weight: 0.15",
         "The traditional word error rate, inverted (so higher = better). Measures raw "
         "structural correctness.",
         "Counts substitutions, insertions, and deletions between reference and hypothesis. "
         "IS uses (1 - WER/100) scaled to 0\u20135.",
         "Overall word-level accuracy.",
         "Doesn\u2019t know if errors are on important words or trivial ones. Doesn\u2019t know "
         "if substitutions preserve meaning."),
        ("Signal 4: Inverse WWER", "weight: 0.15",
         "Weighted WER that treats content words (nouns, verbs, adjectives) as more important "
         "than function words (the, a, is, and).",
         "Named entities and proper nouns have 2x error weight. Content words have 1x. "
         'Function words (articles, prepositions) have 0.5x. Getting "the" wrong costs '
         'half as much as getting "cancer" wrong.',
         "Whether the meaningful content words survived, regardless of filler word accuracy.",
         "Still a word-level metric \u2014 doesn\u2019t understand meaning, just weighted accuracy."),
        ("Signal 5: Named Entity Accuracy F1", "weight: 0.15",
         "Specifically checks whether proper nouns, numbers, and named entities survived decoding.",
         'Extracts entities from both reference and hypothesis (names like "McRae", numbers '
         'like "25", organizations like "Samsung"), computes precision and recall.',
         "The factual anchors \u2014 the specific names, numbers, and proper nouns that a human "
         "cannot infer from context.",
         "Only applies to segments that contain named entities. Short generic sentences "
         "may have no entities to measure."),
        ("Signal 6: Length Ratio", "weight: 0.15",
         "Checks whether the hypothesis is roughly the same length as the reference.",
         "Ratio = len(hypothesis words) / len(reference words). A ratio near 1.0 is ideal. "
         "Too long (>1.5) suggests hallucination. Too short (<0.5) suggests truncation.",
         "Extreme hallucinations where the model generates paragraphs of fabricated text, "
         "or truncations where it outputs just a few words.",
         "A hypothesis can be the right length but entirely wrong in content."),
    ]

    for title, weight, what, how, catches, misses in signals:
        add_heading(doc, f"{title} ({weight})", 2)
        add_bullet_bold_value(doc, "What it does: ", what)
        add_bullet_bold_value(doc, "How: ", how)
        add_bullet_bold_value(doc, "What it catches: ", catches)
        add_bullet_bold_value(doc, "What it misses: ", misses)


def section_6_tiers(doc):
    add_heading(doc, "6. Tier Classification with Examples", 1)

    tiers = [
        ("Tier 5: Excellent (IS 4.0\u20135.0)", "Human fully understands", GREEN_BG, [
            ('WER 3%\nReference: "you\u2019re smart enough to realize that you need a plan"\n'
             'Predicted: "smart enough to realize that you need a plan"\n'
             'Missing only "you\u2019re" at the start. Perfect comprehension.'),
            ('WER 6%\nReference: "encourage a little bit of a deeper thought process about the calories and the weight loss"\n'
             'Predicted: "encourage a little bit of a deeper thought process about the calories in the weight loss"\n'
             'One preposition changed ("and"\u2192"in"). Meaning identical.'),
        ]),
        ("Tier 4: Good (IS 3.0\u20133.99)", "Intelligible, meaning preserved", GREEN_BG, [
            ('WER 23%\nReference: "on today\u2019s the doctor is in the topic is hypertension there is no ideal blood pressure reading"\n'
             'Predicted: "today as the doctor says the topic is hypertension there is no ideal blood pressure reading"\n'
             'Structure reshuffled but meaning perfectly clear: doctor\u2019s show about hypertension.'),
            ('WER 29%\nReference: "return books taken from the main library to the king\u2019s manor library and vice versa"\n'
             'Predicted: "term booked down from the main library to his local library and vice versa"\n'
             '"king\u2019s manor"\u2192"his local" loses the specific name but preserves the concept.'),
        ]),
        ("Tier 3: Fair (IS 2.0\u20132.99)", "Gist recoverable, details wrong", AMBER_BG, [
            ('WER 27%\nReference: "more fabulous queens from around the globe first the lebanese comedian"\n'
             'Predicted: "more fabulous tweets from around the globe first lebanese community"\n'
             'Topic area preserved but the specific subject (drag queens? comedian?) is lost.'),
            ('WER 35%\nReference: "for the sweetest and most all encompassing celebration of this great holiday of purim"\n'
             'Predicted: "for the swedish and most norwegian campus in celebration of this great holiday of jul"\n'
             '"purim"\u2192"jul" and "sweetest"\u2192"swedish" change the cultural context entirely.'),
        ]),
        ("Tier 2: Poor (IS 1.0\u20131.99)", "Fragments only", RED_BG, [
            ('WER 69%\nReference: "then we have the costal cartilages costal always refers to ribs"\n'
             'Predicted: "that we have they cause our hormones to go haywire"\n'
             'An anatomy lesson became a statement about hormones. No meaningful overlap.'),
        ]),
        ("Tier 1: Failed (IS 0.0\u20130.99)", "Unintelligible or hallucinated", RED_BG, [
            ('WER 100%\nReference: "let\u2019s see a half a cup of"\n'
             'Predicted: "i have to say thank you"\n'
             'Zero connection. Cooking \u2192 gratitude.'),
            ('WER 100% (harmful)\nReference: "it doesn\u2019t have a carry strap but you can put it on your shoulder"\n'
             'Predicted: "this is david irving he\u2019s a holocaust denier"\n'
             'Not just wrong \u2014 actively harmful fabrication.'),
        ]),
    ]

    for tier_title, tier_desc, tier_bg, examples in tiers:
        add_heading(doc, f'{tier_title} \u2014 "{tier_desc}"', 2)
        for ex in examples:
            add_quote(doc, ex)


def section_7_formula(doc):
    add_heading(doc, "7. The Composite Formula", 1)

    add_para(doc, "The final Intelligibility Score is a weighted sum of six normalized signals:",
             bold=True)

    # Formula as a styled box
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(
        "IS = 0.25 \u00d7 Semantic_Sim + 0.15 \u00d7 Phonetic_Sim + 0.15 \u00d7 (1\u2212WER) "
        "+ 0.15 \u00d7 (1\u2212WWER) + 0.15 \u00d7 NEA_F1 + 0.15 \u00d7 Length_Ratio"
    )
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    run.bold = True
    run.font.color.rgb = C_PRIMARY

    add_para(doc, "All signals scaled to 0\u20135 before combining. Final IS range: 0.0 to 5.0.")

    add_para(doc, (
        '"Properly captured" = IS \u2265 3.0 (Tier 4: Good + Tier 5: Excellent)'
    ), bold=True, size=Pt(11), color=C_GREEN)

    add_para(doc, (
        "These are the segments where a human could understand what was actually said "
        "from the lip-reading output."
    ))


def section_17_llm_judge_gold_standard(doc):
    add_heading(doc, "17. LLM-as-a-Judge Cross-Validation of IS Framework", 1)

    add_para(doc, (
        "To validate the IS framework against actual language understanding, Claude Opus 4.6 "
        "evaluated all 1,497 hypothesis-reference pairs using holistic LLM reasoning. The "
        "evaluation was blind (no metrics visible), used a 3-level scale (Y/P/N), and included "
        "30 duplicate pairs for intra-rater reliability (86.7% exact agreement)."
    ))

    # Results table
    add_heading(doc, "17.1 Results", 2)

    add_styled_table(doc,
        ["Judgment", "Count", "Rate"],
        [
            ["Y (meaning clearly conveyed)", "345", "23.0%"],
            ["P (partial \u2014 some meaning preserved)", "626", "41.8%"],
            ["N (meaning lost)", "526", "35.1%"],
            ["Y+P (any useful output)", "971", "64.9%"],
        ],
        col_widths=[3.0, 1.0, 1.0],
    )

    # Correlation with IS
    add_heading(doc, "17.2 Correlation with IS Signals", 2)

    add_para(doc, (
        "The LLM judge ratings correlate strongly with IS and its component signals:"
    ))

    add_styled_table(doc,
        ["Signal", "Pearson r", "Spearman \u03c1"],
        [
            ["Intelligibility Score (IS)", "0.850", "0.858"],
            ["Semantic Similarity", "0.680", "0.679"],
            ["Inverse WER", "0.619", "0.626"],
            ["Inverse WWER", "0.614", "0.622"],
            ["Phonetic Similarity", "0.549", "0.555"],
            ["NEA F1", "0.493", "0.512"],
            ["Length Ratio", "0.415", "0.408"],
        ],
        col_widths=[2.5, 1.0, 1.0],
    )

    add_para(doc, (
        "IS achieves the highest correlation (r = 0.850) because it combines all six signals \u2014 "
        "no single signal alone captures what the LLM judge measures. Semantic similarity is the "
        "strongest individual signal (r = 0.680), confirming IS\u2019s 25% semantic weight is appropriate."
    ))

    # Agreement with IS threshold
    add_heading(doc, "17.3 Agreement with IS \u2265 3.0 Threshold", 2)

    add_styled_table(doc,
        ["Comparison", "\u03ba", "Accuracy", "Precision", "Recall"],
        [
            ["LLM Y vs IS \u2265 3.0 (strict)", "0.565", "80.6%", "94.5%", "54.6%"],
            ["LLM Y+P vs IS \u2265 3.0 (lenient)", "0.521", "74.6%", "61.2%", "99.5%"],
        ],
        col_widths=[2.5, 0.6, 0.7, 0.7, 0.6],
    )

    add_para(doc, (
        "Only 22 of 1,497 segments (1.5%) are boundary disagreements. 19 have LLM=Y but IS<3.0 "
        "(validating the salvage analysis), and just 3 have LLM=N but IS\u22653.0. The IS threshold "
        "is well-calibrated."
    ))

    # 3x5 Confusion Matrix
    add_heading(doc, "17.4 LLM Judgment \u00d7 IS Tier Distribution", 2)

    add_para(doc, (
        "The 3\u00d75 confusion matrix shows clean tier separation \u2014 Y concentrates in tiers 4\u20135, "
        "N in tiers 1\u20132, and P spans the middle:"
    ))

    add_styled_table(doc,
        ["IS Tier", "Y", "P", "N"],
        [
            ["5 \u2014 Excellent (4.0\u20135.0)", "214 (62.0%)", "56 (8.9%)", "6 (1.1%)"],
            ["4 \u2014 Good (3.0\u20133.99)", "112 (32.5%)", "179 (28.6%)", "30 (5.7%)"],
            ["3 \u2014 Fair (2.0\u20132.99)", "14 (4.1%)", "206 (32.9%)", "105 (20.0%)"],
            ["2 \u2014 Poor (1.0\u20131.99)", "5 (1.4%)", "152 (24.3%)", "179 (34.0%)"],
            ["1 \u2014 Failed (0.0\u20130.99)", "0 (0.0%)", "33 (5.3%)", "206 (39.2%)"],
        ],
        col_widths=[2.2, 1.2, 1.2, 1.2],
    )

    # Partial judgment
    add_heading(doc, "17.5 Partial Judgment Analysis (626 P Segments)", 2)

    add_styled_table(doc,
        ["Element", "Preserved (% of P)", "Lost (% of P)"],
        [
            ["Structure (word order, grammar)", "88.8%", "0.6%"],
            ["Key content words", "66.6%", "25.6%"],
            ["Semantic meaning / gist", "26.0%", "55.1%"],
            ["Detail / qualifiers", "\u2014", "55.4%"],
            ["Named entities", "0.6%", "20.6%"],
            ["Phonetic resemblance", "16.6%", "\u2014"],
        ],
        col_widths=[2.5, 1.5, 1.5],
    )

    add_para(doc, (
        "The model reliably produces grammatically well-formed output with recognizable content "
        "words, but frequently loses specific meaning, fine-grained detail, and named entities."
    ))

    # Cross-config stability
    add_heading(doc, "17.6 Cross-Configuration Stability", 2)

    add_para(doc, (
        "The llm_context_prob heuristic (a deterministic decision tree, not an LLM API call) "
        "was validated across 16 decode configurations:"
    ))

    add_bullet_bold_value(doc, "Mean correlation with IS: ", "r = 0.925 (std = 0.015)")
    add_bullet_bold_value(doc, "Cohen\u2019s \u03ba range: ", "0.62\u20130.86 across configs")
    add_bullet_bold_value(doc, "Recall for IS \u2265 3.0: ", "97.6\u2013100% across all configs")
    add_bullet_bold_value(doc, "Agreement with IS \u2265 3.0: ", "88.6% (1,325/1,497 segments)")

    add_para(doc, (
        "This stability confirms the IS framework measures a genuine underlying quality dimension "
        "that persists regardless of decode settings, and that the LLM judge\u2019s holistic assessment "
        "tracks this dimension faithfully."
    ))

    # Capture rate comparison
    add_heading(doc, "17.7 Capture Rate Comparison", 2)

    add_styled_table(doc,
        ["Method", "Capture Rate", "What It Measures"],
        [
            ["WER \u2264 20%", "11.4%", "Traditional word accuracy only"],
            ["IS \u2265 3.0", "39.9%", "Multi-signal intelligibility"],
            ["IS + salvage (llm_prob \u2265 0.5)", "50.9%", "IS + recoverable meaning"],
            ["LLM Judge: Y", "23.0%", "Strict holistic meaning preservation"],
            ["LLM Judge: Y+P", "64.9%", "Any useful output (holistic)"],
        ],
        col_widths=[2.3, 1.0, 3.2],
    )

    add_para(doc, (
        "Key implication: WER alone reports 11.4% success; IS raises this to 39.9%; the LLM judge\u2019s "
        "lenient threshold shows 64.9% of outputs contain at least some useful information. The true "
        "capture rate lies between IS (39.9%) and LLM Y+P (64.9%), depending on the downstream "
        "application\u2019s tolerance for partial information."
    ), bold=True)


# =====================================================================
# MAIN
# =====================================================================

def main():
    doc = Document()

    # Page setup: A4, 2.5cm side, 2.0cm top/bottom
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # Cover page
    create_cover_page(doc)

    # TOC
    toc_titles = [
        "1. Why WER Isn\u2019t Enough",
        "2. The 10 Proof Examples",
        "3. Homophene Confusions: The Lip-Reading Problem",
        "4. Claude\u2019s Assessment Process",
        "5. The 6 Signals Explained",
        "6. Tier Classification with Examples",
        "7. The Composite Formula",
        "17. LLM-as-a-Judge Cross-Validation of IS Framework",
    ]
    add_toc(doc, toc_titles)

    # Header / footer (after TOC so page numbering works)
    add_header_footer(doc)

    # Content sections
    section_1_why_wer_isnt_enough(doc)
    doc.add_page_break()
    section_2_proof_examples(doc)
    doc.add_page_break()
    section_3_homophenes(doc)
    doc.add_page_break()
    section_4_assessment_process(doc)
    doc.add_page_break()
    section_5_six_signals(doc)
    doc.add_page_break()
    section_6_tiers(doc)
    doc.add_page_break()
    section_7_formula(doc)
    doc.add_page_break()
    section_17_llm_judge_gold_standard(doc)

    # Save
    OUTPUT_FILE.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_FILE))
    print(f"Saved: {OUTPUT_FILE}")


if __name__ == "__main__":
    main()