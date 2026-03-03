#!/usr/bin/env python3
"""
Argos — The Orchard — Gemini Presentation Maker Instructions (v2)

Generates a branded .docx with 4 sequential Gemini prompts for creating
the Argos VSP project review presentation (30 slides).

v2 changes:
- Structured slide format: LAYOUT / IMAGE / SPEAKER NOTES / [ANIMATION]
- Explicit image insertion directives (no more "leave space for chart")
- Video slides marked [MANUAL: VIDEO] — presenter inserts externally
- Animation checklist section + embedded animation tags in speaker notes
- Image-only Gemini uploads (markdown reports are supplementary reference)

Usage:
    python3 generate_gemini_instructions.py

Output:
    presentation_materials_20260224/gemini_presentation_instructions.docx
"""

from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

# ── Paths ──
SCRIPT_DIR = Path(__file__).parent
ASSETS_DIR = SCRIPT_DIR.parent / "assets"
OUTPUT_DIR = Path("/home/ubuntu/presentation_materials_20260224")
OUTPUT_FILE = OUTPUT_DIR / "gemini_presentation_instructions.docx"

LOGO_ORCHARD = ASSETS_DIR / "logo.png"
LOGO_PEACOCK = ASSETS_DIR / "peacock.png"

# ── Standard Colors ──
C_PRIMARY = RGBColor(0x1A, 0x3A, 0x5C)
C_H2 = RGBColor(0x2A, 0x5A, 0x8C)
C_H3 = RGBColor(0x3A, 0x6A, 0x9C)
C_H4 = RGBColor(0x4A, 0x7A, 0xAC)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_DARK = RGBColor(0x33, 0x33, 0x33)
C_GRAY = RGBColor(0x66, 0x66, 0x66)
C_CODE = RGBColor(0x2D, 0x2D, 0x2D)
C_TEAL = RGBColor(0x00, 0x80, 0x80)
C_CORAL = RGBColor(0xE0, 0x6C, 0x75)

HEADER_BG = "1a3a5c"
ZEBRA_BG = "f0f4f8"
CODE_BG = "f5f5f5"
PROMPT_BG = "eef6ee"   # light green for prompt blocks
TEAL_BG = "e0f2f1"
CORAL_BG = "fce4ec"

LAST_UPDATED = datetime.now().strftime("%B %d, %Y")


# ═══════════════════════════════════════════════════
# STANDARD HELPERS
# ═══════════════════════════════════════════════════

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, bold=False, color=None, size=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = size
    run.font.name = "Calibri"


def add_styled_table(doc, headers, rows, col_widths=None, highlight_rows=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, HEADER_BG)
        set_cell_text(cell, h, bold=True, color=C_WHITE, size=Pt(9))
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            set_cell_text(cell, val, size=Pt(9))
            if highlight_rows and r_idx in highlight_rows:
                set_cell_shading(cell, highlight_rows[r_idx])
            elif r_idx % 2 == 1:
                set_cell_shading(cell, ZEBRA_BG)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def add_heading(doc, text, level, color=None):
    h = doc.add_heading(text, level=level)
    if color is None:
        color = {1: C_PRIMARY, 2: C_H2, 3: C_H3, 4: C_H4}.get(level, C_PRIMARY)
    for run in h.runs:
        run.font.color.rgb = color
        run.font.name = "Calibri"
    return h


def add_para(doc, text, bold=False, italic=False, size=Pt(11), color=None, space_after=Pt(6)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = size
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = space_after
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = "Calibri"
    return p


def add_prompt_block(doc, prompt_text):
    """Add a prompt as a shaded code block with monospace font."""
    for line in prompt_text.strip().split("\n"):
        p = doc.add_paragraph()
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{PROMPT_BG}"/>')
        p._p.get_or_add_pPr().append(shading)
        run = p.add_run(line)
        run.font.size = Pt(9)
        run.font.name = "Consolas"
        run.font.color.rgb = C_CODE
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
    # Add a spacer after the block
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_attachment_table(doc, attachments):
    """Add a table of images to insert manually after export to Google Slides."""
    add_heading(doc, "Images to Insert After Export", 3)
    add_para(doc, (
        "After exporting the Gemini presentation to Google Slides, replace each gray "
        "placeholder rectangle with the corresponding image from the table below. "
        "Do NOT upload these images to Gemini \u2014 Gemini cannot reliably insert uploaded images."
    ), italic=True, color=C_GRAY, size=Pt(10))
    add_styled_table(doc,
        ["#", "File", "Slide", "Replace Placeholder"],
        [[str(i + 1)] + row for i, row in enumerate(attachments)],
        col_widths=[0.3, 2.8, 0.6, 2.9],
    )


def _tight_page_break(doc):
    last_p = doc.paragraphs[-1]
    last_p.add_run().add_break(WD_BREAK.PAGE)


def build_header(doc):
    """Add header with logo and org name."""
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.clear()

    if LOGO_ORCHARD.exists():
        run = hp.add_run()
        run.add_picture(str(LOGO_ORCHARD), width=Inches(0.35))

    text_width = section.page_width - section.left_margin - section.right_margin
    hp.paragraph_format.tab_stops.add_tab_stop(text_width, WD_TAB_ALIGNMENT.RIGHT)
    hp.add_run("\t")
    run = hp.add_run("Argos \u2014 The Orchard")
    run.font.size = Pt(8)
    run.font.name = "Calibri"
    run.italic = True
    run.font.color.rgb = C_GRAY


def build_footer(doc):
    """Add page numbers in footer."""
    section = doc.sections[0]
    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.clear()
    run = fp.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run._r.append(instr)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)
    run.font.size = Pt(8)
    run.font.color.rgb = C_GRAY


def build_cover(doc):
    """Build cover page."""
    doc.add_paragraph()
    doc.add_paragraph()

    if LOGO_PEACOCK.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(str(LOGO_PEACOCK), width=Inches(2.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ARGOS")
    run.bold = True
    run.font.size = Pt(48)
    run.font.color.rgb = C_PRIMARY
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Gemini Presentation Instructions")
    run.font.size = Pt(22)
    run.font.name = "Calibri"
    run.font.color.rgb = C_H2

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Visual Speech Processing \u2014 Project Review")
    run.font.size = Pt(16)
    run.font.name = "Calibri"
    run.font.color.rgb = C_H3

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("The Orchard")
    run.font.size = Pt(20)
    run.font.name = "Calibri"
    run.font.color.rgb = C_PRIMARY

    doc.add_paragraph()

    for label, value in [
        ("Author: ", "Yoad Oxman"),
        ("Target Tool: ", "Google Gemini 3.1 Pro (Slides Generation)"),
        ("Slides: ", "30 slides + appendix"),
        ("Audience: ", "Technical managers and supervisors"),
        ("Last Updated: ", LAST_UPDATED),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        run = p.add_run(value)
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        p.paragraph_format.space_after = Pt(2)

    _tight_page_break(doc)


def build_toc(doc):
    """Build Table of Contents."""
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    add_heading(doc, "Table of Contents", 1)

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-2" \\h \\z \\u '
    run._r.append(instr)
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    toc_entries = [
        "1. How to Use This Document",
        "2. Presentation Overview",
        "3. Prompt 1: Opening & Context (Slides 1-4)",
        "4. Prompt 2: Research Findings (Slides 5-16)",
        "5. Prompt 3: Engineering (Slides 17-23)",
        "6. Prompt 4: Future & Close (Slides 24-30)",
        "7. Appendix Slides",
        "8. Animation & Transition Checklist",
        "9. Post-Generation Manual Steps",
        "10. Material References",
    ]
    for title in toc_entries:
        placeholder = paragraph.add_run(title + "\n")
        placeholder.font.size = Pt(11)
        placeholder.font.name = "Calibri"

    fld_end_run = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    fld_end_run._r.append(fld_end)

    settings = doc.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)

    _tight_page_break(doc)


# ═══════════════════════════════════════════════════
# PROMPTS (v2 — structured format)
# ═══════════════════════════════════════════════════

PROMPT_1 = """Create a presentation in Google Slides with a dark modern theme:
- Background: deep navy (#0D1B2A)
- Text: white (#FFFFFF)
- Accents: teal (#00B4D8) for highlights, coral (#E06C75) for warnings/negatives
- Font: clean sans-serif (Montserrat or similar)
- All slides: consistent style, professional, minimal clutter

CRITICAL IMAGE RULES:
- Do NOT search the internet for images. Do NOT use stock photos. Do NOT add image source slides.
- Where I specify IMAGE PLACEHOLDER, create a dark gray rounded rectangle with the placeholder label in white text inside it. I will replace these with my own images after export.
- Every image placeholder should be a visible, labeled box \u2014 not an empty space.

This is PART 1 of 4 for an internal project review called
"Argos VSP: Research Findings and Production Roadmap"
about a Visual Speech Processing (lip reading) system.
Audience: Technical managers and supervisors.
Create exactly 4 slides with speaker notes for each.

---

SLIDE 1 - Title
LAYOUT: Title slide, clean and minimal
Top-right corner: IMAGE PLACEHOLDER labeled "[Logo]" (small, ~1 inch square, dark gray box)
Title: "Argos VSP: Research Findings and Production Roadmap"
Subtitle: "Visual Speech Processing \u2014 Project Review"
Bottom-right: "February 2026"
SPEAKER NOTES: "Welcome. This presentation covers 3 months of research and engineering on a visual speech processing system \u2014 reading lips from video, no audio. We\u2019ll cover what we found, what we built, and where we go next. [ANIMATION: Title fades in, subtitle follows 0.5s later. Transition: None.]"

---

SLIDE 2 - What is Visual Speech Processing? [VIDEO PLACEHOLDER]
LAYOUT: Minimal center text with a large play button icon
Title: "What is Visual Speech Processing?"
Center text: "A system that reads lips from video \u2014 no audio needed."
Below center: Large triangular play button icon (white outline on dark background) with caption "Opening Demo: 33-Word Perfect Lip Reading"
Do NOT embed any video or search for video images. Just create the play icon as a shape.
SPEAKER NOTES: "PLAY VIDEO: IEa7qEkMvfQ_3__c5447488_with_hyp.mp4 \u2014 33 words about health insurance, WER 0%. Play the video first, then explain: this is the best case. The system perfectly reads 33 consecutive words from lip movement alone. Now let\u2019s see how it works. [ANIMATION: Text appears, play icon pulses. Transition: Fade 0.5s.]"

---

SLIDE 3 - Model Architecture
LAYOUT: Full-width placeholder with text annotations below
Top area: IMAGE PLACEHOLDER labeled "[pipeline_architecture.png]" (large, ~80% slide width, centered)
Title: "How It Works: Three Components"
Below the placeholder, three labeled blocks left-to-right:
"AV-HuBERT (Visual Encoder, frozen, 1024-dim)" \u2192 "Linear Projection (1024\u21924096)" \u2192 "LLaMA-2-7B (4-bit QLoRA, r=16)"
Bottom note in smaller text: "Only 12.6M trainable params (0.19%). LLM is swappable \u2014 Llama 3.1 8B is a drop-in replacement (same 4096 hidden size)."
SPEAKER NOTES: "Three components. Visual encoder (AV-HuBERT) is frozen \u2014 pre-trained on LRS3 lip-reading data. It outputs 1024-dim features per frame. A linear projection maps to 4096-dim (LLM input space). Then LLaMA-2-7B generates text. Key: only the LoRA adapters and projection layer are trained \u2014 12.6M of 7B parameters. And the LLM is swappable: Llama 3.1 8B has the same hidden dimension, making it a trivial 1-2 hour swap. [ANIMATION: Three component blocks fly in left-to-right sequentially. Transition: Fade 0.5s.]"

---

SLIDE 4 - The Benchmark
LAYOUT: Split \u2014 large number on the left, chart on the right
Right half: IMAGE PLACEHOLDER labeled "[P2_paper_vs_reality.png]" (fills the right half of the slide)
Left side: Large "25.4%" in teal accent color, with "WER on LRS3 (TED Talks)" below in white.
Below the number: "Curated data, controlled conditions."
Bottom spanning full width: "Two questions: How does this hold on real-world video? And is WER even the right metric?"
SPEAKER NOTES: "The paper claims 25.4% Word Error Rate on LRS3 \u2014 a curated dataset of TED talks with clear speech, frontal faces, good lighting. Our question: what happens on real-world YouTube video? The chart on the right previews the answer: 64.1% WER, 2.5x worse. And more importantly \u2014 is WER even the right way to measure this? [ANIMATION: Large number appears first, then chart fades in on right. Transition: Fade 0.5s.]"
"""

PROMPT_2 = """Continue the same presentation (dark navy #0D1B2A, white text, teal/coral accents).
This is PART 2 of 4: Research Findings. Create exactly 12 slides with speaker notes.

REMINDER: Do NOT search the internet for images. Do NOT use stock photos. Where I specify IMAGE PLACEHOLDER, create a dark gray rounded rectangle with the label text in white inside it. I will replace these with my own charts after export.

---

SLIDE 5 - The Reality Gap
LAYOUT: Split \u2014 large metric left, chart right
Right half: IMAGE PLACEHOLDER labeled "[P1_quality_tiers.png]"
Left side: Large "64.1%" in coral, with "Mean WER across 1,497 real-world segments" below.
Below, five quality tier bullets with colored dots:
\u2022 11.4% Usable (<30%) \u2014 green dot
\u2022 17.4% Marginal (30-50%) \u2014 yellow dot
\u2022 17.8% Poor (50-75%) \u2014 orange dot
\u2022 32.8% Unusable (75-100%) \u2014 red dot
\u2022 20.6% Hallucinated (>100%) \u2014 dark red dot
Bottom: "But WER overstates failure \u2014 see next slide."
SPEAKER NOTES: "1,497 diverse YouTube segments. 64.1% mean WER \u2014 2.5x worse than the paper\u2019s 25.4%. Only 11.4% usable by WER standards. And 20.6% are hallucinations \u2014 fluent text that\u2019s completely fabricated. This is the most dangerous failure mode. But WER is misleading \u2014 it treats all errors equally. [ANIMATION: Tier bars build top-to-bottom, 0.3s each. Transition: Fade 0.5s.]"

---

SLIDE 6 - WER Is Blind
LAYOUT: Two boxes side by side
Left box (green border): "WER 29%, IS 4.3 \u2014 Fully Intelligible"
  Ref: "allow you to work with the team in a more"
  Hyp: "allow you to work with a team and more"
Right box (red border): "WER 33%, IS 3.0 \u2014 Unintelligible"
  Ref: "today i\u2019m talking with admiral mcrae"
  Hyp: "today i\u2019m talking with animal migratory"
Bottom: "WER says equal. Meaning says opposite. So we built a metric to capture this."
SPEAKER NOTES: "Same WER, opposite meaning. Left: minor grammatical change, meaning fully preserved. Right: a name destroyed \u2014 \u2018admiral mcrae\u2019 becomes \u2018animal migratory.\u2019 WER sees ~30% error in both. But one is usable and the other is garbage. This motivated our Intelligibility Score. [ANIMATION: Left box appears, then right box appears. Transition: Fade 0.5s.]"

---

SLIDE 7 - The Intelligibility Score
LAYOUT: Top section with 6 blocks in a row, bottom section with tier bars
Title: "Intelligibility Score: 39.9% Properly Captured"
Six signal blocks in a row (colored blocks with labels):
  Semantic Sim (25%) | Phonetic Sim (15%) | Inv. WER (15%) | Inv. WWER (15%) | NEA F1 (15%) | Length Ratio (15%)
Key callout in teal: "IS >= 3.0 = Properly Captured: 39.9% \u2014 3.5x more than WER\u2019s 11.4%"
Five tier bars: 18.4% Excellent | 21.4% Good | 21.7% Fair | 22.4% Poor | 16.0% Failed
SPEAKER NOTES: "The Intelligibility Score combines 6 signals into a 0-5 composite. Key insight: 39.9% of segments are properly captured (IS >= 3.0) \u2014 3.5x more than WER\u2019s 11.4% \u2018usable.\u2019 WER dramatically overstates failure. Methodology: LLM-distilled evaluation \u2014 Claude designed the rubric, selected signals and weights, defined tier boundaries. This is \u2018what an expert would score\u2019 made deterministic and free. Validated across 16 decode configs: Claude-designed heuristic (deterministic, no runtime LLM calls) r=0.925 with IS, 88.6% agreement. Phonetic Similarity is the #1 correlate (r=0.943). Six signals collapse into 3 independent dimensions: word accuracy (~60%), meaning preservation (~28%), output sanity (~9%). [ANIMATION: Six signal blocks build one at a time. Tier bars appear below. Transition: Fade 0.5s.]"

---

SLIDE 8 - Failure Mode Taxonomy
LAYOUT: Horizontal bar chart (create this chart directly from the data below \u2014 do NOT use an image)
Title: "10 Classified Failure Modes (900 failed segments)"
Create a horizontal bar chart with these values (sorted by frequency, largest at top):
  Total Topic Drift: 15.9% (143) \u2014 dark red
  Phonetically Similar Wrong Topic: 15.7% (141) \u2014 red
  Accumulated Small Errors: 12.3% (111) \u2014 orange
  Hallucination: 12.3% (111) \u2014 dark red
  High Error Rate: 12.1% (109) \u2014 orange
  Entity Destruction: 12.0% (108) \u2014 red
  Content Word Errors: 10.7% (96) \u2014 yellow
  Empty Output: 7.8% (70) \u2014 gray
  Truncation: 1.1% (10) \u2014 gray
  Over-generation: 0.1% (1) \u2014 gray
Bottom: "Failures are diverse \u2014 no single fix. Each roadmap phase targets specific modes."
SPEAKER NOTES: "900 failed segments classified into 10 failure modes. The top two \u2014 topic drift and phonetically-similar-wrong-topic \u2014 account for 31.6%. Hallucination (12.3%) is the most dangerous: fluent, confident, completely fabricated. Entity destruction (12.0%) loses names and numbers. This taxonomy maps directly to our roadmap: N-best aggregation targets accumulated small errors, confidence scoring flags hallucinations, data scaling reduces topic drift. [ANIMATION: Bars build top-to-bottom, wipe left. Transition: Fade 0.5s.]"

---

SLIDE 9 - Performance Distribution
LAYOUT: Full-width placeholder, minimal text
Center: IMAGE PLACEHOLDER labeled "[09_boxplot_wwer_all_experiments.png]" (large, ~70% slide width)
Title: "Distribution Across 13 Experiments"
Below placeholder: "Most segments: 50-80% WER. Stable core of ~11% always-good segments across all configs."
SPEAKER NOTES: "This boxplot shows WWER distribution across all 13 decode experiments. The box is consistently in the 50-80% range. Important: about 11% of segments are always good regardless of parameters, and about 16% are always bad. The bottleneck is the visual encoder, not decode strategy \u2014 proved by cross-config analysis showing r > 0.92 per-segment ranking stability. [ANIMATION: No build. Transition: Fade 0.5s.]"

---

SLIDE 10 - Why the Gap: Root Causes
LAYOUT: Three columns with data, chart on right
Right half: IMAGE PLACEHOLDER labeled "[01_wer_vs_duration.png]"
Title: "Three Root Causes \u2014 With Data"
Three columns (left two-thirds):
  1. Domain Mismatch \u2014 Business/Finance: IS 3.08, 57% captured (best). DIY/Home: IS 2.13, 30% captured (worst).
  2. Short Segments \u2014 5-10 words: 32% captured, 74% WER. 20+ words: 49% captured, 55% WER.
  3. Hallucination \u2014 LLM prior overwhelms weak visual signal. 15.9% topic drift, 12.3% hallucination.
SPEAKER NOTES: "Three root causes. Domain mismatch: the model was trained on formal TED talks, so business/finance content works best (57% captured). DIY/home improvement is worst (30%). Short segments fail catastrophically \u2014 under 10 words gives only 32% capture rate because there\u2019s not enough visual context. And hallucination: when the visual signal is weak, the LLM\u2019s language prior takes over and generates fluent but fabricated text. The scatter plot shows the duration effect clearly. [ANIMATION: Three columns appear one at a time. Transition: Fade 0.5s.]"

---

SLIDE 11 - Named Entity Accuracy
LAYOUT: Split \u2014 text left, chart right
Right half: IMAGE PLACEHOLDER labeled "[14_nea_vs_wwer_scatter.png]"
Title: "NEA F1: 38.8% \u2014 The Largest Differentiator"
Left side text:
  "Names, numbers, proper nouns \u2014 what context cannot recover."
  Signal gap (success vs failure):
  \u2022 NEA F1: 74% vs 16% (gap: 58pp \u2014 LARGEST)
  \u2022 Semantic: 0.74 vs 0.24
  \u2022 Phonetic: 0.81 vs 0.38
SPEAKER NOTES: "Named Entity Accuracy is the single largest differentiator between success and failure: 74% vs 16%, a 58 percentage point gap. You can guess a missing \u2018the\u2019 from context, but you cannot recover \u2018Admiral McRae\u2019 if the model says \u2018animal migratory.\u2019 NEA contributes 17.3% of IS variance \u2014 disproportionate to its 15% weight because entity accuracy has the highest variance among signals. [ANIMATION: Text appears, then chart fades in. Transition: Fade 0.5s.]"

---

SLIDE 12 - 13 Tuning Experiments
LAYOUT: Two columns
Right half: IMAGE PLACEHOLDER labeled "[10_empty_and_hallucination_rates.png]"
Title: "Systematic Parameter Search: 13 Configurations"
Left column \u2014 "Parameters Tested":
  \u2022 Beam size (5-50)
  \u2022 Length penalty (-0.5 to 2.0)
  \u2022 Temperature (0.3-1.5)
  \u2022 Sampling strategy
  \u2022 Repetition penalty
Right column header: "Best Config (J) \u2014 Full Dataset (1,497 segments)"
  \u2022 IS: 2.60 vs 2.52 baseline (+0.08)
  \u2022 Captured: 622 vs 597 (+25 segments)
  \u2022 Empties: 0 vs 70 (eliminated)
  \u2022 Hallucinations: 348 vs 307 (+41 more)
SPEAKER NOTES: "13 systematic experiments across beam size, length penalty, temperature, and sampling. Config J (beam=20, lenpen=0, temperature=1.0, do_sample=True) achieved the best IS. Key trade-off: eliminated all 70 empty outputs but added 41 hallucinations. Net IS gain: only +0.08. Note: standalone container uses do_sample=True (stochastic), EC2 uses False (deterministic) \u2014 to be unified. The chart shows the empty-vs-hallucination trade-off across configs. [ANIMATION: Left column appears, then right column. Transition: Fade 0.5s.]"

---

SLIDE 13 - Limits of Tuning
LAYOUT: Split \u2014 text left, chart right
Right half: IMAGE PLACEHOLDER labeled "[P4_lenpen_sensitivity.png]"
Title: "Tuning Is Mitigation, Not a Cure"
Left text:
  \u2022 Config J: eliminates empties but doubles hallucinations
  \u2022 Net IS gain: only +0.08 across 1,497 segments
  \u2022 Cross-config proof: per-segment rankings identical (r > 0.92)
  \u2022 "Hard" and "easy" segments stay the same \u2014 bottleneck is the visual encoder
  \u2022 Three levers remain: more data, stronger LLM, smart prompts
SPEAKER NOTES: "Tuning is mitigation, not a cure. Config J\u2019s fundamental trade-off: silent failures (empties) vs noisy failures (hallucinations). The length penalty chart shows the sensitivity \u2014 small changes cause wild swings. Cross-config analysis proves: per-segment IS rankings are nearly identical across all 16 configs (r > 0.92 for most pairs). The bottleneck is the visual encoder, not decode parameters. Data scarcity (1,273 training segments) is the real problem. Three levers: (1) scale data to 20K-50K segments, (2) swap to Llama 3.1 8B (drop-in), (3) smart prompts as force multiplier. [ANIMATION: Text appears, chart fades in. Transition: Fade 0.5s.]"

---

SLIDE 14 - Curated Examples
LAYOUT: Table with color-coded IS column
Title: "Representative Examples"
Create a table with columns: Category | Reference | Hypothesis | WER | IS
Row 1 (green IS): Perfect | "health insurance company they pay..." | [exact match] | 0% | 5.0
Row 2 (green IS): WER Misleads | "work with the team in a more" | "work with a team and more" | 29% | 4.3
Row 3 (yellow IS): Entity Lost | "talking with admiral mcrae" | "talking with animal migratory" | 33% | 3.0
Row 4 (red IS): Hallucinated | "carry strap" | "holocaust denier" | 100% | 0.7
Color the IS column: 5.0 and 4.3 in green, 3.0 in yellow, 0.7 in red.
SPEAKER NOTES: "Four examples spanning the quality range. Row 1: perfect lip-reading. Row 2: WER says 29% error but the meaning is fully preserved \u2014 IS 4.3. Row 3: same WER range but a name destroyed, meaning lost. Row 4: complete hallucination \u2014 \u2018carry strap\u2019 becomes \u2018holocaust denier.\u2019 This is why WER alone is insufficient. [ANIMATION: Table rows build one at a time, top to bottom. Transition: Fade 0.5s.]"

---

SLIDE 15 - Live Demo [VIDEO PLACEHOLDER]
LAYOUT: Minimal dark slide with play button icon
Title: "Demo: Three Videos"
Subtitle: "Perfect \u2192 Near-miss \u2192 Hallucination"
Center: Large triangular play button icon (white outline shape on dark background)
Below: "Three demonstrations played from external files."
Do NOT embed any video or search for images. Just create the play icon as a shape.
SPEAKER NOTES: "PLAY VIDEOS in order: (1) d8BR6hsvzoY_31__2e9546df_with_hyp.mp4 \u2014 \u2018buy one get one free\u2019 (WER 0%, short and punchy). (2) -POZpyVCN8k_9__c7b26ea8_with_hyp.mp4 \u2014 \u2018admiral mcrae\u2019 becomes \u2018animal migratory\u2019 (funny near-miss). (3) 00MUdHQ7GGY_8__b1480c7a_with_hyp.mp4 \u2014 hallucination: fabricates \u2018David Irving\u2019 narrative (WER 100%). [ANIMATION: No build. Transition: Fade 0.5s.]"

---

SLIDE 16 - Research Output
LAYOUT: Numbered list with brief descriptions
Title: "Eight Research Reports + Training Analysis"
Numbered list:
  1. Executive Assessment \u2014 reality gap, quality tiers, 1,497 segments
  2. Hyperparameter Tuning \u2014 13 experiments, parameter sensitivity
  3. Prompt Engineering \u2014 context injection, 7 prompt strategies
  4. Confidence Scoring \u2014 beam score extraction, quality filtering
  5. N-Best Aggregation \u2014 ROVER/MBR hypothesis combination
  6. Fine-Tuning Analysis \u2014 LoRA domain adaptation strategy
  7. Intelligibility Scoring \u2014 6-signal composite metric, failure taxonomy
  8. LLM Upgrade Analysis \u2014 Llama 3.1 8B drop-in, data scaling, GER
  + Exp A Training Report \u2014 overfitting analysis, 10 diagnostic plots
SPEAKER NOTES: "Eight formal research reports plus the Exp A training analysis. Each report is a deep dive with methodology, results, and recommendations. The Intelligibility Score report is the main methodological contribution. The LLM Upgrade Analysis maps the path forward: model alternatives, data scaling projections, prompt strategies, and investment roadmap targeting 27-42% WER. [ANIMATION: List items build one at a time. Transition: Fade 0.5s.]"
"""

PROMPT_3 = """Continue the same presentation (dark navy #0D1B2A, white text, teal/coral accents).
This is PART 3 of 4: Engineering Achievements. Create exactly 7 slides with speaker notes.

REMINDER: Do NOT search the internet for images. Do NOT use stock photos. Where I specify IMAGE PLACEHOLDER, create a dark gray rounded rectangle with the label text in white inside it. I will replace these with my own diagrams after export.

---

SLIDE 17 - Pipeline Architecture
LAYOUT: Full-width placeholder with title
Center: IMAGE PLACEHOLDER labeled "[pipeline_architecture.png]" (large, filling the full slide width below the title)
Title: "8-Stage Automated Pipeline"
Subtitle: "3 research repos \u2192 single orchestrated system"
SPEAKER NOTES: "The pipeline orchestrates three separate research codebases into a single automated system. Eight stages: video normalization, ASR transcription, mouth cropping, LRS3 format conversion, manifest generation, feature clustering, LLM decode, and output generation. Each stage is a separate module in lib/ with its own tests. [ANIMATION: Pipeline diagram fades in. Transition: Fade 0.5s.]"

---

SLIDE 18 - Engineering Challenges
LAYOUT: Two columns
Title: "37 Bugs Fixed: Research Code \u2192 Production"
Left column \u2014 "Critical Fixes":
  \u2022 NVENC silent corruption (destroyed 43% of videos)
  \u2022 HDR/10-bit tone mapping
  \u2022 Cython extension compilation
  \u2022 fairseq patching
  \u2022 spaCy offline installation
  \u2022 Docker networking
  \u2022 Python venv conflicts
Right column \u2014 "Result":
  \u2022 Any format: MP4, MKV, AVI, MOV
  \u2022 Any resolution, any codec
  \u2022 GPU detection with CPU fallback
  \u2022 All 37 bugs documented with root cause
SPEAKER NOTES: "37 bugs found and fixed turning research code into production. The most critical: NVENC silent video corruption that destroyed 43% of processed videos without any error message. Also HDR tone-mapping failures, Cython compilation issues across architectures, fairseq compatibility patches. Every bug documented with root cause analysis. The system now handles any input format, resolution, and codec with automatic GPU/CPU detection. [ANIMATION: Left column appears, then right column. Transition: Fade 0.5s.]"

---

SLIDE 19 - Modular Refactoring
LAYOUT: Before/after comparison
Title: "52% Code Reduction: 823 \u2192 393 Lines"
Left (coral background): "BEFORE" \u2014 "Monolithic 823-line script. Fragile. Untestable."
Right (teal background): "AFTER" \u2014 "393-line orchestrator + 11 modules in lib/. 37 automated tests. Each stage independently testable. Auto-detects EC2 vs container."
SPEAKER NOTES: "The original pipeline was a single 823-line bash script. Refactored into a 393-line orchestrator calling 11 reusable modules. 37 automated tests validate every module. Environment-aware: automatically detects EC2 development vs Docker container deployment. This is Mission 1 in our research backlog \u2014 foundation for everything that followed. [ANIMATION: Before block appears left, then After block appears right. Transition: Morph.]"

---

SLIDE 20 - Deployed Product
LAYOUT: Feature list with icon placeholders
Title: "Standalone Container \u2014 No Cloud Required"
Five feature items:
  \u2022 Docker container with web UI on Linux machine
  \u2022 Drag-and-drop video upload
  \u2022 Automatic end-to-end processing
  \u2022 INSTALL.sh overlay with backup and verification
  \u2022 Currently deployed at client site
SPEAKER NOTES: "The system runs as a standalone Docker container \u2014 no cloud, no internet required. Web UI for drag-and-drop video upload. INSTALL.sh handles deployment with automatic backup and 13-point verification. Currently deployed and running at the client site. [ANIMATION: Feature items build one at a time. Transition: Fade 0.5s.]"

---

SLIDE 21 - Intelligent Features
LAYOUT: Three cards in a row
Title: "Pipeline Intelligence"
Card 1: "Transcription Reuse" \u2014 "Manual corrections persist across runs. Whisper skips matched segments, saving hours."
Card 2: "Golden K-means" \u2014 "1,396-video baseline clustering model for consistent feature extraction."
Card 3: "Smart Segmentation" \u2014 "Configurable overlap for context preservation across segment boundaries."
SPEAKER NOTES: "Three intelligent features. Transcription reuse: if you manually correct a segment\u2019s transcription, it persists across all future pipeline runs. Whisper automatically skips matched segments. Golden k-means: a 1,396-video baseline model ensures consistent clustering. Smart segmentation: configurable overlap ensures context isn\u2019t lost at segment boundaries. [ANIMATION: Three cards fly in from bottom, one at a time. Transition: Fade 0.5s.]"

---

SLIDE 22 - Evaluation Infrastructure
LAYOUT: Two-column list with chart placeholder on right
Right half: IMAGE PLACEHOLDER labeled "[15_cdf_wwer_curated.png]"
Title: "Evaluation: Beyond Standard WER"
Left list:
  \u2022 16 analytical plots per experiment (auto-generated)
  \u2022 Per-segment HTML reports (13 interactive)
  \u2022 Custom NEA metric for entity accuracy
  \u2022 Intelligibility Score pipeline (LLM-distilled, 6 signals)
  \u2022 IS correlation analysis (16 configs, r=0.925)
  \u2022 Failure mode classification (10 modes)
  \u2022 Topic/length analysis (11 categories)
  \u2022 Fine-tuning diagnostics (10 training plots)
SPEAKER NOTES: "Evaluation infrastructure goes far beyond standard WER. 16 plots auto-generated per experiment. Interactive HTML reports. Custom Named Entity Accuracy metric. The full Intelligibility Score pipeline with 6 signals, validated across 16 decode configs (Claude-designed heuristic r=0.925 (no runtime LLM calls)). Automated failure mode classification into 10 categories. Topic and segment length analysis across 11 content categories. Plus 10 fine-tuning diagnostic plots for training analysis. The CDF chart shows actionable quality thresholds. [ANIMATION: List items build one at a time. Transition: Fade 0.5s.]"

---

SLIDE 23 - Process & Documentation
LAYOUT: Bullet list
Title: "Engineering Process"
  \u2022 40+ git versions with semantic tags
  \u2022 EC2/container sync protocol (26 tracked items)
  \u2022 8 formal research reports + methodology docs
  \u2022 Complete architecture, development, and bug documentation
  \u2022 Automated test suite and verification
SPEAKER NOTES: "Rigorous engineering process. Over 40 git versions with semantic tagging. EC2-to-container synchronization protocol tracking 26 items. Eight formal research reports including the LLM Upgrade Analysis with model alternatives, data scaling projections, and investment strategy. Full architecture and development documentation. Automated test suite and deployment verification. [ANIMATION: No build. Transition: Fade 0.5s.]"
"""

PROMPT_4 = """Continue the same presentation (dark navy #0D1B2A, white text, teal/coral accents).
This is PART 4 of 4: Future Directions & Close. Create exactly 7 slides with speaker notes.
Frame the roadmap as a research agenda and resource needs, not a sales pitch.

REMINDER: Do NOT search the internet for images. Do NOT use stock photos. Where I specify IMAGE PLACEHOLDER, create a dark gray rounded rectangle with the label text in white inside it. I will replace these with my own charts after export.

---

SLIDE 24 - Reframing the Starting Point
LAYOUT: Two columns \u2014 red left, green right
Right half: IMAGE PLACEHOLDER labeled "[P1_quality_tiers.png]" (as visual context)
Title: "Starting from 40%, Not 11%"
LEFT column (coral/red tint): "WER Says"
  \u2022 11.4% usable
  \u2022 9 out of 10 segments fail
RIGHT column (teal/green tint): "IS Says"
  \u2022 39.9% properly captured
  \u2022 43-51% context-recoverable
  \u2022 Validated across 16 decode configs
  \u2022 Claude-designed heuristic (no runtime LLM): 88.6% agreement, r=0.925
Bottom: "The gap is real \u2014 but WER dramatically overstates failure. 40% already works."
SPEAKER NOTES: "This is the turning point. WER says 11.4% usable. Our Intelligibility Score says 39.9% properly captured \u2014 3.5x more. And 43-51% is context-recoverable. This isn\u2019t wishful thinking: IS is validated across all 16 decode configs (Claude-designed heuristic r=0.925 (no runtime LLM calls), 88.6% agreement, recall 99.2%). Config J has the highest IS (2.571) despite +14.8pp higher WER than baseline \u2014 IS captures meaning that WER misses. We know WHY it works: 41.5% of successes are phonetically preserved (r=0.943 with IS). Per-segment rankings stable across configs (r > 0.92) \u2014 the bottleneck is the encoder, not decode. [ANIMATION: Left (red) column appears, then right (green) column. Transition: Fade 0.5s.]"

---

SLIDE 25 - Research Roadmap
LAYOUT: Staircase/waterfall descending from 67% to 27-42%
Right half: IMAGE PLACEHOLDER labeled "[P3_wer_trajectory.png]"
Title: "Five Phases \u2014 From 67% to Target 27-42% WER"
Five phases as descending steps (left side):
  Phase 1: Surface the good 40% \u2014 confidence scoring (2-4 hours)
  Phase 2: Improve the fair 22% \u2014 N-best aggregation (-5 to -15%)
  Phase 3: LLM swap + smart prompts \u2014 Llama 3.1 8B + prompts (-8 to -18pp)
  Phase 4: Scale data to 20K-50K segments + fine-tune (-15 to -25pp)
  Phase 5: GER post-processing \u2014 no retraining needed (-8 to -15pp)
Bottom: "Combined target: 27-42% WER. Multiplicative scaling law (ICLR 2024)."
SPEAKER NOTES: "Five phases, each targeting different bottlenecks. Phase 1 is immediate: confidence scoring to surface the 40% that\u2019s already good (2-4 hours). Phase 2: N-best aggregation exploiting all 20 beam hypotheses. Phase 3: swap LLM to Llama 3.1 8B (drop-in, 1-2 hours) plus smart prompts as force multiplier. Phase 4: the biggest gain \u2014 scale training data from 1.3K to 20K-50K segments. Phase 5: GER post-processing using an external correction LLM. Key: ICLR 2024 shows these gains are multiplicative \u2014 stronger LLM extracts MORE from the same data. [ANIMATION: Five phases build as descending staircase, one at a time. Transition: Fade 0.5s.]"

---

SLIDE 26 - Phase 1: Confidence Scoring
LAYOUT: Feature description with bullet points
Title: "Phase 1: Automatic Quality Flagging"
  \u2022 IS proves 40% is good \u2014 beam scores can flag it at decode time
  \u2022 No additional inference needed \u2014 scores already computed
  \u2022 Effort: 2-4 hours implementation
  \u2022 Priority: Business/Finance segments (57% captured) most reliable
  \u2022 Short segments (<10 words, 32%) need lower confidence thresholds
  \u2022 This is the immediate next step
SPEAKER NOTES: "Phase 1 is the quick win. We know 40% is already good; beam scores computed during decode can flag quality automatically. Business/Finance content (57% captured) should get highest confidence. Short segments need lower thresholds. 2-4 hours of implementation. [ANIMATION: Bullet items build one at a time. Transition: Fade 0.5s.]"

---

SLIDE 27 - Phase 2: N-Best Aggregation
LAYOUT: Feature description with key numbers
Title: "Phase 2: Exploit All 20 Hypotheses"
  \u2022 Currently discarding 19 of 20 beam candidates
  \u2022 ROVER/MBR \u2014 established technique, consistent gains
  \u2022 Expected: 5-15% relative improvement
  \u2022 Targets: Accumulated Small Errors (12.3%) and Content Word Errors (10.7%)
  \u2022 Moves "fair" tier segments up to "good"
SPEAKER NOTES: "We\u2019re currently throwing away 19 of 20 hypotheses. N-best aggregation (ROVER/MBR) is an established technique that combines multiple hypotheses to reduce errors. Targets the \u2018death by 1000 cuts\u2019 failure mode \u2014 accumulated small errors (12.3%) and content word errors (10.7%). These segments have the right words scattered across multiple hypotheses. [ANIMATION: Bullet items build one at a time. Transition: Fade 0.5s.]"

---

SLIDE 28 - Fine-Tuning + Data Scaling
LAYOUT: Split \u2014 dashboard top, projections bottom
Top half: IMAGE PLACEHOLDER labeled "[FT_10_summary_dashboard.png]" (large, filling top half of slide)
Title: "Domain Adaptation: Data Is the Bottleneck"
Below the placeholder, two columns:
Left \u2014 "Exp A/B Results":
  \u2022 Exp A (r=16): best val acc 62.94% at epoch 2, severe overfitting
  \u2022 Exp B (r=64): 3.1pp WORSE \u2014 more params = faster overfitting
  \u2022 KEY: 1,273 segments too small for LoRA generalization
Right \u2014 "Data Scaling (ICLR 2024)":
  \u2022 5K segments: 55-60% WER
  \u2022 20K segments: 45-50% / 40-45% (Llama 3.1)
  \u2022 50K segments: 40-45% / 35-40%
  \u2022 AVSpeech: 290K videos available
SPEAKER NOTES: "Exp A: 1,273 segments, 17 hours on Tesla T4. Best validation accuracy 62.94% at epoch 2 out of 19 \u2014 severe overfitting. 36.5pp gap between train (95.5%) and val (59.0%) by end. Exp B with r=64: 3.1pp WORSE, not better. More parameters = faster overfitting on tiny data. This is NOT a model capacity problem. 1,273 segments is below the ~1K minimum for LoRA generalization (ICLR 2024 scaling laws). Data scaling projections: 20K segments brings WER to 45-50% with Llama-2, or 40-45% with Llama 3.1 8B. AVSpeech has 290K videos \u2014 data curation is tractable. Key: multiplicative scaling law means stronger LLM extracts MORE from same data. [ANIMATION: Dashboard fades in top, then columns appear below. Transition: Fade 0.5s.]"

---

SLIDE 29 - LLM Upgrade + Advanced Capabilities
LAYOUT: Three columns
Title: "Stronger LLM + Smart Prompts = Force Multiplier"
Column 1 \u2014 "LLM Swap (immediate)":
  \u2022 Llama 3.1 8B: drop-in (same hidden_size 4096)
  \u2022 Quality \u2248 Llama-2 70B, 128K vocab, 128K context
  \u2022 Setup: 1-2 hours
  \u2022 Alone: -3 to -8pp WER
Column 2 \u2014 "Smart Prompts (force multiplier)":
  \u2022 7 strategies: topic context, word count, anti-hallucination, GER
  \u2022 Llama-2: +5-10pp | Llama 3.1: +12-20pp
  \u2022 GER post-processing: +8-15pp, no retraining
Column 3 \u2014 "Future":
  \u2022 Arabic (K-means model exists)
  \u2022 VALLR: 18.7% WER with 3B model
  \u2022 Multi-speaker, streaming
SPEAKER NOTES: "Three columns of future capability. Left: LLM swap to Llama 3.1 8B is trivial \u2014 same hidden dimension, 1-2 hours of setup. Alone it gives -3 to -8pp WER. Llama-3 8B quality is comparable to Llama-2 70B but with 4x vocabulary and 32x context window. Center: 7 prompt strategies are a force multiplier \u2014 more effective on stronger models. Llama-2 gets +5-10pp, Llama 3.1 gets +12-20pp from the same strategies. GER (Generative Error Correction) uses N-best hypotheses + a correction LLM (even an external API like Claude) for +8-15pp with no retraining. Right: Arabic support exists (k-means model trained), VALLR achieves 18.7% WER on LRS3 with a 3B model \u2014 architecture innovation matters more than model size. [ANIMATION: Three columns appear one at a time, left to right. Transition: Fade 0.5s.]"

---

SLIDE 30 - Summary
LAYOUT: Four numbered points
Title: "Key Takeaways"
Four points (teal numbers):
  1. Rigorous assessment: 2.5x WER gap on 1,497 segments. Novel IS metric reveals 40% properly captured. 10 classified failure modes.
  2. Production system delivered: standalone container, 37 bugs fixed, 8-stage pipeline, 37 tests, 8 research reports.
  3. Data is the bottleneck: Exp A/B proved 1,273 segments too small. Multiplicative scaling law: stronger LLM + more data compounds.
  4. Actionable roadmap to 27-42% WER: LLM swap + smart prompts + data scaling + GER. Each targets a different bottleneck.
SPEAKER NOTES: "Four takeaways. One: we know exactly where we stand \u2014 rigorous assessment with a novel metric that reveals the true picture. Two: production system delivered and deployed. Three: data, not model capacity, is the bottleneck \u2014 and we know the scaling law. Four: clear roadmap from 67% to 27-42% WER, with each phase targeting a different bottleneck and gains that compound multiplicatively. [ANIMATION: Four points build one at a time. Transition: Fade 0.5s.]"
"""


# ═══════════════════════════════════════════════════
# IMAGE INSERTION TABLES (insert manually in Google
# Slides after exporting from Gemini — NOT uploaded)
# ═══════════════════════════════════════════════════

ATTACHMENTS_1 = [
    ["08_branding/BlackLogo300x300-W-BG.png", "1", "Replace [Logo] placeholder, top-right ~1 inch"],
    ["09_pipeline_diagram/pipeline_architecture.png", "3", "Replace [pipeline_architecture.png] placeholder, full width"],
    ["01_plots_for_slides/P2_paper_vs_reality.png", "4", "Replace [P2_paper_vs_reality.png] placeholder, right half"],
]

ATTACHMENTS_2 = [
    ["01_plots_for_slides/P1_quality_tiers.png", "5", "Replace [P1_quality_tiers.png] placeholder, right half"],
    ["01_plots_for_slides/09_boxplot_wwer_all_experiments.png", "9", "Replace [09_boxplot_wwer...] placeholder, center ~70% width"],
    ["01_plots_for_slides/01_wer_vs_duration.png", "10", "Replace [01_wer_vs_duration.png] placeholder, right half"],
    ["01_plots_for_slides/14_nea_vs_wwer_scatter.png", "11", "Replace [14_nea_vs_wwer_scatter.png] placeholder, right half"],
    ["01_plots_for_slides/10_empty_and_hallucination_rates.png", "12", "Replace [10_empty_and_hallucination...] placeholder, right half"],
    ["01_plots_for_slides/P4_lenpen_sensitivity.png", "13", "Replace [P4_lenpen_sensitivity.png] placeholder, right half"],
    ["01_plots_for_slides/P5_tuning_before_after.png", "12", "Optional: add to left half or keep as speaker notes backup"],
]

ATTACHMENTS_3 = [
    ["09_pipeline_diagram/pipeline_architecture.png", "17", "Replace [pipeline_architecture.png] placeholder, full width"],
    ["01_plots_for_slides/15_cdf_wwer_curated.png", "22", "Replace [15_cdf_wwer_curated.png] placeholder, right half"],
]

ATTACHMENTS_4 = [
    ["01_plots_for_slides/P1_quality_tiers.png", "24", "Replace [P1_quality_tiers.png] placeholder, right half"],
    ["01_plots_for_slides/P3_wer_trajectory.png", "25", "Replace [P3_wer_trajectory.png] placeholder, right half"],
    ["01_plots_for_slides/finetune/FT_10_summary_dashboard.png", "28", "Replace [FT_10_summary_dashboard.png] placeholder, top half"],
    ["01_plots_for_slides/finetune/FT_01_loss_curves.png", "28", "Appendix backup \u2014 loss curves (speaker notes ref)"],
    ["01_plots_for_slides/finetune/FT_02_accuracy_curves.png", "28", "Appendix backup \u2014 accuracy curves (speaker notes ref)"],
    ["01_plots_for_slides/finetune/FT_03_overfitting_gap.png", "28", "Appendix backup \u2014 overfitting gap (speaker notes ref)"],
]


# ═══════════════════════════════════════════════════
# ANIMATION CHECKLIST DATA
# ═══════════════════════════════════════════════════

ANIMATION_CHECKLIST = [
    ["1", "Title, subtitle", "Fade in sequence (0.5s delay)", "None (first slide)"],
    ["2", "Text, play icon", "Appear", "Fade 0.5s"],
    ["3", "Three component blocks", "Fly in left-to-right", "Fade 0.5s"],
    ["4", "Number, chart", "Number appears, chart fades in", "Fade 0.5s"],
    ["5", "5 quality tier bars", "Build top-to-bottom (0.3s each)", "Fade 0.5s"],
    ["6", "Left/right comparison boxes", "Left appears, then right", "Fade 0.5s"],
    ["7", "6 IS signal blocks + tier bars", "Blocks build one at a time", "Fade 0.5s"],
    ["8", "10 failure mode bars", "Wipe left, top-to-bottom", "Fade 0.5s"],
    ["9", "Chart", "No build animation", "Fade 0.5s"],
    ["10", "Three root cause columns", "Columns appear one at a time", "Fade 0.5s"],
    ["11", "Text, chart", "Text appears, chart fades in", "Fade 0.5s"],
    ["12", "Left/right columns", "Left appears, then right", "Fade 0.5s"],
    ["13", "Text, chart", "Text appears, chart fades in", "Fade 0.5s"],
    ["14", "Table rows", "Build one row at a time", "Fade 0.5s"],
    ["15", "Play icon", "No build animation", "Fade 0.5s"],
    ["16", "List items", "Build one at a time", "Fade 0.5s"],
    ["17", "Pipeline diagram", "Fade in", "Fade 0.5s"],
    ["18", "Left/right columns", "Left appears, then right", "Fade 0.5s"],
    ["19", "Before/after blocks", "Left appears, then right", "Morph"],
    ["20", "Feature items", "Build one at a time", "Fade 0.5s"],
    ["21", "Three cards", "Fly in from bottom, one at a time", "Fade 0.5s"],
    ["22", "List items, chart", "List builds, chart fades in", "Fade 0.5s"],
    ["23", "Bullet list", "No build animation", "Fade 0.5s"],
    ["24", "Red/green columns", "Left (red) appears, then right (green)", "Fade 0.5s"],
    ["25", "5 roadmap phases", "Staircase build, one at a time", "Fade 0.5s"],
    ["26", "Bullet items", "Build one at a time", "Fade 0.5s"],
    ["27", "Bullet items", "Build one at a time", "Fade 0.5s"],
    ["28", "Dashboard, columns", "Dashboard fades in, columns appear below", "Fade 0.5s"],
    ["29", "Three columns", "Appear one at a time, left to right", "Fade 0.5s"],
    ["30", "Four takeaway points", "Build one at a time", "Fade 0.5s"],
]


# ═══════════════════════════════════════════════════
# DOCUMENT SECTIONS
# ═══════════════════════════════════════════════════

def build_intro(doc):
    """Section 1: How to Use This Document."""
    add_heading(doc, "1. How to Use This Document", 1)

    add_para(doc, (
        "This document contains 4 sequential prompts for Google Gemini 3.1 Pro "
        "to generate the Argos VSP project review presentation (30 slides). "
        "Gemini creates the slide structure, text, layout, and speaker notes. "
        "You then insert images and apply animations manually after export."
    ))

    add_heading(doc, "Step-by-Step Workflow", 2)
    add_bullet(doc, 'Open Google Gemini (gemini.google.com) and start a new conversation')
    add_bullet(doc, 'Paste Prompt 1 (NO image uploads needed \u2014 Gemini creates placeholders)')
    add_bullet(doc, 'In the SAME conversation, paste Prompts 2, 3, and 4 sequentially')
    add_bullet(doc, 'Export the presentation to Google Slides (Gemini offers this option)')
    add_bullet(doc, 'In Google Slides: replace each gray placeholder box with the real image (see tables below each prompt)')
    add_bullet(doc, 'Apply animations using the Animation Checklist (Section 8) \u2014 ~10-15 min')
    add_bullet(doc, 'Insert demo videos on slides 2 and 15 from 06_demo_videos/')

    add_heading(doc, "Why Not Upload Images to Gemini?", 2)
    add_para(doc, (
        "Gemini\u2019s presentation generator does not reliably insert uploaded images. "
        "When told to use uploaded files, it either ignores them, fetches stock photos "
        "from the internet, or creates broken image placeholders. Instead, the prompts "
        "tell Gemini to create labeled dark gray placeholder rectangles (e.g., "
        "\"[pipeline_architecture.png]\"). You replace these in Google Slides after export."
    ))

    add_heading(doc, "How Animations Work", 2)
    add_para(doc, (
        "Each slide\u2019s speaker notes contain an [ANIMATION: ...] tag describing "
        "the recommended animation. Gemini cannot add animations programmatically, "
        "so apply them manually after generation using the Animation Checklist "
        "(Section 8). Estimated time: 10-15 minutes."
    ))

    add_heading(doc, "Video Slides", 2)
    add_para(doc, (
        "Slides 2 and 15 are marked [VIDEO PLACEHOLDER]. Gemini creates "
        "slides with play button icons. You play the demo videos from "
        "06_demo_videos/ externally during the presentation."
    ))

    add_heading(doc, "Design Theme", 2)
    add_para(doc, (
        "All prompts specify: dark navy background (#0D1B2A), white text (#FFFFFF), "
        "teal accents (#00B4D8), coral accents (#E06C75). "
        "Prompt 1 sets the full theme; subsequent prompts reference it briefly."
    ))

    _tight_page_break(doc)


def build_overview(doc):
    """Section 2: Presentation Overview."""
    add_heading(doc, "2. Presentation Overview", 1)

    add_styled_table(doc,
        ["Section", "Prompt", "Slides", "Duration", "Images to Insert After"],
        [
            ["Opening & Context", "Prompt 1", "1-4", "~5 min", "3"],
            ["Research Findings", "Prompt 2", "5-16", "~20 min", "7"],
            ["Engineering", "Prompt 3", "17-23", "~10 min", "2"],
            ["Future & Close", "Prompt 4", "24-30", "~15 min", "6"],
        ],
        col_widths=[1.3, 0.8, 0.7, 0.8, 1.0],
    )

    add_heading(doc, "Key Numbers", 2)
    add_styled_table(doc,
        ["Metric", "Value", "Source"],
        [
            ["Paper WER (LRS3)", "25.4%", "VSP-LLM paper, Table 1"],
            ["Real-world mean WER", "64.1%", "Report 1, 1,497 segments"],
            ["Reality gap", "2.5x worse", "Report 1"],
            ["Properly captured (IS \u2265 3.0)", "39.9% (597 segments)", "Intelligibility Score"],
            ["WER overstatement factor", "3.5x", "39.9% vs 11.4%"],
            ["Failure modes classified", "10 modes", "900 failed segments"],
            ["#1 failure: Topic Drift", "15.9% (143 segments)", "Failure taxonomy"],
            ["#1 success: Phonetic Preservation", "41.5%", "248 of 597 successes"],
            ["Bugs fixed", "37", "Bug tracking docs"],
            ["Pipeline: lines before \u2192 after", "823 \u2192 393", "52% reduction"],
            ["Total experiments", "13 (A-M)", "Tuning dataset"],
            ["Exp A: best val accuracy", "62.94% at epoch 2", "r=16, 320 updates"],
            ["Exp A: overfitting gap", "36.5 pp", "Train 95.5% vs val 59.0%"],
            ["IS: top correlate", "Phonetic Sim r=0.943", "Correlation analysis"],
            ["IS: Claude-designed heuristic (16 configs)", "r=0.925, 88.6% agreement", "No runtime LLM calls"],
            ["LLM: current model", "Llama-2-7B (4-bit QLoRA)", "Hidden size 4096"],
            ["LLM: recommended swap", "Llama 3.1 8B (drop-in)", "1-2 hours setup"],
            ["Prompts: force multiplier", "Llama-2 +5-10pp / 3.1 +12-20pp", "7 strategies"],
            ["GER post-processing", "+8-15pp, no retraining", "N-best + correction LLM"],
            ["Combined target", "27-42% WER (from 67%)", "LLM + data + prompts + GER"],
        ],
        col_widths=[2.5, 2.2, 2.0],
    )

    _tight_page_break(doc)


def build_prompt_section(doc, number, title, slides, prompt_text, attachments=None):
    """Build a prompt section with description, image table, and copyable prompt block."""
    add_heading(doc, f"{number + 2}. Prompt {number}: {title} ({slides})", 1)

    add_para(doc, (
        f"Paste this prompt into Gemini (no file uploads needed). "
        f"This prompt generates {slides}."
    ), bold=True)

    add_heading(doc, "Prompt Text (copy and paste into Gemini)", 2)
    add_prompt_block(doc, prompt_text)

    if attachments:
        add_attachment_table(doc, attachments)

    _tight_page_break(doc)


def build_appendix(doc):
    """Section 7: Appendix Slides."""
    add_heading(doc, "7. Appendix Slides", 1)

    add_para(doc, (
        "These slides are NOT generated by the Gemini prompts. "
        "Add them manually to the end of the deck as needed during Q&A."
    ))

    add_styled_table(doc,
        ["#", "Slide", "When to Use"],
        [
            ["A1", "Homophenes table", "If asked about visual confusion patterns"],
            ["A2", "Exp A Fine-Tuning Deep Dive \u2014 6-panel dashboard (FT_10), "
                   "loss/accuracy curves, overfitting gap. Key: 1,273 train segments, "
                   "17h Tesla T4, best at epoch 2/19",
                   "If asked for training details or overfitting analysis"],
            ["A3", "Catastrophic lenpen=2.0 (WER 171%)", "If asked about parameter space boundaries"],
            ["A4", "CDF WWER chart", "If asked about quality thresholds"],
            ["A5", "Topic analysis detail (11 categories)", "If asked about domain-specific performance"],
            ["A6", "Signal comparison table (success vs failure)", "If asked about what makes IS work"],
            ["A7", "Config J full-dataset comparison", "If asked about tuning details on full data"],
            ["A8", "IS Correlation & Cross-Config Analysis", "If asked about IS validation methodology"],
            ["A9", "LLM Upgrade Analysis \u2014 model alternatives, "
                   "data scaling projections (1.3K\u2192100K), "
                   "7 prompt strategies, improvement waterfall (67%\u219227-42%)",
                   "If asked about LLM alternatives or investment strategy"],
        ],
        col_widths=[0.4, 3.0, 3.2],
    )

    _tight_page_break(doc)


def build_animation_checklist(doc):
    """Section 8: Animation & Transition Checklist."""
    add_heading(doc, "8. Animation & Transition Checklist", 1)

    add_para(doc, (
        "Gemini cannot add animations programmatically. Apply these manually after "
        "generating the slides. Each slide\u2019s speaker notes also contain an "
        "[ANIMATION: ...] tag as a reminder. Estimated time: 10-15 minutes."
    ))

    add_para(doc, (
        "In Google Slides: select elements \u2192 Insert \u2192 Animation \u2192 choose effect. "
        "For transitions: Slide \u2192 Transition \u2192 Fade (0.5s) for all slides."
    ), italic=True, color=C_GRAY, size=Pt(10))

    add_styled_table(doc,
        ["Slide", "Elements", "Animation", "Transition"],
        ANIMATION_CHECKLIST,
        col_widths=[0.5, 2.0, 2.5, 1.5],
    )

    _tight_page_break(doc)


def build_manual_steps(doc):
    """Section 9: Post-Generation Manual Steps."""
    add_heading(doc, "9. Post-Generation Manual Steps", 1)

    add_para(doc, (
        "After all 4 prompts are complete, do these steps to finalize the deck:"
    ))

    steps = [
        ("Replace image placeholders", (
            "Each slide with a chart or diagram has a labeled gray placeholder rectangle. "
            "In Google Slides: click the placeholder \u2192 delete it \u2192 Insert \u2192 Image \u2192 Upload from computer. "
            "Use the image tables below each prompt (Sections 3-6) to find the correct file and placement."
        )),
        ("Apply animations", (
            "Use the Animation Checklist (Section 8) to apply animations and transitions. "
            "Default transition: Fade 0.5s for all slides. Build animations for data-heavy slides."
        )),
        ("Insert demo videos", (
            "Slides 2 and 15 have play button placeholders. Either: "
            "(a) link video files if presenting from Google Slides, or "
            "(b) play from external media player during presentation (recommended). "
            "Video files are in 06_demo_videos/."
        )),
        ("Add logo to slide master", (
            "If Gemini didn\u2019t carry the logo through all slides, add it to the slide master "
            "layout (View \u2192 Theme builder) for consistent branding."
        )),
        ("Add appendix slides", (
            "Add appendix slides (A1-A9) manually as needed for Q&A backup. "
            "See Section 7 for the list and when to use each."
        )),
        ("Review speaker notes", (
            "Each slide has detailed speaker notes generated by Gemini. "
            "Review for accuracy and add any personal notes or talking points."
        )),
    ]

    for i, (title, description) in enumerate(steps, 1):
        add_para(doc, f"{i}. {title}", bold=True, size=Pt(11))
        add_para(doc, description, size=Pt(10), color=C_GRAY)


def build_materials(doc):
    """Section 10: Material References."""
    add_heading(doc, "10. Material References", 1)

    add_para(doc, "All materials are in the presentation_materials_20260224/ folder:")

    add_styled_table(doc,
        ["Folder", "Contents", "Used In"],
        [
            ["01_plots_for_slides/", "17 presentation graphs (P1-P5, existing, finetune/)",
             "Insert into slides 4-13, 22, 24-28 (replace placeholders)"],
            ["01_plots_for_slides/finetune/", "4 key fine-tuning plots (FT_01, FT_02, FT_03, FT_10)",
             "Insert into slide 28 (replace placeholder)"],
            ["02_plots_boss_deep_dive/", "14 technical deep-dive graphs",
             "Appendix, Q&A backup"],
            ["03_reports_md/supplementary/", "10 research reports in Markdown",
             "Presenter reference only"],
            ["04_reports_docx/", "18 formatted reports (Word + PDF)",
             "Handouts"],
            ["05_data/", "18+ data files (CSV, JSON, HTML reports)",
             "Reference data"],
            ["06_demo_videos/", "8 burned videos with subtitle overlays",
             "Play externally on slides 2, 15"],
            ["07_paper/", "VSP-LLM paper + 2025 presentation",
             "Background reference"],
            ["08_branding/", "3 logo files",
             "Insert into slide 1 (replace placeholder)"],
            ["09_pipeline_diagram/", "8-stage pipeline flow diagram",
             "Insert into slides 3, 17 (replace placeholders)"],
            ["10_examples/", "Curated example data",
             "Content source for slide 14"],
        ],
        col_widths=[1.8, 2.8, 2.0],
    )


# ═══════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════

def main():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # Build document
    build_header(doc)
    build_footer(doc)
    build_cover(doc)
    build_toc(doc)
    build_intro(doc)
    build_overview(doc)

    # 4 Gemini prompts with image attachments
    prompts = [
        (1, "Opening & Context", "Slides 1-4", PROMPT_1, ATTACHMENTS_1),
        (2, "Research Findings", "Slides 5-16", PROMPT_2, ATTACHMENTS_2),
        (3, "Engineering", "Slides 17-23", PROMPT_3, ATTACHMENTS_3),
        (4, "Future & Close", "Slides 24-30", PROMPT_4, ATTACHMENTS_4),
    ]
    for num, title, slides, text, attachments in prompts:
        build_prompt_section(doc, num, title, slides, text, attachments)

    build_appendix(doc)
    build_animation_checklist(doc)
    build_manual_steps(doc)
    build_materials(doc)

    # Save
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_FILE))
    print(f"Generated: {OUTPUT_FILE}")
    print(f"Size: {OUTPUT_FILE.stat().st_size / 1024:.0f} KB")


if __name__ == "__main__":
    main()
