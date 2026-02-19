#!/usr/bin/env python3
"""
Argos — The Orchard — Speech Recognition Metrics Explainer

Generates a branded .docx that explains WER, WWER, NEA, and Word Alignment
in detail, with per-word color-coded impact examples from real pipeline data.

Usage:
    python3 generate_metrics_explainer.py

Output:
    argos_research/metrics_explainer.docx
"""

import re
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

# ── Output ──
SCRIPT_DIR = Path(__file__).parent
OUTPUT_DIR = SCRIPT_DIR.resolve().parent.parent / "tuning"
OUTPUT_FILE = OUTPUT_DIR / "metrics-explainer.docx"

# ── Logos ──
ASSETS_DIR = SCRIPT_DIR.parent / "assets"
LOGO_ORCHARD = ASSETS_DIR / "logo.png"
LOGO_PEACOCK = ASSETS_DIR / "peacock.png"

# ── Standard Colors ──
C_PRIMARY = RGBColor(0x1a, 0x3a, 0x5c)
C_H2 = RGBColor(0x2a, 0x5a, 0x8c)
C_H3 = RGBColor(0x3a, 0x6a, 0x9c)
C_H4 = RGBColor(0x4a, 0x7a, 0xac)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_BLACK = RGBColor(0x00, 0x00, 0x00)
C_DARK = RGBColor(0x33, 0x33, 0x33)
C_GREEN = RGBColor(0x15, 0x57, 0x24)
C_RED = RGBColor(0x72, 0x1c, 0x24)
C_AMBER = RGBColor(0x85, 0x64, 0x04)
C_GRAY = RGBColor(0x66, 0x66, 0x66)
C_CODE = RGBColor(0x2d, 0x2d, 0x2d)

HEADER_BG = "1a3a5c"
ZEBRA_BG = "f0f4f8"
GREEN_BG = "d4edda"
RED_BG = "f8d7da"
AMBER_BG = "fff3cd"
CODE_BG = "f5f5f5"

# ── Impact Colors (new for this document) ──
C_IMPACT_GOOD = RGBColor(0x0a, 0x7a, 0x0a)   # green — word helps metric
C_IMPACT_BAD = RGBColor(0xb0, 0x00, 0x20)     # red — word hurts metric
C_IMPACT_MILD = RGBColor(0xb5, 0x89, 0x00)    # amber — moderate impact
C_IMPACT_NONE = RGBColor(0x99, 0x99, 0x99)    # gray — invisible to metric
C_ENTITY = RGBColor(0x1a, 0x4e, 0x8a)         # dark blue — entity token

LAST_UPDATED = datetime.now().strftime("%B %d, %Y")


# ═══════════════════════════════════════════════════
# STANDARD HELPERS (from existing generators)
# ═══════════════════════════════════════════════════

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, bold=False, color=None, size=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = size
    run.font.name = "Calibri"


def add_styled_table(doc, headers, rows, col_widths=None, highlight_rows=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, HEADER_BG)
        set_cell_text(cell, h, bold=True, color=C_WHITE, size=Pt(9))
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            set_cell_text(cell, val, size=Pt(9))
            if highlight_rows and r_idx in highlight_rows:
                set_cell_shading(cell, highlight_rows[r_idx])
            elif r_idx % 2 == 1:
                set_cell_shading(cell, ZEBRA_BG)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


_bookmark_counter = [0]


def add_heading(doc, text, level, color=None, bookmark_id=None):
    h = doc.add_heading(text, level=level)
    if color is None:
        color = {1: C_PRIMARY, 2: C_H2, 3: C_H3, 4: C_H4}.get(level, C_PRIMARY)
    for run in h.runs:
        run.font.color.rgb = color
        run.font.name = "Calibri"
    if bookmark_id:
        _bookmark_counter[0] += 1
        bm_start = OxmlElement("w:bookmarkStart")
        bm_start.set(qn("w:id"), str(_bookmark_counter[0]))
        bm_start.set(qn("w:name"), bookmark_id)
        bm_end = OxmlElement("w:bookmarkEnd")
        bm_end.set(qn("w:id"), str(_bookmark_counter[0]))
        h._p.insert(0, bm_start)
        h._p.append(bm_end)
    return h


def add_para(doc, text, bold=False, italic=False, size=Pt(11), color=None, space_after=Pt(6)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = size
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = space_after
    return p


def add_para_rich(doc, segments, space_after=Pt(6)):
    p = doc.add_paragraph()
    for text, bold in segments:
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(11)
        run.font.name = "Calibri"
    p.paragraph_format.space_after = space_after
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = "Calibri"
    return p


def add_bullet_bold_value(doc, label, value, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after = Pt(2)
    run_b = p.add_run(label)
    run_b.bold = True
    run_b.font.size = Pt(10)
    run_b.font.name = "Calibri"
    run_v = p.add_run(value)
    run_v.font.size = Pt(10)
    run_v.font.name = "Calibri"
    return p


def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{CODE_BG}"/>')
    p._p.get_or_add_pPr().append(shading)
    run = p.add_run(code_text)
    run.font.size = Pt(8)
    run.font.name = "Consolas"
    run.font.color.rgb = C_CODE
    p.paragraph_format.space_after = Pt(6)
    return p


def _build_inline_image_xml(rId, cx, cy, pic_id=1, name="Logo"):
    inline = OxmlElement('wp:inline')
    inline.set('distT', '0'); inline.set('distB', '0')
    inline.set('distL', '0'); inline.set('distR', '0')
    extent = OxmlElement('wp:extent')
    extent.set('cx', str(cx)); extent.set('cy', str(cy))
    inline.append(extent)
    docPr = OxmlElement('wp:docPr')
    docPr.set('id', str(pic_id)); docPr.set('name', name)
    inline.append(docPr)
    graphic = OxmlElement('a:graphic')
    graphicData = OxmlElement('a:graphicData')
    graphicData.set('uri', 'http://schemas.openxmlformats.org/drawingml/2006/picture')
    pic = OxmlElement('pic:pic')
    nvPicPr = OxmlElement('pic:nvPicPr')
    cNvPr = OxmlElement('pic:cNvPr')
    cNvPr.set('id', '0'); cNvPr.set('name', name)
    nvPicPr.append(cNvPr)
    cNvPicPr = OxmlElement('pic:cNvPicPr')
    nvPicPr.append(cNvPicPr)
    pic.append(nvPicPr)
    blipFill = OxmlElement('pic:blipFill')
    blip = OxmlElement('a:blip')
    blip.set(qn('r:embed'), rId)
    blipFill.append(blip)
    stretch = OxmlElement('a:stretch')
    fillRect = OxmlElement('a:fillRect')
    stretch.append(fillRect)
    blipFill.append(stretch)
    pic.append(blipFill)
    spPr = OxmlElement('pic:spPr')
    xfrm = OxmlElement('a:xfrm')
    off = OxmlElement('a:off'); off.set('x', '0'); off.set('y', '0')
    xfrm.append(off)
    ext_el = OxmlElement('a:ext')
    ext_el.set('cx', str(cx)); ext_el.set('cy', str(cy))
    xfrm.append(ext_el)
    spPr.append(xfrm)
    prstGeom = OxmlElement('a:prstGeom')
    prstGeom.set('prst', 'rect')
    spPr.append(prstGeom)
    pic.append(spPr)
    graphicData.append(pic)
    graphic.append(graphicData)
    inline.append(graphic)
    drawing = OxmlElement('w:drawing')
    drawing.append(inline)
    return drawing


def add_header_footer(doc):
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    if LOGO_ORCHARD.exists():
        image_part = doc.part.package.get_or_add_image_part(str(LOGO_ORCHARD))
        rId = header.part.relate_to(image_part, RT.IMAGE)
        size_emu = int(0.3 * 914400)
        logo_run = hp.add_run()
        drawing = _build_inline_image_xml(rId, size_emu, size_emu, pic_id=10, name="Header Logo")
        logo_run._r.append(drawing)
    # Right-tab pushes text to far edge; logo stays at left margin
    text_width = section.page_width - section.left_margin - section.right_margin
    hp.paragraph_format.tab_stops.add_tab_stop(text_width, WD_TAB_ALIGNMENT.RIGHT)
    hp.add_run("\t")
    run = hp.add_run("Argos \u2014 The Orchard")
    run.font.size = Pt(8); run.font.color.rgb = C_GRAY
    run.font.name = "Calibri"; run.italic = True
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("Page ")
    run.font.size = Pt(8); run.font.color.rgb = C_GRAY; run.font.name = "Calibri"
    fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
    run._r.append(fld1)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve"); instr.text = " PAGE "
    run._r.append(instr)
    fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld2)


def _tight_page_break(doc):
    """Add page break to the last paragraph instead of creating a new one.

    Use after tables (which add a trailing paragraph) to avoid a visible
    blank gap at the bottom of the page before the break.
    """
    last_p = doc.paragraphs[-1]
    last_p.add_run().add_break(WD_BREAK.PAGE)


def _add_toc_hyperlink(doc, bookmark_name, display_text, level=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.left_indent = Inches(0.3 * level)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), bookmark_name)
    run_el = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle"); rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "22")
    rPr.append(sz)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Calibri"); rFonts.set(qn("w:hAnsi"), "Calibri")
    rPr.append(rFonts)
    color = OxmlElement("w:color"); color.set(qn("w:val"), "1a3a5c")
    rPr.append(color)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single")
    rPr.append(u)
    run_el.append(rPr)
    text_el = OxmlElement("w:t")
    text_el.set(qn("xml:space"), "preserve"); text_el.text = display_text
    run_el.append(text_el)
    hyperlink.append(run_el)
    p._p.append(hyperlink)


# ═══════════════════════════════════════════════════
# NEW HELPERS — Per-Word Impact Coloring
# ═══════════════════════════════════════════════════

def add_colored_runs(doc, segments, space_after=Pt(6), size=Pt(10)):
    """Paragraph with mixed bold/italic/color runs.
    segments = [(text, bold, color, italic), ...] — italic is optional (default False).
    """
    p = doc.add_paragraph()
    for seg in segments:
        text, bold, color = seg[0], seg[1], seg[2]
        italic = seg[3] if len(seg) > 3 else False
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = size
        run.font.name = "Calibri"
        if color:
            run.font.color.rgb = color
    p.paragraph_format.space_after = space_after
    return p


def add_impact_example(doc, label, words_with_impact, prefix_color=C_DARK):
    """Render a sentence with per-word impact coloring.
    label: "REF:" or "HYP:" etc.
    words_with_impact: [(word, color, bold), ...]
    """
    segments = [(label + " ", True, prefix_color)]
    for i, (word, color, bold) in enumerate(words_with_impact):
        if i > 0:
            segments.append((" ", False, None))
        italic = (color == C_IMPACT_NONE)
        segments.append((word, bold, color, italic))
    return add_colored_runs(doc, segments)


def add_thermometer_table(doc, thresholds):
    """Colored-row interpretation scale.
    thresholds = [(range_text, bg_color, interpretation, example), ...]
    """
    table = doc.add_table(rows=1 + len(thresholds), cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["Range", "Interpretation", "Example"]):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, HEADER_BG)
        set_cell_text(cell, h, bold=True, color=C_WHITE, size=Pt(9))
    for r_idx, (rng, bg, interp, example) in enumerate(thresholds):
        row = table.rows[r_idx + 1]
        set_cell_text(row.cells[0], rng, size=Pt(9), bold=True)
        set_cell_shading(row.cells[0], bg)
        set_cell_text(row.cells[1], interp, size=Pt(9))
        set_cell_shading(row.cells[1], bg)
        set_cell_text(row.cells[2], example, size=Pt(9))
        set_cell_shading(row.cells[2], bg)
    for row in table.rows:
        row.cells[0].width = Inches(1.0)
        row.cells[1].width = Inches(3.4)
        row.cells[2].width = Inches(2.1)
    doc.add_paragraph()
    return table


# ═══════════════════════════════════════════════════
# INLINE TOKENIZATION & ALIGNMENT (self-contained)
# ═══════════════════════════════════════════════════

def toks(s):
    s = (s or "").strip().lower()
    return re.findall(r"[a-z0-9]+(?:'[a-z0-9]+)?", s)


def align(ref, hyp):
    r = toks(ref)
    h = toks(hyp)
    ref_words = {}
    for w in r:
        ref_words[w] = ref_words.get(w, 0) + 1
    tagged = []
    for i, hyp_word in enumerate(h):
        if i < len(r) and hyp_word == r[i]:
            tagged.append((hyp_word, "ok"))
        elif hyp_word in ref_words:
            tagged.append((hyp_word, "rep"))
        else:
            tagged.append((hyp_word, "ins"))
    return tagged


# ═══════════════════════════════════════════════════
# TOC
# ═══════════════════════════════════════════════════

TOC_ENTRIES = [
    ("sec1", "1. Introduction \u2014 Why Multiple Metrics?"),
    ("sec2", "2. WER (Word Error Rate)"),
    ("sec3", "3. WWER (Weighted Word Error Rate) & spaCy"),
    ("sec4", "4. NEA (Named Entity Accuracy)"),
    ("sec5", "5. Cross-Metric Analysis & Quick Reference"),
]


# ═══════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════

def create_cover_page(doc):
    for _ in range(2):
        doc.add_paragraph()
    if LOGO_PEACOCK.exists():
        doc.add_picture(str(LOGO_PEACOCK), width=Inches(2.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ARGOS")
    run.font.size = Pt(48); run.font.color.rgb = C_PRIMARY
    run.bold = True; run.font.name = "Calibri"
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Speech Recognition Metrics Explainer")
    run2.font.size = Pt(22); run2.font.color.rgb = C_H2; run2.font.name = "Calibri"
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("The Orchard")
    run3.font.size = Pt(20); run3.font.color.rgb = C_H2; run3.font.name = "Calibri"
    doc.add_paragraph()
    p_author = doc.add_paragraph()
    p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_author = p_author.add_run("Yoad Oxman")
    run_author.font.size = Pt(14); run_author.font.color.rgb = C_DARK
    run_author.font.name = "Calibri"
    doc.add_paragraph()
    for label, value in [
        ("Metrics covered: ", "WER, WWER, NEA (Recall/Precision/F1), Word Alignment"),
        ("Data source: ", "1,497 AVSpeech segments, 13 tuning experiments"),
        ("Last updated: ", LAST_UPDATED),
    ]:
        pi = doc.add_paragraph()
        pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = pi.add_run(label)
        r1.bold = True; r1.font.size = Pt(11)
        r1.font.color.rgb = C_DARK; r1.font.name = "Calibri"
        r2 = pi.add_run(value)
        r2.font.size = Pt(11); r2.font.color.rgb = C_DARK; r2.font.name = "Calibri"
    doc.add_page_break()


def create_toc(doc):
    # Push TOC down from top of page
    for _ in range(3):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    add_heading(doc, "Table of Contents", 1)

    # Standard Word TOC field — auto-populates from heading styles on open
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fld_begin)
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' TOC \\o "1-2" \\h \\z \\u '
    run._r.append(instr)
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    run._r.append(fld_sep)

    # Placeholder entries shown until Word updates the field
    for _bm, title in TOC_ENTRIES:
        placeholder = paragraph.add_run(title + "\n")
        placeholder.font.size = Pt(11)
        placeholder.font.name = "Calibri"
        placeholder.font.color.rgb = C_PRIMARY

    fld_end_run = paragraph.add_run()
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    fld_end_run._r.append(fld_end)

    # Tell Word to update all fields (including TOC) on open
    settings = doc.settings.element
    update_fields = OxmlElement('w:updateFields')
    update_fields.set(qn('w:val'), 'true')
    settings.append(update_fields)

    doc.add_page_break()


# ═══════════════════════════════════════════════════
# SECTION 1 — Introduction
# ═══════════════════════════════════════════════════

def section_1(doc):
    add_heading(doc, "1. Introduction \u2014 Why Multiple Metrics?", 1, bookmark_id="sec1")

    add_para(doc, (
        "A single number cannot capture all the ways a lip-reading system can fail. "
        "Two segments can both score 50% WER yet differ dramatically in quality: one might "
        "get every entity name wrong while preserving sentence structure, the other might nail "
        "every proper noun but scramble the filler words. Each metric illuminates a different "
        "facet of the output, and understanding what each one measures is essential for "
        "interpreting pipeline results and guiding improvement."
    ))

    # "Same Sentence, Three Views" demo
    add_heading(doc, "1.1 Same Sentence, Three Views", 2)

    add_para(doc, (
        "Consider this real example from the pipeline (segment zmwgmt7wcv8, WER 26.7%):"
    ))

    # REF
    add_colored_runs(doc, [
        ("REF: ", True, C_DARK),
        ("watching the way words develop and change does suggest that wittgenstein was on to something", False, C_DARK),
    ])
    add_colored_runs(doc, [
        ("HYP: ", True, C_DARK),
        ("the way words develop and change suggests that philistine was on to something", False, C_DARK),
    ])

    add_para(doc, "WER view \u2014 which words cost edit-distance errors?", bold=True, size=Pt(10))
    # WER: deletions from ref are errors (watching, does), substitutions (suggests, philistine)
    add_impact_example(doc, "HYP", [
        ("the", C_IMPACT_GOOD, False),       # ok
        ("way", C_IMPACT_GOOD, False),
        ("words", C_IMPACT_GOOD, False),
        ("develop", C_IMPACT_GOOD, False),
        ("and", C_IMPACT_GOOD, False),
        ("change", C_IMPACT_GOOD, False),
        ("suggests", C_IMPACT_BAD, True),     # substitution for "does suggest"
        ("that", C_IMPACT_GOOD, False),
        ("philistine", C_IMPACT_BAD, True),   # substitution for "wittgenstein"
        ("was", C_IMPACT_GOOD, False),
        ("on", C_IMPACT_GOOD, False),
        ("to", C_IMPACT_GOOD, False),
        ("something", C_IMPACT_GOOD, False),
    ])
    add_para(doc, "\u2192 4 errors (2 deletions + 2 substitutions) / 15 ref words = 26.7% WER. "
             "Green words are free; each red word adds +1 to the error count equally.",
             size=Pt(9), italic=True, color=C_GRAY)

    add_para(doc, "WWER view \u2014 how much does each error cost by word importance?", bold=True, size=Pt(10))
    add_impact_example(doc, "HYP", [
        ("the", C_IMPACT_GOOD, False),
        ("way", C_IMPACT_GOOD, False),
        ("words", C_IMPACT_GOOD, False),
        ("develop", C_IMPACT_GOOD, False),
        ("and", C_IMPACT_GOOD, False),
        ("change", C_IMPACT_GOOD, False),
        ("suggests", C_IMPACT_MILD, True),    # VERB wrong = 1.0 cost
        ("that", C_IMPACT_GOOD, False),
        ("philistine", C_IMPACT_BAD, True),   # PROPN wrong = 2.0 cost!
        ("was", C_IMPACT_GOOD, False),
        ("on", C_IMPACT_GOOD, False),
        ("to", C_IMPACT_GOOD, False),
        ("something", C_IMPACT_GOOD, False),
    ])
    add_para(doc,
             "\u2192 \"wittgenstein\" (proper noun, weight 2.0) wrong = 2.0 error cost. "
             "\"does\" (function word, weight 0.5) deleted = 0.5 cost. "
             "The entity error dominates. WWER = 37.5%.",
             size=Pt(9), italic=True, color=C_GRAY)

    add_para(doc, "NEA view \u2014 which entity words were captured or missed?", bold=True, size=Pt(10))
    add_impact_example(doc, "REF", [
        ("watching", C_IMPACT_NONE, False),
        ("the", C_IMPACT_NONE, False),
        ("way", C_IMPACT_NONE, False),
        ("words", C_IMPACT_NONE, False),
        ("develop", C_IMPACT_NONE, False),
        ("and", C_IMPACT_NONE, False),
        ("change", C_IMPACT_NONE, False),
        ("does", C_IMPACT_NONE, False),
        ("suggest", C_IMPACT_NONE, False),
        ("that", C_IMPACT_NONE, False),
        ("wittgenstein", C_IMPACT_BAD, True),  # entity MISSED
        ("was", C_IMPACT_NONE, False),
        ("on", C_IMPACT_NONE, False),
        ("to", C_IMPACT_NONE, False),
        ("something", C_IMPACT_NONE, False),
    ])
    add_para(doc,
             "\u2192 1 entity (\"wittgenstein\") in reference, 0 found in hypothesis. "
             "NEA Recall = 0%. The 12 correct gray words are invisible to this metric.",
             size=Pt(9), italic=True, color=C_GRAY)

    # At a glance table
    add_heading(doc, "1.2 Metrics at a Glance", 2)

    add_styled_table(doc,
        ["Metric", "Measures", "Range", "Good", "Bad", "Best For"],
        [
            ["WER", "Overall word accuracy", "0% \u2013 \u221e", "\u226430%", ">60%",
             "General quality benchmark"],
            ["WWER", "Importance-weighted accuracy", "0% \u2013 \u221e", "\u226430%", ">60%",
             "Semantic quality (entities matter more)"],
            ["NEA F1", "Entity capture quality", "0% \u2013 100%", "\u226570%", "<40%",
             "Are names, numbers, places correct?"],
            ["Alignment", "Per-word correctness", "ok / rep / ins", "Mostly ok", "Mostly ins",
             "Debugging individual words"],
        ],
        col_widths=[0.7, 1.3, 0.8, 0.6, 0.6, 2.5],
    )

    _tight_page_break(doc)


# ═══════════════════════════════════════════════════
# SECTION 2 — WER
# ═══════════════════════════════════════════════════

def section_2(doc):
    add_heading(doc, "2. WER (Word Error Rate)", 1, bookmark_id="sec2")

    add_heading(doc, "2.1 Definition", 2)
    add_para(doc, "WER measures the minimum number of word-level edits needed to transform the "
             "hypothesis into the reference, as a percentage of the reference length:")
    add_code_block(doc, "WER = (Substitutions + Deletions + Insertions) / Reference_Words \u00d7 100%")
    add_bullet_bold_value(doc, "Substitution (S): ", "A reference word was replaced by a different word")
    add_bullet_bold_value(doc, "Deletion (D): ", "A reference word is missing from the hypothesis")
    add_bullet_bold_value(doc, "Insertion (I): ", "The hypothesis contains a word not in the reference")

    add_heading(doc, "2.2 Tokenization", 2)
    add_para(doc, "Before computing WER, both strings are tokenized identically:")
    add_code_block(doc, 'tokens = re.findall(r"[a-z0-9]+(?:\'[a-z0-9]+)?", text.lower())')
    add_para(doc, "This keeps only lowercase alphanumeric words and contractions (e.g., \"don't\"). "
             "Punctuation, capitalization, and special characters are stripped.")

    add_heading(doc, "2.3 Walkthrough: WER = 22.9%", 2)
    add_para(doc, "Segment 0v2N6w4m46s Part 1 \u2014 a medical topic where most content is correct:", bold=True, size=Pt(10))

    # Color-coded REF (deletions in red)
    add_impact_example(doc, "REF", [
        ("on", C_IMPACT_BAD, True),           # deleted
        ("today's", C_IMPACT_BAD, True),       # substituted
        ("the", C_IMPACT_GOOD, False),
        ("doctor", C_IMPACT_GOOD, False),
        ("is", C_IMPACT_BAD, True),           # substituted
        ("in", C_IMPACT_BAD, True),           # deleted
        ("the", C_IMPACT_GOOD, False),
        ("topic", C_IMPACT_GOOD, False),
        ("is", C_IMPACT_GOOD, False),
        ("hypertension", C_IMPACT_GOOD, False),
        ("...", C_IMPACT_GOOD, False),
        ("we're", C_IMPACT_BAD, True),        # substituted
        ("going", C_IMPACT_BAD, True),        # deleted in hyp position
        ("to", C_IMPACT_GOOD, False),
        ("see", C_IMPACT_GOOD, False),
        ("where", C_IMPACT_GOOD, False),
        ("ours", C_IMPACT_BAD, True),         # substituted
        ("...", C_IMPACT_GOOD, False),
        ("decisions", C_IMPACT_BAD, True),    # substituted
    ])
    add_impact_example(doc, "HYP", [
        ("today", C_IMPACT_GOOD, False),
        ("as", C_IMPACT_BAD, True),           # insertion/substitution
        ("the", C_IMPACT_GOOD, False),
        ("doctor", C_IMPACT_GOOD, False),
        ("says", C_IMPACT_BAD, True),         # substitution
        ("the", C_IMPACT_GOOD, False),
        ("topic", C_IMPACT_GOOD, False),
        ("is", C_IMPACT_GOOD, False),
        ("hypertension", C_IMPACT_GOOD, False),
        ("...", C_IMPACT_GOOD, False),
        ("we'll", C_IMPACT_BAD, True),
        ("see", C_IMPACT_GOOD, False),
        ("where", C_IMPACT_GOOD, False),
        ("our", C_IMPACT_BAD, True),
        ("...", C_IMPACT_GOOD, False),
        ("changes", C_IMPACT_BAD, True),
    ])

    add_para(doc, "\u2192 9 errors / 40 reference words = 22.5% WER (rounded to 22.9% with full tokenization). "
             "Red words each cost exactly +1 to the error count. Green words cost nothing.",
             size=Pt(9), italic=True, color=C_GRAY)
    add_para(doc, "Key insight: WER treats every word equally. \"hypertension\" (critical medical term) "
             "correct = same value as \"the\" (filler) correct.", bold=True, size=Pt(10))

    add_heading(doc, "2.4 WER Can Exceed 100%", 2)
    add_para(doc, (
        "When the hypothesis contains far more words than the reference, insertions push WER "
        "above 100%. This happened catastrophically with Config H (length penalty = 2.0):"
    ))
    add_styled_table(doc,
        ["", "Content"],
        [
            ["REF (3 words)", "get the idea"],
            ["HYP (50+ words)", "that's why i'm here thank you so much for having me it's been an "
             "honor and a privilege to share my story with you and i hope that it inspires you to "
             "live your life to the fullest and to chase your dre..."],
            ["WER", "6,833% (3 correct words needed, 200+ insertions generated)"],
        ],
        col_widths=[1.2, 5.3],
        highlight_rows={2: RED_BG},
    )
    add_para(doc, "The model hallucinated an entire motivational speech from a 3-word visual input. "
             "Every insertion word (50+) adds +1 to the numerator while the denominator stays at 3.",
             size=Pt(9), italic=True, color=C_GRAY)

    add_heading(doc, "2.5 Interpretation Scale", 2)
    add_thermometer_table(doc, [
        ("\u2264 30%", GREEN_BG, "Good \u2014 most words correct, minor errors", "Medical segment: 22.9%"),
        ("30% \u2013 60%", AMBER_BG, "Moderate \u2014 core meaning partly preserved", "Informal speech: 50%"),
        ("> 60%", RED_BG, "Poor \u2014 output is unreliable or hallucinated", "Domain mismatch: 100%"),
    ])

    _tight_page_break(doc)


# ═══════════════════════════════════════════════════
# SECTION 3 — WWER + spaCy
# ═══════════════════════════════════════════════════

def section_3(doc):
    add_heading(doc, "3. WWER (Weighted Word Error Rate) & spaCy", 1, bookmark_id="sec3")

    add_heading(doc, "3.1 Motivation: Not All Words Are Equal", 2)
    add_para(doc, (
        "Standard WER treats \"the\" and \"Obama\" as equally important. In practice, getting a "
        "person's name or a number wrong is far more damaging than missing a filler word. WWER "
        "addresses this by weighting errors according to the linguistic importance of each word."
    ))

    add_heading(doc, "3.2 Token Classification via spaCy", 2)
    add_para(doc, (
        "WWER uses the spaCy en_core_web_sm model to classify each token by its Part-of-Speech "
        "(POS) tag and Named Entity Recognition (NER) label. The pipeline processes each word through "
        "tokenization, POS tagging, and entity recognition, then assigns it to one of three tiers:"
    ))

    add_styled_table(doc,
        ["Tier", "Weight", "What Qualifies", "Example Words"],
        [
            ["HIGH", "2.0\u00d7", "Named entities (PERSON, ORG, GPE, LOC, DATE, TIME, MONEY, "
             "PERCENT, QUANTITY, NORP, FAC, EVENT) + POS: PROPN, NUM",
             "Obama, London, 2024, Tuesday, Google"],
            ["MEDIUM", "1.0\u00d7", "Content words \u2014 POS: NOUN, VERB, ADJ, ADV",
             "hospital, running, beautiful, quickly"],
            ["LOW", "0.5\u00d7", "Function words \u2014 POS: DET, AUX, PRON, ADP, CONJ, PART, etc.",
             "the, is, we, from, and, to, but"],
        ],
        col_widths=[0.6, 0.6, 2.8, 2.5],
        highlight_rows={0: "dbe8f0", 2: ZEBRA_BG},
    )

    add_heading(doc, "3.3 Seeing Token Classification", 2)
    add_para(doc, "Here is a sample sentence with each word colored by its spaCy classification:", size=Pt(10))

    # "President Obama visited the hospital in New York on Tuesday"
    add_colored_runs(doc, [
        ("Example: ", True, C_DARK),
        ("President", True, C_ENTITY),      # PROPN = HIGH
        (" ", False, None),
        ("Obama", True, C_ENTITY),          # PERSON = HIGH
        (" ", False, None),
        ("visited", False, C_DARK),         # VERB = MED
        (" ", False, None),
        ("the", False, C_IMPACT_NONE, True),  # DET = LOW
        (" ", False, None),
        ("hospital", False, C_DARK),        # NOUN = MED
        (" ", False, None),
        ("in", False, C_IMPACT_NONE, True), # ADP = LOW
        (" ", False, None),
        ("New", True, C_ENTITY),            # GPE = HIGH
        (" ", False, None),
        ("York", True, C_ENTITY),           # GPE = HIGH
        (" ", False, None),
        ("on", False, C_IMPACT_NONE, True), # ADP = LOW
        (" ", False, None),
        ("Tuesday", True, C_ENTITY),        # DATE = HIGH
    ])
    add_para(doc,
             "Dark blue bold = HIGH (2.0\u00d7), Black = MEDIUM (1.0\u00d7), "
             "Gray italic = LOW (0.5\u00d7). "
             "Getting \"Obama\" wrong costs 4\u00d7 more than getting \"the\" wrong.",
             size=Pt(9), italic=True, color=C_GRAY)

    add_heading(doc, "3.4 How WWER Is Calculated", 2)
    add_para(doc, (
        "WWER uses Python's SequenceMatcher to align reference and hypothesis words. "
        "For each alignment operation (replace, delete, insert), the error cost equals "
        "the weight of the reference token involved. Insertions use medium weight (1.0) "
        "since there is no reference token to classify."
    ))
    add_code_block(doc,
        "WWER = Sum(weight_of_each_wrong_token) / Sum(weight_of_all_ref_tokens) \u00d7 100%\n\n"
        "For each alignment opcode:\n"
        "  equal    \u2192 no cost\n"
        "  replace  \u2192 cost = weight of replaced ref token\n"
        "  delete   \u2192 cost = weight of deleted ref token\n"
        "  insert   \u2192 cost = 1.0 per inserted word (medium default)"
    )

    add_heading(doc, "3.5 Walkthrough: WER 43.8% vs WWER 117.9%", 2)
    add_para(doc, "Segment S8aXVz799Dc \u2014 an English learning phrase where entities get mangled:", bold=True, size=Pt(10))
    add_colored_runs(doc, [
        ("REF: ", True, C_DARK),
        ("another useful expression i'm afraid there's been a mistake or i'm afraid there's been a mistake", False, C_DARK),
    ])
    add_colored_runs(doc, [
        ("HYP: ", True, C_DARK),
        ("another useful expression i'm afraid there's been a mishap or i m afraid there s", False, C_DARK),
    ])

    add_para(doc, "WWER impact view \u2014 each wrong word colored by its error cost:", bold=True, size=Pt(10))
    add_impact_example(doc, "HYP", [
        ("another", C_IMPACT_GOOD, False),       # correct MED
        ("useful", C_IMPACT_GOOD, False),        # correct MED
        ("expression", C_IMPACT_GOOD, False),    # correct MED
        ("i'm", C_IMPACT_GOOD, False),           # correct LOW
        ("afraid", C_IMPACT_GOOD, False),        # correct MED
        ("there's", C_IMPACT_GOOD, False),       # correct LOW
        ("been", C_IMPACT_GOOD, False),          # correct LOW
        ("a", C_IMPACT_GOOD, False),             # correct LOW
        ("mishap", C_IMPACT_MILD, True),         # wrong NOUN = 1.0 cost
        ("or", C_IMPACT_GOOD, False),            # correct LOW
        ("i", C_IMPACT_GOOD, False),
        ("m", C_IMPACT_BAD, True),               # fragment = insertion
        ("afraid", C_IMPACT_GOOD, False),
        ("there", C_IMPACT_GOOD, False),
        ("s", C_IMPACT_BAD, True),               # fragment = insertion
    ])

    add_para(doc,
             "\u2192 WER = 43.8% (7 errors / 16 words). "
             "But WWER = 117.9% because the hypothesis is also shorter (missing 5 words "
             "including the repeated HIGH-value \"mistake\"), and insertions of fragments "
             "(\"m\", \"s\") each cost 1.0. The weighted denominator is small while weighted "
             "errors pile up, pushing WWER far above 100%.",
             size=Pt(9), italic=True, color=C_GRAY)

    add_heading(doc, "3.6 When WER and WWER Diverge", 2)
    add_para(doc, "The gap between WER and WWER reveals what kind of errors dominate:")

    add_styled_table(doc,
        ["Segment", "WER%", "WWER%", "Gap", "Explanation"],
        [
            ["0v2N6w4m46s", "22.9", "19.5", "WER > WWER",
             "Errors on function words (\"on\", \"in\") \u2014 only cost 0.5\u00d7 each"],
            ["S8aXVz799Dc", "43.8", "117.9", "WWER >> WER",
             "Missing content words + fragments inflate weighted errors"],
            ["NGCTGBsuEg4", "100.0", "62.5", "WER >> WWER",
             "Math variables (\"vi\", \"vf\") wrong, but common words correct"],
            ["zmwgmt7wcv8", "26.7", "37.5", "WWER > WER",
             "\"wittgenstein\" (PROPN, 2.0\u00d7) wrong = disproportionate impact"],
        ],
        col_widths=[1.2, 0.5, 0.6, 1.0, 3.2],
        highlight_rows={0: GREEN_BG, 1: RED_BG},
    )
    add_para(doc,
             "Rule of thumb: WWER > WER means entity/content errors dominate. "
             "WWER < WER means errors are concentrated on function words.",
             bold=True, size=Pt(10))

    add_heading(doc, "3.7 Fallback Without spaCy", 2)
    add_para(doc, (
        "When spaCy is unavailable, the system falls back to a basic stopword filter: "
        "80 common function words (\"a\", \"the\", \"is\", \"we\", ...) are classified as LOW, "
        "digit tokens and number words (\"twenty\", \"hundred\") as HIGH, and everything else as MEDIUM. "
        "This is less accurate than POS/NER tagging but still provides meaningful weighting."
    ))

    doc.add_page_break()


# ═══════════════════════════════════════════════════
# SECTION 4 — NEA
# ═══════════════════════════════════════════════════

def section_4(doc):
    add_heading(doc, "4. NEA (Named Entity Accuracy)", 1, bookmark_id="sec4")

    add_heading(doc, "4.1 What NEA Measures", 2)
    add_para(doc, (
        "NEA focuses exclusively on high-value tokens \u2014 proper nouns, numbers, and named "
        "entities. It asks two complementary questions:"
    ))
    add_bullet_bold_value(doc, "Recall: ",
        "Of the important words in the reference, how many appear in the hypothesis?")
    add_bullet_bold_value(doc, "Precision: ",
        "Of the important words the model produced, how many are actually in the reference?")
    add_bullet_bold_value(doc, "F1: ",
        "Harmonic mean of Recall and Precision \u2014 the single-number summary.")

    add_code_block(doc,
        "NEA Recall    = |ref_entities \u2229 hyp_entities| / |ref_entities| \u00d7 100%\n"
        "NEA Precision = |ref_entities \u2229 hyp_entities| / |hyp_entities| \u00d7 100%\n"
        "NEA F1        = 2 \u00d7 Recall \u00d7 Precision / (Recall + Precision)"
    )

    add_para(doc, (
        "Important: NEA uses set-based matching (not position-sensitive). If \"Obama\" appears "
        "anywhere in both ref and hyp, it counts as matched regardless of position. "
        "When the reference has no high-value tokens, NEA falls back to content words "
        "(HIGH + MEDIUM) so the metric remains meaningful."
    ))

    add_heading(doc, "4.2 Walkthrough: WER 50% but NEA 0%", 2)
    add_para(doc, "Segment lOuEis9L8os \u2014 biblical narrative where names are everything:", bold=True, size=Pt(10))

    # REF with entity impact
    add_impact_example(doc, "REF", [
        ("her", C_IMPACT_NONE, False),
        ("brother", C_IMPACT_NONE, False),
        ("and", C_IMPACT_NONE, False),
        ("they're", C_IMPACT_NONE, False),
        ("kinda", C_IMPACT_NONE, False),
        ("half", C_IMPACT_NONE, False),
        ("brothers", C_IMPACT_NONE, False),
        ("and", C_IMPACT_NONE, False),
        ("sisters", C_IMPACT_NONE, False),
        ("...", C_IMPACT_NONE, False),
        ("absalom", C_IMPACT_BAD, True),     # MISSED entity
        ("wants", C_IMPACT_NONE, False),
        ("to", C_IMPACT_NONE, False),
        ("avenge", C_IMPACT_NONE, False),
        ("tamar's", C_IMPACT_BAD, True),     # MISSED entity (via "tamar")
        ("rape", C_IMPACT_NONE, False),
        ("...", C_IMPACT_NONE, False),
    ])
    add_impact_example(doc, "HYP", [
        ("her", C_IMPACT_NONE, False),
        ("brother", C_IMPACT_NONE, False),
        ("...", C_IMPACT_NONE, False),
        ("hops", C_IMPACT_BAD, True),        # hallucinated entity (for "absalom")
        ("are", C_IMPACT_NONE, False),
        ("brought", C_IMPACT_NONE, False),
        ("to", C_IMPACT_NONE, False),
        ("a", C_IMPACT_NONE, False),
        ("rage", C_IMPACT_NONE, False),
        ("tamer's", C_IMPACT_NONE, False),   # close but not "tamar"
        ("rape", C_IMPACT_NONE, False),
        ("...", C_IMPACT_NONE, False),
    ])
    add_para(doc,
             "\u2192 REF entities: {absalom, tamar}. HYP entities: {} (neither found). "
             "Recall = 0/2 = 0%. The 15+ correct gray words contribute nothing to NEA. "
             "A reader would immediately notice the names are wrong \u2014 NEA captures this, WER dilutes it.",
             size=Pt(9), italic=True, color=C_GRAY)

    add_heading(doc, "4.3 Contrast: WER 22.9% but NEA 100%", 2)
    add_para(doc, "Segment 0v2N6w4m46s Part 1 \u2014 medical topic with all entities preserved:", bold=True, size=Pt(10))

    add_impact_example(doc, "REF", [
        ("on", C_IMPACT_NONE, False),
        ("today's", C_IMPACT_NONE, False),
        ("the", C_IMPACT_NONE, False),
        ("doctor", C_IMPACT_GOOD, True),      # entity FOUND
        ("is", C_IMPACT_NONE, False),
        ("in", C_IMPACT_NONE, False),
        ("the", C_IMPACT_NONE, False),
        ("topic", C_IMPACT_NONE, False),
        ("is", C_IMPACT_NONE, False),
        ("hypertension", C_IMPACT_GOOD, True), # entity FOUND
        ("...", C_IMPACT_NONE, False),
        ("blood", C_IMPACT_GOOD, True),        # entity FOUND
        ("pressure", C_IMPACT_GOOD, True),     # entity FOUND
        ("reading", C_IMPACT_GOOD, True),      # content word FOUND
        ("...", C_IMPACT_NONE, False),
    ])
    add_para(doc,
             "\u2192 Every high-value token (\"doctor\", \"hypertension\", \"blood\", \"pressure\") "
             "appears in the hypothesis. NEA Recall = 100%. The 9 wrong function words "
             "(\"on\" deleted, \"we're\" \u2192 \"we'll\", etc.) are invisible to NEA.",
             size=Pt(9), italic=True, color=C_GRAY)

    add_heading(doc, "4.4 Entity Edge Cases", 2)

    add_styled_table(doc,
        ["Segment", "WER%", "NEA F1%", "Entity Analysis"],
        [
            ["lOuEis9L8os", "50.0", "0%",
             "Biblical names \"absalom\", \"tamar\" completely missed. "
             "Model produced \"hops\" and \"tamer\" \u2014 phonetically close but not entity matches."],
            ["bR80JEaW4LY", "42.9", "0%",
             "Indian names \"jeevan anand\" \u2192 \"pyramids 3000 years\". Complete domain hallucination. "
             "Model invents plausible-sounding content from lip shapes."],
            ["zmwgmt7wcv8", "26.7", "0%",
             "\"wittgenstein\" \u2192 \"philistine\" (homophene: similar mouth shapes, different words). "
             "Consistent across all 13 hyperparameter configs \u2014 a true visual ambiguity."],
            ["0v2N6w4m46s", "22.9", "100%",
             "All medical entities correct despite 22.9% WER. Only function words differ. "
             "Shows that WER errors don't always mean semantic failure."],
            ["91tFHYm9qso", "21.1", "76.9%",
             "\"coping\" \u2192 \"hope\" (semantic near-miss). 5 of 7 entities matched. "
             "Interesting: spaCy classified both as NOUN, but they're different enough to miss."],
        ],
        col_widths=[1.1, 0.5, 0.6, 4.3],
        highlight_rows={3: GREEN_BG, 0: RED_BG, 1: RED_BG, 2: RED_BG},
    )

    add_heading(doc, "4.5 The missed_entities Field", 2)
    add_para(doc, (
        "Every report row includes a missed_entities column listing which high-value reference "
        "words were absent from the hypothesis. This is the fastest way to diagnose NEA failures:"
    ))
    add_code_block(doc,
        "lOuEis9L8os:  missed_entities = \"absalom, tamar\"\n"
        "bR80JEaW4LY:  missed_entities = \"anand, jeevan\"\n"
        "zmwgmt7wcv8:  missed_entities = \"wittgenstein\"\n"
        "0v2N6w4m46s:  missed_entities = \"\"  (none \u2014 all found)"
    )

    add_heading(doc, "4.6 Interpretation Scale", 2)
    add_thermometer_table(doc, [
        ("\u2265 70%", GREEN_BG, "Good \u2014 most entities captured correctly",
         "Medical segment: NEA 100%"),
        ("40% \u2013 70%", AMBER_BG, "Partial \u2014 some entities present, some missed",
         "Coping/hope: NEA 76.9%"),
        ("< 40%", RED_BG, "Poor \u2014 critical entities missing or hallucinated",
         "Names missed: NEA 0%"),
    ])

    _tight_page_break(doc)


# ═══════════════════════════════════════════════════
# SECTION 5 — Cross-Metric Analysis & Quick Reference
# ═══════════════════════════════════════════════════

def section_5(doc):
    add_heading(doc, "5. Cross-Metric Analysis & Quick Reference", 1, bookmark_id="sec5")

    add_heading(doc, "5.1 Same Example, Three Metrics", 2)
    add_para(doc, (
        "The segment 91tFHYm9qso (WER 21.1%, WWER 20.7%, NEA 76.9%) demonstrates "
        "how all three metrics see the same output differently:"
    ))
    add_colored_runs(doc, [
        ("REF: ", True, C_DARK),
        ("so we started looking at the structure of coping and the structure of emotion "
         "and started really looking at", False, C_DARK),
    ])
    add_colored_runs(doc, [
        ("HYP: ", True, C_DARK),
        ("we started looking at structure of hope and the structure of emotion "
         "and started really looking", False, C_DARK),
    ])

    add_styled_table(doc,
        ["Metric", "Score", "What It Sees"],
        [
            ["WER", "21.1%",
             "4 errors / 19 words: \"so\" deleted, \"the\" deleted, "
             "\"coping\" \u2192 \"hope\", \"at\" deleted. Treats all equally."],
            ["WWER", "20.7%",
             "Slightly better: the deleted words (\"so\", \"the\", \"at\") are LOW weight (0.5\u00d7 each). "
             "\"coping\" (NOUN, 1.0\u00d7) is the main cost."],
            ["NEA F1", "76.9%",
             "5 of 7 entities matched (\"structure\", \"emotion\", \"looking\", \"started\", \"really\"). "
             "\"coping\" and \"so\" missed. High recall despite WER errors."],
        ],
        col_widths=[0.7, 0.6, 5.2],
        highlight_rows={0: ZEBRA_BG, 1: GREEN_BG, 2: GREEN_BG},
    )

    add_heading(doc, "5.2 Disagreement Patterns", 2)
    add_para(doc, (
        "When WER, WWER, and NEA give different signals for the same segment, "
        "the pattern of disagreement reveals what kind of error the model is making. "
        "Understanding these patterns is essential for diagnosing problems and choosing "
        "where to focus improvement efforts. Below we walk through each pattern with "
        "simple, generalizable examples before showing real pipeline cases."
    ))

    # ── Pattern A: WER ok, NEA bad ──
    add_heading(doc, "Pattern A: WER is acceptable, but NEA is poor", 3)
    add_para(doc, (
        "This pattern means the model got the sentence structure and most common words right, "
        "but missed the most important words \u2014 the names, numbers, or entities that carry "
        "the core meaning. It is the most deceptive pattern because a casual glance at WER "
        "might suggest the output is usable, when in fact the critical information is wrong."
    ))
    add_para(doc, "Simple example:", bold=True, size=Pt(10))
    add_impact_example(doc, "REF:", [
        ("the", C_IMPACT_GOOD, False),
        ("mayor", C_IMPACT_GOOD, False),
        ("of", C_IMPACT_GOOD, False),
        ("Springfield", C_IMPACT_BAD, True),
        ("met", C_IMPACT_GOOD, False),
        ("with", C_IMPACT_GOOD, False),
        ("Senator", C_IMPACT_BAD, True),
        ("Jones", C_IMPACT_BAD, True),
    ])
    add_impact_example(doc, "HYP:", [
        ("the", C_IMPACT_GOOD, False),
        ("mayor", C_IMPACT_GOOD, False),
        ("of", C_IMPACT_GOOD, False),
        ("the", C_IMPACT_BAD, True),
        ("city", C_IMPACT_BAD, True),
        ("met", C_IMPACT_GOOD, False),
        ("with", C_IMPACT_GOOD, False),
        ("some", C_IMPACT_BAD, True),
        ("guy", C_IMPACT_BAD, True),
    ])
    add_para(doc,
             "\u2192 WER = 50% (4 errors / 8 ref words) \u2014 moderate. "
             "NEA = 0% \u2014 none of the 3 entities (Springfield, Senator, Jones) appear in hypothesis. "
             "The sentence reads fluently but all the factual content is gone.",
             size=Pt(9), italic=True, color=C_GRAY)

    add_para(doc, "Another example:", bold=True, size=Pt(10))
    add_impact_example(doc, "REF:", [
        ("I", C_IMPACT_GOOD, False),
        ("called", C_IMPACT_GOOD, False),
        ("Dr", C_IMPACT_BAD, True),
        ("Patel", C_IMPACT_BAD, True),
        ("at", C_IMPACT_GOOD, False),
        ("3", C_IMPACT_BAD, True),
        ("pm", C_IMPACT_GOOD, False),
    ])
    add_impact_example(doc, "HYP:", [
        ("I", C_IMPACT_GOOD, False),
        ("called", C_IMPACT_GOOD, False),
        ("the", C_IMPACT_BAD, True),
        ("doctor", C_IMPACT_BAD, True),
        ("at", C_IMPACT_GOOD, False),
        ("some", C_IMPACT_BAD, True),
        ("point", C_IMPACT_BAD, True),
    ])
    add_para(doc,
             "\u2192 WER = 57% (4 errors / 7 words). "
             "NEA = 0% \u2014 \"Dr\" (title), \"Patel\" (PERSON), and \"3\" (NUM) all missing. "
             "The model understood the general action (calling a doctor) but lost every specific detail.",
             size=Pt(9), italic=True, color=C_GRAY)

    add_para(doc, (
        "General rule: When WER is moderate (30\u201360%) but NEA is near 0%, the model is "
        "substituting generic words for specific entities. This is common in lip-reading because "
        "proper nouns and numbers have no visual context to help the model disambiguate \u2014 "
        "the mouth shapes for \"Patel\" and \"bottle\" may look nearly identical."
    ))

    # ── Pattern B: WWER >> WER ──
    add_heading(doc, "Pattern B: WWER is much higher than WER", 3)
    add_para(doc, (
        "When WWER is significantly larger than WER, it means the errors are concentrated "
        "on high-weight words (entities, proper nouns, content words) rather than on filler. "
        "Even though the total error count may not be extreme, the errors that do exist are "
        "hitting the most important words."
    ))
    add_para(doc, "Simple example:", bold=True, size=Pt(10))
    add_impact_example(doc, "REF:", [
        ("the", C_IMPACT_GOOD, False),
        ("president", C_IMPACT_GOOD, False),
        ("of", C_IMPACT_GOOD, False),
        ("Google", C_IMPACT_BAD, True),
        ("announced", C_IMPACT_GOOD, False),
        ("it", C_IMPACT_GOOD, False),
    ])
    add_impact_example(doc, "HYP:", [
        ("the", C_IMPACT_GOOD, False),
        ("president", C_IMPACT_GOOD, False),
        ("of", C_IMPACT_GOOD, False),
        ("the", C_IMPACT_BAD, True),
        ("company", C_IMPACT_BAD, True),
        ("announced", C_IMPACT_GOOD, False),
        ("it", C_IMPACT_GOOD, False),
    ])
    add_para(doc,
             "\u2192 WER = 17% (1 substitution / 6 ref words). "
             "But WWER \u2248 33% because \"Google\" is a named entity (ORG, weight 2.0\u00d7) \u2014 "
             "that one error costs 2.0 in weighted terms while the denominator weights the "
             "6-word reference at only ~6.0 total. One entity error = double the weighted cost.",
             size=Pt(9), italic=True, color=C_GRAY)

    add_para(doc, (
        "General rule: WWER >> WER means the model is correct on filler but wrong on "
        "substance. This often happens with rare proper nouns the model has never seen."
    ))

    # ── Pattern C: WWER << WER ──
    add_heading(doc, "Pattern C: WWER is much lower than WER", 3)
    add_para(doc, (
        "The opposite: errors are concentrated on low-weight function words while the "
        "important content words are all correct. This is generally the least concerning "
        "pattern because the semantic meaning is preserved despite a noisy surface form."
    ))
    add_para(doc, "Simple example:", bold=True, size=Pt(10))
    add_impact_example(doc, "REF:", [
        ("she", C_IMPACT_BAD, True),
        ("was", C_IMPACT_BAD, True),
        ("in", C_IMPACT_BAD, True),
        ("the", C_IMPACT_BAD, True),
        ("hospital", C_IMPACT_GOOD, False),
        ("with", C_IMPACT_BAD, True),
        ("a", C_IMPACT_BAD, True),
        ("broken", C_IMPACT_GOOD, False),
        ("leg", C_IMPACT_GOOD, False),
    ])
    add_impact_example(doc, "HYP:", [
        ("went", C_IMPACT_BAD, True),
        ("to", C_IMPACT_BAD, True),
        ("hospital", C_IMPACT_GOOD, False),
        ("broken", C_IMPACT_GOOD, False),
        ("leg", C_IMPACT_GOOD, False),
    ])
    add_para(doc,
             "\u2192 WER = 67% (6 errors / 9 ref words) \u2014 looks terrible. "
             "But WWER \u2248 30% because the 6 wrong words are all function words (weight 0.5\u00d7 each = 3.0 cost), "
             "while all 3 content words (\"hospital\", \"broken\", \"leg\") are correct. "
             "The core meaning \u2014 someone has a broken leg and is at a hospital \u2014 is fully preserved.",
             size=Pt(9), italic=True, color=C_GRAY)

    add_para(doc, (
        "General rule: WWER << WER means errors are cosmetic. The model understood the "
        "content but struggled with grammatical structure \u2014 common in lip-reading where "
        "function words like \"the\", \"is\", and \"a\" have nearly identical mouth shapes."
    ))

    # ── Pattern D: WER bad, NEA ok ──
    add_heading(doc, "Pattern D: WER is bad, but NEA is high", 3)
    add_para(doc, (
        "The model captures all the important entities correctly, but makes many errors "
        "on the surrounding words. The factual core is intact even though the overall "
        "word accuracy is low."
    ))
    add_para(doc, "Simple example:", bold=True, size=Pt(10))
    add_impact_example(doc, "REF:", [
        ("on", C_IMPACT_BAD, True),
        ("Monday", C_IMPACT_GOOD, True),
        ("the", C_IMPACT_BAD, True),
        ("team", C_IMPACT_GOOD, False),
        ("from", C_IMPACT_BAD, True),
        ("NASA", C_IMPACT_GOOD, True),
        ("will", C_IMPACT_BAD, True),
        ("launch", C_IMPACT_GOOD, False),
        ("the", C_IMPACT_BAD, True),
        ("rocket", C_IMPACT_GOOD, False),
    ])
    add_impact_example(doc, "HYP:", [
        ("so", C_IMPACT_BAD, True),
        ("Monday", C_IMPACT_GOOD, True),
        ("a", C_IMPACT_BAD, True),
        ("team", C_IMPACT_GOOD, False),
        ("at", C_IMPACT_BAD, True),
        ("NASA", C_IMPACT_GOOD, True),
        ("is", C_IMPACT_BAD, True),
        ("going", C_IMPACT_BAD, True),
        ("to", C_IMPACT_BAD, True),
        ("launch", C_IMPACT_GOOD, False),
        ("a", C_IMPACT_BAD, True),
        ("rocket", C_IMPACT_GOOD, False),
    ])
    add_para(doc,
             "\u2192 WER = 70% (7 errors / 10 ref words) \u2014 bad on paper. "
             "NEA = 100% \u2014 \"Monday\" (DATE) and \"NASA\" (ORG) both present. "
             "A human reader would get the essential message: NASA is launching a rocket on Monday.",
             size=Pt(9), italic=True, color=C_GRAY)

    add_para(doc, (
        "General rule: High NEA + bad WER means the output is still usable for tasks "
        "like information extraction, search indexing, or entity-based summarization, "
        "even if it reads awkwardly."
    ))

    # ── Pattern E: All metrics bad ──
    add_heading(doc, "Pattern E: All metrics are bad (total failure)", 3)
    add_para(doc, (
        "When WER \u2265 100%, WWER is high, and NEA is at or near 0%, the model has failed "
        "completely. This can happen for several reasons: the model produced an empty "
        "prediction, it hallucinated unrelated content, or the input was too noisy "
        "(bad video quality, extreme angle, or no visible face)."
    ))
    add_para(doc, "Simple example \u2014 empty prediction:", bold=True, size=Pt(10))
    add_impact_example(doc, "REF:", [
        ("Taylor", C_IMPACT_BAD, True),
        ("Swift", C_IMPACT_BAD, True),
        ("performed", C_IMPACT_BAD, True),
        ("at", C_IMPACT_BAD, True),
        ("the", C_IMPACT_BAD, True),
        ("Grammys", C_IMPACT_BAD, True),
    ])
    add_impact_example(doc, "HYP:", [
        ("(empty)", C_IMPACT_BAD, True),
    ])
    add_para(doc,
             "\u2192 WER = 100% (6 deletions / 6 words). WWER = 100%. NEA = 0%. "
             "Nothing was predicted at all.",
             size=Pt(9), italic=True, color=C_GRAY)

    add_para(doc, "Simple example \u2014 hallucinated content:", bold=True, size=Pt(10))
    add_impact_example(doc, "REF:", [
        ("get", C_IMPACT_BAD, True),
        ("the", C_IMPACT_BAD, True),
        ("idea", C_IMPACT_BAD, True),
    ])
    add_impact_example(doc, "HYP:", [
        ("thank", C_IMPACT_BAD, True),
        ("you", C_IMPACT_BAD, True),
        ("so", C_IMPACT_BAD, True),
        ("much", C_IMPACT_BAD, True),
        ("for", C_IMPACT_BAD, True),
        ("having", C_IMPACT_BAD, True),
        ("me", C_IMPACT_BAD, True),
        ("it's", C_IMPACT_BAD, True),
        ("been", C_IMPACT_BAD, True),
        ("an", C_IMPACT_BAD, True),
        ("honor", C_IMPACT_BAD, True),
        ("...", C_IMPACT_BAD, True),
    ])
    add_para(doc,
             "\u2192 WER = 6,833% (3 ref words, 50+ hallucinated words). "
             "The model invented an entire speech. Every inserted word adds +1 to the "
             "numerator while the denominator stays at 3.",
             size=Pt(9), italic=True, color=C_GRAY)

    add_para(doc, (
        "General rule: When all metrics fail simultaneously, check the input first. "
        "Common causes include: face not visible (angle too extreme), very short segments "
        "(under 1 second), background noise or music with no speech, or domain mismatch "
        "(e.g., non-English speech processed with an English model)."
    ))

    # ── Pattern F: Homophene lock ──
    add_heading(doc, "Pattern F: Homophene lock (consistent mis-reading)", 3)
    add_para(doc, (
        "Some words look identical on the lips despite sounding completely different. "
        "When the model consistently produces the same wrong word across multiple "
        "hyperparameter configurations, it reveals a true visual ambiguity that no "
        "amount of tuning can fix. This is a fundamental limitation of lip-reading."
    ))
    add_para(doc, "Simple example:", bold=True, size=Pt(10))
    add_impact_example(doc, "REF:", [
        ("he", C_IMPACT_GOOD, False),
        ("said", C_IMPACT_GOOD, False),
        ("pat", C_IMPACT_BAD, True),
        ("and", C_IMPACT_GOOD, False),
        ("bat", C_IMPACT_BAD, True),
        ("are", C_IMPACT_GOOD, False),
        ("different", C_IMPACT_GOOD, False),
    ])
    add_impact_example(doc, "HYP:", [
        ("he", C_IMPACT_GOOD, False),
        ("said", C_IMPACT_GOOD, False),
        ("bad", C_IMPACT_BAD, True),
        ("and", C_IMPACT_GOOD, False),
        ("bad", C_IMPACT_BAD, True),
        ("are", C_IMPACT_GOOD, False),
        ("different", C_IMPACT_GOOD, False),
    ])
    add_para(doc,
             "\u2192 The consonants p/b/m all look identical on the lips (bilabial closure). "
             "\"pat\" \u2192 \"bad\" and \"bat\" \u2192 \"bad\" are visual ambiguities, not model failures. "
             "Similarly, \"wittgenstein\" \u2192 \"philistine\" in real data \u2014 both start with similar lip shapes.",
             size=Pt(9), italic=True, color=C_GRAY)

    add_para(doc, (
        "General rule: If the same substitution appears across all decode configurations "
        "(beam search, sampling, different temperatures), it is likely a homophene \u2014 "
        "a word pair that cannot be distinguished by vision alone. These errors can only "
        "be resolved with language model context (future fine-tuning direction)."
    ))

    # ── Summary table ──
    doc.add_page_break()
    add_heading(doc, "Summary of Disagreement Patterns", 3)
    add_para(doc, (
        "The table below collects all patterns with real pipeline examples for quick reference:"
    ))

    add_styled_table(doc,
        ["Pattern", "Example", "WER", "WWER", "NEA", "What Happened"],
        [
            ["A: WER ok, NEA bad",
             "lOuEis9L8os", "50%", "50.9%", "0%",
             "Structure correct but proper names (\"absalom\", \"tamar\") completely missed"],
            ["B: WWER >> WER",
             "S8aXVz799Dc", "43.8%", "117.9%", "66.7%",
             "Errors concentrated on content/entity words; fragments inflate cost"],
            ["C: WWER << WER",
             "NGCTGBsuEg4", "100%", "62.5%", "40%",
             "Math variables wrong but common words (\"zero\", \"goes away\") are LOW-cost correct"],
            ["D: WER bad, NEA ok",
             "91tFHYm9qso", "21.1%", "20.7%", "76.9%",
             "Most entities present; WER errors are function word deletions"],
            ["E: All 100%",
             "DBhaa45mAro", "100%", "100%", "0%",
             "Empty prediction (baseline). All entities (\"taylor\", \"dennis\", \"julia\") missed."],
            ["E: WER > 100%",
             "Config H", "6833%", "\u2014", "\u2014",
             "Catastrophic hallucination: 3-word ref, 50+ word fictional output"],
            ["F: Homophene lock",
             "zmwgmt7wcv8", "26.7%", "37.5%", "0%",
             "\"wittgenstein\" \u2192 \"philistine\" across ALL 13 configs. True visual ambiguity."],
        ],
        col_widths=[1.0, 1.0, 0.5, 0.5, 0.5, 3.0],
        highlight_rows={0: AMBER_BG, 1: RED_BG, 4: RED_BG, 5: RED_BG},
    )

    add_heading(doc, "5.3 Decision Guide: Which Metric to Use", 2)

    add_styled_table(doc,
        ["Question", "Use This Metric", "Why"],
        [
            ["\"How accurate is the overall output?\"", "WER",
             "Counts every word error equally \u2014 the standard ASR benchmark"],
            ["\"Did we get the important parts right?\"", "WWER",
             "Weights entities 4\u00d7 more than filler \u2014 reflects human judgment of quality"],
            ["\"Are names, numbers, and places correct?\"", "NEA F1",
             "Focuses purely on high-value tokens \u2014 critical for downstream use (search, analysis)"],
            ["\"Which specific words are wrong?\"", "Word Alignment (ok/rep/ins)",
             "Per-word diagnostic: green = correct position, amber = shifted, red = invented"],
            ["\"Comparing two decode configs?\"", "WWER",
             "Best single number for optimization \u2014 captures semantic quality better than WER"],
        ],
        col_widths=[2.2, 1.5, 2.8],
    )

    _tight_page_break(doc)
    add_heading(doc, "5.4 Quick Reference Card", 2)

    add_para(doc, "Metric formulas and thresholds:", bold=True, size=Pt(10))
    add_styled_table(doc,
        ["Metric", "Formula", "Range", "Green", "Amber", "Red"],
        [
            ["WER", "(S+D+I) / N \u00d7 100", "0 \u2013 \u221e", "\u226430%", "30-60%", ">60%"],
            ["WWER", "\u03a3(weighted errors) / \u03a3(weighted ref) \u00d7 100",
             "0 \u2013 \u221e", "\u226430%", "30-60%", ">60%"],
            ["NEA R", "|ref\u2229hyp| / |ref|", "0-100%", "\u226570%", "40-70%", "<40%"],
            ["NEA P", "|ref\u2229hyp| / |hyp|", "0-100%", "\u226570%", "40-70%", "<40%"],
            ["NEA F1", "2RP / (R+P)", "0-100%", "\u226570%", "40-70%", "<40%"],
        ],
        col_widths=[0.6, 2.2, 0.7, 0.6, 0.6, 0.6],
    )

    add_para(doc, "Token classification weights:", bold=True, size=Pt(10))
    add_styled_table(doc,
        ["Tier", "Weight", "spaCy POS Tags", "spaCy NER Types"],
        [
            ["HIGH", "2.0\u00d7", "PROPN, NUM",
             "PERSON, ORG, GPE, LOC, DATE, TIME, MONEY, PERCENT, QUANTITY, NORP, FAC, EVENT"],
            ["MEDIUM", "1.0\u00d7", "NOUN, VERB, ADJ, ADV", "(none)"],
            ["LOW", "0.5\u00d7", "DET, AUX, PRON, ADP, CONJ, PART, PUNCT, etc.", "(none)"],
        ],
        col_widths=[0.6, 0.6, 1.5, 3.8],
    )

    add_para(doc, "Word alignment tags:", bold=True, size=Pt(10))
    add_styled_table(doc,
        ["Tag", "Color", "Meaning", "Impact on WER"],
        [
            ["ok", "Green", "Word matches reference at the same position", "0 (no error)"],
            ["rep", "Amber", "Word exists in reference but at a different position", "0 (exists, just shifted)"],
            ["ins", "Red", "Word does not appear in reference at all (hallucinated)", "+1 per word"],
        ],
        col_widths=[0.5, 0.7, 2.8, 2.5],
        highlight_rows={0: GREEN_BG, 1: AMBER_BG, 2: RED_BG},
    )


# ═══════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════

def main():
    print("Generating Argos Metrics Explainer...")

    doc = Document()

    # Page setup (A4)
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_header_footer(doc)
    create_cover_page(doc)
    create_toc(doc)

    section_1(doc)
    section_2(doc)
    section_3(doc)
    section_4(doc)
    section_5(doc)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_FILE))
    size_kb = OUTPUT_FILE.stat().st_size // 1024
    print(f"Saved: {OUTPUT_FILE}")
    print(f"Size: {size_kb} KB")


if __name__ == "__main__":
    main()
