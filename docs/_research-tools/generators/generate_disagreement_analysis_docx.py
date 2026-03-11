#!/usr/bin/env python3
"""
Argos -- IS vs Opus Judge Disagreement Analysis (Word Document)

Summarises the phenomenon of disagreements between the Intelligibility Score (IS)
metric and the Claude Opus LLM-as-a-Judge verdicts, for both blind and
context-aware evaluation modes. Includes curated examples, root-cause
explanations, and mitigation strategies.

Usage:
    python3 generate_disagreement_analysis_docx.py

Output:
    /home/ubuntu/docs/evaluation/llm_judge/disagreement_analysis.docx
"""

import csv
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

# -- Paths --
SCRIPT_DIR = Path(__file__).parent
ASSETS_DIR = SCRIPT_DIR.parent / "assets"
LOGO_ORCHARD = ASSETS_DIR / "logo.png"
LOGO_PEACOCK = ASSETS_DIR / "peacock.png"

CSV_FILE = Path("/home/ubuntu/docs/evaluation/llm_judge/context_eval/context_eval_results.csv")
OUTPUT_FILE = Path("/home/ubuntu/docs/evaluation/llm_judge/disagreement_analysis.docx")

# -- Colors --
C_PRIMARY = RGBColor(0x1A, 0x3A, 0x5C)
C_H2 = RGBColor(0x2A, 0x5A, 0x8C)
C_H3 = RGBColor(0x3A, 0x6A, 0x9C)
C_H4 = RGBColor(0x4A, 0x7A, 0xAC)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_DARK = RGBColor(0x33, 0x33, 0x33)
C_GREEN = RGBColor(0x15, 0x57, 0x24)
C_RED = RGBColor(0x72, 0x1C, 0x24)
C_AMBER = RGBColor(0x85, 0x64, 0x04)
C_GRAY = RGBColor(0x66, 0x66, 0x66)

HEADER_BG = "1a3a5c"
ZEBRA_BG = "f0f4f8"
GREEN_BG = "d4edda"
RED_BG = "f8d7da"
AMBER_BG = "fff3cd"

LAST_UPDATED = datetime.now().strftime("%B %d, %Y")


# =====================================================================
# HELPER FUNCTIONS
# =====================================================================

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, bold=False, color=None, size=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = size
    run.font.name = "Calibri"


def add_styled_table(doc, headers, rows, col_widths=None, highlight_rows=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, HEADER_BG)
        set_cell_text(cell, h, bold=True, color=C_WHITE, size=Pt(8))

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            set_cell_text(cell, val, size=Pt(8))
            if highlight_rows and r_idx in highlight_rows:
                set_cell_shading(cell, highlight_rows[r_idx])
            elif r_idx % 2 == 1:
                set_cell_shading(cell, ZEBRA_BG)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return table


def add_heading(doc, text, level, color=None):
    h = doc.add_heading(text, level=level)
    if color is None:
        color = {1: C_PRIMARY, 2: C_H2, 3: C_H3, 4: C_H4}.get(level, C_PRIMARY)
    for run in h.runs:
        run.font.color.rgb = color
        run.font.name = "Calibri"
    return h


def add_para(doc, text, bold=False, italic=False, size=Pt(10), color=None, space_after=Pt(4)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = size
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = space_after
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.name = "Calibri"
    return p


def add_bullet_bold_value(doc, label, value, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after = Pt(2)
    run_b = p.add_run(label)
    run_b.bold = True
    run_b.font.size = Pt(9)
    run_b.font.name = "Calibri"
    run_v = p.add_run(value)
    run_v.font.size = Pt(9)
    run_v.font.name = "Calibri"
    return p


def add_example_block(doc, num, ref, hyp, is_score, wer, judge_verdict, explanation):
    """Add a formatted example block with REF/HYP pair and explanation."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"Example {num}")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    run.font.color.rgb = C_H3
    meta = p.add_run(f"  |  IS = {is_score}  |  WER = {wer}%  |  Judge: {judge_verdict}")
    meta.font.size = Pt(9)
    meta.font.name = "Calibri"
    meta.font.color.rgb = C_GRAY

    # REF
    p_ref = doc.add_paragraph()
    p_ref.paragraph_format.left_indent = Inches(0.3)
    p_ref.paragraph_format.space_after = Pt(1)
    r1 = p_ref.add_run("REF: ")
    r1.bold = True
    r1.font.size = Pt(9)
    r1.font.name = "Calibri"
    r2 = p_ref.add_run(ref)
    r2.font.size = Pt(9)
    r2.font.name = "Calibri"
    r2.italic = True

    # HYP
    p_hyp = doc.add_paragraph()
    p_hyp.paragraph_format.left_indent = Inches(0.3)
    p_hyp.paragraph_format.space_after = Pt(1)
    r3 = p_hyp.add_run("HYP: ")
    r3.bold = True
    r3.font.size = Pt(9)
    r3.font.name = "Calibri"
    r4 = p_hyp.add_run(hyp)
    r4.font.size = Pt(9)
    r4.font.name = "Calibri"
    r4.italic = True

    # Explanation
    p_exp = doc.add_paragraph()
    p_exp.paragraph_format.left_indent = Inches(0.3)
    p_exp.paragraph_format.space_after = Pt(6)
    r5 = p_exp.add_run(explanation)
    r5.font.size = Pt(9)
    r5.font.name = "Calibri"
    r5.font.color.rgb = C_DARK


def _tight_page_break(doc):
    last_p = doc.paragraphs[-1]
    last_p.add_run().add_break(WD_BREAK.PAGE)


def add_header_footer(doc):
    section = doc.sections[0]
    # Header
    header = section.header
    hp = header.paragraphs[0]
    hp.paragraph_format.space_after = Pt(0)
    text_width = section.page_width - section.left_margin - section.right_margin
    hp.paragraph_format.tab_stops.add_tab_stop(text_width, WD_TAB_ALIGNMENT.RIGHT)
    if LOGO_ORCHARD.exists():
        run_img = hp.add_run()
        run_img.add_picture(str(LOGO_ORCHARD), height=Inches(0.3))
    hp.add_run("\t")
    run_brand = hp.add_run("Argos \u2014 The Orchard")
    run_brand.font.size = Pt(8)
    run_brand.font.name = "Calibri"
    run_brand.italic = True
    run_brand.font.color.rgb = C_GRAY

    # Footer
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f = fp.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run_f._r.append(fld_begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run_f._r.append(instr)
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run_f._r.append(fld_sep)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run_f._r.append(fld_end)
    run_f.font.size = Pt(8)
    run_f.font.color.rgb = C_GRAY


def add_cover_page(doc):
    for _ in range(2):
        doc.add_paragraph()
    if LOGO_PEACOCK.exists():
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_logo.add_run()
        run.add_picture(str(LOGO_PEACOCK), width=Inches(2.0))

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = p_title.add_run("ARGOS")
    run_t.bold = True
    run_t.font.size = Pt(48)
    run_t.font.color.rgb = C_PRIMARY
    run_t.font.name = "Calibri"

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_s = p_sub.add_run("IS vs Opus Judge: Disagreement Analysis")
    run_s.font.size = Pt(22)
    run_s.font.name = "Calibri"
    run_s.font.color.rgb = C_H2

    p_brand = doc.add_paragraph()
    p_brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_b = p_brand.add_run("The Orchard")
    run_b.font.size = Pt(20)
    run_b.font.name = "Calibri"

    p_author = doc.add_paragraph()
    p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_a = p_author.add_run("Yoad Oxman")
    run_a.font.size = Pt(14)
    run_a.font.name = "Calibri"

    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_d = p_date.add_run(f"Last updated: {LAST_UPDATED}")
    run_d.font.size = Pt(10)
    run_d.font.name = "Calibri"
    run_d.font.color.rgb = C_GRAY

    _tight_page_break(doc)


def add_toc(doc):
    for _ in range(3):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(0)
        sp.paragraph_format.space_before = Pt(0)

    add_heading(doc, "Table of Contents", 1)

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-2" \\h \\z \\u '
    run._r.append(instr)
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)
    fld_end_run = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    fld_end_run._r.append(fld_end)

    settings = doc.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)

    _tight_page_break(doc)


# =====================================================================
# LOAD DATA
# =====================================================================

def load_csv():
    rows = []
    with open(CSV_FILE) as f:
        reader = csv.DictReader(f)
        for r in reader:
            rows.append(r)
    return rows


# =====================================================================
# DOCUMENT SECTIONS
# =====================================================================

def section_overview(doc):
    add_heading(doc, "1. Overview", 1)
    add_para(doc, (
        "This document analyses disagreements between the Intelligibility Score (IS) metric "
        "and the Claude Opus LLM-as-a-Judge gold standard. Disagreements are cases where one "
        "system rates a segment as successful but the other rates it as failed. Understanding "
        "these edge cases is critical for calibrating the IS framework and identifying its "
        "blind spots."
    ))
    add_para(doc, (
        "Two evaluation modes are analysed: (1) Blind \u2014 the judge sees only the reference "
        "and hypothesis text, and (2) Context-aware \u2014 the judge infers the domain/topic from "
        "the reference and applies domain expectations. In both cases, the judge used a "
        "3-level Y/P/N scale (Yes = meaning clearly conveyed, Partial = some meaning preserved, "
        "No = meaning lost or misleading)."
    ))
    add_para(doc, (
        "The IS metric is a deterministic composite of 6 signals (semantic similarity, phonetic "
        "similarity, inverse WER, inverse WWER, named entity accuracy, length ratio), scored "
        "0\u20135.0. The legacy threshold was IS \u2265 3.0; the current NIV thresholds are "
        "IS \u2265 3.80 (for Y) and IS \u2265 2.00 (for Y+P)."
    ))

    add_heading(doc, "Disagreement Counts", 2)
    add_styled_table(doc,
        ["Mode", "Judge Y but IS < 3.0", "Judge N but IS \u2265 3.0", "Total Disagreements", "Agreement Rate"],
        [
            ["Blind", "19", "3", "22 / 1,497", "98.5%"],
            ["Context-aware", "5", "11", "16 / 1,497", "98.9%"],
        ],
        col_widths=[1.2, 1.4, 1.4, 1.3, 1.0],
    )
    add_para(doc, (
        "Disagreements are rare in both modes: 1.5% (blind) and 1.1% (context). The "
        "asymmetry differs: blind mode has more false negatives (IS misses useful output), "
        "while context mode has more false positives (IS overrates domain-confused output)."
    ), italic=True, size=Pt(9), color=C_GRAY)


def section_blind_y_is_fail(doc):
    _tight_page_break(doc)
    add_heading(doc, "2. Blind: Judge Says Y, but IS < 3.0", 1)
    add_para(doc, (
        "These 19 segments convey the reference meaning despite poor metric scores. The judge "
        "recovers meaning through holistic reasoning that the IS formula cannot replicate. "
        "These validate the LLM salvage hypothesis: a human (or LLM) viewer can fill gaps that "
        "word-level metrics punish."
    ))

    examples = [
        ("one really nice thing about this is",
         "what a brilliant idea this is",
         "1.841", "71.4", "Y",
         "Core meaning preserved via paraphrase. 'Brilliant idea' captures 'really nice thing'. "
         "IS penalises because no words overlap \u2014 but the semantic intent is identical."),

        ("about the being the living in space part that",
         "we're basically about the business of sending kids to live in space",
         "1.984", "111.1", "Y",
         "Core topic (living in space) preserved despite hallucinated expansion. WER > 100% "
         "triggers IS hard failure, but the judge recognises the elaboration as harmless."),

        ("real human application that we can pull from that",
         "what are the human implications of that but because",
         "2.059", "100.0", "Y",
         "'Human implications' captures 'human application' \u2014 a valid synonym substitution. "
         "IS penalises 100% WER; the judge sees the semantic bridge."),

        ("rotates all the way around once the dial that looks like a clock it means that you've used 7 48",
         "process all the way around once you get to the end of the clock it means that you viewed seven or seven point four a",
         "2.256", "80.0", "Y",
         "Temporal structure and numerical content (7, 7.4) are recoverable. The hypothesis "
         "preserves the instructional flow even though individual words differ."),

        ("to the next level",
         "to the next level and they have tried",
         "2.318", "100.0", "Y",
         "Core phrase perfectly intact. The trailing addition ('and they have tried') is minor "
         "elaboration that does not change meaning. WER=100% due to length inflation."),
    ]

    for i, (ref, hyp, is_s, wer, verdict, expl) in enumerate(examples, 1):
        add_example_block(doc, i, ref, hyp, is_s, wer, verdict, expl)

    add_heading(doc, "Root Cause Pattern", 2)
    add_bullet(doc, "Semantic preservation under high WER: meaning conveyed via paraphrase or synonym, but no word overlap to reward")
    add_bullet(doc, "Harmless hallucination: extra words inflate length ratio and WER, but the core content is intact")
    add_bullet(doc, "Phonetic bridging: lip-reading confusions produce phonetically close but lexically different words")
    add_bullet(doc, "Short-segment amplification: in short references (4\u20136 words), a single word change causes disproportionate WER")


def section_blind_n_is_pass(doc):
    _tight_page_break(doc)
    add_heading(doc, "3. Blind: Judge Says N, but IS \u2265 3.0", 1)
    add_para(doc, (
        "Only 3 segments (0.2%) are IS false positives under blind evaluation. The IS metric "
        "rates them highly because structural or phonetic overlap is strong, but meaning is "
        "fundamentally broken."
    ))

    examples = [
        ("all you have to do is unscrew",
         "all you have to do is not to",
         "3.419", "28.6", "N",
         "Semantic reversal: 'unscrew' \u2192 'not to' inverts the instruction. Structure matches "
         "perfectly, WER is low, but the actionable content is opposite. IS trusts structural "
         "alignment and cannot detect valence inversion."),

        ("blood extraction first before going to x ray because i wish i did that",
         "cut your hair first before going to ashram because i wish i did that",
         "3.142", "35.7", "N",
         "Complete domain swap: medical (blood extraction, x-ray) \u2192 wellness (cut hair, ashram). "
         "Phonetically plausible substitutions fool the IS phonetic and structure signals. "
         "The judge recognises the domain is completely wrong."),

        ("one twitch is all you do",
         "one to rich is all the",
         "3.011", "66.7", "N",
         "Phonetic garbage: 'to rich' and 'all the' are not meaningful phrases. IS scores it "
         "at 3.0 due to partial word overlap and acceptable length ratio. The judge sees "
         "complete semantic breakdown."),
    ]

    for i, (ref, hyp, is_s, wer, verdict, expl) in enumerate(examples, 1):
        add_example_block(doc, i, ref, hyp, is_s, wer, verdict, expl)

    add_heading(doc, "Root Cause Pattern", 2)
    add_bullet(doc, "Semantic reversal: IS structure-matches without catching that the action is inverted")
    add_bullet(doc, "Domain confusion: phonetically plausible but semantically wrong substitutions score well on phonetic/structure")
    add_bullet(doc, "Word salad: scattered correct words create acceptable overlap ratios despite no coherent meaning")


def section_context_y_is_fail(doc):
    _tight_page_break(doc)
    add_heading(doc, "4. Context-Aware: Judge Says Y, but IS < 3.0", 1)
    add_para(doc, (
        "With domain context, only 5 segments show this pattern (down from 19 blind). "
        "Context makes the judge stricter overall (Y drops 8pp), so fewer segments get "
        "upgraded to Y. The surviving cases are strong semantic matches where domain "
        "knowledge actually helps the judge accept approximate output."
    ))

    examples = [
        ("so um",
         "so i kind of",
         "2.059", "150.0", "Y (was N blind)",
         "The only N\u2192Y transition in the entire dataset. Both utterances are filler phrases; "
         "context lets the judge recognise they serve the same conversational function. "
         "IS penalises 150% WER, but both are semantically empty."),

        ("ho where to begin with this",
         "o to begin with this",
         "2.939", "33.3", "Y",
         "Minor word drop ('ho' \u2192 'o') preserves meaning completely. IS scores just below 3.0 "
         "because the missing word inflates WER for this short segment."),

        ("to the next level",
         "to the next level and they have tried",
         "2.318", "100.0", "Y",
         "Same as blind Example 5 \u2014 core phrase intact. Context confirms the added words "
         "are harmless elaboration."),

        ("is also very heavy and is seven and a half pounds over 11 and a half inches",
         "also very heavy at 7 5 pounds and over 11 5 inches",
         "2.600", "64.7", "Y",
         "Numerical content (7.5 pounds, 11.5 inches) preserved with different formatting. "
         "Context (product review) confirms the measurements are correct."),

        ("to put around it and kind of we're all getting a lot more knowledge about kind of where we came from",
         "like to poke around and get to know more about where we came from",
         "2.618", "61.9", "Y",
         "Core theme (learning about origins) intact. Rewording is significant but semantically "
         "faithful. Context confirms the topic (ancestry/exploration) is preserved."),
    ]

    for i, (ref, hyp, is_s, wer, verdict, expl) in enumerate(examples, 1):
        add_example_block(doc, i, ref, hyp, is_s, wer, verdict, expl)

    add_heading(doc, "Root Cause Pattern", 2)
    add_bullet(doc, "Context rescues almost nothing: only 1 N\u2192Y in 1,497 pairs")
    add_bullet(doc, "Surviving Y cases are paraphrases or numerical equivalences that IS penalises on word overlap")
    add_bullet(doc, "Context makes the Y threshold stricter, not more lenient \u2014 fewer cases qualify")


def section_context_n_is_pass(doc):
    _tight_page_break(doc)
    add_heading(doc, "5. Context-Aware: Judge Says N, but IS \u2265 3.0", 1)
    add_para(doc, (
        "This is the most revealing category: 11 segments (up from 3 blind). Context "
        "awareness exposes domain-specific failures that blind evaluation missed. The judge "
        "knows what the speaker was supposed to be talking about and catches when the "
        "hypothesis drifts into the wrong domain."
    ))

    examples = [
        ("i just learned something recently because i'm a lover of",
         "i just learned something recently because i'm not a lover of",
         "4.750", "10.0", "N (was P blind)",
         "Negation insertion: 'a lover of' \u2192 'not a lover of'. IS=4.75 (near perfect!) "
         "because only one word is added. But the meaning is reversed. The context-aware "
         "judge catches that this changes the speaker's stance on the topic completely."),

        ("i will let you guys know i am a very lazy natural which i'm trying",
         "will let you guys know i am a very lazy astronaut with some training",
         "3.602", "33.3", "N (was P blind)",
         "Domain swap: 'natural' (hair care) \u2192 'astronaut'. Without context, 'lazy astronaut' "
         "seems plausible. With context (beauty/hair video), it is absurd. IS scores high "
         "because structure and most words match."),

        ("i now have 50 stitches on my needle and so now we're ready to begin the decreases",
         "now have 50 stitches on my neck and so now we're ready to begin the skin grafting on samuel",
         "3.278", "39.1", "N (was P blind)",
         "Domain swap: knitting (needle, decreases) \u2192 medical (neck, skin grafting). Without "
         "context, both are plausible. With context (crafting tutorial), the medical interpretation "
         "is completely wrong."),

        ("in my opinion one of the most difficult things about dealing with an arc",
         "it might be one of the most difficult things about raising daughters",
         "3.083", "50.0", "N (was P blind)",
         "Topic drift: welding/technical ('dealing with an arc') \u2192 parenting ('raising daughters'). "
         "Structure preserved, emotional tone similar, but the actual subject is wrong."),

        ("and growing in student loan debt you could hire all of the us marshals in the entire country",
         "and growing in south korea to allow all of the us marketers in the entire country",
         "3.238", "30.8", "N (was P blind)",
         "Domain confusion: US education policy (student loan debt, US marshals) \u2192 international "
         "business (South Korea, US marketers). Context reveals the economics topic is "
         "completely lost despite low WER."),
    ]

    for i, (ref, hyp, is_s, wer, verdict, expl) in enumerate(examples, 1):
        add_example_block(doc, i, ref, hyp, is_s, wer, verdict, expl)

    add_heading(doc, "Root Cause Pattern", 2)
    add_bullet(doc, "Negation/valence inversion: a single word ('not') reverses meaning; IS cannot detect this")
    add_bullet(doc, "Domain-specific vocabulary swaps: phonetically similar words from the wrong domain fool IS but not a domain-aware judge")
    add_bullet(doc, "Context raises the bar: segments that seemed 'close enough' blind are revealed as wrong-domain when the judge knows the topic")
    add_bullet(doc, "This is the strongest argument for domain-aware fine-tuning \u2014 the model resolves lip movements to the wrong vocabulary")


def section_transition_matrix(doc):
    _tight_page_break(doc)
    add_heading(doc, "6. Blind \u2192 Context Transition Matrix", 1)
    add_para(doc, (
        "The transition matrix shows how judgments shift when the judge gains domain context. "
        "Context is a quality tool, not a rescue tool: it mainly downgrades shallow successes."
    ))

    add_styled_table(doc,
        ["Blind \u2192 Context", "\u2192 Y", "\u2192 P", "\u2192 N", "Total"],
        [
            ["Y", "207", "138", "0", "345"],
            ["P", "17", "517", "92", "626"],
            ["N", "1", "50", "475", "526"],
        ],
        col_widths=[1.2, 0.8, 0.8, 0.8, 0.8],
        highlight_rows={0: GREEN_BG},
    )

    add_bullet_bold_value(doc, "230 downgrades ", "vs 68 upgrades \u2192 net \u2212162 (\u221210.8%)")
    add_bullet_bold_value(doc, "80.1% stable: ", "1,199 pairs received the same judgment in both modes")
    add_bullet_bold_value(doc, "Dominant transition \u2014 Y\u2192P (138 cases): ", "hypothesis seemed complete without domain context, but domain knowledge revealed vocabulary failures")
    add_bullet_bold_value(doc, "P\u2192N (92 cases): ", "partial successes revealed to be deeper failures under domain scrutiny")
    add_bullet_bold_value(doc, "N\u2192Y (1 case only): ", "context essentially never rescues complete failures")


def section_mitigations(doc):
    _tight_page_break(doc)
    add_heading(doc, "7. Mitigation Strategies", 1)

    add_heading(doc, "7.1 For IS False Negatives (Judge Y, IS < 3.0)", 2)
    add_para(doc, "These are segments where IS is too harsh. The model output is useful but metrics penalise it.")

    add_bullet_bold_value(doc, "Adopt NIV Y+P threshold (IS \u2265 2.00): ",
        "already captures most of these cases. At IS \u2265 2.00, \u03ba=0.818 with judge Y+P (922 segments, 61.6%).")
    add_bullet_bold_value(doc, "Deploy llm_context_prob heuristic: ",
        "the 15-rule decision tree identifies 165 salvageable segments (18.3% of metric failures) "
        "with llm_context_prob \u2265 0.5, lifting useful rate to 61.6% (IS \u2265 2.00), confirmed at 64.9% by Opus judge.")
    add_bullet_bold_value(doc, "Accept the residual gap: ",
        "19 of 1,497 (1.3%) is a small false negative rate. Perfect alignment would require "
        "runtime LLM evaluation, which is expensive and non-deterministic.")

    add_heading(doc, "7.2 For IS False Positives (Judge N, IS \u2265 3.0)", 2)
    add_para(doc, "These are segments where IS is too generous. The model output is misleading but metrics rate it highly.")

    add_bullet_bold_value(doc, "NIV Y threshold (IS \u2265 3.80): ",
        "the stricter Y threshold reduces false positive exposure. Only 83 false positives at IS \u2265 3.80 vs 271 at IS \u2265 3.0.")
    add_bullet_bold_value(doc, "Cross-check with semantic similarity: ",
        "semantic sim \u2265 0.70 (\u03ba=0.714 vs judge Y) catches domain reversals better than IS alone.")
    add_bullet_bold_value(doc, "Domain-aware fine-tuning: ",
        "the root cause is the model resolving lip movements to wrong-domain vocabulary. "
        "Fine-tuning with 20K+ segments from diverse domains would reduce domain confusion at the source.")

    add_heading(doc, "7.3 For Context-Specific False Positives (11 cases)", 2)
    add_para(doc, "These are the most actionable \u2014 blind evaluation misses them, but domain knowledge catches them.")

    add_bullet_bold_value(doc, "Negation detection: ",
        "the IS=4.75 false positive ('lover of' \u2192 'not a lover of') shows IS cannot detect "
        "single-word negations. A dedicated negation-aware signal could catch these.")
    add_bullet_bold_value(doc, "Domain vocabulary validation: ",
        "checking whether hypothesis key terms belong to the expected topic vocabulary "
        "would flag 'astronaut' in a hair care video or 'skin grafting' in a knitting tutorial.")
    add_bullet_bold_value(doc, "Topic-conditioned decoding: ",
        "providing the model with a topic label at decode time was tested and failed (naive injection). "
        "Topic-aware fine-tuning is the viable path.")


def section_key_takeaways(doc):
    _tight_page_break(doc)
    add_heading(doc, "8. Key Takeaways", 1)

    takeaways = [
        ("Disagreements are rare.",
         "22 of 1,497 (1.5%) blind, 16 of 1,497 (1.1%) context. IS and the Opus judge "
         "agree on 98%+ of segments."),
        ("NIV thresholds resolve most blind disagreements.",
         "IS \u2265 3.80 for Y (\u03ba=0.690) and IS \u2265 2.00 for Y+P (\u03ba=0.818) align IS with "
         "the judge's natural decision boundaries."),
        ("IS is better at ranking than classifying.",
         "Pearson r=0.85 between IS and judge scores. The ranking is excellent; the edge cases "
         "are at classification boundaries."),
        ("Context makes the judge stricter, not more lenient.",
         "Y drops 8pp, Y+P drops 2.7pp. The dominant transition is Y\u2192P (138 cases): domain "
         "knowledge reveals vocabulary failures invisible to blind evaluation."),
        ("The 11 context false positives point to domain-aware fine-tuning.",
         "IS cannot detect wrong-domain vocabulary swaps. The solution is not better metrics "
         "but a better model that resolves lip movements to the correct vocabulary domain."),
    ]

    for i, (title, body) in enumerate(takeaways, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        run_n = p.add_run(f"{i}. {title}")
        run_n.bold = True
        run_n.font.size = Pt(10)
        run_n.font.name = "Calibri"
        run_n.font.color.rgb = C_PRIMARY

        p_body = doc.add_paragraph()
        p_body.paragraph_format.left_indent = Inches(0.3)
        p_body.paragraph_format.space_after = Pt(6)
        run_b = p_body.add_run(body)
        run_b.font.size = Pt(9)
        run_b.font.name = "Calibri"


# =====================================================================
# MAIN
# =====================================================================

def main():
    doc = Document()

    # Page setup: A4
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    add_header_footer(doc)
    add_cover_page(doc)
    add_toc(doc)

    section_overview(doc)
    section_blind_y_is_fail(doc)
    section_blind_n_is_pass(doc)
    section_context_y_is_fail(doc)
    section_context_n_is_pass(doc)
    section_transition_matrix(doc)
    section_mitigations(doc)
    section_key_takeaways(doc)

    OUTPUT_FILE.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_FILE))
    print(f"Saved: {OUTPUT_FILE}")


if __name__ == "__main__":
    main()