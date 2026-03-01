#!/usr/bin/env python3
"""
Argos -- Intelligibility Assessment Report

Generates a branded .docx report with all conclusions from the
intelligibility scoring analysis of 1497 lip-reading decoded segments.

Reads:
    - intelligibility_scores.csv (augmented segment data)
    - intelligibility_summary.json (aggregate stats)

Output:
    docs/evaluation/intelligibility/intelligibility_report.docx

Usage:
    python3 generate_intelligibility_report.py
"""

import csv
import json
import statistics
from pathlib import Path
from datetime import datetime
from collections import Counter

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

# -- Paths --
SCRIPT_DIR = Path(__file__).parent
ASSETS_DIR = SCRIPT_DIR.parent / "assets"
LOGO_ORCHARD = ASSETS_DIR / "logo.png"
LOGO_PEACOCK = ASSETS_DIR / "peacock.png"

DATA_DIR = Path("/home/ubuntu/docs/evaluation/intelligibility")
SCORES_CSV = DATA_DIR / "intelligibility_scores.csv"
SUMMARY_JSON = DATA_DIR / "intelligibility_summary.json"

OUTPUT_FILE = DATA_DIR / "intelligibility_report.docx"

LAST_UPDATED = datetime.now().strftime("%B %d, %Y")

# -- Colors --
C_PRIMARY = RGBColor(0x1A, 0x3A, 0x5C)
C_H2 = RGBColor(0x2A, 0x5A, 0x8C)
C_H3 = RGBColor(0x3A, 0x6A, 0x9C)
C_H4 = RGBColor(0x4A, 0x7A, 0xAC)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_DARK = RGBColor(0x33, 0x33, 0x33)
C_GREEN = RGBColor(0x15, 0x57, 0x24)
C_RED = RGBColor(0x72, 0x1C, 0x24)
C_AMBER = RGBColor(0x85, 0x64, 0x04)
C_GRAY = RGBColor(0x66, 0x66, 0x66)

HEADER_BG = "1a3a5c"
ZEBRA_BG = "f0f4f8"
GREEN_BG = "d4edda"
RED_BG = "f8d7da"
AMBER_BG = "fff3cd"
BLUE_BG = "d1ecf1"
TIER5_BG = "c3e6cb"
TIER4_BG = "d4edda"
TIER3_BG = "fff3cd"
TIER2_BG = "f8d7da"
TIER1_BG = "f5c6cb"


# =====================================================================
# DATA LOADING
# =====================================================================

def load_scores():
    rows = []
    with open(SCORES_CSV, newline="", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for r in reader:
            for col in ["wer_%", "wwer_%", "nea_f1_%", "semantic_sim",
                        "phonetic_sim", "length_ratio", "intelligibility_score"]:
                try:
                    r[col] = float(r.get(col, 0) or 0)
                except (ValueError, TypeError):
                    r[col] = 0.0
            for col in ["intelligibility_tier", "phonetic_matches", "phonetic_near_misses"]:
                try:
                    r[col] = int(r.get(col, 0) or 0)
                except (ValueError, TypeError):
                    r[col] = 0
            try:
                r["llm_context_prob"] = float(r.get("llm_context_prob", 0) or 0)
            except (ValueError, TypeError):
                r["llm_context_prob"] = 0.0
            r["context_recoverable"] = str(r.get("context_recoverable", "")).strip().lower() == "true"
            rows.append(r)
    return rows


def load_summary():
    with open(SUMMARY_JSON, encoding="utf-8") as f:
        return json.load(f)


# =====================================================================
# HELPER FUNCTIONS
# =====================================================================

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, bold=False, color=None, size=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = size
    run.font.name = "Calibri"


def add_styled_table(doc, headers, rows_data, col_widths=None, highlight_rows=None):
    table = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, HEADER_BG)
        set_cell_text(cell, h, bold=True, color=C_WHITE, size=Pt(8))

    for r_idx, row_data in enumerate(rows_data):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            set_cell_text(cell, val, size=Pt(8))
            if highlight_rows and r_idx in highlight_rows:
                set_cell_shading(cell, highlight_rows[r_idx])
            elif r_idx % 2 == 1:
                set_cell_shading(cell, ZEBRA_BG)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return table


_bookmark_counter = [0]


def add_heading(doc, text, level, color=None):
    h = doc.add_heading(text, level=level)
    if color is None:
        color = {1: C_PRIMARY, 2: C_H2, 3: C_H3, 4: C_H4}.get(level, C_PRIMARY)
    for run in h.runs:
        run.font.color.rgb = color
        run.font.name = "Calibri"
    return h


def add_para(doc, text, bold=False, italic=False, size=Pt(10), color=None, space_after=Pt(4)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = size
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = space_after
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.name = "Calibri"
    return p


def add_bullet_bold_value(doc, label, value, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after = Pt(2)
    run_b = p.add_run(label)
    run_b.bold = True
    run_b.font.size = Pt(9)
    run_b.font.name = "Calibri"
    run_v = p.add_run(value)
    run_v.font.size = Pt(9)
    run_v.font.name = "Calibri"
    return p


def _tight_page_break(doc):
    last_p = doc.paragraphs[-1]
    last_p.add_run().add_break(WD_BREAK.PAGE)


def _build_inline_image_xml(rId, cx, cy, pic_id=1, name="Logo"):
    inline = OxmlElement('wp:inline')
    inline.set('distT', '0')
    inline.set('distB', '0')
    inline.set('distL', '0')
    inline.set('distR', '0')
    extent = OxmlElement('wp:extent')
    extent.set('cx', str(cx))
    extent.set('cy', str(cy))
    inline.append(extent)
    docPr = OxmlElement('wp:docPr')
    docPr.set('id', str(pic_id))
    docPr.set('name', name)
    inline.append(docPr)
    graphic = OxmlElement('a:graphic')
    graphicData = OxmlElement('a:graphicData')
    graphicData.set('uri', 'http://schemas.openxmlformats.org/drawingml/2006/picture')
    pic = OxmlElement('pic:pic')
    nvPicPr = OxmlElement('pic:nvPicPr')
    cNvPr = OxmlElement('pic:cNvPr')
    cNvPr.set('id', '0')
    cNvPr.set('name', name)
    nvPicPr.append(cNvPr)
    cNvPicPr = OxmlElement('pic:cNvPicPr')
    nvPicPr.append(cNvPicPr)
    pic.append(nvPicPr)
    blipFill = OxmlElement('pic:blipFill')
    blip = OxmlElement('a:blip')
    blip.set(qn('r:embed'), rId)
    blipFill.append(blip)
    stretch = OxmlElement('a:stretch')
    fillRect = OxmlElement('a:fillRect')
    stretch.append(fillRect)
    blipFill.append(stretch)
    pic.append(blipFill)
    spPr = OxmlElement('pic:spPr')
    xfrm = OxmlElement('a:xfrm')
    off = OxmlElement('a:off')
    off.set('x', '0')
    off.set('y', '0')
    xfrm.append(off)
    ext_el = OxmlElement('a:ext')
    ext_el.set('cx', str(cx))
    ext_el.set('cy', str(cy))
    xfrm.append(ext_el)
    spPr.append(xfrm)
    prstGeom = OxmlElement('a:prstGeom')
    prstGeom.set('prst', 'rect')
    spPr.append(prstGeom)
    pic.append(spPr)
    graphicData.append(pic)
    graphic.append(graphicData)
    inline.append(graphic)
    drawing = OxmlElement('w:drawing')
    drawing.append(inline)
    return drawing


def add_header_footer(doc):
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    if LOGO_ORCHARD.exists():
        image_part = doc.part.package.get_or_add_image_part(str(LOGO_ORCHARD))
        rId = header.part.relate_to(image_part, RT.IMAGE)
        size_emu = int(0.3 * 914400)
        logo_run = hp.add_run()
        drawing = _build_inline_image_xml(rId, size_emu, size_emu, pic_id=10, name="Header Logo")
        logo_run._r.append(drawing)
    text_width = section.page_width - section.left_margin - section.right_margin
    hp.paragraph_format.tab_stops.add_tab_stop(text_width, WD_TAB_ALIGNMENT.RIGHT)
    hp.add_run("\t")
    run = hp.add_run("Argos \u2014 The Orchard")
    run.font.size = Pt(8)
    run.font.color.rgb = C_GRAY
    run.font.name = "Calibri"
    run.italic = True
    # Footer with page numbers
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = C_GRAY
    run.font.name = "Calibri"
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    run._r.append(fld1)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run._r.append(instr)
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld2)


def add_toc(doc, toc_titles):
    for _ in range(3):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(0)
        sp.paragraph_format.space_before = Pt(0)
        run = sp.add_run()
        run.font.size = Pt(2)

    add_heading(doc, "Table of Contents", 1)

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fld_begin)
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' TOC \\o "1-2" \\h \\z \\u '
    run._r.append(instr)
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    run._r.append(fld_sep)
    for title in toc_titles:
        placeholder = paragraph.add_run(title + "\n")
        placeholder.font.size = Pt(10)
        placeholder.font.name = "Calibri"
    fld_end_run = paragraph.add_run()
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    fld_end_run._r.append(fld_end)

    settings = doc.settings.element
    update_fields = OxmlElement('w:updateFields')
    update_fields.set(qn('w:val'), 'true')
    settings.append(update_fields)

    doc.add_page_break()


def fmt(val, decimals=1):
    return f"{val:.{decimals}f}"


# =====================================================================
# COVER PAGE
# =====================================================================

def create_cover_page(doc, summary):
    for _ in range(2):
        doc.add_paragraph()

    if LOGO_PEACOCK.exists():
        doc.add_picture(str(LOGO_PEACOCK), width=Inches(2.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ARGOS")
    run.font.size = Pt(48)
    run.font.color.rgb = C_PRIMARY
    run.bold = True
    run.font.name = "Calibri"

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Intelligibility Assessment Report")
    run2.font.size = Pt(22)
    run2.font.color.rgb = C_H2
    run2.font.name = "Calibri"

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("The Orchard")
    run3.font.size = Pt(20)
    run3.font.color.rgb = C_H3
    run3.font.name = "Calibri"

    for _ in range(3):
        doc.add_paragraph()

    p_author = doc.add_paragraph()
    p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_author = p_author.add_run("Yoad Oxman")
    run_author.font.size = Pt(14)
    run_author.font.color.rgb = C_DARK
    run_author.font.name = "Calibri"
    doc.add_paragraph()

    meta = [
        ("Segments Analyzed", str(summary["n_segments"])),
        ("Properly Captured", f'{summary["properly_captured_count"]} ({summary["properly_captured_pct"]}%)'),
        ("Date", LAST_UPDATED),
    ]
    for k, v in meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rb = p.add_run(f"{k}: ")
        rb.bold = True
        rb.font.size = Pt(11)
        rb.font.name = "Calibri"
        rb.font.color.rgb = C_DARK
        rv = p.add_run(v)
        rv.font.size = Pt(11)
        rv.font.name = "Calibri"
        rv.font.color.rgb = C_H2

    doc.add_page_break()


# =====================================================================
# REPORT SECTIONS
# =====================================================================

def section_executive_summary(doc, summary, rows):
    add_heading(doc, "Executive Summary", 1)

    n = summary["n_segments"]
    captured = summary["properly_captured_count"]
    captured_pct = summary["properly_captured_pct"]
    ctx_rule = summary["context_recoverable_count"]
    ctx_rule_pct = summary["context_recoverable_pct"]
    ctx_llm = summary["llm_context_recoverable_count"]
    ctx_llm_pct = summary["llm_context_recoverable_pct"]

    add_para(doc,
        f"This report assesses how many of the {n} lip-reading decoded segments "
        f"were 'properly captured' \u2014 meaning a human could understand what was actually "
        f"said from reading the model's output. We move beyond Word Error Rate (WER) to "
        f"a composite Intelligibility Score (IS) that combines semantic similarity, "
        f"phonetic analysis, and entity preservation.",
        size=Pt(10))

    add_heading(doc, "Key Findings", 2)

    # Headline metrics table
    add_styled_table(doc,
        ["Metric", "Count", "Percentage"],
        [
            ["Properly Captured (IS >= 3.0)", str(captured), f"{captured_pct}%"],
            ["Context-Recoverable (rule-based)", str(ctx_rule), f"{ctx_rule_pct}%"],
            ["Context-Recoverable (LLM-based)", str(ctx_llm), f"{ctx_llm_pct}%"],
            ["Empty Hypotheses", str(summary["n_empty_hyp"]), f'{summary["n_empty_hyp"]/n*100:.1f}%'],
        ],
        col_widths=[3.0, 1.2, 1.2],
        highlight_rows={0: GREEN_BG, 1: BLUE_BG, 2: BLUE_BG, 3: RED_BG}
    )

    add_para(doc,
        f"Of {n} segments, {captured} ({captured_pct}%) are intelligible enough that a human "
        f"could understand what was said. When we add topic context (a viewer who knows the "
        f"general subject), the rule-based estimate rises to {ctx_rule} ({ctx_rule_pct}%) "
        f"and the LLM-knowledge-based estimate reaches {ctx_llm} ({ctx_llm_pct}%).",
        size=Pt(10))

    add_para(doc,
        f"Mean Intelligibility Score: {summary['mean_is']:.2f} / 5.0, "
        f"Median: {summary['median_is']:.2f} / 5.0.",
        bold=True, size=Pt(10))

    _tight_page_break(doc)


def section_why_wer_isnt_enough(doc):
    add_heading(doc, "Why WER Isn't Enough", 1)

    add_para(doc,
        "Word Error Rate (WER) counts how many words differ between the reference "
        "and hypothesis. But it treats all errors equally \u2014 it doesn't distinguish "
        "whether errors destroy meaning or merely change surface form.",
        size=Pt(10))

    add_heading(doc, "Same WER, Opposite Intelligibility", 2)

    add_para(doc,
        "Consider these two predictions, both with ~30% WER:",
        size=Pt(10))

    # Example A
    add_para(doc, "Prediction A \u2014 WER 29%, FULLY INTELLIGIBLE:", bold=True, size=Pt(10))
    add_bullet(doc, 'Reference: "that are going to allow you to work with the team in a more"')
    add_bullet(doc, 'Predicted: "are going to allow you to work with a team and more"')
    add_bullet(doc, "A human reading the prediction completely understands the message. "
               "Minor articles changed, trivial restructuring.")

    add_para(doc, "Prediction B \u2014 WER 33%, COMPLETELY UNINTELLIGIBLE:", bold=True, size=Pt(10))
    add_bullet(doc, 'Reference: "today i\'m talking with admiral mcrae"')
    add_bullet(doc, 'Predicted: "today i\'m talking with animal migratory"')
    add_bullet(doc, "The structure is similar but the meaning is destroyed. "
               '"Admiral McRae" (a person) became "animal migratory" (a biology concept).')

    add_para(doc,
        "WER says these are almost equally good (29% vs 33%). "
        "A human says one is perfect and the other is nonsense. "
        "This is why we need the Intelligibility Score.",
        italic=True, size=Pt(10))

    doc.add_page_break()


def section_methodology(doc, summary):
    add_heading(doc, "Methodology: The 6-Signal Composite", 1)

    add_para(doc,
        "The Intelligibility Score (IS) combines six automated signals that collectively "
        "approximate how a human expert would judge whether a lip-reading output preserves "
        "the speaker's meaning.",
        size=Pt(10))

    signals = [
        ("Semantic Similarity", "0.25",
         "Compares sentence meaning using AI embeddings (all-MiniLM-L6-v2). "
         "Captures whether ref and hyp are about the same topic."),
        ("Phonetic Similarity", "0.15",
         "Double Metaphone encoding checks if substituted words SOUND alike. "
         "Captures natural lip-reading confusions (p/b/m, t/d/n groups)."),
        ("Inverse WER", "0.15",
         "Traditional word error rate, inverted (lower WER = higher score). "
         "Raw structural correctness."),
        ("Inverse WWER", "0.15",
         "Weighted WER that penalizes content word errors more than function words. "
         "Content words 1x, function words 0.5x, entities 2x."),
        ("NEA F1", "0.15",
         "Named Entity Accuracy F1 \u2014 whether proper nouns, numbers, and names survived. "
         "These are the words humans cannot recover from context."),
        ("Length Ratio", "0.15",
         "Ratio of hypothesis to reference length. Detects hallucination (too long) "
         "and truncation (too short)."),
    ]

    add_styled_table(doc,
        ["Signal", "Weight", "Description"],
        [[s[0], s[1], s[2]] for s in signals],
        col_widths=[1.5, 0.6, 4.5])

    add_heading(doc, "Composite Formula", 2)
    add_para(doc,
        "IS = 0.25 \u00d7 Semantic + 0.15 \u00d7 Phonetic + 0.15 \u00d7 (1\u2212WER) "
        "+ 0.15 \u00d7 (1\u2212WWER) + 0.15 \u00d7 NEA_F1 + 0.15 \u00d7 Length_Ratio",
        bold=True, size=Pt(10))
    add_para(doc,
        "All signals scaled to 0\u20135 before combining. Final IS range: 0.0 to 5.0. "
        '"Properly captured" = IS >= 3.0 (Tier 4: Good + Tier 5: Excellent).',
        size=Pt(10))

    add_heading(doc, "Signal Statistics", 2)
    ss = summary["signal_stats"]
    add_styled_table(doc,
        ["Signal", "Mean", "Median"],
        [
            ["Semantic Similarity", fmt(ss["semantic_sim"]["mean"], 3), fmt(ss["semantic_sim"]["median"], 3)],
            ["Phonetic Similarity", fmt(ss["phonetic_sim"]["mean"], 3), fmt(ss["phonetic_sim"]["median"], 3)],
            ["Inverse WER", fmt(ss["inv_wer_mean"], 3), "\u2014"],
            ["Inverse WWER", fmt(ss["inv_wwer_mean"], 3), "\u2014"],
            ["NEA F1 (%)", fmt(ss["nea_f1_mean"], 1), "\u2014"],
            ["Length Ratio", fmt(ss["length_ratio"]["mean"], 3), fmt(ss["length_ratio"]["median"], 3)],
        ],
        col_widths=[2.0, 1.5, 1.5])

    _tight_page_break(doc)


def section_tier_distribution(doc, summary, rows):
    add_heading(doc, "Tier Distribution", 1)

    tier_data = summary["tier_distribution"]
    n = summary["n_segments"]

    tier_desc = {
        "5_excellent": "Human fully understands. Minor word differences only.",
        "4_good": "Meaning clearly preserved. Some wrong words but recoverable.",
        "3_fair": "Gist recoverable. General topic correct, details wrong.",
        "2_poor": "Only fragments survive. Different message conveyed.",
        "1_failed": "No meaningful connection. Hallucinated or empty.",
    }

    tier_rows = []
    tier_highlights = {}
    tier_bgs = {"5_excellent": TIER5_BG, "4_good": TIER4_BG, "3_fair": TIER3_BG,
                "2_poor": TIER2_BG, "1_failed": TIER1_BG}
    for i, key in enumerate(["5_excellent", "4_good", "3_fair", "2_poor", "1_failed"]):
        td = tier_data[key]
        tier_rows.append([
            key.replace("_", " ").title(),
            str(td["count"]),
            f'{td["pct"]}%',
            tier_desc.get(key, ""),
        ])
        tier_highlights[i] = tier_bgs[key]

    add_styled_table(doc,
        ["Tier", "Count", "Pct", "Description"],
        tier_rows,
        col_widths=[1.2, 0.7, 0.6, 4.1],
        highlight_rows=tier_highlights)

    # Interpretation
    excellent = tier_data["5_excellent"]["count"]
    good = tier_data["4_good"]["count"]
    fair = tier_data["3_fair"]["count"]
    poor = tier_data["2_poor"]["count"]
    failed = tier_data["1_failed"]["count"]

    add_para(doc,
        f"The top two tiers (Excellent + Good) total {excellent + good} segments ({(excellent+good)/n*100:.1f}%) "
        f"\u2014 these are the outputs a human could understand without additional context.",
        size=Pt(10))
    add_para(doc,
        f"The Fair tier ({fair} segments, {fair/n*100:.1f}%) represents partial intelligibility: "
        f"the general topic survives but key details are wrong.",
        size=Pt(10))
    add_para(doc,
        f"The bottom two tiers (Poor + Failed) total {poor + failed} segments ({(poor+failed)/n*100:.1f}%) "
        f"\u2014 these are outputs where meaning is largely or completely lost.",
        size=Pt(10))

    doc.add_page_break()

    # Tier examples
    add_heading(doc, "Tier Examples from Data", 2)

    tier_examples = {5: [], 4: [], 3: [], 2: [], 1: []}
    for r in rows:
        t = r["intelligibility_tier"]
        if len(tier_examples.get(t, [])) < 2:
            if r.get("ref", "").strip() and r.get("hyp", "").strip():
                tier_examples.setdefault(t, []).append(r)

    tier_names = {5: "Excellent", 4: "Good", 3: "Fair", 2: "Poor", 1: "Failed"}
    for tier_num in [5, 4, 3, 2, 1]:
        examples = tier_examples.get(tier_num, [])
        if not examples:
            continue
        add_heading(doc, f"Tier {tier_num}: {tier_names[tier_num]}", 3)
        for ex in examples[:2]:
            ref_text = ex.get("ref", "")[:120]
            hyp_text = ex.get("hyp", "")[:120]
            wer = ex["wer_%"]
            is_score = ex["intelligibility_score"]
            add_para(doc, f"IS: {is_score:.2f}, WER: {wer:.0f}%", bold=True, size=Pt(9))
            add_bullet(doc, f'Ref: "{ref_text}"')
            add_bullet(doc, f'Hyp: "{hyp_text}"')

    _tight_page_break(doc)


def add_bar_chart_table(doc, title, items, color_map=None, max_count=None):
    """
    Create a color-coded horizontal bar chart using a Word table.

    items: list of (label, count, pct) tuples
    color_map: dict mapping label -> hex color for the bar, or None for default
    """
    if not items:
        return

    if max_count is None:
        max_count = max(c for _, c, _ in items) if items else 1

    # Default colors cycle
    default_colors = [
        "2a5a8c", "3a8c3a", "c47f17", "8c2a2a", "5a2a8c",
        "178c8c", "8c5a1a", "2a2a8c", "1a6b3a", "8c1a5a",
    ]

    table = doc.add_table(rows=len(items), cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (label, count, pct) in enumerate(items):
        # Label column
        cell_label = table.rows[i].cells[0]
        set_cell_text(cell_label, label, size=Pt(8), bold=True)
        cell_label.width = Inches(2.5)

        # Bar column - use shading to create visual bar effect
        cell_bar = table.rows[i].cells[1]
        bar_text = "\u2588" * max(1, int(count * 20 / max(1, max_count)))
        cell_bar.text = ""
        p = cell_bar.paragraphs[0]
        run = p.add_run(bar_text)
        run.font.size = Pt(8)
        run.font.name = "Calibri"
        bar_color = None
        if color_map and label in color_map:
            bar_color = color_map[label]
        else:
            bar_color = default_colors[i % len(default_colors)]
        run.font.color.rgb = RGBColor(
            int(bar_color[:2], 16),
            int(bar_color[2:4], 16),
            int(bar_color[4:6], 16),
        )
        cell_bar.width = Inches(2.5)

        # Count column
        cell_count = table.rows[i].cells[2]
        set_cell_text(cell_count, f"{count} ({pct:.1f}%)", size=Pt(8))
        cell_count.width = Inches(1.2)

        # Zebra striping
        if i % 2 == 1:
            for cell in [cell_label, cell_bar, cell_count]:
                set_cell_shading(cell, ZEBRA_BG)

    doc.add_paragraph()
    return table


def section_is_histogram(doc, summary):
    add_heading(doc, "Intelligibility Score Distribution", 1)

    add_para(doc,
        "The histogram below shows how IS scores are distributed across the 1497 segments. "
        "A roughly uniform distribution means the model doesn't cluster around any particular "
        "quality level \u2014 it produces the full range from excellent to failed.",
        size=Pt(10))

    histogram = summary.get("is_histogram", {})
    n = summary["n_segments"]

    # Color-code by tier: 0-1 red, 1-2 orange, 2-3 amber, 3-4 green, 4-5 dark green
    hist_colors = {
        "0.0-0.5": "c0392b", "0.5-1.0": "e74c3c",  # Failed - red
        "1.0-1.5": "d35400", "1.5-2.0": "e67e22",  # Poor - orange
        "2.0-2.5": "f39c12", "2.5-3.0": "f1c40f",  # Fair - amber
        "3.0-3.5": "27ae60", "3.5-4.0": "2ecc71",  # Good - green
        "4.0-4.5": "1a8c4e", "4.5-5.0": "155724",  # Excellent - dark green
    }

    items = [(bucket, count, count / n * 100) for bucket, count in histogram.items()]
    add_bar_chart_table(doc, "IS Distribution", items, color_map=hist_colors)

    # Interpretation
    low = sum(histogram.get(k, 0) for k in ["0.0-0.5", "0.5-1.0"])
    mid_low = sum(histogram.get(k, 0) for k in ["1.0-1.5", "1.5-2.0"])
    mid = sum(histogram.get(k, 0) for k in ["2.0-2.5", "2.5-3.0"])
    mid_high = sum(histogram.get(k, 0) for k in ["3.0-3.5", "3.5-4.0"])
    high = sum(histogram.get(k, 0) for k in ["4.0-4.5", "4.5-5.0"])

    add_para(doc,
        f"The distribution is roughly bell-shaped with a slight negative skew: "
        f"{high} segments ({high/n*100:.0f}%) score IS 4.0\u20135.0 (excellent), "
        f"{mid_high} ({mid_high/n*100:.0f}%) score 3.0\u20134.0 (good), "
        f"{mid} ({mid/n*100:.0f}%) score 2.0\u20133.0 (fair), "
        f"{mid_low} ({mid_low/n*100:.0f}%) score 1.0\u20132.0 (poor), and "
        f"{low} ({low/n*100:.0f}%) score 0.0\u20131.0 (failed).",
        size=Pt(10))

    _tight_page_break(doc)


def section_failure_analysis(doc, summary, rows):
    add_heading(doc, "Failure Mode Analysis", 1)

    fm_dist = summary.get("failure_mode_distribution", {})
    n_failed = sum(v["count"] for v in fm_dist.values())

    add_para(doc,
        f"Of the {n_failed} segments that failed to be properly captured (IS < 3.0), "
        f"what went wrong? The chart below classifies each failure into its dominant mode.",
        size=Pt(10))

    # Color map for failure modes: severity gradient
    fm_colors = {
        "Hallucination": "c0392b",
        "Empty Output": "922b21",
        "Total Topic Drift": "e74c3c",
        "Phonetically Similar but Wrong Topic": "d35400",
        "Entity Destruction": "e67e22",
        "High Error Rate": "f39c12",
        "Content Word Errors": "f1c40f",
        "Accumulated Small Errors": "d4ac0d",
        "Truncation": "7d3c98",
        "Over-generation": "6c3483",
    }

    items = [(mode, v["count"], v["pct"]) for mode, v in fm_dist.items()]
    add_bar_chart_table(doc, "Failure Modes", items, color_map=fm_colors)

    add_heading(doc, "Failure Mode Descriptions", 2)

    descriptions = [
        ("Hallucination",
         "WER > 100%. The model generated fluent but completely fabricated text, longer "
         "than the original. This is the most dangerous failure mode because the output "
         "looks confident and plausible."),
        ("Empty Output",
         "The model produced nothing. The decoding process failed silently, producing "
         "zero tokens. These represent complete system failures."),
        ("Total Topic Drift",
         "Neither semantic similarity nor phonetic similarity could find any connection "
         "between reference and hypothesis. The model produced text about a completely "
         "different subject."),
        ("Phonetically Similar but Wrong Topic",
         "The words sound similar to the reference but convey a different meaning. "
         "The lip movements were correctly read but mapped to wrong word choices. "
         'Example: "admiral mcrae" \u2192 "animal migratory".'),
        ("Entity Destruction",
         "The general topic is vaguely preserved but all proper nouns, numbers, and "
         "names are destroyed. These are unrecoverable because names cannot be guessed "
         "from context."),
        ("High Error Rate",
         "WER > 70% with no compensating semantic or phonetic similarity. "
         "Too many words are wrong for any signal to survive."),
        ("Content Word Errors",
         "The sentence structure is intact but key content words (nouns, verbs) are "
         "substituted with unrelated words. Function words are correct but meaning is lost."),
        ("Accumulated Small Errors",
         "No single catastrophic failure, but many small errors add up to destroy "
         "the overall meaning. Each error is individually minor but collectively fatal."),
        ("Truncation",
         "The model produced a much shorter output than expected (< 30% of reference length). "
         "The decoding ended too early."),
        ("Over-generation",
         "The model produced a much longer output than expected (> 180% of reference length) "
         "without WER exceeding 100%. Extra content was inserted."),
    ]

    for mode, desc in descriptions:
        if mode in fm_dist:
            count = fm_dist[mode]["count"]
            pct = fm_dist[mode]["pct"]
            add_para(doc, f"{mode} ({count} segments, {pct}%)", bold=True, size=Pt(9))
            add_para(doc, desc, size=Pt(9), space_after=Pt(6))

    _tight_page_break(doc)


def section_success_analysis(doc, summary, rows):
    add_heading(doc, "Success Pattern Analysis", 1)

    sp_dist = summary.get("success_pattern_distribution", {})
    n_success = sum(v["count"] for v in sp_dist.values())

    add_para(doc,
        f"Of the {n_success} segments that were properly captured (IS >= 3.0), "
        f"what made them succeed? The chart below reveals the dominant pattern "
        f"that drove each segment above the intelligibility threshold.",
        size=Pt(10))

    # Color map for success patterns: positive gradient
    sp_colors = {
        "Near-Perfect Output": "155724",
        "Minor Errors, High Semantic Match": "1a8c4e",
        "Phonetically Preserved": "27ae60",
        "Entities Preserved": "2ecc71",
        "Good Semantic + Correct Length": "58d68d",
        "Low-Moderate WER": "82e0aa",
        "Combined Semantic + Phonetic Bridge": "abebc6",
        "Multiple Partial Signals": "d5f5e3",
    }

    items = [(pat, v["count"], v["pct"]) for pat, v in sp_dist.items()]
    add_bar_chart_table(doc, "Success Patterns", items, color_map=sp_colors)

    add_heading(doc, "Key Insight: Phonetic Preservation Is the #1 Success Driver", 2)

    phon_count = sp_dist.get("Phonetically Preserved", {}).get("count", 0)
    phon_pct = sp_dist.get("Phonetically Preserved", {}).get("pct", 0)

    add_para(doc,
        f"The single largest success category is 'Phonetically Preserved' ({phon_count} "
        f"segments, {phon_pct}%). These are segments where the WER is moderate (> 25%) "
        f"but the wrong words SOUND like the right words. The model correctly read the "
        f"lip movements and mapped them to phonetically plausible words.",
        size=Pt(10))

    add_para(doc,
        "This confirms a key property of lip-reading: the visual signal preserves the "
        "phonetic structure of speech even when specific words are wrong. A system that "
        "corrects phonetic near-misses (e.g., via acoustic language model rescoring) "
        "could potentially recover a large fraction of these segments to near-perfect accuracy.",
        size=Pt(10))

    add_heading(doc, "What Successful Segments Look Like (Signal Comparison)", 2)

    sig_cmp = summary.get("signal_comparison", {})
    success_stats = sig_cmp.get("success", {})
    failure_stats = sig_cmp.get("failure", {})

    if success_stats and failure_stats:
        cmp_rows = [
            ["Semantic Similarity",
             fmt(success_stats.get("semantic_sim", 0), 3),
             fmt(failure_stats.get("semantic_sim", 0), 3),
             f'{success_stats.get("semantic_sim", 0) - failure_stats.get("semantic_sim", 0):+.3f}'],
            ["Phonetic Similarity",
             fmt(success_stats.get("phonetic_sim", 0), 3),
             fmt(failure_stats.get("phonetic_sim", 0), 3),
             f'{success_stats.get("phonetic_sim", 0) - failure_stats.get("phonetic_sim", 0):+.3f}'],
            ["WER (%)",
             fmt(success_stats.get("wer", 0), 1),
             fmt(failure_stats.get("wer", 0), 1),
             f'{success_stats.get("wer", 0) - failure_stats.get("wer", 0):+.1f}'],
            ["WWER (%)",
             fmt(success_stats.get("wwer", 0), 1),
             fmt(failure_stats.get("wwer", 0), 1),
             f'{success_stats.get("wwer", 0) - failure_stats.get("wwer", 0):+.1f}'],
            ["NEA F1 (%)",
             fmt(success_stats.get("nea_f1", 0), 1),
             fmt(failure_stats.get("nea_f1", 0), 1),
             f'{success_stats.get("nea_f1", 0) - failure_stats.get("nea_f1", 0):+.1f}'],
            ["Length Ratio",
             fmt(success_stats.get("length_ratio", 0), 3),
             fmt(failure_stats.get("length_ratio", 0), 3),
             f'{success_stats.get("length_ratio", 0) - failure_stats.get("length_ratio", 0):+.3f}'],
        ]

        add_styled_table(doc,
            ["Signal", "Success Mean", "Failure Mean", "Gap"],
            cmp_rows,
            col_widths=[2.0, 1.3, 1.3, 1.0],
            highlight_rows={0: GREEN_BG, 1: GREEN_BG, 4: GREEN_BG})

        add_para(doc,
            "The largest differentiators are NEA F1 (entity preservation, +58 gap), "
            "WER (-56 gap), and semantic similarity (+0.50 gap). Successful segments "
            "preserve entities, have lower error rates, and maintain semantic coherence. "
            "Length ratio is notably similar between groups, confirming that length alone "
            "is not a strong predictor of quality.",
            size=Pt(10))

    _tight_page_break(doc)


def section_wer_thresholds(doc, rows, summary):
    add_heading(doc, "WER/WWER Threshold Analysis", 1)

    add_para(doc,
        "While the Intelligibility Score is a composite metric, practitioners often need "
        "simple threshold rules. The table below maps WER ranges to intelligibility expectations.",
        size=Pt(10))

    n = len(rows)
    wer_vals = [r["wer_%"] for r in rows]

    threshold_rows = []
    thresholds = [
        (0, 15, "Almost Certainly Understandable", GREEN_BG),
        (15, 30, "Probably Understandable", TIER4_BG),
        (30, 50, "Coin Flip \u2014 Depends on WHICH Words", TIER3_BG),
        (50, 70, "Probably NOT Understandable", TIER2_BG),
        (70, 100, "Almost Certainly NOT Understandable", TIER1_BG),
        (100, 999, "Hallucinated / Guaranteed NOT Understandable", RED_BG),
    ]

    highlights = {}
    for i, (lo, hi, label, bg) in enumerate(thresholds):
        if hi == 999:
            count = sum(1 for w in wer_vals if w > lo)
        else:
            count = sum(1 for w in wer_vals if lo < w <= hi)
            if lo == 0:
                count = sum(1 for w in wer_vals if w <= hi)
        pct = count / n * 100
        if hi == 999:
            range_str = f"WER > {lo}%"
        elif lo == 0:
            range_str = f"WER \u2264 {hi}%"
        else:
            range_str = f"WER {lo}\u2013{hi}%"
        threshold_rows.append([range_str, str(count), f"{pct:.1f}%", label])
        highlights[i] = bg

    add_styled_table(doc,
        ["WER Range", "Count", "Pct", "Intelligibility Expectation"],
        threshold_rows,
        col_widths=[1.2, 0.7, 0.6, 4.1],
        highlight_rows=highlights)

    add_heading(doc, "What the Thresholds Mean", 2)

    add_para(doc, "WER \u2264 15% (Almost Certainly Understandable):", bold=True, size=Pt(10))
    add_para(doc,
        "At 15% WER on a 10-word sentence, only ~1.5 words are wrong. In practice this means "
        "one minor substitution \u2014 typically a function word like 'the'\u2192'a' or a morphological "
        "variant like 'going'\u2192'gonna'. Example: 'is 100 in our control simply by thinking positive' "
        "(WER 10%) \u2014 only 'wellness' dropped.",
        size=Pt(9))

    add_para(doc, "WER 15\u201330% (Probably Understandable):", bold=True, size=Pt(10))
    add_para(doc,
        "2\u20133 wrong words out of 10. Usually function words or near-synonyms. The sentence "
        "structure survives and content words are mostly intact.",
        size=Pt(9))

    add_para(doc, "WER 30\u201350% (Coin Flip):", bold=True, size=Pt(10))
    add_para(doc,
        "This is where it depends entirely on WHICH words are wrong. If errors hit function words "
        "and phonetic near-misses, it's still understandable. If they hit content words and entities, "
        "it's not. Same WER, completely different outcomes.",
        size=Pt(9))

    add_para(doc, "WER 50\u201370% (Probably NOT Understandable):", bold=True, size=Pt(10))
    add_para(doc,
        "More than half the words are wrong. Only fragments of the original survive. Occasionally "
        "recoverable if surviving fragments happen to be key content words.",
        size=Pt(9))

    add_para(doc, "WER > 70% (Almost Certainly NOT Understandable):", bold=True, size=Pt(10))
    add_para(doc,
        "Nearly every word is wrong. Even if a word or two survives, there's not enough signal.",
        size=Pt(9))

    add_para(doc, "WER > 100% (Hallucinated):", bold=True, size=Pt(10))
    add_para(doc,
        "WER can exceed 100% when the hypothesis is longer than the reference (insertions "
        "counted as errors). This always indicates hallucination \u2014 the model generated "
        "fluent but fabricated text.",
        size=Pt(9))

    add_para(doc,
        "WWER is usually more informative because it weights content words more heavily. "
        "A sentence with WWER < 40% usually has its content words intact even if articles "
        "and prepositions are garbled.",
        italic=True, size=Pt(9))

    _tight_page_break(doc)


def section_context_recovery(doc, summary, rows):
    add_heading(doc, "Context Recovery Analysis", 1)

    add_para(doc,
        "A viewer who already knows the general topic (e.g., 'this is a cooking video') can "
        "fill in gaps that a cold reader cannot. We estimate this recoverability using two "
        "independent approaches.",
        size=Pt(10))

    add_heading(doc, "Method 1: Rule-Based Context Recovery", 2)

    add_para(doc,
        "A decision-tree approach using semantic similarity, WER, WWER, and length ratio "
        "thresholds to estimate whether context helps.",
        size=Pt(10))

    rule_reasons = summary["context_recovery_reasons"]
    n = summary["n_segments"]
    rule_rows = sorted(rule_reasons.items(), key=lambda x: -x[1])
    add_styled_table(doc,
        ["Reason", "Count", "Pct"],
        [[r, str(c), f"{c/n*100:.1f}%"] for r, c in rule_rows],
        col_widths=[3.0, 1.0, 1.0])

    add_para(doc,
        f'Result: {summary["context_recoverable_count"]} / {n} segments '
        f'({summary["context_recoverable_pct"]}%) are recoverable with topic context.',
        bold=True, size=Pt(10))

    add_heading(doc, "Method 2: LLM-Knowledge-Based Context Recovery", 2)

    add_para(doc,
        "A multi-factor analysis that approximates how an expert would assess recoverability. "
        "Instead of hard thresholds, it considers: content word overlap, sequence preservation, "
        "phonetic plausibility, length sanity, semantic domain coherence, and information density.",
        size=Pt(10))

    llm_reasons = summary["llm_context_reasons"]
    llm_rows = sorted(llm_reasons.items(), key=lambda x: -x[1])
    add_styled_table(doc,
        ["Reasoning Category", "Count", "Pct"],
        [[r, str(c), f"{c/n*100:.1f}%"] for r, c in llm_rows],
        col_widths=[3.0, 1.0, 1.0])

    add_para(doc,
        f'Result: {summary["llm_context_recoverable_count"]} / {n} segments '
        f'({summary["llm_context_recoverable_pct"]}%) are recoverable (probability >= 0.5).',
        bold=True, size=Pt(10))

    add_para(doc,
        f'Mean recovery probability: {summary["llm_context_mean_prob"]:.3f}, '
        f'Median: {summary["llm_context_median_prob"]:.3f}.',
        size=Pt(10))

    add_heading(doc, "Comparison of Approaches", 2)

    add_styled_table(doc,
        ["Approach", "Recoverable", "Pct", "Notes"],
        [
            ["Rule-Based", str(summary["context_recoverable_count"]),
             f'{summary["context_recoverable_pct"]}%', "Conservative; hard thresholds"],
            ["LLM-Knowledge", str(summary["llm_context_recoverable_count"]),
             f'{summary["llm_context_recoverable_pct"]}%', "More nuanced; multi-factor"],
            ["Difference", str(summary["llm_context_recoverable_count"] - summary["context_recoverable_count"]),
             f'{summary["llm_context_recoverable_pct"] - summary["context_recoverable_pct"]:.1f}%',
             "LLM approach recovers more via structural + phonetic bridges"],
        ],
        col_widths=[1.5, 1.0, 0.7, 3.4],
        highlight_rows={2: BLUE_BG})

    _tight_page_break(doc)


def section_topic_analysis(doc, summary):
    add_heading(doc, "Topic Analysis", 1)

    topic_data = summary.get("topic_analysis", {})
    if not topic_data:
        add_para(doc, "No topic analysis data available.", size=Pt(10))
        return

    add_para(doc,
        "Topics are assigned by keyword matching on reference text. This reveals which "
        "content domains the lip-reading model handles best and worst. All three assessment "
        "methods (IS threshold, rule-based context, LLM-based context) are shown per topic.",
        size=Pt(10))

    add_heading(doc, "Intelligibility by Topic (sorted by Mean IS)", 2)

    # Color gradient: green for high IS, red for low IS
    topic_colors = {}
    sorted_topics = list(topic_data.items())
    for i, (topic, _) in enumerate(sorted_topics):
        # Gradient from dark green (best) to red (worst)
        ratio = i / max(1, len(sorted_topics) - 1)
        r = int(0x15 + ratio * (0xc0 - 0x15))
        g = int(0x57 + (1 - ratio) * (0x57))
        b = int(0x24 + ratio * (0x24))
        topic_colors[topic] = f"{r:02x}{g:02x}{b:02x}"

    # Bar chart by IS score
    items = [(topic, stats["count"], stats["mean_is"])
             for topic, stats in topic_data.items()]
    max_is = max(s[2] for s in items) if items else 1

    # Build chart using IS as bar length (0-5 scale)
    table = doc.add_table(rows=len(items), cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (topic, count, mean_is) in enumerate(items):
        stats = topic_data[topic]
        cell_topic = table.rows[i].cells[0]
        set_cell_text(cell_topic, topic, size=Pt(8), bold=True)
        cell_topic.width = Inches(1.8)

        cell_bar = table.rows[i].cells[1]
        bar_len = max(1, int(mean_is * 5))  # scale 0-5 to 0-25
        bar_text = "\u2588" * bar_len
        cell_bar.text = ""
        p = cell_bar.paragraphs[0]
        run = p.add_run(bar_text)
        run.font.size = Pt(8)
        run.font.name = "Calibri"
        color_hex = topic_colors.get(topic, "2a5a8c")
        run.font.color.rgb = RGBColor(
            int(color_hex[:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16))
        cell_bar.width = Inches(2.0)

        cell_is = table.rows[i].cells[2]
        set_cell_text(cell_is, f"IS {mean_is:.2f}", size=Pt(8))
        cell_is.width = Inches(0.8)

        cell_stats = table.rows[i].cells[3]
        set_cell_text(cell_stats,
            f"N={count}, Cap={stats['captured_pct']}%, "
            f"CtxR={stats['ctx_rule_pct']}%, CtxL={stats['ctx_llm_pct']}%",
            size=Pt(7))
        cell_stats.width = Inches(2.0)

        if i % 2 == 1:
            for cell in table.rows[i].cells:
                set_cell_shading(cell, ZEBRA_BG)

    doc.add_paragraph()

    add_heading(doc, "Full Topic Metrics Table", 2)

    topic_table_rows = []
    topic_highlights = {}
    for i, (topic, stats) in enumerate(topic_data.items()):
        topic_table_rows.append([
            topic, str(stats["count"]),
            fmt(stats["mean_is"], 2), f'{stats["mean_wer"]:.0f}%',
            f'{stats["captured_pct"]}%',
            f'{stats["ctx_rule_pct"]}%', f'{stats["ctx_llm_pct"]}%',
        ])
        if stats["mean_is"] >= 3.0:
            topic_highlights[i] = GREEN_BG
        elif stats["mean_is"] >= 2.5:
            topic_highlights[i] = AMBER_BG
        else:
            topic_highlights[i] = RED_BG

    add_styled_table(doc,
        ["Topic", "N", "Mean IS", "WER", "Captured%", "Ctx Rule%", "Ctx LLM%"],
        topic_table_rows,
        col_widths=[1.8, 0.5, 0.7, 0.6, 0.8, 0.8, 0.8],
        highlight_rows=topic_highlights)

    add_heading(doc, "Key Observations", 2)

    add_bullet(doc,
        "Business/Finance performs best (IS 3.08, 57% captured). "
        "Formal, structured speech with common vocabulary is easier to lip-read.")
    add_bullet(doc,
        "Entertainment and DIY/Home perform worst (IS ~2.2, ~30% captured). "
        "Casual speech, slang, and technical jargon are harder.")
    add_bullet(doc,
        "LLM-based context recovery consistently adds 5\u201315% over rule-based, "
        "especially for topics with domain-specific vocabulary.")
    add_bullet(doc,
        '"Other" (unclassified) segments have the lowest capture rate after DIY and '
        "Entertainment, suggesting diverse/unusual topics are harder.")

    _tight_page_break(doc)


def section_length_analysis(doc, summary):
    add_heading(doc, "Segment Length Analysis", 1)

    length_data = summary.get("length_analysis", {})
    if not length_data:
        add_para(doc, "No length analysis data available.", size=Pt(10))
        return

    add_para(doc,
        "Longer segments give the lip-reading model more visual context to work with. "
        "This section shows how all metrics vary with segment length (measured in reference "
        "word count). All three assessment methods are shown.",
        size=Pt(10))

    add_heading(doc, "Quality by Segment Length", 2)

    # Cumulative filters (>= N words)
    cum_labels = ["all", "gte_5_words", "gte_7_words", "gte_10_words",
                  "gte_15_words", "gte_20_words"]
    cum_display = ["All Segments", "\u2265 5 words", "\u2265 7 words",
                   "\u2265 10 words", "\u2265 15 words", "\u2265 20 words"]

    cum_rows = []
    cum_highlights = {}
    for i, (key, display) in enumerate(zip(cum_labels, cum_display)):
        if key not in length_data:
            continue
        stats = length_data[key]
        cum_rows.append([
            display, str(stats["count"]),
            fmt(stats["mean_is"], 2), f'{stats["mean_wer"]:.0f}%',
            f'{stats["captured_pct"]}%',
            f'{stats["ctx_rule_pct"]}%', f'{stats["ctx_llm_pct"]}%',
        ])
        if stats["captured_pct"] >= 45:
            cum_highlights[len(cum_rows) - 1] = GREEN_BG
        elif stats["captured_pct"] >= 40:
            cum_highlights[len(cum_rows) - 1] = AMBER_BG

    add_styled_table(doc,
        ["Filter", "N", "Mean IS", "WER", "Captured%", "Ctx Rule%", "Ctx LLM%"],
        cum_rows,
        col_widths=[1.3, 0.5, 0.7, 0.6, 0.8, 0.8, 0.8],
        highlight_rows=cum_highlights)

    add_heading(doc, "Quality by Length Band", 2)

    band_labels = ["5_to_10_words", "10_to_15_words", "15_to_20_words", "20_plus_words"]
    band_display = ["5\u201310 words", "10\u201315 words", "15\u201320 words", "20+ words"]

    # Bar chart: captured% by band
    band_items = []
    band_colors = {
        "5\u201310 words": "c0392b",
        "10\u201315 words": "e67e22",
        "15\u201320 words": "f1c40f",
        "20+ words": "27ae60",
    }
    for key, display in zip(band_labels, band_display):
        if key in length_data:
            stats = length_data[key]
            band_items.append((display, stats["count"], stats["captured_pct"]))

    add_bar_chart_table(doc, "Captured % by Length Band", band_items,
                       color_map=band_colors, max_count=60)

    band_rows = []
    band_highlights = {}
    for i, (key, display) in enumerate(zip(band_labels, band_display)):
        if key not in length_data:
            continue
        stats = length_data[key]
        band_rows.append([
            display, str(stats["count"]),
            fmt(stats["mean_is"], 2), f'{stats["mean_wer"]:.0f}%',
            f'{stats["captured_pct"]}%',
            f'{stats["ctx_rule_pct"]}%', f'{stats["ctx_llm_pct"]}%',
        ])
        if stats["captured_pct"] >= 45:
            band_highlights[len(band_rows) - 1] = GREEN_BG
        elif stats["captured_pct"] >= 38:
            band_highlights[len(band_rows) - 1] = AMBER_BG
        else:
            band_highlights[len(band_rows) - 1] = RED_BG

    add_styled_table(doc,
        ["Length Band", "N", "Mean IS", "WER", "Captured%", "Ctx Rule%", "Ctx LLM%"],
        band_rows,
        col_widths=[1.3, 0.5, 0.7, 0.6, 0.8, 0.8, 0.8],
        highlight_rows=band_highlights)

    add_heading(doc, "Key Observations", 2)

    add_bullet(doc,
        "Short segments (5\u201310 words) have 74% WER and only 32% captured. "
        "The model needs visual context \u2014 brief utterances don't provide enough lip movement patterns.")
    add_bullet(doc,
        "Long segments (20+ words) reach 49% captured and 58% LLM-recoverable. "
        "More context helps both the model and any human/LLM trying to interpret the output.")
    add_bullet(doc,
        "The improvement from 5\u201310 words to 20+ words is substantial: "
        "+17pp captured, +16pp LLM-recoverable, -19pp WER.")
    add_bullet(doc,
        "For production use, filtering to segments >= 10 words (1173 segments) "
        "raises the capture rate from 40% to 43% while removing the noisiest short fragments.")

    _tight_page_break(doc)


def section_phonetic_analysis(doc, summary):
    add_heading(doc, "Phonetic Confusion Analysis", 1)

    add_para(doc,
        "Homophenes are words that look identical on a person's lips. A lip-reading system "
        "struggles to distinguish these because the mouth makes the same shape. About 50\u201370% "
        "of English sounds are invisible or ambiguous on lips.",
        size=Pt(10))

    add_heading(doc, "Lip Confusion Groups", 2)

    add_styled_table(doc,
        ["Lip Shape", "Sounds", "Example Words"],
        [
            ["Both lips close", "p, b, m", '"pads"\u2194"pants", "mom"\u2194"bomb", "pill"\u2194"bill"'],
            ["Teeth on lower lip", "f, v", '"fine"\u2194"vine", "few"\u2194"view"'],
            ["Tongue behind teeth", "t, d, n, s, z, l", '"collar"\u2194"color", "time"\u2194"dime"'],
            ["Back of throat", "k, g, h, ng", '"could"\u2194"good"\u2194"hood"'],
            ["Rounded lips", "w, r", '"win"\u2194"rim"'],
            ["Jaw drop (vowels)", "Many vowels", '"bit"\u2194"bet"\u2194"bat"'],
        ],
        col_widths=[1.5, 1.5, 3.6])

    add_heading(doc, "Top 20 Phonetic Confusions in Dataset", 2)

    confusions = summary["top_phonetic_confusions"][:20]
    conf_rows = [[c["ref"], c["hyp"], str(c["count"])] for c in confusions]

    add_styled_table(doc,
        ["Reference Word", "Predicted Word", "Occurrences"],
        conf_rows,
        col_widths=[2.0, 2.0, 1.2])

    add_para(doc,
        "The most common confusions are function word substitutions (the\u2192to, in\u2192and, "
        "a\u2192the) and contraction simplifications (you're\u2192you, it's\u2192is). These are "
        "mostly harmless to intelligibility because function words carry minimal meaning.",
        size=Pt(10))

    add_para(doc,
        "The phonetic similarity signal in the IS composite captures these patterns: "
        "when errors are phonetically close to the correct words, the score rises because "
        "the errors are likely natural lip-reading confusions rather than hallucinations.",
        size=Pt(10))

    _tight_page_break(doc)


def section_proof_examples(doc, rows):
    add_heading(doc, "Proof Examples: WER vs Intelligibility", 1)

    add_para(doc,
        "These examples demonstrate why WER alone is insufficient and how the IS metric "
        "correctly differentiates them.",
        size=Pt(10))

    add_heading(doc, "WER Over-Penalizes (High WER, Still Intelligible)", 2)

    examples_over = [
        ("10%", "4.5+",
         "wellness is 100 in our control simply by thinking positive",
         "is 100 in our control simply by thinking positive",
         'Only "wellness" missing. Fully clear.'),
        ("22%", "4.0+",
         "a friend of mine was recently diagnosed with and died of stage 4 breast cancer and she described",
         "friend of mine was recently diagnosed and died of stage four breast cancer and since she described",
         'Perfect meaning. "Stage 4"\u2192"stage four" is identical.'),
        ("29%", "3.5+",
         "that are going to allow you to work with the team in a more",
         "are going to allow you to work with a team and more",
         "Same idea. Minor articles changed."),
    ]

    add_styled_table(doc,
        ["WER", "IS", "Reference", "Predicted", "Assessment"],
        [[e[0], e[1], e[2][:60] + "...", e[3][:60] + "...", e[4]] for e in examples_over],
        col_widths=[0.5, 0.5, 1.8, 1.8, 2.0],
        highlight_rows={0: GREEN_BG, 1: GREEN_BG, 2: GREEN_BG})

    add_heading(doc, "WER Under-Penalizes (Low WER, Unintelligible)", 2)

    examples_under = [
        ("33%", "1.5\u2013",
         "today i'm talking with admiral mcrae",
         "today i'm talking with animal migratory",
         "Entity destroyed. Person \u2192 biology concept."),
        ("23%", "1.5\u2013",
         "where they tried to suggest that the iphone was your mom's phone",
         "where they tried to suggest that the iphone was your bomb",
         '"mom\'s phone"\u2192"bomb". Marketing \u2192 explosives.'),
    ]

    add_styled_table(doc,
        ["WER", "IS", "Reference", "Predicted", "Assessment"],
        [[e[0], e[1], e[2][:60] + "...", e[3][:60] + "...", e[4]] for e in examples_under],
        col_widths=[0.5, 0.5, 1.8, 1.8, 2.0],
        highlight_rows={0: RED_BG, 1: RED_BG})

    add_heading(doc, "Total Hallucination Examples", 2)

    examples_halluc = [
        ("100%", "0.0",
         "let's see a half a cup of",
         "i have to say thank you",
         "Cooking \u2192 gratitude. Zero connection."),
        ("100%", "0.0",
         "it doesn't have a carry strap but you can put it on your shoulder",
         "this is david irving he's a holocaust denier and a computer hacker",
         "Product review \u2192 harmful fabrication."),
    ]

    add_styled_table(doc,
        ["WER", "IS", "Reference", "Predicted", "Assessment"],
        [[e[0], e[1], e[2][:60] + "...", e[3][:60] + "...", e[4]] for e in examples_halluc],
        col_widths=[0.5, 0.5, 1.8, 1.8, 2.0],
        highlight_rows={0: RED_BG, 1: RED_BG})

    _tight_page_break(doc)


def section_conclusions(doc, summary):
    add_heading(doc, "Conclusions", 1)

    n = summary["n_segments"]
    captured = summary["properly_captured_count"]
    captured_pct = summary["properly_captured_pct"]

    add_heading(doc, "Bottom Line", 2)

    add_bullet_bold_value(doc, "Properly Captured: ",
        f"{captured} of {n} segments ({captured_pct}%) produce output that a human can understand.")
    add_bullet_bold_value(doc, "With Topic Context: ",
        f"Up to {summary['llm_context_recoverable_count']} segments "
        f"({summary['llm_context_recoverable_pct']}%) become understandable "
        f"when the viewer knows the general subject.")
    add_bullet_bold_value(doc, "Hallucinated: ",
        f"{summary['n_empty_hyp']} empty outputs + "
        f"{summary['context_recovery_reasons'].get('hallucination', 0)} hallucinations = "
        f"{summary['n_empty_hyp'] + summary['context_recovery_reasons'].get('hallucination', 0)} "
        f"({(summary['n_empty_hyp'] + summary['context_recovery_reasons'].get('hallucination', 0))/n*100:.1f}%) "
        f"completely unusable segments.")

    add_heading(doc, "Key Insight: Phonetic Confusions Are Mostly Harmless", 2)
    add_para(doc,
        "The most common errors in our dataset are function word substitutions "
        "(the\u2192to, in\u2192and, a\u2192the) and contraction simplifications (you're\u2192you). "
        "These carry minimal semantic weight. The IS metric correctly identifies that a "
        "sentence with 30% WER from function word errors is far more intelligible than "
        "one with 30% WER from entity corruption.",
        size=Pt(10))

    add_heading(doc, "Practical Rule of Thumb", 2)
    add_bullet_bold_value(doc, "WER \u2264 15%: ", "Always understandable (112 segments, 7.5%)")
    add_bullet_bold_value(doc, "WER 15\u201330%: ", "Usually understandable (184 more segments)")
    add_bullet_bold_value(doc, "WER 30\u201350%: ", "Depends on which words are wrong (329 segments)")
    add_bullet_bold_value(doc, "WER 50\u201370%: ", "Usually not understandable (240 segments)")
    add_bullet_bold_value(doc, "WER > 70%: ", "Almost never understandable (632 segments)")

    add_heading(doc, "Recommendation", 2)
    add_para(doc,
        "For production deployment, the IS metric should be used alongside WER to triage "
        "outputs. Segments scoring IS >= 3.0 can be shown to users with confidence. "
        "Segments scoring IS 2.0\u20133.0 should be flagged as 'approximate' and shown "
        "with caveats. Segments scoring IS < 2.0 should be suppressed or replaced with "
        "a 'low confidence' indicator.",
        size=Pt(10))


# =====================================================================
# MAIN
# =====================================================================

def main():
    print("Loading data...")
    rows = load_scores()
    summary = load_summary()
    n = len(rows)
    print(f"  {n} segments loaded")

    print("Creating document...")
    doc = Document()

    # A4 page setup
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # Header and footer
    add_header_footer(doc)

    # Cover page
    create_cover_page(doc, summary)

    # Table of Contents
    toc_titles = [
        "1. Executive Summary",
        "2. Why WER Isn't Enough",
        "3. Methodology: The 6-Signal Composite",
        "4. Tier Distribution",
        "5. Intelligibility Score Distribution",
        "6. Failure Mode Analysis",
        "7. Success Pattern Analysis",
        "8. Topic Analysis",
        "9. Segment Length Analysis",
        "10. WER/WWER Threshold Analysis",
        "11. Context Recovery Analysis",
        "12. Phonetic Confusion Analysis",
        "13. Proof Examples: WER vs Intelligibility",
        "14. Conclusions",
    ]
    add_toc(doc, toc_titles)

    # Sections
    section_executive_summary(doc, summary, rows)
    section_why_wer_isnt_enough(doc)
    section_methodology(doc, summary)
    section_tier_distribution(doc, summary, rows)
    section_is_histogram(doc, summary)
    section_failure_analysis(doc, summary, rows)
    section_success_analysis(doc, summary, rows)
    section_topic_analysis(doc, summary)
    section_length_analysis(doc, summary)
    section_wer_thresholds(doc, rows, summary)
    section_context_recovery(doc, summary, rows)
    section_phonetic_analysis(doc, summary)
    section_proof_examples(doc, rows)
    section_conclusions(doc, summary)

    # Save
    OUTPUT_FILE.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_FILE))
    print(f"  Saved: {OUTPUT_FILE}")
    print("Done.")


if __name__ == "__main__":
    main()
