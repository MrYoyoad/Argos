#!/usr/bin/env python3
"""
Convert analysis markdown files to branded Argos .docx documents.

Parses markdown generically (headings, tables, code blocks, bullets, paragraphs)
and applies the Argos -- The Orchard style guide (cover page, header/footer,
Calibri font, colored headings, styled tables with zebra striping).

Usage:
    python3 convert_analysis_docx.py           # Convert all missing analysis files
    python3 convert_analysis_docx.py --file X  # Convert a single file

Output:
    .docx files alongside their source .md files
"""

import re
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

# -- Paths --
SCRIPT_DIR = Path(__file__).parent
ASSETS_DIR = SCRIPT_DIR.parent / "assets"
LOGO_ORCHARD = ASSETS_DIR / "logo.png"
LOGO_PEACOCK = ASSETS_DIR / "peacock.png"

# -- Colors --
C_PRIMARY = RGBColor(0x1A, 0x3A, 0x5C)
C_H2 = RGBColor(0x2A, 0x5A, 0x8C)
C_H3 = RGBColor(0x3A, 0x6A, 0x9C)
C_H4 = RGBColor(0x4A, 0x7A, 0xAC)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_DARK = RGBColor(0x33, 0x33, 0x33)
C_GRAY = RGBColor(0x66, 0x66, 0x66)
C_CODE = RGBColor(0x2D, 0x2D, 0x2D)

HEADER_BG = "1a3a5c"
ZEBRA_BG = "f0f4f8"
CODE_BG = "f5f5f5"

LAST_UPDATED = datetime.now().strftime("%B %d, %Y")

# -- Analysis files to convert --
ANALYSIS_FILES = [
    {
        "md_path": Path("/home/ubuntu/docs/evaluation/llm_upgrade_analysis.md"),
        "docx_path": Path("/home/ubuntu/docs/evaluation/llm_upgrade_analysis.docx"),
        "title": "Impact of a Stronger LLM on VSP-LLM Pipeline Performance",
        "subtitle": "LLM Alternatives, Data Scaling, Multilingual Models & Prompt Strategies",
        "version": "Analysis Report v1.0",
    },
    {
        "md_path": Path("/home/ubuntu/docs/evaluation/is_correlation_analysis.md"),
        "docx_path": Path("/home/ubuntu/docs/evaluation/is_correlation_analysis.docx"),
        "title": "IS Metric — Component Correlation Analysis",
        "subtitle": "Variance Decomposition, Signal Redundancy & Cross-Config Stability",
        "version": "Analysis Report v1.0",
    },
    {
        "md_path": Path("/home/ubuntu/docs/finetuning/experiments/comparison_report.md"),
        "docx_path": Path("/home/ubuntu/docs/finetuning/experiments/comparison_report.docx"),
        "title": "Weekend Fine-Tuning Comparison Report",
        "subtitle": "Baseline vs Exp A (r=16) vs Exp B (r=64) — AVSpeech Domain Adaptation",
        "version": "Experiment Report v1.0",
    },
    {
        "md_path": Path("/home/ubuntu/docs/evaluation/baseline_vs_J_vs_C_intelligibility.md"),
        "docx_path": Path("/home/ubuntu/docs/evaluation/baseline_vs_J_vs_C_intelligibility.docx"),
        "title": "Intelligibility Comparison: Baseline vs Config J vs Config C",
        "subtitle": "Decode Configuration Impact on Intelligibility Score Distribution",
        "version": "Analysis Report v1.0",
    },
]


# =====================================================================
# HELPER FUNCTIONS (matching Argos style guide)
# =====================================================================

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, bold=False, color=None, size=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    _add_formatted_runs(p, str(text))
    for run in p.runs:
        if bold:
            run.bold = True
        if color:
            run.font.color.rgb = color
        if size:
            run.font.size = size
        run.font.name = "Calibri"


def _add_formatted_runs(paragraph, text):
    """Add text with inline formatting (bold, code, italic) to a paragraph."""
    parts = re.split(r'(\*\*.*?\*\*|`[^`]+`|\*[^*]+\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = "Calibri"
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x66, 0x33, 0x99)
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            run.font.name = "Calibri"
        else:
            run = paragraph.add_run(part)
            run.font.name = "Calibri"


def add_styled_table(doc, headers, rows, col_widths=None):
    if not rows:
        return None
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, HEADER_BG)
        set_cell_text(cell, h, bold=True, color=C_WHITE, size=Pt(8))

    for r_idx, row_data in enumerate(rows):
        for c_idx in range(min(len(row_data), n_cols)):
            cell = table.rows[r_idx + 1].cells[c_idx]
            set_cell_text(cell, row_data[c_idx], size=Pt(8))
            if r_idx % 2 == 1:
                set_cell_shading(cell, ZEBRA_BG)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return table


def add_heading_styled(doc, text, level, color=None):
    h = doc.add_heading(text, level=level)
    if color is None:
        color = {1: C_PRIMARY, 2: C_H2, 3: C_H3, 4: C_H4}.get(level, C_PRIMARY)
    for run in h.runs:
        run.font.color.rgb = color
        run.font.name = "Calibri"
    return h


def add_para(doc, text, bold=False, italic=False, size=Pt(10),
             color=None, space_after=Pt(4)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = size
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = space_after
    return p


def add_formatted_para(doc, text, size=Pt(10), space_after=Pt(4)):
    """Add paragraph with inline markdown formatting (bold, code, italic)."""
    p = doc.add_paragraph()
    _add_formatted_runs(p, text)
    for run in p.runs:
        if run.font.size is None:
            run.font.size = size
        if run.font.name is None:
            run.font.name = "Calibri"
    p.paragraph_format.space_after = space_after
    return p


def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.2)

    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{CODE_BG}" w:val="clear"/>')
    pPr.append(shading)

    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(7.5)
    run.font.color.rgb = C_CODE
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.text = ""
    _add_formatted_runs(p, text)
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        if run.font.size is None:
            run.font.size = Pt(9)
        if run.font.name is None or run.font.name == "Calibri":
            run.font.name = "Calibri"
    return p


def _build_inline_image_xml(rId, cx, cy, pic_id=1, name="Logo"):
    """Build inline image XML for header embedding."""
    nsmap = (
        'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
        'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
        'xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"'
    )
    xml = (
        f'<w:drawing {nsdecls("w")} {nsmap}>'
        f'<wp:inline distT="0" distB="0" distL="0" distR="0">'
        f'<wp:extent cx="{cx}" cy="{cy}"/>'
        f'<wp:docPr id="{pic_id}" name="{name}"/>'
        f'<a:graphic>'
        f'<a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        f'<pic:pic>'
        f'<pic:nvPicPr><pic:cNvPr id="{pic_id}" name="{name}"/><pic:cNvPicPr/></pic:nvPicPr>'
        f'<pic:blipFill><a:blip r:embed="{rId}"/><a:stretch><a:fillRect/></a:stretch></pic:blipFill>'
        f'<pic:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm>'
        f'<a:prstGeom prst="rect"><a:avLst/></a:prstGeom></pic:spPr>'
        f'</pic:pic>'
        f'</a:graphicData>'
        f'</a:graphic>'
        f'</wp:inline>'
        f'</w:drawing>'
    )
    return parse_xml(xml)


def add_header_footer(doc):
    """Add header with logo + 'Argos -- The Orchard' and footer with page numbers."""
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    section = doc.sections[0]
    section.different_first_page_header_footer = True

    # -- Header (non-first pages) --
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]

    if LOGO_ORCHARD.exists():
        try:
            image_part = doc.part.package.get_or_add_image_part(str(LOGO_ORCHARD))
            rId = header.part.relate_to(image_part, RT.IMAGE)
            size_emu = int(0.3 * 914400)
            logo_run = hp.add_run()
            drawing = _build_inline_image_xml(rId, size_emu, size_emu,
                                              pic_id=10, name="Header Logo")
            logo_run._r.append(drawing)
        except Exception:
            pass

    text_width = section.page_width - section.left_margin - section.right_margin
    hp.paragraph_format.tab_stops.add_tab_stop(text_width, WD_TAB_ALIGNMENT.RIGHT)
    hp.add_run("\t")
    run = hp.add_run("Argos \u2014 The Orchard")
    run.font.size = Pt(8)
    run.font.color.rgb = C_GRAY
    run.font.name = "Calibri"
    run.italic = True

    # -- Footer --
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = fp.add_run("Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = C_GRAY
    run.font.name = "Calibri"

    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    run._r.append(fld1)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run._r.append(instr)
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld2)


def create_cover_page(doc, title, subtitle, version):
    """Create branded Argos cover page."""
    for _ in range(2):
        doc.add_paragraph()

    if LOGO_PEACOCK.exists():
        doc.add_picture(str(LOGO_PEACOCK), width=Inches(2.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ARGOS title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ARGOS")
    run.font.size = Pt(48)
    run.font.color.rgb = C_PRIMARY
    run.bold = True
    run.font.name = "Calibri"

    # Document title
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(title)
    run2.font.size = Pt(18)
    run2.font.color.rgb = C_H2
    run2.font.name = "Calibri"

    # The Orchard
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("The Orchard")
    run3.font.size = Pt(20)
    run3.font.color.rgb = C_H2
    run3.font.name = "Calibri"

    # Version
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run(version)
    run4.font.size = Pt(14)
    run4.font.color.rgb = C_H3
    run4.italic = True
    run4.font.name = "Calibri"

    for _ in range(2):
        doc.add_paragraph()

    # Metadata
    info_lines = [
        ("Subtitle:", subtitle),
        ("Lead Engineer:", "Yoad Oxman"),
        ("Last Updated:", LAST_UPDATED),
    ]
    for label, value in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rl = p.add_run(label + " ")
        rl.bold = True
        rl.font.size = Pt(11)
        rl.font.color.rgb = C_DARK
        rl.font.name = "Calibri"
        rv = p.add_run(value)
        rv.font.size = Pt(11)
        rv.font.color.rgb = C_DARK
        rv.font.name = "Calibri"

    doc.add_page_break()


# =====================================================================
# MARKDOWN PARSER
# =====================================================================

def parse_markdown_to_docx(doc, md_text):
    """Parse markdown text and add content to a docx Document."""
    lines = md_text.split('\n')
    i = 0
    in_code_block = False
    code_block_lines = []
    in_table = False
    table_header = []
    table_rows = []

    while i < len(lines):
        line = lines[i]
        stripped = line.rstrip()

        # --- Code blocks ---
        if stripped.startswith('```'):
            if in_code_block:
                code_text = '\n'.join(code_block_lines)
                add_code_block(doc, code_text)
                code_block_lines = []
                in_code_block = False
            else:
                _flush_table(doc, table_header, table_rows)
                table_header = []
                table_rows = []
                in_table = False
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_block_lines.append(line.rstrip())
            i += 1
            continue

        # --- Table rows ---
        if '|' in stripped and stripped.startswith('|'):
            # Skip separator rows (|---|---|)
            if re.match(r'^\|[\s\-:|]+\|$', stripped):
                i += 1
                continue
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            if cells:
                if not in_table:
                    in_table = True
                    table_header = cells
                else:
                    table_rows.append(cells)
            i += 1
            continue
        else:
            if in_table:
                _flush_table(doc, table_header, table_rows)
                table_header = []
                table_rows = []
                in_table = False

        # --- Empty lines ---
        if not stripped:
            i += 1
            continue

        # --- Horizontal rules ---
        if stripped in ('---', '***', '___'):
            # Only add a thin separator, not a full page break
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run('\u2500' * 60)
            run.font.size = Pt(6)
            run.font.color.rgb = C_GRAY
            i += 1
            continue

        # --- Headings ---
        heading_match = re.match(r'^(#{1,4})\s+(.*)', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            # Remove inline bold markers from heading text
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            add_heading_styled(doc, text, level=level)
            i += 1
            continue

        # --- Bullet/list items ---
        bullet_match = re.match(r'^(\s*)[*\-+]\s+(.*)', stripped)
        if bullet_match:
            indent = len(bullet_match.group(1))
            text = bullet_match.group(2)
            level = min(indent // 2, 2)
            add_bullet(doc, text, level=level)
            i += 1
            continue

        # --- Numbered list items ---
        numbered_match = re.match(r'^(\s*)\d+[.)]\s+(.*)', stripped)
        if numbered_match:
            text = numbered_match.group(2)
            indent = len(numbered_match.group(1))
            level = min(indent // 2, 2)
            add_bullet(doc, text, level=level)
            i += 1
            continue

        # --- Regular paragraph ---
        add_formatted_para(doc, stripped, size=Pt(10), space_after=Pt(4))
        i += 1

    # Flush remaining table
    if in_table:
        _flush_table(doc, table_header, table_rows)


def _flush_table(doc, header, rows):
    """Flush accumulated table rows into a styled table."""
    if not header:
        return
    add_styled_table(doc, header, rows)


# =====================================================================
# MAIN CONVERSION
# =====================================================================

def convert_md_to_docx(md_path, docx_path, title, subtitle, version):
    """Convert a single markdown file to a branded Argos docx."""
    print(f"  Reading: {md_path}")
    md_text = md_path.read_text(encoding='utf-8')

    doc = Document()

    # Page setup (A4)
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    # Header and footer
    add_header_footer(doc)

    # Cover page
    create_cover_page(doc, title, subtitle, version)

    # Parse markdown content
    parse_markdown_to_docx(doc, md_text)

    # Save
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    size_kb = docx_path.stat().st_size / 1024
    print(f"  Saved:   {docx_path} ({size_kb:.0f} KB)")


def main():
    print(f"Argos Analysis DOCX Converter")
    print(f"{'=' * 50}")
    print(f"Date: {LAST_UPDATED}\n")

    # Filter to specific file if --file argument provided
    targets = ANALYSIS_FILES
    if len(sys.argv) > 2 and sys.argv[1] == '--file':
        file_filter = sys.argv[2]
        targets = [f for f in ANALYSIS_FILES if file_filter in str(f['md_path'])]
        if not targets:
            print(f"No matching file found for: {file_filter}")
            sys.exit(1)

    converted = 0
    for entry in targets:
        md_path = entry['md_path']
        if not md_path.exists():
            print(f"  SKIP: {md_path} not found")
            continue

        print(f"\nConverting: {md_path.name}")
        try:
            convert_md_to_docx(
                md_path=md_path,
                docx_path=entry['docx_path'],
                title=entry['title'],
                subtitle=entry['subtitle'],
                version=entry['version'],
            )
            converted += 1
        except Exception as e:
            print(f"  FAILED: {e}")
            import traceback
            traceback.print_exc()

    print(f"\n{'=' * 50}")
    print(f"Done! Converted {converted}/{len(targets)} files.")


if __name__ == '__main__':
    main()
