#!/usr/bin/env python3
"""
Argos — The Orchard — Teammate Briefing Docx Generator

Renders docs/guides/teammate-briefing-aug2026.md into a branded .docx
following the house docx conventions in STYLE_GUIDE.md (Calibri, A4,
header logo + right-tab brand line, cover page, Word TOC field, styled
tables, page-number footer).

The generator is a generic markdown renderer — it parses the .md at run
time (headings, paragraphs with bold/italic/inline-code, tables, bullet
and numbered lists, fenced code blocks, links) so it can simply be re-run
after any briefing edit. No content is hardcoded.

Relative markdown links are rendered as visible text with the target path
in parentheses (resolved to a repo-root-relative path where possible),
because the .docx travels off-repo where hyperlinks would be dead.

Interpreter: system `python3` (Python 3.x with python-docx 1.2.0 installed
site-wide on this EC2 box). Verified 2026-08-03; neither project venv is
required (/home/ubuntu/vsp-llm-yoad-venv also has python-docx 1.2.0 and
works, auto_avsr/pre-process-venv does NOT have python-docx).

Usage:
    python3 generate_teammate_briefing.py [--src <briefing.md>] [--out <out.docx>]

Default output:
    docs/guides/teammate-briefing-aug2026.docx
"""

import argparse
import datetime
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

# ── Paths ──
SCRIPT_DIR = Path(__file__).resolve().parent
REPO_ROOT = SCRIPT_DIR.parent.parent.parent            # /home/ubuntu
DEFAULT_SRC = REPO_ROOT / "docs" / "guides" / "teammate-briefing-aug2026.md"
DEFAULT_OUT = REPO_ROOT / "docs" / "guides" / "teammate-briefing-aug2026.docx"

# ── Logos ──
ASSETS_DIR = SCRIPT_DIR.parent / "assets"
LOGO_ORCHARD = ASSETS_DIR / "logo.png"
LOGO_PEACOCK = ASSETS_DIR / "peacock.png"

# ── Colors (house palette, matches generate_q2_summary.py) ──
C_PRIMARY = RGBColor(0x1a, 0x3a, 0x5c)
C_H2 = RGBColor(0x2a, 0x5a, 0x8c)
C_H3 = RGBColor(0x3a, 0x6a, 0x9c)
C_H4 = RGBColor(0x4a, 0x7a, 0xac)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_DARK = RGBColor(0x33, 0x33, 0x33)
C_GRAY = RGBColor(0x66, 0x66, 0x66)
C_CODE = RGBColor(0x2d, 0x2d, 0x2d)
C_LINK = RGBColor(0x2a, 0x5a, 0x8c)

HEADER_BG = "1a3a5c"
ZEBRA_BG = "f0f4f8"
CODE_BG = "f5f5f5"

MONO_FONT = "Consolas"
TEXT_WIDTH_IN = 6.3   # A4 (21.0 cm) minus 2 x 2.5 cm side margins ≈ 16 cm


# ═══════════════════════════════════════════════════
# MARKDOWN PARSING
# ═══════════════════════════════════════════════════

def parse_table_block(lines):
    """Return (headers, rows) from a run of '|'-prefixed lines."""
    def split_row(line):
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        return cells

    sep_re = re.compile(r"^:?-{3,}:?$")
    rows = [split_row(l) for l in lines]
    headers = rows[0]
    body = []
    for r in rows[1:]:
        if all(sep_re.match(c) for c in r if c != "") and any(r):
            continue  # alignment separator row
        body.append(r)
    # pad ragged rows
    ncol = len(headers)
    body = [(r + [""] * ncol)[:ncol] for r in body]
    return headers, body


def parse_markdown(text):
    """Parse markdown text into a list of (kind, payload) blocks.

    Kinds: heading (level, text), para (text), table (headers, rows),
    ul [(indent_level, text)...], ol [(number, text)...], code [lines], hr.
    """
    lines = text.split("\n")
    blocks = []
    para_buf = []
    i = 0

    def flush_para():
        if para_buf:
            blocks.append(("para", " ".join(para_buf).strip()))
            para_buf.clear()

    ul_re = re.compile(r"^(\s*)[-*+]\s+(.*)$")
    ol_re = re.compile(r"^(\s*)(\d+)[.)]\s+(.*)$")

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            flush_para()
            i += 1
            code = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code.append(lines[i])
                i += 1
            blocks.append(("code", code))
            i += 1
            continue

        if not stripped:
            flush_para()
            i += 1
            continue

        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            flush_para()
            blocks.append(("heading", (len(m.group(1)), m.group(2).strip())))
            i += 1
            continue

        if re.match(r"^(-{3,}|\*{3,}|_{3,})$", stripped):
            flush_para()
            blocks.append(("hr", None))
            i += 1
            continue

        if stripped.startswith("|"):
            flush_para()
            tbl = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                tbl.append(lines[i])
                i += 1
            blocks.append(("table", parse_table_block(tbl)))
            continue

        m = ul_re.match(line)
        if m:
            flush_para()
            items = []
            while i < len(lines):
                mm = ul_re.match(lines[i])
                if not mm:
                    break
                items.append((len(mm.group(1)) // 2, mm.group(2).strip()))
                i += 1
            blocks.append(("ul", items))
            continue

        m = ol_re.match(line)
        if m:
            flush_para()
            items = []
            while i < len(lines):
                mm = ol_re.match(lines[i])
                if not mm:
                    break
                items.append((mm.group(2), mm.group(3).strip()))
                i += 1
            blocks.append(("ol", items))
            continue

        para_buf.append(stripped)
        i += 1

    flush_para()
    return blocks


def strip_inline_md(text):
    """Plain-text version of an inline-markdown string (for headings/TOC)."""
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    return text.strip()


# ═══════════════════════════════════════════════════
# INLINE RENDERING (bold / italic / code / links)
# ═══════════════════════════════════════════════════

INLINE_RE = re.compile(
    r"(?P<code>`[^`]+`)"
    r"|(?P<link>\[[^\]]+\]\([^)]+\))"
    r"|(?P<bold>\*\*.+?\*\*)"
    r"|(?P<italic>\*[^*\s].*?\*)"
)
LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def display_target(target, md_dir):
    """Human-readable link target for the parenthesized suffix.

    Relative paths resolve against the md file's directory and are shown
    repo-root-relative when possible (the docx travels off-repo, so a
    readable path beats '../../..'). http/mailto shown as-is; pure
    in-document anchors are dropped.
    """
    target = target.strip()
    if re.match(r"^(https?:|mailto:)", target):
        return target
    if target.startswith("#"):
        return None
    path_part = target.split("#", 1)[0]
    if not path_part:
        return None
    try:
        resolved = (md_dir / path_part).resolve()
        return str(resolved.relative_to(REPO_ROOT))
    except (ValueError, OSError):
        return path_part


def add_run(p, text, size, bold=False, italic=False, color=None, font=None):
    if not text:
        return
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = size
    run.font.name = font or "Calibri"
    if color:
        run.font.color.rgb = color


def add_inline(p, text, size, md_dir, bold=False, italic=False, color=None):
    """Render inline markdown into runs on paragraph p."""
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            add_run(p, text[pos:m.start()], size, bold, italic, color)
        if m.group("code"):
            add_run(p, m.group("code")[1:-1], size, bold, italic, C_CODE, MONO_FONT)
        elif m.group("link"):
            lm = LINK_RE.match(m.group("link"))
            label, target = lm.group(1), lm.group(2)
            add_inline(p, label, size, md_dir, bold=bold, italic=italic, color=C_LINK)
            disp = display_target(target, md_dir)
            if disp and disp != strip_inline_md(label):
                add_run(p, f" ({disp})", size, False, True, C_GRAY)
        elif m.group("bold"):
            add_inline(p, m.group("bold")[2:-2], size, md_dir, bold=True,
                       italic=italic, color=color)
        elif m.group("italic"):
            add_inline(p, m.group("italic")[1:-1], size, md_dir, bold=bold,
                       italic=True, color=color)
        pos = m.end()
    add_run(p, text[pos:], size, bold, italic, color)


# ═══════════════════════════════════════════════════
# DOCX HELPERS (house style)
# ═══════════════════════════════════════════════════

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def shade_paragraph(p, color_hex):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color_hex)
    pPr.append(shd)


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    color = {1: C_PRIMARY, 2: C_H2, 3: C_H3, 4: C_H4}.get(level, C_PRIMARY)
    for run in h.runs:
        run.font.color.rgb = color
        run.font.name = "Calibri"
    return h


def add_styled_table(doc, headers, rows, md_dir):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Column widths proportional to content length, within the text width.
    weights = []
    for c in range(len(headers)):
        longest = max([len(strip_inline_md(headers[c]))] +
                      [len(strip_inline_md(r[c])) for r in rows] or [1])
        weights.append(max(6, min(longest, 60)))
    total = sum(weights)
    widths = [Inches(TEXT_WIDTH_IN * w / total) for w in weights]

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, HEADER_BG)
        cell.text = ""
        p = cell.paragraphs[0]
        add_inline(p, strip_inline_md(h), Pt(9), md_dir, bold=True, color=C_WHITE)

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            add_inline(p, val, Pt(9), md_dir)
            if r_idx % 2 == 1:
                set_cell_shading(cell, ZEBRA_BG)

    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = w

    doc.add_paragraph()
    return table


def add_code_block(doc, lines):
    for line in (lines or [""]):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent = Inches(0.2)
        shade_paragraph(p, CODE_BG)
        run = p.add_run(line if line else " ")
        run.font.name = MONO_FONT
        run.font.size = Pt(9)
        run.font.color.rgb = C_CODE
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _build_inline_image_xml(rId, cx, cy, pic_id=1, name="Logo"):
    inline = OxmlElement("wp:inline")
    for attr in ("distT", "distB", "distL", "distR"):
        inline.set(attr, "0")
    extent = OxmlElement("wp:extent")
    extent.set("cx", str(cx))
    extent.set("cy", str(cy))
    inline.append(extent)
    docPr = OxmlElement("wp:docPr")
    docPr.set("id", str(pic_id))
    docPr.set("name", name)
    inline.append(docPr)
    graphic = OxmlElement("a:graphic")
    graphicData = OxmlElement("a:graphicData")
    graphicData.set("uri", "http://schemas.openxmlformats.org/drawingml/2006/picture")
    pic = OxmlElement("pic:pic")
    nvPicPr = OxmlElement("pic:nvPicPr")
    cNvPr = OxmlElement("pic:cNvPr")
    cNvPr.set("id", "0")
    cNvPr.set("name", name)
    nvPicPr.append(cNvPr)
    nvPicPr.append(OxmlElement("pic:cNvPicPr"))
    pic.append(nvPicPr)
    blipFill = OxmlElement("pic:blipFill")
    blip = OxmlElement("a:blip")
    blip.set(qn("r:embed"), rId)
    blipFill.append(blip)
    stretch = OxmlElement("a:stretch")
    stretch.append(OxmlElement("a:fillRect"))
    blipFill.append(stretch)
    pic.append(blipFill)
    spPr = OxmlElement("pic:spPr")
    xfrm = OxmlElement("a:xfrm")
    off = OxmlElement("a:off")
    off.set("x", "0")
    off.set("y", "0")
    xfrm.append(off)
    ext_el = OxmlElement("a:ext")
    ext_el.set("cx", str(cx))
    ext_el.set("cy", str(cy))
    xfrm.append(ext_el)
    spPr.append(xfrm)
    prstGeom = OxmlElement("a:prstGeom")
    prstGeom.set("prst", "rect")
    spPr.append(prstGeom)
    pic.append(spPr)
    graphicData.append(pic)
    graphic.append(graphicData)
    inline.append(graphic)
    drawing = OxmlElement("w:drawing")
    drawing.append(inline)
    return drawing


def add_header_footer(doc, brand_suffix):
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    section = doc.sections[0]
    section.different_first_page_header_footer = True

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]

    if LOGO_ORCHARD.exists():
        image_part = doc.part.package.get_or_add_image_part(str(LOGO_ORCHARD))
        rId = header.part.relate_to(image_part, RT.IMAGE)
        size_emu = int(0.3 * 914400)
        logo_run = hp.add_run()
        drawing = _build_inline_image_xml(rId, size_emu, size_emu,
                                          pic_id=10, name="Header Logo")
        logo_run._r.append(drawing)
    text_width = section.page_width - section.left_margin - section.right_margin
    hp.paragraph_format.tab_stops.add_tab_stop(text_width, WD_TAB_ALIGNMENT.RIGHT)
    hp.add_run("\t")
    run = hp.add_run(f"Argos — The Orchard  ·  {brand_suffix}")
    run.font.size = Pt(8)
    run.font.color.rgb = C_GRAY
    run.font.name = "Calibri"
    run.italic = True

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = C_GRAY
    run.font.name = "Calibri"
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    run._r.append(fld1)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run._r.append(instr)
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld2)


def create_cover_page(doc, title, src_path):
    for _ in range(2):
        doc.add_paragraph()

    if LOGO_PEACOCK.exists():
        doc.add_picture(str(LOGO_PEACOCK), width=Inches(2.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ARGOS")
    run.font.size = Pt(48)
    run.font.color.rgb = C_PRIMARY
    run.bold = True
    run.font.name = "Calibri"

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Visual Speech Processing System")
    run2.font.size = Pt(22)
    run2.font.color.rgb = C_H2
    run2.font.name = "Calibri"

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("The Orchard")
    run3.font.size = Pt(20)
    run3.font.color.rgb = C_H2
    run3.font.name = "Calibri"

    doc.add_paragraph()

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run(title)
    run4.font.size = Pt(16)
    run4.font.color.rgb = C_H3
    run4.italic = True
    run4.font.name = "Calibri"

    for _ in range(2):
        doc.add_paragraph()

    info_lines = [
        ("Author:", "Yoad Oxman"),
        ("Source:", str(src_path.name)),
        ("Generated:", datetime.date.today().strftime("%B %d, %Y")),
    ]
    for label, value in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rl = p.add_run(label + " ")
        rl.bold = True
        rl.font.size = Pt(11)
        rl.font.color.rgb = C_DARK
        rl.font.name = "Calibri"
        rv = p.add_run(value)
        rv.font.size = Pt(11)
        rv.font.color.rgb = C_DARK
        rv.font.name = "Calibri"

    doc.add_page_break()


def create_toc(doc, toc_titles):
    for _ in range(3):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    add_heading(doc, "Table of Contents", 1)

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-2" \\h \\z \\u '
    run._r.append(instr)
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    for title in toc_titles:
        placeholder = paragraph.add_run(title + "\n")
        placeholder.font.size = Pt(11)
        placeholder.font.name = "Calibri"
        placeholder.font.color.rgb = C_PRIMARY

    fld_end_run = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    fld_end_run._r.append(fld_end)

    settings = doc.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)

    doc.add_page_break()


# ═══════════════════════════════════════════════════
# BODY RENDERING
# ═══════════════════════════════════════════════════

def render_blocks(doc, blocks, md_dir):
    first_h1_skipped = False
    for kind, payload in blocks:
        if kind == "heading":
            level, text = payload
            text_plain = strip_inline_md(text)
            if level == 1 and not first_h1_skipped:
                first_h1_skipped = True   # doc title lives on the cover page
                continue
            # md ## -> docx H1 etc. (md's single H1 is the cover title)
            add_heading(doc, text_plain, max(1, min(level - 1, 4)))
        elif kind == "para":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            add_inline(p, payload, Pt(11), md_dir)
        elif kind == "table":
            headers, rows = payload
            add_styled_table(doc, headers, rows, md_dir)
        elif kind == "ul":
            for indent, text in payload:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.left_indent = Inches(0.25 + indent * 0.25)
                p.paragraph_format.space_after = Pt(2)
                add_inline(p, text, Pt(10), md_dir)
        elif kind == "ol":
            for number, text in payload:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.45)
                p.paragraph_format.first_line_indent = Inches(-0.2)
                p.paragraph_format.space_after = Pt(2)
                add_run(p, f"{number}.  ", Pt(10), bold=True, color=C_PRIMARY)
                add_inline(p, text, Pt(10), md_dir)
        elif kind == "code":
            add_code_block(doc, payload)
        elif kind == "hr":
            continue  # section rules add no value in the docx


# ═══════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════

def main():
    ap = argparse.ArgumentParser(description=__doc__.splitlines()[1])
    ap.add_argument("--src", type=Path, default=DEFAULT_SRC,
                    help="Source markdown file (default: teammate-briefing-aug2026.md)")
    ap.add_argument("--out", type=Path, default=DEFAULT_OUT,
                    help="Output .docx path (default: alongside the source)")
    args = ap.parse_args()

    text = args.src.read_text(encoding="utf-8")
    md_dir = args.src.resolve().parent
    blocks = parse_markdown(text)

    title = args.src.stem
    for k, p in blocks:
        if k == "heading" and p[0] == 1:
            title = strip_inline_md(p[1])
            break
    toc_titles = [strip_inline_md(p[1]) for k, p in blocks
                  if k == "heading" and p[0] == 2]

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)

    create_cover_page(doc, title, args.src)
    create_toc(doc, toc_titles)
    add_header_footer(doc, "Teammate Briefing")

    render_blocks(doc, blocks, md_dir)

    args.out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(args.out))
    print(f"✓ Saved: {args.out}")
    print(f"  Paragraphs: {len(doc.paragraphs)}  Tables: {len(doc.tables)}")


if __name__ == "__main__":
    main()
