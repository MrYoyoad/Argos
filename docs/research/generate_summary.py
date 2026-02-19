#!/usr/bin/env python3
"""
Argos — The Orchard — Condensed Project Summary Generator

Generates a 10-page branded .docx summarizing the entire Argos project
in 4 parts: Research, Operations, Technical Development, Lessons & To-Do.

Usage:
    python3 generate_summary.py

Output:
    argos_research/project_summary.docx
"""

import os
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

# ── Output ──
OUTPUT_DIR = Path(__file__).parent
OUTPUT_FILE = OUTPUT_DIR / "project_summary.docx"

# ── Logos ──
LOGO_ORCHARD = OUTPUT_DIR / "logo.png"
LOGO_PEACOCK = OUTPUT_DIR / "peacock.png"

# ── Colors ──
C_PRIMARY = RGBColor(0x1a, 0x3a, 0x5c)
C_H2 = RGBColor(0x2a, 0x5a, 0x8c)
C_H3 = RGBColor(0x3a, 0x6a, 0x9c)
C_H4 = RGBColor(0x4a, 0x7a, 0xac)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_BLACK = RGBColor(0x00, 0x00, 0x00)
C_DARK = RGBColor(0x33, 0x33, 0x33)
C_GREEN = RGBColor(0x15, 0x57, 0x24)
C_RED = RGBColor(0x72, 0x1c, 0x24)
C_AMBER = RGBColor(0x85, 0x64, 0x04)
C_GRAY = RGBColor(0x66, 0x66, 0x66)
C_CODE = RGBColor(0x2d, 0x2d, 0x2d)

HEADER_BG = "1a3a5c"
ZEBRA_BG = "f0f4f8"
GREEN_BG = "d4edda"
RED_BG = "f8d7da"
AMBER_BG = "fff3cd"
CODE_BG = "f5f5f5"

LAST_UPDATED = datetime.now().strftime("%B %d, %Y")


# ═══════════════════════════════════════════════════
# HELPER FUNCTIONS (reused from generate_research_journal.py)
# ═══════════════════════════════════════════════════

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, bold=False, color=None, size=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = size
    run.font.name = "Calibri"


def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, HEADER_BG)
        set_cell_text(cell, h, bold=True, color=C_WHITE, size=Pt(9))

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            set_cell_text(cell, val, size=Pt(9))
            if r_idx % 2 == 1:
                set_cell_shading(cell, ZEBRA_BG)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return table


_bookmark_counter = [0]

def add_heading(doc, text, level, color=None, bookmark_id=None):
    h = doc.add_heading(text, level=level)
    if color is None:
        color = {1: C_PRIMARY, 2: C_H2, 3: C_H3, 4: C_H4}.get(level, C_PRIMARY)
    for run in h.runs:
        run.font.color.rgb = color
        run.font.name = "Calibri"
    if bookmark_id:
        _bookmark_counter[0] += 1
        bm_start = OxmlElement("w:bookmarkStart")
        bm_start.set(qn("w:id"), str(_bookmark_counter[0]))
        bm_start.set(qn("w:name"), bookmark_id)
        bm_end = OxmlElement("w:bookmarkEnd")
        bm_end.set(qn("w:id"), str(_bookmark_counter[0]))
        h._p.insert(0, bm_start)
        h._p.append(bm_end)
    return h


def add_para(doc, text, bold=False, italic=False, size=Pt(11), color=None, space_after=Pt(6)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = size
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = space_after
    return p


def add_para_rich(doc, segments, space_after=Pt(6)):
    p = doc.add_paragraph()
    for text, bold in segments:
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(11)
        run.font.name = "Calibri"
    p.paragraph_format.space_after = space_after
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = "Calibri"
    return p


def add_bullet_bold_value(doc, label, value, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after = Pt(2)
    run_b = p.add_run(label)
    run_b.bold = True
    run_b.font.size = Pt(10)
    run_b.font.name = "Calibri"
    run_v = p.add_run(value)
    run_v.font.size = Pt(10)
    run_v.font.name = "Calibri"
    return p


def _build_inline_image_xml(rId, cx, cy, pic_id=1, name="Logo"):
    inline = OxmlElement('wp:inline')
    inline.set('distT', '0')
    inline.set('distB', '0')
    inline.set('distL', '0')
    inline.set('distR', '0')

    extent = OxmlElement('wp:extent')
    extent.set('cx', str(cx))
    extent.set('cy', str(cy))
    inline.append(extent)

    docPr = OxmlElement('wp:docPr')
    docPr.set('id', str(pic_id))
    docPr.set('name', name)
    inline.append(docPr)

    graphic = OxmlElement('a:graphic')
    graphicData = OxmlElement('a:graphicData')
    graphicData.set('uri', 'http://schemas.openxmlformats.org/drawingml/2006/picture')

    pic = OxmlElement('pic:pic')
    nvPicPr = OxmlElement('pic:nvPicPr')
    cNvPr = OxmlElement('pic:cNvPr')
    cNvPr.set('id', '0')
    cNvPr.set('name', name)
    nvPicPr.append(cNvPr)
    cNvPicPr = OxmlElement('pic:cNvPicPr')
    nvPicPr.append(cNvPicPr)
    pic.append(nvPicPr)

    blipFill = OxmlElement('pic:blipFill')
    blip = OxmlElement('a:blip')
    blip.set(qn('r:embed'), rId)
    blipFill.append(blip)
    stretch = OxmlElement('a:stretch')
    fillRect = OxmlElement('a:fillRect')
    stretch.append(fillRect)
    blipFill.append(stretch)
    pic.append(blipFill)

    spPr = OxmlElement('pic:spPr')
    xfrm = OxmlElement('a:xfrm')
    off = OxmlElement('a:off')
    off.set('x', '0')
    off.set('y', '0')
    xfrm.append(off)
    ext_el = OxmlElement('a:ext')
    ext_el.set('cx', str(cx))
    ext_el.set('cy', str(cy))
    xfrm.append(ext_el)
    spPr.append(xfrm)
    prstGeom = OxmlElement('a:prstGeom')
    prstGeom.set('prst', 'rect')
    spPr.append(prstGeom)
    pic.append(spPr)

    graphicData.append(pic)
    graphic.append(graphicData)
    inline.append(graphic)

    drawing = OxmlElement('w:drawing')
    drawing.append(inline)
    return drawing


def add_header_footer(doc):
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    section = doc.sections[0]
    section.different_first_page_header_footer = True

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]

    if LOGO_ORCHARD.exists():
        image_part = doc.part.package.get_or_add_image_part(str(LOGO_ORCHARD))
        rId = header.part.relate_to(image_part, RT.IMAGE)
        size_emu = int(0.3 * 914400)
        logo_run = hp.add_run()
        drawing = _build_inline_image_xml(rId, size_emu, size_emu, pic_id=10, name="Header Logo")
        logo_run._r.append(drawing)
        hp.add_run("  ")

    run = hp.add_run("Argos \u2014 The Orchard  |  INTERNAL")
    run.font.size = Pt(8)
    run.font.color.rgb = C_GRAY
    run.font.name = "Calibri"
    run.italic = True
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = C_GRAY
    run.font.name = "Calibri"
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    run._r.append(fld1)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run._r.append(instr)
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld2)


def _add_toc_hyperlink(doc, bookmark_name, display_text, level=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.left_indent = Inches(0.3 * level)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), bookmark_name)

    run_el = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "22")
    rPr.append(sz)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:hAnsi"), "Calibri")
    rPr.append(rFonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1a3a5c")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    run_el.append(rPr)

    text_el = OxmlElement("w:t")
    text_el.set(qn("xml:space"), "preserve")
    text_el.text = display_text
    run_el.append(text_el)
    hyperlink.append(run_el)
    p._p.append(hyperlink)


# ═══════════════════════════════════════════════════
# TOC ENTRIES
# ═══════════════════════════════════════════════════

TOC_ENTRIES = [
    ("part1", "Part 1. Research"),
    ("part2", "Part 2. Operations"),
    ("part3", "Part 3. Technical Development"),
    ("part4", "Part 4. Project Insights, Lessons Learned & To-Do"),
]


# ═══════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════

def create_cover_page(doc):
    for _ in range(2):
        doc.add_paragraph()

    if LOGO_PEACOCK.exists():
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_logo = p_logo.add_run()
        run_logo.add_picture(str(LOGO_PEACOCK), width=Inches(2.5))

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ARGOS")
    run.font.size = Pt(48)
    run.font.color.rgb = C_PRIMARY
    run.bold = True
    run.font.name = "Calibri"

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Visual Speech Processing System")
    run2.font.size = Pt(22)
    run2.font.color.rgb = C_H2
    run2.font.name = "Calibri"

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("The Orchard")
    run3.font.size = Pt(20)
    run3.font.color.rgb = C_H2
    run3.font.name = "Calibri"

    doc.add_paragraph()

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run("Project Summary")
    run4.font.size = Pt(16)
    run4.font.color.rgb = C_H3
    run4.italic = True
    run4.font.name = "Calibri"

    for _ in range(2):
        doc.add_paragraph()

    info_lines = [
        ("Project Start:", "October 18, 2024"),
        ("Lead Engineer:", "Yoad Oxman (Solo Developer)"),
        ("Last Updated:", LAST_UPDATED),
        ("Classification:", "Internal R&D Documentation"),
    ]
    for label, value in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rl = p.add_run(label + " ")
        rl.bold = True
        rl.font.size = Pt(11)
        rl.font.color.rgb = C_DARK
        rl.font.name = "Calibri"
        rv = p.add_run(value)
        rv.font.size = Pt(11)
        rv.font.color.rgb = C_DARK
        rv.font.name = "Calibri"

    doc.add_page_break()


# ═══════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════

def create_toc(doc):
    add_heading(doc, "Table of Contents", 1)
    doc.add_paragraph()

    for bm_id, title in TOC_ENTRIES:
        _add_toc_hyperlink(doc, bm_id, title, level=0)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Click any section title above to jump directly to that part.")
    run.font.size = Pt(9)
    run.font.color.rgb = C_GRAY
    run.italic = True
    run.font.name = "Calibri"

    doc.add_page_break()


# ═══════════════════════════════════════════════════
# PART 1: RESEARCH (~3 pages)
# ═══════════════════════════════════════════════════

def part_1_research(doc):
    add_heading(doc, "Part 1. Research", 1, bookmark_id="part1")

    # 1.1 Model Overview
    add_heading(doc, "1.1 VSP-LLM Model Architecture", 2)
    add_para(doc, (
        "Argos is built on VSP-LLM (Yeo et al., arXiv:2402.15151, May 2024), the first framework "
        "to combine visual speech recognition with large language models. The system takes raw video "
        "of a speaker's face and produces text transcriptions purely from lip movements \u2014 no audio required."
    ))

    add_styled_table(doc,
        ["Component", "Implementation", "Role"],
        [
            ["Visual Encoder", "AV-HuBERT Large (315M params)", "Extracts visual speech features from mouth crops"],
            ["Deduplication", "200 K-means clusters", "Reduces sequence length ~50% with no quality loss"],
            ["LLM Backbone", "LLaMA-2-7B (4-bit quantized)", "Language generation from visual features"],
            ["Adaptation", "QLoRA (r=16, alpha=32, Q/K/V)", "Efficient fine-tuning (~4.2M trainable params)"],
            ["Bridge", "avfeat_to_llm linear layer", "Maps 1024-dim visual features to LLaMA embedding space"],
        ],
        col_widths=[1.3, 2.0, 3.2]
    )

    add_para(doc, (
        "The paper reports 25.4% WER on LRS3 (433h TED talks, studio conditions) using the finetune "
        "configuration. The model uniquely improves with longer videos: 27.3% WER at 0-2s down to 12.9% "
        "WER at 6+ seconds, thanks to LLM context modeling."
    ))

    # 1.2 Real-World Evaluation
    add_heading(doc, "1.2 Real-World Evaluation Results", 2)
    add_para(doc, (
        "We evaluated the model on real-world YouTube content across two datasets:"
    ))

    add_styled_table(doc,
        ["Metric", "Paper (LRS3)", "english_1k (860 seg)", "english_full (1,497 seg)"],
        [
            ["Mean WER", "~25-30%", "67.0%", "64.1%"],
            ["Corpus WER", "\u2014", "125.5%", "\u2014"],
            ["NEA F1", "Not reported", "38.8%", "39%"],
            ["WWER", "\u2014", "\u2014", "61.9%"],
            ["Usable (WER \u226420%)", "\u2014", "11.4%", "\u2014"],
            ["Hallucinated (WER \u2265100%)", "\u2014", "20.6%", "\u2014"],
        ],
        col_widths=[1.8, 1.3, 1.7, 1.7]
    )

    add_para(doc, (
        "The model performs 2.5x worse on real-world video than the paper's benchmark. "
        "Only 11.4% of segments achieve acceptable quality (WER \u226420%), while 20.6% are catastrophically "
        "hallucinated (the model generates fluent but completely fabricated text). Corpus WER exceeds 100%, "
        "meaning the model inserts more error words than exist in the reference."
    ))

    # 1.3 Quality Distribution
    add_heading(doc, "1.3 Quality Distribution", 2)
    add_styled_table(doc,
        ["Quality Tier", "WER Range", "Segments", "% of Total"],
        [
            ["Usable", "0-20%", "98", "11.4%"],
            ["Marginal", "21-40%", "150", "17.4%"],
            ["Poor", "41-60%", "153", "17.8%"],
            ["Unusable", "61-99%", "282", "32.8%"],
            ["Hallucinated", "100%+", "177", "20.6%"],
        ],
        col_widths=[1.5, 1.2, 1.2, 1.2]
    )

    # 1.4 Root Cause
    add_heading(doc, "1.4 Root Cause: Domain Mismatch", 2)
    add_para(doc, (
        "The performance gap is primarily caused by domain mismatch between training data (LRS3: "
        "TED talks, studio lighting, professional speakers, frontal view) and test data (YouTube: "
        "variable lighting, head pose, casual speech, diverse accents, brands/slang). "
        "Length distribution shift is secondary \u2014 the model handles variable lengths natively via "
        "RoPE and attention mechanisms."
    ))

    add_styled_table(doc,
        ["Dimension", "LRS3 (Training)", "YouTube (Test)", "Impact"],
        [
            ["Head pose", "Frontal, static", "Variable, moving", "High"],
            ["Lighting", "Professional", "Variable, poor", "High"],
            ["Speaker diversity", "TED presenters", "Global creators", "Medium"],
            ["Vocabulary", "Academic", "Casual, slang, brands", "Medium"],
            ["Speech style", "Rehearsed, clear", "Natural, fast", "Medium"],
        ],
        col_widths=[1.3, 1.5, 1.5, 0.8]
    )

    # 1.5 Improvement Roadmap
    add_heading(doc, "1.5 Improvement Roadmap", 2)
    add_para(doc, (
        "Six detailed analysis reports identified multiple improvement paths. "
        "The table below summarizes expected impact by strategy:"
    ))

    add_styled_table(doc,
        ["Strategy", "Expected WER Reduction", "Complexity", "Status"],
        [
            ["Hyperparameter tuning (lenpen, beam, max_len)", "-1.8 WWER pts (tested)", "Low", "DONE \u2014 13 experiments run"],
            ["Fine-tuning on AVSpeech (encoder unfreeze)", "-15 to -25 pts", "High", "Infrastructure ready"],
            ["N-best aggregation (ROVER / MBR)", "-5 to -10 pts", "Medium", "Code design complete"],
            ["Prompt engineering (topic, word count)", "-5 to -10 pts", "Low", "Needs testing"],
            ["Confidence scoring (token/sequence)", "Quality filtering", "Medium", "Code design complete"],
            ["LoRA rank increase (r=16 \u2192 r=64)", "-3 to -8 pts", "Low", "Config change ready"],
            ["Combined realistic target", "~35-45% WER", "\u2014", "\u2014"],
        ],
        col_widths=[2.8, 1.5, 0.9, 1.3]
    )

    # 1.6 Fine-Tuning Strategy
    add_heading(doc, "1.6 Fine-Tuning Strategy", 2)
    add_para(doc, (
        "Domain adaptation fine-tuning on AVSpeech data is the highest-impact single intervention. "
        "Key parameters for the recommended configuration:"
    ))
    add_bullet_bold_value(doc, "Training data: ", "AVSpeech (~4,700h, ~290K YouTube videos, 3-10s clips)")
    add_bullet_bold_value(doc, "Length strategy: ", "Mixed (3-12s) \u2014 don't filter by length; data quantity trumps length matching")
    add_bullet_bold_value(doc, "LoRA rank: ", "Increase r=16 \u2192 r=64, alpha=128 (current r=16 is underfit for TED\u2192YouTube shift)")
    add_bullet_bold_value(doc, "Encoder: ", "Unfreeze after 5K steps (encoder is the bottleneck \u2014 never saw YouTube conditions)")
    add_bullet_bold_value(doc, "Data curation: ", "80% clean frontal + 20% moderate difficulty; exclude extreme angles/lighting")
    add_bullet_bold_value(doc, "GPU requirement: ", "p3.16xlarge (8x V100, 128GB) \u2014 ~3-5 hours training, ~$24/hr")
    add_bullet_bold_value(doc, "Expected outcome: ", "WER from 67% \u2192 42-52%, usable segments from 11% \u2192 18-28%")

    # 1.7 Hyperparameter Tuning Results
    add_heading(doc, "1.7 Hyperparameter Tuning Results (13 Experiments)", 2)
    add_para(doc, (
        "A systematic sweep of 13 decode configurations was run on a 107-segment test set. "
        "Parameters tested: length penalty, beam width, sampling, temperature, repetition penalty."
    ))

    add_styled_table(doc,
        ["Exp", "Config", "WWER%", "NEA F1%", "Verdict"],
        [
            ["A", "baseline (beam=20, lenpen=0)", "59.5", "40.9", "Reference"],
            ["C", "lenpen=1.0", "58.6", "41.1", "Best single param"],
            ["J", "lenpen=1.0 + sample temp=0.5", "57.7", "41.0", "BEST WWER"],
            ["F", "sample temp=1.0", "59.4", "42.4", "Best NEA F1"],
            ["G", "greedy (beam=1)", "63.2", "37.7", "Worse"],
            ["D", "lenpen=-0.5", "75.3", "26.1", "FAIL (45% empty)"],
            ["H", "lenpen=2.0", "171.5", "39.1", "FAIL (over-generation)"],
        ],
        col_widths=[0.3, 2.3, 0.6, 0.6, 1.5]
    )

    add_para(doc, (
        "Key finding: lenpen=1.0 + sampling at temp=0.5 (Config J) achieves best WWER (57.7%, "
        "-1.8 vs baseline). Tuning alone yields modest gains (~1.8 pts); "
        "fine-tuning remains essential for meaningful improvement. Full decode with Config J on "
        "1,497 segments is in progress. See tuning_experiments.docx for detailed analysis."
    ), bold=True)


# ═══════════════════════════════════════════════════
# PART 2: OPERATIONS (~1.5 pages)
# ═══════════════════════════════════════════════════

def part_2_operations(doc):
    add_heading(doc, "Part 2. Operations", 1, bookmark_id="part2")

    # 2.1 Dual-Environment Architecture
    add_heading(doc, "2.1 Dual-Environment Architecture", 2)
    add_para(doc, (
        "The system operates across two primary environments with strict synchronization requirements. "
        "Every EC2 change must be replicated to the Linux container with path translation."
    ))

    add_styled_table(doc,
        ["Aspect", "EC2 (Development)", "Linux Container (Production)"],
        [
            ["Base path", "/home/ubuntu/", "/workspace/ or /host/galaxy_export/"],
            ["Purpose", "Development, testing, evaluation", "End-user deployment"],
            ["GPU", "T4 (16GB)", "Variable (12-24GB)"],
            ["Internet", "Available", "May be restricted (HORIZON)"],
            ["Sync protocol", "Source of truth", "28 pending sync items documented"],
        ],
        col_widths=[1.3, 2.5, 2.7]
    )

    # 2.2 Deployment Evolution
    add_heading(doc, "2.2 Deployment Packages", 2)
    add_para(doc, (
        "Deployment evolved from 13+ scattered update folders into a single, verified package:"
    ))

    add_styled_table(doc,
        ["Release", "Date", "Size", "Key Changes"],
        [
            ["v1.0.0", "Feb 3, 2026", "1.8 MB", "12 critical fixes, 7 UI features, automated install"],
            ["v1.0.36+", "Feb 17, 2026", "2.9 MB", "37 total bugs fixed, NVENC fix, segment naming"],
            ["Current", "Feb 18, 2026", "\u2014", "28 EC2 changes pending container sync"],
        ],
        col_widths=[0.8, 1.2, 0.7, 3.8]
    )

    add_para(doc, (
        "Each package includes automated INSTALL.sh and VERIFY.sh scripts. "
        "Size reduction: 15 GB scattered files \u2192 1.8 MB verified package (99.99% reduction)."
    ))

    # 2.3 Critical Bugs Fixed
    add_heading(doc, "2.3 Critical Bugs Fixed (37 Total)", 2)
    add_para(doc, "Bugs were tracked across three phases:")

    add_styled_table(doc,
        ["Phase", "Count", "Key Highlights"],
        [
            ["Installation (1-13)", "13", "Missing scripts, module imports, hardcoded EC2 paths, python3 compat"],
            ["Deployment (14-25)", "12", "GPU OOM on 12GB, Cython extensions, max_len config, NVENC corruption"],
            ["Final (26-37)", "12", "Upload limits, terminal encoding, inference errors, metric calculation"],
        ],
        col_widths=[1.5, 0.6, 4.4]
    )

    add_para(doc, (
        "Most impactful fixes: Fairseq Cython auto-build (decode failed without it), "
        "max_len configuration (predictions truncated mid-sentence), transcription persistence "
        "(saved hours of Whisper re-processing), and NVENC silent corruption (GPU encoder "
        "produced corrupt output that passed all checks but was visually broken)."
    ))

    # 2.4 Testing
    add_heading(doc, "2.4 Testing & Verification", 2)
    add_styled_table(doc,
        ["Test Suite", "Assertions", "Coverage"],
        [
            ["Module tests (test_all_modules.sh)", "37 tests", "All 12 lib/ modules"],
            ["Environment tests (EC2)", "35+ assertions", "Path detection, env vars, config"],
            ["Smoke tests", "60+ checks", "Dependencies, imports, CLI tools"],
            ["Pre/post-deployment checklists", "Comprehensive", "All fixes verified per release"],
        ],
        col_widths=[2.5, 1.2, 2.8]
    )


# ═══════════════════════════════════════════════════
# PART 3: TECHNICAL DEVELOPMENT (~2 pages)
# ═══════════════════════════════════════════════════

def part_3_technical(doc):
    add_heading(doc, "Part 3. Technical Development", 1, bookmark_id="part3")

    # 3.1 Pipeline Architecture
    add_heading(doc, "3.1 Pipeline Architecture (9 Stages)", 2)
    add_para(doc, (
        "The main pipeline (run_flat_english_pipeline.sh) orchestrates video processing "
        "from raw input to final reports through 9 stages:"
    ))

    add_styled_table(doc,
        ["Stage", "Description", "Key Technology"],
        [
            ["0.1 Segment", "Split videos into 12s segments (codec copy)", "FFmpeg stream copy"],
            ["0.5 Normalize", "HDR/10-bit tone mapping on segments", "NVENC GPU / CPU fallback"],
            ["0.6 Transcriptions", "Copy existing .wrd files (skip Whisper)", "Bash file matching"],
            ["2. Mouth Crop", "Face detection + 88x88 grayscale crops", "MediaPipe, 25fps"],
            ["3. ASR", "Generate word transcriptions", "OpenAI Whisper"],
            ["4-5. LRS3 + Manifests", "Format conversion, TSV/split generation", "Custom Python"],
            ["6. K-means", "200-cluster feature extraction", "AV-HuBERT layer 12 (1024-dim)"],
            ["7. Decode", "LLM inference (segment-level)", "VSP-LLM + LLaMA-2-7B"],
            ["8. Outputs", "JSON reports + burned subtitle videos", "Custom Python + FFmpeg"],
        ],
        col_widths=[1.3, 2.8, 2.4]
    )

    # 3.2 Key Innovation: Segment-First Normalization
    add_heading(doc, "3.2 Segment-First Normalization", 2)
    add_para(doc, (
        "A key architectural innovation: videos are segmented FIRST using fast codec copy, "
        "then normalization runs on 12-second segments instead of full-length videos. "
        "For a 60-minute video, this means normalizing 300 small segments instead of one "
        "large file \u2014 resulting in 90% faster processing, lower memory usage, and better "
        "parallelization potential."
    ))

    # 3.3 Modular Refactoring
    add_heading(doc, "3.3 Modular Refactoring", 2)
    add_para(doc, (
        "The pipeline was refactored from a monolithic 823-line script into 11 reusable "
        "modules under lib/ (January 2026). Main orchestrator reduced to 393 lines (52% reduction). "
        "Each module is independently testable, with environment-aware configuration that "
        "auto-detects EC2 vs container paths."
    ))

    add_styled_table(doc,
        ["Module", "Responsibility"],
        [
            ["common.sh", "Logging (log_info, log_error, log_stage), validation"],
            ["config.sh", "Environment detection (EC2/container), path configuration"],
            ["archive.sh", "Archive management with transcription preservation"],
            ["normalization.sh", "Video normalization (HDR/10-bit, GPU/CPU encoding)"],
            ["asr.sh", "Whisper ASR with transcription reuse (Step 0.6 + 1.5 + 3)"],
            ["decode.sh", "VSP-LLM decode with Cython auto-build check"],
            ["outputs.sh", "Client reports (JSON/HTML/CSV) and burned videos"],
        ],
        col_widths=[1.5, 5.0]
    )

    # 3.4 Transcription Persistence
    add_heading(doc, "3.4 Transcription Persistence", 2)
    add_para(doc, (
        "Whisper ASR is the slowest pipeline stage. A persistence system ensures it runs only once "
        "per video across all pipeline runs. Transcriptions are stored in .transcriptions/ alongside "
        "input videos, survive pipeline archiving, and are automatically copied before Whisper runs. "
        "User-edited transcriptions (marked [MANUAL]) are preserved and never overwritten."
    ))

    # 3.5 Web UI
    add_heading(doc, "3.5 Web UI", 2)
    add_para(doc, (
        "A zero-dependency web interface built entirely on Python standard library (http.server, "
        "subprocess, threading) with vanilla HTML/CSS/JavaScript frontend."
    ))

    add_styled_table(doc,
        ["Feature", "Details"],
        [
            ["Video validation", "FFprobe checks: format, duration, resolution, K-means viability"],
            ["Drag-and-drop upload", "File upload with progress tracking"],
            ["9-stage progress", "Real-time pipeline tracking with stage-level timeouts"],
            ["Transcription management", "Modal dialog for manual transcription entry/editing"],
            ["Video exclusion", "Move videos to .excluded/ without deletion"],
            ["K-means toggle", "Skip retraining when reusing existing cluster model"],
        ],
        col_widths=[1.8, 4.7]
    )

    # 3.6 Technology Stack
    add_heading(doc, "3.6 Technology Stack", 2)
    add_styled_table(doc,
        ["Layer", "Technologies"],
        [
            ["ML/AI", "PyTorch 2.5.1, Fairseq (custom), LLaMA-2-7B (4-bit), AV-HuBERT Large"],
            ["ASR", "OpenAI Whisper (offline), SentencePiece 0.1.96"],
            ["Vision", "MediaPipe (face detection), RetinaFace (alternative)"],
            ["Video", "FFmpeg/FFprobe (encode/decode/segment/probe)"],
            ["Backend", "Python 3.11 stdlib (http.server, subprocess, threading)"],
            ["Frontend", "Vanilla HTML5 / CSS / JavaScript (no frameworks)"],
            ["Infrastructure", "Docker (nvidia/cuda base), EC2, standalone Ubuntu, HORIZON"],
        ],
        col_widths=[1.0, 5.5]
    )

    # 3.7 Code Inventory
    add_heading(doc, "3.7 Code Inventory", 2)
    add_styled_table(doc,
        ["Category", "Files", "Lines"],
        [
            ["Pipeline scripts", "4", "1,299"],
            ["Library modules (lib/)", "12", "1,562"],
            ["Custom Python scripts", "13", "3,342"],
            ["Web UI + server", "11", "5,780"],
            ["Testing scripts", "5", "1,731"],
            ["Grand total", "~55", "~14,882"],
        ],
        col_widths=[2.5, 0.7, 0.7]
    )


# ═══════════════════════════════════════════════════
# PART 4: LESSONS LEARNED & TO-DO (~2 pages)
# ═══════════════════════════════════════════════════

def part_4_lessons_and_todo(doc):
    add_heading(doc, "Part 4. Project Insights, Lessons Learned & To-Do", 1, bookmark_id="part4")

    # 4.1 Key Insights
    add_heading(doc, "4.1 Key Insights", 2)

    add_bullet_bold_value(doc, "Benchmark \u2260 Reality: ",
        "Paper reports 25.4% WER on LRS3; real-world YouTube yields 64-67% WER (2.5x worse). "
        "Domain mismatch is the dominant factor, not model architecture.")

    add_bullet_bold_value(doc, "Hallucination is Architectural: ",
        "When visual signal is ambiguous, the LLM generates fluent but fabricated text from its "
        "language prior. This affects 20.6% of segments and cannot be fully solved by tuning alone \u2014 "
        "it requires confidence scoring to detect and filter.")

    add_bullet_bold_value(doc, "Short Segments Fail: ",
        "Segments with <5 words produce 128.5% WER. The model needs context to disambiguate "
        "homophenes (e.g., /p/, /b/, /m/ look identical). Minimum 5-second segment length is recommended.")

    add_bullet_bold_value(doc, "No Quality Signal Exists: ",
        "Best heuristics (word count, duration) achieve only 24% precision for identifying good outputs. "
        "Confidence scoring from model internals is essential for production use.")

    add_bullet_bold_value(doc, "Corpus WER > 100%: ",
        "The model generates MORE error words than reference words exist, primarily due to "
        "length penalty being disabled (lenpen=0.0). This is the single easiest fix available.")

    # 4.2 Operational Lessons
    add_heading(doc, "4.2 Operational Lessons", 2)

    add_bullet_bold_value(doc, "Container Sync Discipline: ",
        "Every EC2 change must be replicated to the production container. Forgetting this caused "
        "repeated production bugs. The sync changelog protocol was essential \u2014 28 items documented.")

    add_bullet_bold_value(doc, "Transcription Persistence Saves Hours: ",
        "Whisper is the slowest pipeline stage. Running it once and persisting results across "
        "pipeline runs was a high-impact optimization that fundamentally changed the workflow.")

    add_bullet_bold_value(doc, "Modular > Monolithic: ",
        "The refactoring from 823\u2192393 lines made the pipeline testable and debuggable. "
        "Each module can be tested independently (37 tests). Worth the investment.")

    add_bullet_bold_value(doc, "GPU Encoding Pitfalls: ",
        "NVENC (GPU video encoder) produced silently corrupt output that passed all automated checks "
        "but was visually broken. Hardware acceleration requires explicit validation of output quality.")

    # 4.3 Timeline
    add_heading(doc, "4.3 Development Timeline", 2)
    add_styled_table(doc,
        ["Date", "Milestone"],
        [
            ["Oct 18, 2024", "Project start \u2014 literature review begins"],
            ["Oct-Nov 2024", "First pipeline version (~285 lines)"],
            ["Nov-Dec 2024", "Pipeline iterations, face detection, AVSpeech processing"],
            ["Dec 29, 2025", "Presentation: \"Argos \u2014 A Working Pipeline\" (19 slides)"],
            ["Jan 2026", "Modular refactoring (823 \u2192 501+1,562 lines)"],
            ["Jan 2026", "Web UI development (5,780 lines)"],
            ["Feb 3, 2026", "First container deployment (v1.0.0)"],
            ["Feb 3-17, 2026", "Five container releases, 37 bugs fixed"],
            ["Feb 2026", "Performance evaluation (860 segments, 6 reports)"],
            ["Feb 2026", "Fine-tuning infrastructure ready"],
        ],
        col_widths=[1.3, 5.2]
    )

    # 4.4 Project To-Do List
    add_heading(doc, "4.4 Project To-Do List", 2)

    add_heading(doc, "Research & Model Improvement", 3)
    add_styled_table(doc,
        ["#", "Task", "Priority", "Expected Impact"],
        [
            ["1", "Fine-tune on AVSpeech with encoder unfreeze", "Critical", "-15 to -25 WER pts"],
            ["2", "Increase LoRA rank r=16 \u2192 r=64 (alpha=128)", "High", "-3 to -8 WER pts"],
            ["3", "Test length penalty lenpen=0.5-1.0", "High", "-5 to -15 WER pts"],
            ["4", "Reduce beam width 20 \u2192 5-10", "Medium", "-1 to -3 WER pts"],
            ["5", "Tighten max_len_a=1.0, max_len_b=50", "Medium", "Reduce hallucinations"],
            ["6", "Implement confidence scoring", "High", "Quality filtering"],
            ["7", "Implement N-best aggregation (ROVER/MBR)", "Medium", "-5 to -10 WER pts"],
            ["8", "Evaluate prompt engineering strategies", "Medium", "-5 to -10 WER pts"],
            ["9", "Curate training data (80% clean / 20% moderate)", "High", "Better training signal"],
            ["10", "Evaluate RetinaFace vs MediaPipe", "Low", "-2 to -5 WER pts"],
        ],
        col_widths=[0.3, 3.0, 0.8, 1.5]
    )

    add_heading(doc, "Operations & Deployment", 3)
    add_styled_table(doc,
        ["#", "Task", "Priority"],
        [
            ["11", "Sync all 28 pending EC2 \u2192 container changes", "Critical"],
            ["12", "Deploy latest container package to HORIZON server", "High"],
            ["13", "Automate container sync verification (CI/CD or script)", "Medium"],
            ["14", "Add minimum segment duration filter (\u22655s) to pipeline", "Medium"],
        ],
        col_widths=[0.3, 4.5, 0.8]
    )

    add_heading(doc, "Technical Development", 3)
    add_styled_table(doc,
        ["#", "Task", "Priority"],
        [
            ["15", "Add confidence scores to JSON report output", "High"],
            ["16", "Add N-best hypotheses to decode output format", "Medium"],
            ["17", "Investigate multi-language support (Arabic \u2014 needs multilingual LLM)", "Low"],
            ["18", "Consider newer model architectures (LLaMA-3, Mistral)", "Low"],
        ],
        col_widths=[0.3, 4.5, 0.8]
    )

    add_heading(doc, "Documentation & Process", 3)
    add_styled_table(doc,
        ["#", "Task", "Priority"],
        [
            ["19", "Update research_documentation.docx after each fine-tuning experiment", "Ongoing"],
            ["20", "Document HORIZON server setup and constraints", "Medium"],
            ["21", "Create runbook for fine-tuning on cloud GPUs (p3.16xlarge)", "Medium"],
        ],
        col_widths=[0.3, 4.5, 0.8]
    )

    # Final note
    doc.add_paragraph()
    add_para(doc, (
        "This summary condenses the full 60-page Argos R&D documentation into the most critical "
        "findings, decisions, and next steps. For detailed technical analysis, code diffs, and "
        "configuration specifics, refer to the complete research_documentation.docx."
    ), italic=True, color=C_GRAY)


# ═══════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════

def main():
    print("Generating Argos Project Summary...")
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Header/Footer
    add_header_footer(doc)

    # Build document
    create_cover_page(doc)
    create_toc(doc)
    part_1_research(doc)
    part_2_operations(doc)
    part_3_technical(doc)
    part_4_lessons_and_todo(doc)

    # Save
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_FILE))
    print(f"Saved: {OUTPUT_FILE}")
    print(f"Size: {OUTPUT_FILE.stat().st_size / 1024:.0f} KB")


if __name__ == "__main__":
    main()
