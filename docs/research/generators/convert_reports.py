#!/usr/bin/env python3
"""Convert markdown reports to PDF and Word (DOCX) format."""

import os
import sys
import markdown
from pathlib import Path

ANALYSIS_DIR = Path(__file__).parent
REPORTS = [
    ("report_1_executive_assessment.md", "Report 1 - Executive Technical Assessment"),
    ("report_2_hyperparameter_tuning.md", "Report 2 - Hyperparameter Tuning Analysis"),
    ("report_3_prompt_engineering.md", "Report 3 - Prompt Engineering Analysis"),
    ("report_4_confidence_scoring.md", "Report 4 - Word-Level Confidence Scoring"),
    ("report_5_beam_search_aggregation.md", "Report 5 - Beam Search Aggregation"),
    ("report_6_finetuning_analysis.md", "Report 6 - Fine-Tuning Analysis"),
]

CSS = """
@page {
    size: A4;
    margin: 2cm;
}
body {
    font-family: 'Segoe UI', Calibri, Arial, Helvetica, sans-serif;
    font-size: 11pt;
    line-height: 1.5;
    color: #1a1a1a;
    max-width: 100%;
}
h1 {
    font-size: 20pt;
    color: #1a3a5c;
    border-bottom: 3px solid #1a3a5c;
    padding-bottom: 8px;
    margin-top: 0;
    page-break-before: avoid;
}
h2 {
    font-size: 16pt;
    color: #2a5a8c;
    border-bottom: 1px solid #ccc;
    padding-bottom: 4px;
    margin-top: 24pt;
}
h3 {
    font-size: 13pt;
    color: #3a6a9c;
    margin-top: 18pt;
}
h4 {
    font-size: 11pt;
    color: #4a7aac;
    margin-top: 14pt;
}
table {
    border-collapse: collapse;
    width: 100%;
    margin: 12px 0;
    font-size: 10pt;
}
th {
    background-color: #1a3a5c;
    color: white;
    padding: 8px 10px;
    text-align: left;
    font-weight: 600;
}
td {
    padding: 6px 10px;
    border-bottom: 1px solid #ddd;
}
tr:nth-child(even) td {
    background-color: #f5f8fc;
}
code {
    background-color: #f0f3f8;
    padding: 2px 5px;
    border-radius: 3px;
    font-family: 'Consolas', 'Courier New', monospace;
    font-size: 9.5pt;
}
pre {
    background-color: #f0f3f8;
    padding: 12px 16px;
    border-radius: 5px;
    border-left: 4px solid #1a3a5c;
    overflow-x: auto;
    font-size: 9pt;
    line-height: 1.4;
}
pre code {
    background: none;
    padding: 0;
}
blockquote {
    border-left: 4px solid #2a5a8c;
    margin: 12px 0;
    padding: 8px 16px;
    background-color: #f5f8fc;
    color: #555;
}
strong {
    color: #1a1a1a;
}
hr {
    border: none;
    border-top: 2px solid #1a3a5c;
    margin: 20px 0;
}
.header-meta {
    color: #666;
    font-size: 10pt;
}
"""


def convert_to_pdf(md_path, pdf_path, title):
    """Convert markdown file to PDF using weasyprint."""
    from weasyprint import HTML

    with open(md_path, 'r') as f:
        md_content = f.read()

    html_body = markdown.markdown(
        md_content,
        extensions=['tables', 'fenced_code', 'codehilite', 'toc'],
        extension_configs={'codehilite': {'css_class': 'highlight', 'guess_lang': False}}
    )

    full_html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>{title}</title>
    <style>{CSS}</style>
</head>
<body>
{html_body}
</body>
</html>"""

    HTML(string=full_html).write_pdf(str(pdf_path))
    print(f"  PDF: {pdf_path} ({os.path.getsize(pdf_path) // 1024} KB)")


def convert_to_docx(md_path, docx_path, title):
    """Convert markdown file to Word DOCX."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    import re

    with open(md_path, 'r') as f:
        lines = f.readlines()

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    in_code_block = False
    code_block_lines = []
    in_table = False
    table_rows = []
    i = 0

    while i < len(lines):
        line = lines[i].rstrip('\n')

        # Code blocks
        if line.startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_block_lines)
                p = doc.add_paragraph()
                p.style = doc.styles['Normal']
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                pf = p.paragraph_format
                pf.space_before = Pt(4)
                pf.space_after = Pt(4)
                code_block_lines = []
                in_code_block = False
            else:
                # Flush any pending table
                if in_table and table_rows:
                    _add_table(doc, table_rows)
                    table_rows = []
                    in_table = False
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue

        # Table rows
        if '|' in line and line.strip().startswith('|'):
            stripped = line.strip()
            # Skip separator rows
            if re.match(r'^\|[\s\-:|]+\|$', stripped):
                i += 1
                continue
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            if cells:
                if not in_table:
                    in_table = True
                table_rows.append(cells)
            i += 1
            continue
        else:
            if in_table and table_rows:
                _add_table(doc, table_rows)
                table_rows = []
                in_table = False

        # Empty lines
        if not line.strip():
            i += 1
            continue

        # Horizontal rule
        if line.strip() in ('---', '***', '___'):
            doc.add_paragraph('─' * 70)
            i += 1
            continue

        # Headings
        heading_match = re.match(r'^(#{1,4})\s+(.*)', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Remove bold markers
            h = doc.add_heading(text, level=level)
            if level == 1:
                for run in h.runs:
                    run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5c)
            elif level == 2:
                for run in h.runs:
                    run.font.color.rgb = RGBColor(0x2a, 0x5a, 0x8c)
            i += 1
            continue

        # Regular paragraph with inline formatting
        p = doc.add_paragraph()
        _add_formatted_text(p, line)
        i += 1

    # Flush remaining table
    if in_table and table_rows:
        _add_table(doc, table_rows)

    doc.save(str(docx_path))
    print(f"  DOCX: {docx_path} ({os.path.getsize(docx_path) // 1024} KB)")


def _add_table(doc, rows):
    """Add a table to the document."""
    from docx.shared import Pt, RGBColor, Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    if not rows:
        return

    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = 'Light Grid Accent 1'

    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx]
        for c_idx, cell_text in enumerate(row_data):
            if c_idx < n_cols:
                cell = row.cells[c_idx]
                cell.text = ''
                p = cell.paragraphs[0]
                _add_formatted_text(p, cell_text)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.size = Pt(9)
                    if r_idx == 0:
                        run.bold = True


def _add_formatted_text(paragraph, text):
    """Add text with basic inline formatting (bold, code, italic)."""
    from docx.shared import Pt, RGBColor
    import re

    # Split by formatting markers
    parts = re.split(r'(\*\*.*?\*\*|`[^`]+`|\*[^*]+\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x66, 0x33, 0x99)
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def main():
    print("Converting reports to PDF and DOCX...\n")

    pdf_dir = ANALYSIS_DIR / "pdf"
    docx_dir = ANALYSIS_DIR / "docx"
    pdf_dir.mkdir(exist_ok=True)
    docx_dir.mkdir(exist_ok=True)

    for md_file, title in REPORTS:
        md_path = ANALYSIS_DIR / md_file
        if not md_path.exists():
            print(f"  SKIP: {md_file} not found")
            continue

        base = md_file.replace('.md', '')
        print(f"Converting: {md_file}")

        try:
            convert_to_pdf(md_path, pdf_dir / f"{base}.pdf", title)
        except Exception as e:
            print(f"  PDF FAILED: {e}")

        try:
            convert_to_docx(md_path, docx_dir / f"{base}.docx", title)
        except Exception as e:
            print(f"  DOCX FAILED: {e}")

        print()

    print("Done!")
    print(f"\nPDF files:  {pdf_dir}/")
    print(f"DOCX files: {docx_dir}/")


if __name__ == '__main__':
    main()
