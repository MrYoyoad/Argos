#!/usr/bin/env python3
"""
Argos — The Orchard — Comprehensive R&D Document Generator (v2)

Generates a detailed, chapter-based research documentation Word document
covering ALL work done on the Argos visual speech processing project.

14 chapters, ~60 pages, content-rich with code examples and config snippets.

Usage:
    python3 generate_research_journal.py

Output:
    argos_research/research_documentation.docx
"""

import os
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

# ── Output ──
OUTPUT_DIR = Path(__file__).parent
OUTPUT_FILE = OUTPUT_DIR / "research_documentation.docx"

# ── Logos ──
LOGO_ORCHARD = OUTPUT_DIR / "logo.png"       # HaPardes tree logo (headers)
LOGO_PEACOCK = OUTPUT_DIR / "peacock.png"     # Peacock mascot (cover page)

# ── Colors ──
C_PRIMARY = RGBColor(0x1a, 0x3a, 0x5c)
C_H2 = RGBColor(0x2a, 0x5a, 0x8c)
C_H3 = RGBColor(0x3a, 0x6a, 0x9c)
C_H4 = RGBColor(0x4a, 0x7a, 0xac)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_BLACK = RGBColor(0x00, 0x00, 0x00)
C_DARK = RGBColor(0x33, 0x33, 0x33)
C_GREEN = RGBColor(0x15, 0x57, 0x24)
C_RED = RGBColor(0x72, 0x1c, 0x24)
C_AMBER = RGBColor(0x85, 0x64, 0x04)
C_GRAY = RGBColor(0x66, 0x66, 0x66)
C_CODE = RGBColor(0x2d, 0x2d, 0x2d)

HEADER_BG = "1a3a5c"
ZEBRA_BG = "f0f4f8"
GREEN_BG = "d4edda"
RED_BG = "f8d7da"
AMBER_BG = "fff3cd"
CODE_BG = "f5f5f5"

LAST_UPDATED = datetime.now().strftime("%B %d, %Y")


# ═══════════════════════════════════════════════════
# HELPER FUNCTIONS
# ═══════════════════════════════════════════════════

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, bold=False, color=None, size=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = size
    run.font.name = "Calibri"


def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, HEADER_BG)
        set_cell_text(cell, h, bold=True, color=C_WHITE, size=Pt(9))

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            set_cell_text(cell, val, size=Pt(9))
            if r_idx % 2 == 1:
                set_cell_shading(cell, ZEBRA_BG)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return table


def add_heading(doc, text, level, color=None):
    h = doc.add_heading(text, level=level)
    if color is None:
        color = {1: C_PRIMARY, 2: C_H2, 3: C_H3, 4: C_H4}.get(level, C_PRIMARY)
    for run in h.runs:
        run.font.color.rgb = color
        run.font.name = "Calibri"
    return h


def add_para(doc, text, bold=False, italic=False, size=Pt(11), color=None, space_after=Pt(6)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = size
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = space_after
    return p


def add_para_rich(doc, segments, space_after=Pt(6)):
    """Add a paragraph with mixed bold/normal text.
    segments is a list of (text, bold) tuples."""
    p = doc.add_paragraph()
    for text, bold in segments:
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(11)
        run.font.name = "Calibri"
    p.paragraph_format.space_after = space_after
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = "Calibri"
    return p


def add_bullet_bold_value(doc, label, value, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after = Pt(2)
    run_b = p.add_run(label)
    run_b.bold = True
    run_b.font.size = Pt(10)
    run_b.font.name = "Calibri"
    run_v = p.add_run(value)
    run_v.font.size = Pt(10)
    run_v.font.name = "Calibri"
    return p


def add_code_block(doc, code_text, language=""):
    """Add a code block with monospace font and gray background."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.3)

    # Background shading
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{CODE_BG}" w:val="clear"/>')
    pPr.append(shading)

    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(8)
    run.font.color.rgb = C_CODE
    return p


def _build_inline_image_xml(rId, cx, cy, pic_id=1, name="Logo"):
    """Build the WordprocessingML XML for an inline image."""
    inline = OxmlElement('wp:inline')
    inline.set('distT', '0')
    inline.set('distB', '0')
    inline.set('distL', '0')
    inline.set('distR', '0')

    extent = OxmlElement('wp:extent')
    extent.set('cx', str(cx))
    extent.set('cy', str(cy))
    inline.append(extent)

    docPr = OxmlElement('wp:docPr')
    docPr.set('id', str(pic_id))
    docPr.set('name', name)
    inline.append(docPr)

    graphic = OxmlElement('a:graphic')
    graphicData = OxmlElement('a:graphicData')
    graphicData.set('uri', 'http://schemas.openxmlformats.org/drawingml/2006/picture')

    pic = OxmlElement('pic:pic')
    nvPicPr = OxmlElement('pic:nvPicPr')
    cNvPr = OxmlElement('pic:cNvPr')
    cNvPr.set('id', '0')
    cNvPr.set('name', name)
    nvPicPr.append(cNvPr)
    cNvPicPr = OxmlElement('pic:cNvPicPr')
    nvPicPr.append(cNvPicPr)
    pic.append(nvPicPr)

    blipFill = OxmlElement('pic:blipFill')
    blip = OxmlElement('a:blip')
    blip.set(qn('r:embed'), rId)
    blipFill.append(blip)
    stretch = OxmlElement('a:stretch')
    fillRect = OxmlElement('a:fillRect')
    stretch.append(fillRect)
    blipFill.append(stretch)
    pic.append(blipFill)

    spPr = OxmlElement('pic:spPr')
    xfrm = OxmlElement('a:xfrm')
    off = OxmlElement('a:off')
    off.set('x', '0')
    off.set('y', '0')
    xfrm.append(off)
    ext_el = OxmlElement('a:ext')
    ext_el.set('cx', str(cx))
    ext_el.set('cy', str(cy))
    xfrm.append(ext_el)
    spPr.append(xfrm)
    prstGeom = OxmlElement('a:prstGeom')
    prstGeom.set('prst', 'rect')
    spPr.append(prstGeom)
    pic.append(spPr)

    graphicData.append(pic)
    graphic.append(graphicData)
    inline.append(graphic)

    drawing = OxmlElement('w:drawing')
    drawing.append(inline)
    return drawing


def add_header_footer(doc):
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    section = doc.sections[0]
    section.different_first_page_header_footer = True

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]

    # Add Orchard logo to header (left side) if file exists
    if LOGO_ORCHARD.exists():
        # Get image part from package, relate to header part
        image_part = doc.part.package.get_or_add_image_part(str(LOGO_ORCHARD))
        rId = header.part.relate_to(image_part, RT.IMAGE)
        # Logo is 300x300px square — render at 0.3 inches in header
        size_emu = int(0.3 * 914400)  # 0.3 inches in EMU
        logo_run = hp.add_run()
        drawing = _build_inline_image_xml(rId, size_emu, size_emu, pic_id=10, name="Header Logo")
        logo_run._r.append(drawing)
        hp.add_run("  ")  # spacer

    run = hp.add_run("Argos — The Orchard  |  INTERNAL")
    run.font.size = Pt(8)
    run.font.color.rgb = C_GRAY
    run.font.name = "Calibri"
    run.italic = True

    # Right-align the whole header paragraph
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = C_GRAY
    run.font.name = "Calibri"
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    run._r.append(fld1)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run._r.append(instr)
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld2)


# ═══════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════

def create_cover_page(doc):
    for _ in range(2):
        doc.add_paragraph()

    # Peacock mascot logo on cover
    if LOGO_PEACOCK.exists():
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_logo = p_logo.add_run()
        run_logo.add_picture(str(LOGO_PEACOCK), width=Inches(2.5))

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ARGOS")
    run.font.size = Pt(48)
    run.font.color.rgb = C_PRIMARY
    run.bold = True
    run.font.name = "Calibri"

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Visual Speech Processing System")
    run2.font.size = Pt(22)
    run2.font.color.rgb = C_H2
    run2.font.name = "Calibri"

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("The Orchard")
    run3.font.size = Pt(20)
    run3.font.color.rgb = C_H2
    run3.font.name = "Calibri"

    doc.add_paragraph()

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run("Detailed Edition v2.0")
    run4.font.size = Pt(14)
    run4.font.color.rgb = C_H3
    run4.italic = True
    run4.font.name = "Calibri"

    for _ in range(2):
        doc.add_paragraph()

    info_lines = [
        ("Project Start:", "October 18, 2024"),
        ("Lead Engineer:", "Yoad Oxman (Solo Developer)"),
        ("Last Updated:", LAST_UPDATED),
        ("Classification:", "Internal R&D Documentation"),
    ]
    for label, value in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rl = p.add_run(label + " ")
        rl.bold = True
        rl.font.size = Pt(11)
        rl.font.color.rgb = C_DARK
        rl.font.name = "Calibri"
        rv = p.add_run(value)
        rv.font.size = Pt(11)
        rv.font.color.rgb = C_DARK
        rv.font.name = "Calibri"

    doc.add_page_break()


# ═══════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════

def create_toc(doc):
    add_heading(doc, "Table of Contents", 1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(
        "This table of contents will update when you open the document in Microsoft Word "
        "and press Ctrl+A then F9, or right-click and select 'Update Field'."
    )
    run.font.size = Pt(9)
    run.font.color.rgb = C_GRAY
    run.italic = True
    run.font.name = "Calibri"

    # Insert TOC field
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = ' TOC \\o "1-3" \\h \\z \\u '
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    p2 = doc.add_paragraph()
    r = p2.add_run()
    r._r.append(fld_char_begin)
    r2 = p2.add_run()
    r2._r.append(instr_text)
    r3 = p2.add_run()
    r3._r.append(fld_char_separate)
    r4 = p2.add_run("[Table of Contents — Update in Word]")
    r4.font.color.rgb = C_GRAY
    r4.font.size = Pt(10)
    r5 = p2.add_run()
    r5._r.append(fld_char_end)

    doc.add_page_break()


# ═══════════════════════════════════════════════════
# CHAPTER 1: Research Foundation & Literature Review
# ═══════════════════════════════════════════════════

def chapter_1(doc):
    add_heading(doc, "1. Research Foundation & Literature Review", 1)

    # 1.1 Team
    add_heading(doc, "1.1 Team & Project Inception", 2)
    add_para(doc, (
        "The Argos project began on October 18, 2024, with the goal of building a production-ready "
        "visual speech processing (lip-reading) system. The project team consists of three members, "
        "though all engineering and code development was carried out by a single developer:"
    ))
    add_styled_table(doc,
        ["Team Member", "Role", "Responsibilities"],
        [
            ["Yoad Oxman", "Lead Engineer", "All code development, pipeline architecture, UI, container deployment, testing, documentation"],
            ["Omer Leibovitch", "Data & Research", "Data sourcing, early conceptual samples, initial experiments"],
            ["Ido Springer", "Supervisor", "Management, supervision, project direction"],
        ],
        col_widths=[1.5, 1.2, 3.8]
    )
    add_para(doc, (
        "This is fundamentally a solo engineering effort. Every script, module, UI component, Docker container, "
        "and deployment artifact documented in this journal was designed and implemented by Yoad Oxman. "
        "The total codebase comprises approximately 14,882 lines of custom production code across ~55 files."
    ))

    # 1.2 VSP-LLM Paper
    add_heading(doc, "1.2 The VSP-LLM Paper", 2)
    add_para(doc, (
        "The core model is based on the VSP-LLM paper by Yeo et al. (arXiv:2402.15151v2, May 2024), "
        "titled 'Visual Speech Processing Incorporated with Large Language Models.' This paper presents "
        "a novel approach that combines visual speech features extracted by AV-HuBERT with the language "
        "understanding capabilities of LLaMA-2 through a parameter-efficient fine-tuning approach."
    ))

    add_heading(doc, "Architecture", 3)
    add_para(doc, "The VSP-LLM architecture consists of three main components connected in sequence:")
    add_code_block(doc, (
        "┌─────────────────────┐    ┌────────────────────┐    ┌──────────────────────┐\n"
        "│   AV-HuBERT Encoder │    │  Linear Projection │    │   LLaMA-2-7B + QLoRA │\n"
        "│   (315M parameters) │ -> │  (1024 -> 4096)    │ -> │   (7B params, ~0.1%  │\n"
        "│   Video -> Features │    │  Dim. alignment    │    │   trainable via LoRA) │\n"
        "└─────────────────────┘    └────────────────────┘    └──────────────────────┘\n"
        "\n"
        "Input: 88x88 mouth crops at 25 FPS\n"
        "Output: Text transcription (beam search or sampling decode)"
    ))

    add_para(doc, "The paper reports the following key results on the LRS3 benchmark dataset:")
    add_styled_table(doc,
        ["Metric", "Value", "Dataset", "Notes"],
        [
            ["Word Error Rate (WER)", "25.4%", "LRS3 (test)", "Table 1, visual-only setting"],
            ["Training Data", "433 hours", "LRS3 (labeled)", "Fine-tuning subset"],
            ["Pre-training", "1,759 hours", "LRS3 + VoxCeleb2", "AV-HuBERT self-supervised"],
            ["LLM", "LLaMA-2-7B", "—", "4-bit NF4 quantization with QLoRA"],
            ["LoRA Config", "rank=16, alpha=32", "—", "Applied to all linear layers"],
        ],
        col_widths=[1.8, 1.2, 1.5, 2.0]
    )

    # 1.3 Papers Studied
    add_heading(doc, "1.3 Literature Review (~2-3 Weeks)", 2)
    add_para(doc, (
        "Before any code was written, approximately two to three weeks were spent studying the foundational "
        "papers and understanding the full technical stack. This was essential because the VSP-LLM system "
        "sits at the intersection of multiple research areas: audio-visual speech recognition, self-supervised "
        "representation learning, large language models, and parameter-efficient fine-tuning."
    ))
    add_styled_table(doc,
        ["Paper / Framework", "Authors", "Key Contribution to Argos"],
        [
            ["VSP-LLM (arXiv:2402.15151)", "Yeo et al., 2024", "Core model architecture, decode strategy, QLoRA approach"],
            ["AV-HuBERT (arXiv:2201.02184)", "Shi et al., 2022", "Self-supervised audio-visual encoder, clustering approach"],
            ["Auto-AVSR (arXiv:2303.14307)", "Ma et al., 2023", "Preprocessing pipeline, face detection, mouth ROI extraction"],
            ["fairseq", "Ott et al. (Meta)", "Training framework, sequence generation, Hydra configs"],
            ["LRS3-TED Dataset", "Afouras et al., 2018", "Benchmark dataset structure, evaluation protocol"],
            ["LLaMA-2", "Touvron et al. (Meta), 2023", "Base LLM decoder, tokenizer, 7B architecture"],
            ["QLoRA (arXiv:2305.14314)", "Dettmers et al., 2023", "4-bit quantization + LoRA for efficient fine-tuning"],
        ],
        col_widths=[2.2, 1.5, 2.8]
    )
    add_para(doc, (
        "Understanding these papers was non-trivial. Each paper introduces its own terminology, assumptions, "
        "and dependencies. The AV-HuBERT paper alone requires understanding masked prediction pre-training, "
        "k-means clustering of speech features, and iterative refinement of cluster assignments. The fairseq "
        "framework adds its own configuration system (Hydra/OmegaConf) with deeply nested YAML files."
    ))

    # 1.4 Key Insight
    add_heading(doc, "1.4 Key Insight: Research-Grade vs. Production-Ready", 2)
    add_para(doc, (
        "A critical realization emerged early in the literature review, later formalized in the December 29, 2025 "
        "project presentation (Slide 9):"
    ))
    add_para(doc, (
        '"The paper assumes clean data, multi-GPU infrastructure, and internet connectivity. '
        'The code is research-grade, not plug-and-play."'
    ), bold=True, italic=True, color=C_H2)
    add_para(doc, (
        "This insight shaped the entire Argos project. The published code provides model definitions and training "
        "scripts, but assumes the user has already preprocessed data in LRS3 format, has multi-GPU clusters "
        "available, and can download dependencies on demand. None of these assumptions hold in the production "
        "environments where Argos needed to operate."
    ))

    # 1.5 Paper vs Argos
    add_heading(doc, "1.5 What the Paper Provides vs. What Argos Built", 2)
    add_styled_table(doc,
        ["Aspect", "Paper Provides", "Argos Built"],
        [
            ["Model Code", "PyTorch model definitions", "Same (used as-is)"],
            ["End-to-End Pipeline", "None — 'follow AV-HuBERT docs'", "Complete 8-stage pipeline (501 lines + 1,562 lib)"],
            ["Preprocessing", "Assumes LRS3 format", "'Flat' adaptation for arbitrary videos"],
            ["Video Handling", "None", "Normalization, segmentation, codec handling"],
            ["Face Detection", "References Auto-AVSR", "MediaPipe two-tier + RetinaFace integration"],
            ["ASR Ground Truth", "Assumes pre-existing labels", "Whisper integration with transcription reuse"],
            ["Evaluation", "Standard WER", "WER + WWER + NEA (809-line report generator)"],
            ["User Interface", "None (CLI only)", "Full web UI (5,780 lines, 6 screens, 12 API endpoints)"],
            ["Deployment", "None", "Docker container, 5 versions, 37 bugs fixed"],
            ["Multi-Environment", "Single machine", "4 environments (EC2, 2 standalone, HORIZON)"],
            ["Offline Operation", "Requires internet", "Full offline capability with local wheels"],
        ],
        col_widths=[1.5, 2.0, 3.0]
    )

    # 1.6 Three Upstream Repos
    add_heading(doc, "1.6 Three Upstream Repositories", 2)
    add_para(doc, (
        "Argos integrates three separate open-source repositories, each serving a distinct role in the pipeline. "
        "These repositories were forked, modified, and orchestrated into a unified system:"
    ))
    add_styled_table(doc,
        ["Repository", "Origin", "Role in Argos", "Key Modifications"],
        [
            ["auto_avsr", "Ma et al. (mpc001)", "Preprocessing: face detection, mouth cropping, video transforms",
             "MediaPipe integration, flat dataset support, fast_segment.py, overlapping_segmentation.py"],
            ["VSP-LLM", "Yeo et al. (Sally-SH)", "Model: inference, decode, training, clustering",
             "12 custom Python scripts (3,342 lines), decode configs, training configs"],
            ["av_hubert", "Meta (facebookresearch)", "Feature extraction: HuBERT encoder, clustering preparation",
             "Flat-to-LRS3 conversion, avhubert_flat variant for arbitrary datasets"],
        ],
        col_widths=[1.2, 1.3, 2.0, 2.0]
    )

    # 1.7 Architecture Overview
    add_heading(doc, "1.7 Full Pipeline Architecture", 2)
    add_para(doc, (
        "The complete Argos pipeline processes raw video files through eight sequential stages, "
        "managed by a master orchestrator script. Each stage is implemented as a reusable module "
        "in the lib/ directory:"
    ))
    add_code_block(doc, (
        "Stage 0:   Input Validation & Archive\n"
        "           └── Validate input videos, archive previous outputs\n"
        "           └── Preserve manual transcriptions across runs\n"
        "\n"
        "Stage 0.5: Video Segmentation (fast_segment.py)\n"
        "           └── Split long videos into 12-second segments\n"
        "           └── Ultra-fast codec-copy (no re-encoding)\n"
        "\n"
        "Stage 0.6: Transcription Reuse\n"
        "           └── Copy existing transcriptions from .transcriptions/\n"
        "           └── Match by exact filename\n"
        "\n"
        "Stage 1:   Video Normalization (lib/normalization.sh)\n"
        "           └── HDR/10-bit tone mapping, FPS standardization\n"
        "           └── GPU (NVENC) or CPU (libx264) encoding\n"
        "\n"
        "Stage 2:   Face Detection & Mouth Cropping (auto_avsr)\n"
        "           └── MediaPipe or RetinaFace detection\n"
        "           └── 88x88 pixel mouth ROI extraction at 25 FPS\n"
        "\n"
        "Stage 3:   ASR Transcription (Whisper)\n"
        "           └── Run Whisper on segments without transcriptions\n"
        "           └── Save new auto-transcriptions for future reuse\n"
        "\n"
        "Stage 4:   LRS3 Format Conversion\n"
        "           └── Convert flat dataset structure to LRS3 format\n"
        "           └── Generate manifests, TSVs, and metadata\n"
        "\n"
        "Stage 5:   Feature Clustering (K-Means)\n"
        "           └── Extract AV-HuBERT features\n"
        "           └── K-means clustering (200 clusters) or use golden model\n"
        "\n"
        "Stage 6:   VSP-LLM Decode\n"
        "           └── Beam search (beam=20) or sampling decode\n"
        "           └── Automatic Cython extension build for containers\n"
        "\n"
        "Stage 7:   Client Outputs\n"
        "           └── Reports (CSV, HTML, JSON, TXT, ANSI)\n"
        "           └── Burned videos with subtitle overlays\n"
        "           └── Lip crop video copies"
    ))

    # 1.8 Decode Config
    add_heading(doc, "1.8 Decode Configuration", 2)
    add_para(doc, (
        "The decode stage is controlled by a Hydra YAML configuration file (s2s_decode.yaml) that defines "
        "beam search parameters, length penalties, and generation constraints:"
    ))
    add_code_block(doc, (
        "# s2s_decode.yaml — VSP-LLM Decode Configuration\n"
        "generation:\n"
        "  beam: 20                  # Beam width for search\n"
        "  lenpen: 0.0               # Length penalty (paper's value)\n"
        "  max_len_a: 2.0            # Dynamic max_len = max_len_a * src_clusters + max_len_b\n"
        "  max_len_b: 200            # Buffer for short-input segments\n"
        "  max_len: 2048             # Hard cap on generated length\n"
        "  no_repeat_ngram_size: 3   # Prevent degenerate repetitions\n"
        "  repetition_penalty: 1.2   # Mild penalty for repeating tokens\n"
        "  do_sample: false          # false = deterministic beam search\n"
        "  temperature: 1.0          # Sampling temperature (only with do_sample=true)\n"
        "  top_p: 0.9                # Nucleus sampling threshold"
    ))

    # 1.9 Presentation Milestone
    add_heading(doc, "1.9 December 29, 2025 Presentation Milestone", 2)
    add_para(doc, (
        'On December 29, 2025, a 19-slide presentation titled "Argos — A Working Pipeline" was delivered '
        "to the team. This marked the transition from research and development to a functioning, demonstrable "
        "system. Key accomplishments highlighted in the presentation include:"
    ))
    add_bullet(doc, "HORIZON server operational — pipeline running on a closed network with no internet access")
    add_bullet(doc, "Labeled data preparation — AVSpeech dataset imported and preprocessed")
    add_bullet(doc, "Improved lip cropping — MediaPipe two-tier detection strategy")
    add_bullet(doc, "Linux standalone Docker — containerized deployment for production machines")
    add_bullet(doc, "AVSpeech support — pipeline processes arbitrary YouTube videos, not just LRS3")
    add_bullet(doc, "Whisper built in — ASR ground truth generation integrated into pipeline")
    add_bullet(doc, "Demo-ready system with web UI for non-technical users")

    doc.add_page_break()


# ═══════════════════════════════════════════════════
# CHAPTER 2: Environment & Infrastructure Setup
# ═══════════════════════════════════════════════════

def chapter_2(doc):
    add_heading(doc, "2. Environment & Infrastructure Setup", 1)

    add_para(doc, (
        "Setting up the infrastructure for Argos was one of the most time-consuming aspects of the project. "
        "The system requires GPU-accelerated computing with specific CUDA versions, two separate Python virtual "
        "environments with conflicting dependencies, and the ability to operate in environments with no internet "
        "connectivity. This chapter documents the approximately six weeks spent on environment setup alone."
    ))

    # 2.1 Four Environments
    add_heading(doc, "2.1 Four Deployment Environments", 2)
    add_styled_table(doc,
        ["Environment", "Type", "Internet", "GPU", "Purpose"],
        [
            ["AWS EC2", "Cloud instance", "Full", "NVIDIA (varies)", "Development, testing, experiments"],
            ["Standalone Ubuntu #1", "Physical machine", "Limited", "NVIDIA GPU", "Production deployment target"],
            ["Standalone Ubuntu #2", "Physical machine", "Limited", "NVIDIA GPU", "Secondary production machine"],
            ["HORIZON Server", "Closed network", "None", "NVIDIA GPU", "Secure deployment (no internet)"],
        ],
        col_widths=[1.4, 1.2, 0.8, 1.2, 1.9]
    )
    add_para(doc, (
        "Each environment presented unique challenges. The EC2 instance served as the primary development "
        "environment where all code was first written and tested. Changes were then replicated to the standalone "
        "machines and HORIZON server, accounting for path differences (/home/ubuntu/ on EC2 vs /workspace/ in "
        "containers) and the absence of internet connectivity."
    ))
    add_para(doc, (
        "Making the HORIZON server operational was a significant milestone (highlighted in the December 2025 "
        "presentation). This closed-network server has zero internet access, meaning every Python package, "
        "model weight, and system dependency had to be pre-packaged and transferred physically."
    ))

    # 2.2 Failed Venv
    add_heading(doc, "2.2 Failed Virtual Environment Build Attempts (~2 Weeks)", 2)
    add_para(doc, (
        "The first attempt to set up the software environment on the standalone computers consumed approximately "
        "two weeks and ultimately failed. The core issue was attempting to build Python virtual environments "
        "directly on machines with limited internet access and non-standard system configurations."
    ))
    add_para(doc, "Key problems encountered during the failed venv builds:")
    add_bullet(doc, "System Python versions incompatible with required packages (needed Python 3.9+ for PyTorch)")
    add_bullet(doc, "pip failing to resolve dependency conflicts between fairseq and transformers")
    add_bullet(doc, "Compilation failures for packages requiring CUDA headers (torch, triton, bitsandbytes)")
    add_bullet(doc, "Network timeouts on machines with restricted or slow internet access")
    add_bullet(doc, "Version pinning conflicts — omegaconf==2.0.6 requires old pip for metadata compatibility")
    add_para(doc, (
        "The lesson learned was that building venvs from scratch on production machines was not viable. "
        "The solution that eventually worked was to build complete virtual environments on the EC2 instance "
        "and transfer them as part of the container deployment package."
    ))

    # 2.3 NVIDIA Driver
    add_heading(doc, "2.3 NVIDIA Driver & CUDA Installation (~1 Month)", 2)
    add_para(doc, (
        "After the venv build failures, approximately one month was spent on getting GPU acceleration working "
        "on the standalone machines. This involved installing NVIDIA drivers, CUDA toolkit, and verifying "
        "that PyTorch could communicate with the GPU."
    ))
    add_para(doc, "The installation process involved multiple iterations:")
    add_bullet(doc, "Identifying the correct NVIDIA driver version for the installed GPU hardware")
    add_bullet(doc, "Installing the NVIDIA driver (NVIDIA-Linux-x86_64-535.183.01.run or similar)")
    add_bullet(doc, "Installing CUDA toolkit (multiple versions tried: 12.1, 12.4, 12.8)")
    add_bullet(doc, "Resolving kernel module conflicts with existing nouveau drivers")
    add_bullet(doc, "Verifying GPU visibility with nvidia-smi and CUDA availability in Python")
    add_bullet(doc, "Testing that PyTorch CUDA operations work correctly (tensor operations, memory allocation)")
    add_para(doc, (
        "A particular challenge was that the pipeline requires two different CUDA versions: CUDA 12.8 for "
        "the preprocessing venv (with triton 3.4.0) and CUDA 12.4 compatibility for the decode venv "
        "(with triton 3.1.0). This was resolved through careful library path management within each "
        "virtual environment."
    ))

    # 2.4 Dual Venv
    add_heading(doc, "2.4 Dual Virtual Environment Strategy", 2)
    add_para(doc, (
        "The most significant infrastructure decision was adopting a dual virtual environment architecture. "
        "The preprocessing and decoding stages have fundamentally incompatible dependency requirements "
        "that cannot coexist in a single environment:"
    ))
    add_styled_table(doc,
        ["Aspect", "pre-process-venv (ASR)", "vsp-llm-yoad-venv (Decode)"],
        [
            ["PyTorch", "2.8.0+cu128", "2.5.1+cu124"],
            ["Triton", "3.4.0", "3.1.0"],
            ["CUDA Runtime", "12.8.*", "12.4.*"],
            ["Key Packages", "Whisper, MediaPipe, OpenCV, JAX", "fairseq, transformers, peft, bitsandbytes"],
            ["Package Count", "~135 packages", "~256 packages"],
            ["Purpose", "Face detection, mouth cropping, ASR", "K-means, decode, training"],
            ["Python", "3.9+", "3.9+"],
        ],
        col_widths=[1.3, 2.6, 2.6]
    )
    add_para(doc, (
        "The root conflict is between triton versions. The preprocessing pipeline requires triton 3.4.0 "
        "(which needs CUDA 12.8 runtime libraries), while fairseq's Cython extensions and the decode pipeline "
        "require triton 3.1.0 (CUDA 12.4). Installing both in one environment leads to import errors "
        "and GPU operation failures."
    ))

    # 2.5 Dependency Hell
    add_heading(doc, "2.5 Dependency Conflicts & Resolution", 2)
    add_para(doc, "Several specific dependency conflicts required creative solutions:")
    add_bullet_bold_value(doc, "omegaconf==2.0.6: ",
        "fairseq requires this exact version, but its package metadata format is incompatible with modern pip. "
        "Solution: install using an older pip version, then upgrade pip afterward.")
    add_bullet_bold_value(doc, "fairseq installation: ",
        "Required as an editable install from the forked repository (Sally-SH/VSP-LLM) because the model "
        "code adds custom tasks and criteria that must be importable.")
    add_bullet_bold_value(doc, "numpy version: ",
        "Preprocessing needs numpy 1.26.4 (for MediaPipe compatibility), while decode works with numpy 2.x. "
        "Resolved by pinning in each venv independently.")
    add_bullet_bold_value(doc, "Cython extensions: ",
        "fairseq's data_utils_fast Cython module must be compiled on the target machine's architecture. "
        "The decode module includes automatic detection and compilation on first run.")

    # 2.6 HORIZON
    add_heading(doc, "2.6 HORIZON: Operating Without Internet", 2)
    add_para(doc, (
        "The HORIZON server operates on a completely isolated network with no internet connectivity. "
        "This required pre-packaging every dependency:"
    ))
    add_bullet(doc, "All Python packages as wheel files (.whl) — transferred physically to the server")
    add_bullet(doc, "spaCy language model (en_core_web_sm) pre-downloaded with offline wheels")
    add_bullet(doc, "Whisper model files (medium.pt) pre-cached in the whisper_cache directory")
    add_bullet(doc, "NVIDIA CUDA runtime libraries bundled within the container")
    add_bullet(doc, "Model weights (checkpoint_finetune.pt, LLaMA-2-7B) pre-loaded")
    add_para(doc, (
        "The outputs module (lib/outputs.sh) implements a graceful degradation strategy for spaCy: "
        "it first tries to install from local wheels, falls back to online pip if available, and "
        "finally degrades to basic WER metrics if neither option succeeds. This ensures the pipeline "
        "produces useful output even when advanced NER metrics are unavailable."
    ))

    # 2.7 Model Weights
    add_heading(doc, "2.7 Model Weights Acquisition", 2)
    add_para(doc, (
        "The pre-trained model weights were obtained from the VSP-LLM paper's official Git repository "
        "(Sally-SH/VSP-LLM). These weights represent the model after the authors' own fine-tuning on "
        "the LRS3 dataset — specifically the checkpoint_finetune.pt file containing the AV-HuBERT encoder "
        "weights and linear projection layer. The LLaMA-2-7B base model was separately downloaded "
        "from Meta's distribution."
    ))

    # 2.8 AVSpeech Data
    add_heading(doc, "2.8 AVSpeech Data Sourcing", 2)
    add_para(doc, (
        "Training and evaluation data was sourced from the AVSpeech dataset, which contains short clips "
        "of people speaking extracted from YouTube videos. Unlike the LRS3 dataset (TED talks with clean "
        "audio and frontal faces), AVSpeech represents real-world conditions with diverse camera angles, "
        "lighting, and background noise. The data was imported from an old server that had been previously "
        "mounted, making it accessible for processing through the pipeline."
    ))

    doc.add_page_break()


# ═══════════════════════════════════════════════════
# CHAPTER 3: Building the Pipeline from Scratch
# ═══════════════════════════════════════════════════

def chapter_3(doc):
    add_heading(doc, "3. Building the Pipeline from Scratch", 1)

    # 3.1 The Gap
    add_heading(doc, "3.1 The Gap: No Pipeline Existed", 2)
    add_para(doc, (
        "The single biggest engineering challenge in the Argos project was that no end-to-end pipeline existed. "
        "The VSP-LLM paper provides model code (PyTorch modules, training scripts, decode scripts) but assumes "
        "the user has already preprocessed their data into LRS3 format. The paper's instructions effectively say: "
        "'For preprocessing, follow the AV-HuBERT documentation. For face detection, follow Auto-AVSR.' "
        "This means three separate research codebases needed to be understood, modified, and orchestrated "
        "into a unified pipeline that could process raw video files into text transcriptions."
    ))
    add_para(doc, (
        "Building this pipeline from scratch — connecting auto_avsr's preprocessing with av_hubert's feature "
        "extraction and VSP-LLM's decode — required deep understanding of each repository's data formats, "
        "directory conventions, configuration systems, and Python import structures."
    ))

    # 3.2 Flat Dataset Adaptation
    add_heading(doc, '3.2 "Flat" Dataset Adaptation', 2)
    add_para(doc, (
        "The paper's code exclusively supports the LRS3 dataset format, which has a rigid directory structure "
        "organized by TED talk ID with specific naming conventions. Real-world deployment requires processing "
        "arbitrary video files that don't follow this structure."
    ))
    add_para(doc, (
        'The "flat" adaptation was a fundamental modification to the preprocessing pipeline that allows '
        "it to accept any collection of video files in a single directory. This required changes to:"
    ))
    add_bullet(doc, "Directory traversal logic — scan a flat directory instead of nested LRS3 paths")
    add_bullet(doc, "Manifest generation — create TSV files mapping flat filenames to features")
    add_bullet(doc, "Naming conventions — establish a segment naming system for arbitrary videos")
    add_bullet(doc, "LRS3 compatibility layer — convert flat dataset to LRS3 format for downstream stages")
    add_para(doc, (
        "This adaptation is what enabled Argos to process AVSpeech data and any other video source, "
        "rather than being limited to the LRS3 dataset the paper was designed for."
    ))

    # 3.3 First Pipeline
    add_heading(doc, "3.3 First Pipeline Version (Oct-Nov 2024)", 2)
    add_para(doc, (
        "The first pipeline version was approximately 285 lines of bash script, created in the galaxy_export "
        "directory. It implemented the basic 8-stage flow: face detection, mouth cropping, LRS3 conversion, "
        "manifest generation, k-means clustering, and VSP-LLM decode. This version was functional but "
        "fragile — it assumed specific directory structures, had hardcoded paths, and offered no error "
        "recovery or transcription reuse."
    ))

    # 3.4 Pipeline Evolution
    add_heading(doc, "3.4 Pipeline Evolution: 7 Versions", 2)
    add_para(doc, (
        "The pipeline script underwent seven major revisions, each adding significant capabilities. "
        "The following table tracks the evolution by line count and key features added:"
    ))
    add_styled_table(doc,
        ["Version", "Lines", "Key Changes"],
        [
            ["v1 (Initial)", "~240", "Basic sequential flow, hardcoded paths"],
            ["v2 (galaxy_export)", "~285", "Galaxy export packaging, input validation"],
            ["v3 (Segmentation)", "~410", "Video segmentation (4s windows), overlap support"],
            ["v4 (Normalization)", "~612", "Video normalization, ASR integration, HDR handling"],
            ["v5 (Transcription)", "~673", "Transcription reuse system, manual transcription support"],
            ["v6 (Monolithic)", "823", "Full feature set in single file, all stages complete"],
            ["v7 (Modular)", "501 + 1,562", "Refactored into 12 lib/ modules (see Chapter 8)"],
        ],
        col_widths=[1.5, 1.0, 4.0]
    )
    add_para(doc, (
        "The transition from v6 (823 lines) to v7 (501 + 1,562 lines) represents the modular refactoring "
        "detailed in Chapter 8. While the total line count increased, the main orchestrator script was reduced "
        "by 52%, and the extracted modules became independently testable and reusable across environments."
    ))

    # 3.5 Integration Challenges
    add_heading(doc, "3.5 Key Integration Challenges", 2)
    add_para(doc, (
        "Connecting three research repositories into a unified pipeline required solving several "
        "cross-cutting concerns:"
    ))
    add_bullet_bold_value(doc, "Path Management: ",
        "Each repository assumes its own directory structure. The pipeline must translate between "
        "auto_avsr's flat video paths, av_hubert's LRS3-style paths, and VSP-LLM's dataset directory "
        "conventions. This is handled by lib/config.sh which provides environment-aware path functions.")
    add_bullet_bold_value(doc, "Virtual Environment Switching: ",
        "The preprocessing stages (face detection, ASR) use the pre-process-venv, while stages 5-8 "
        "(clustering, decode, outputs) use vsp-llm-yoad-venv. The pipeline must activate and deactivate "
        "venvs at the correct boundaries without contaminating the shell environment.")
    add_bullet_bold_value(doc, "Data Format Translation: ",
        "auto_avsr produces mouth crop videos as .mp4 files. av_hubert expects features as .npy arrays. "
        "VSP-LLM expects fairseq-style manifests with .tsv and .wrd files. Each transition requires "
        "format conversion scripts.")
    add_bullet_bold_value(doc, "Error Propagation: ",
        "A failure in face detection (e.g., no face found in a segment) must be handled gracefully "
        "without stopping the entire pipeline. The modular architecture uses log_warn for non-fatal "
        "issues and log_error for stage-level failures.")

    # 3.6 Custom Python Scripts
    add_heading(doc, "3.6 Custom Python Scripts (13 Files, 3,342 Lines)", 2)
    add_para(doc, (
        "In addition to the pipeline orchestrator and library modules, 13 custom Python scripts were written "
        "to handle specific processing tasks. These scripts are the workhorses of the pipeline, handling "
        "everything from video segmentation to quality assessment:"
    ))
    add_styled_table(doc,
        ["Script", "Lines", "Location", "Purpose"],
        [
            ["make_report.py", "809", "VSP-LLM/scripts/", "WER/WWER/NEA metrics, color-coded HTML reports"],
            ["make_burn.py", "510", "VSP-LLM/scripts/", "Burn subtitle overlays onto videos with ffmpeg"],
            ["overlapping_segmentation.py", "390", "auto_avsr/preparation/", "Time-based overlap splitting algorithm"],
            ["transcribe_segments.py", "353", "VSP-LLM/scripts/", "Per-segment transcription management tool"],
            ["calculate_per_video_wer.py", "300", "VSP-LLM/scripts/", "Per-video WER with overlap deduplication"],
            ["merge_overlapping_predictions.py", "289", "VSP-LLM/scripts/", "Conflict resolution for overlapping segments"],
            ["fast_segment.py", "197", "auto_avsr/preparation/", "Ultra-fast codec-copy video segmentation"],
            ["asr_to_words_notime.py", "151", "auto_avsr/", "Whisper ASR wrapper — audio to word tokens"],
            ["generate_conflict_report.py", "149", "VSP-LLM/scripts/", "Analyze and report prediction conflicts"],
            ["build_flat_train_tsv.py", "93", "VSP-LLM/scripts/", "Generate training TSV manifests"],
            ["make_flat_tsv.py", "56", "VSP-LLM/scripts/", "Create flat-format TSV files"],
            ["build_flat_cluster_counts.py", "45", "VSP-LLM/scripts/", "Generate cluster count statistics"],
            ["inspect_km_hist.py", "38", "VSP-LLM/scripts/", "K-means cluster histogram visualization"],
        ],
        col_widths=[2.0, 0.5, 1.8, 2.2]
    )

    # 3.7 Transcription Reuse
    add_heading(doc, "3.7 Transcription Reuse System", 2)
    add_para(doc, (
        "One of the most impactful features for production use is the transcription reuse system, "
        "implemented across three pipeline steps (0.6, 3, and 1.5) in lib/asr.sh. This system ensures "
        "that Whisper ASR only runs on segments that don't already have transcriptions, saving hours "
        "of processing time on pipeline re-runs."
    ))
    add_code_block(doc, (
        "Step 0.6: Copy Existing Transcriptions\n"
        "  └── Check .transcriptions/ directory for existing .wrd files\n"
        "  └── Match transcription to segment by exact filename\n"
        "  └── Copy matching transcriptions to segment_wrd_tmp/\n"
        "  └── Whisper will skip segments that already have .wrd files\n"
        "\n"
        "Step 3: Run Whisper ASR\n"
        "  └── python3 asr_to_words_notime.py \\\n"
        "        --in_videos <segment_dir> --out_wrd <wrd_dir> \\\n"
        "        --model medium --lang en --tokenize alnum --lower\n"
        "  └── Only processes segments WITHOUT existing .wrd files\n"
        "\n"
        "Step 1.5: Save New Auto-Transcriptions\n"
        "  └── Copy new Whisper outputs to .transcriptions/ directory\n"
        "  └── Mark as 'auto' type in metadata.json\n"
        "  └── Skip if manual transcription already exists (preserve user edits)\n"
        "  └── Next run, Step 0.6 will reuse these transcriptions"
    ))
    add_para(doc, (
        "This system means that a user can manually transcribe specific segments via the web UI, "
        "and those manual transcriptions will persist across all future pipeline runs. Whisper never "
        "overwrites manual work, and auto-transcriptions are cached for reuse."
    ))

    doc.add_page_break()


# ═══════════════════════════════════════════════════
# PLACEHOLDER — remaining chapters will be added
# ═══════════════════════════════════════════════════

def chapter_4(doc):
    add_heading(doc, "4. Video Normalization", 1)

    add_para(doc, (
        "Real-world video files arrive in an enormous variety of formats: different codecs (H.264, H.265, VP9, AV1), "
        "color spaces (SDR, HDR, BT.2020), bit depths (8-bit, 10-bit), frame rates (24-60 FPS), and resolutions. "
        "The face detection and model inference stages require standardized input: 8-bit SDR video at a consistent "
        "frame rate. The normalization module handles this conversion, and its development uncovered one of the "
        "most challenging production bugs in the entire project."
    ))

    # 4.1 Why Needed
    add_heading(doc, "4.1 Why Normalization Is Needed", 2)
    add_para(doc, (
        "Without normalization, the pipeline fails in multiple ways. HDR (High Dynamic Range) and 10-bit video "
        "causes face detection to produce incorrect landmarks because the pixel value ranges differ from what "
        "the model expects. Variable frame rates cause the model to receive inconsistent temporal information, "
        "degrading transcription quality. Non-standard codecs may not be compatible with the preprocessing "
        "pipeline's ffmpeg-based extraction."
    ))
    add_para(doc, "The normalization module standardizes all input to:")
    add_bullet(doc, "Codec: H.264 (universally compatible)")
    add_bullet(doc, "Color space: SDR (BT.709), 8-bit")
    add_bullet(doc, "Frame rate: 25 FPS (matching LRS3 training data)")
    add_bullet(doc, "Pixel format: yuv420p")

    # 4.2 The Module
    add_heading(doc, "4.2 The Normalization Module (lib/normalization.sh, 235 Lines)", 2)
    add_para(doc, (
        "The normalization module is the largest single module in the lib/ directory. It implements "
        "HDR detection, GPU-accelerated encoding, CPU fallback, and post-encode validation."
    ))

    add_heading(doc, "HDR Detection", 3)
    add_para(doc, (
        "The needs_tonemap() function uses ffprobe to detect videos that require tone mapping "
        "from HDR to SDR. It checks two indicators:"
    ))
    add_code_block(doc, (
        "needs_tonemap() {\n"
        '  local pix_fmt=$(ffprobe -v error -select_streams v:0 \\\n'
        '    -show_entries stream=pix_fmt -of csv=p=0 "$1")\n'
        '  local color_space=$(ffprobe -v error -select_streams v:0 \\\n'
        '    -show_entries stream=color_space -of csv=p=0 "$1")\n'
        '\n'
        '  # 10-bit pixel formats indicate HDR\n'
        '  if [[ "$pix_fmt" == *"10le"* ]] || [[ "$pix_fmt" == *"10be"* ]]; then\n'
        '    return 0  # needs tone mapping\n'
        '  fi\n'
        '  # BT.2020 color space indicates HDR\n'
        '  if [[ "$color_space" == *"bt2020"* ]]; then\n'
        '    return 0\n'
        '  fi\n'
        '  return 1  # standard SDR\n'
        '}'
    ))

    add_heading(doc, "GPU vs. CPU Encoding Paths", 3)
    add_styled_table(doc,
        ["Aspect", "GPU Path (NVENC)", "CPU Path (libx264)"],
        [
            ["Encoder", "h264_nvenc", "libx264"],
            ["Preset", "-preset p7", "-preset slow"],
            ["Quality", "-rc vbr -cq 18", "-crf 18"],
            ["Scaling", "scale_cuda filter", "scale filter"],
            ["HDR Support", "Falls back to CPU", "zscale tonemap"],
            ["Speed", "~5-10x real-time", "~1-2x real-time"],
            ["Requirement", "NVIDIA GPU with NVENC", "Any CPU"],
        ],
        col_widths=[1.3, 2.6, 2.6]
    )

    # 4.3 NVENC Bug
    add_heading(doc, "4.3 NVENC Silent Corruption Bug", 2)
    add_para(doc, (
        "The NVENC silent corruption bug was the most severe production issue encountered in the entire project. "
        "It caused 43% of normalized video segments to be silently corrupted — the ffmpeg command exited with "
        "code 0 (success), the output file existed and had non-zero size, but the video was undecodable. "
        "Face detection would then fail on these segments, and no transcription would be produced."
    ), bold=True)

    add_para(doc, "Three root causes were identified:")

    add_heading(doc, "Root Cause 1: NVENC Hardware Encoder Corruption", 3)
    add_para(doc, (
        "NVIDIA's hardware H.264 encoder (NVENC) silently produces corrupt output streams when given certain "
        "input codec combinations. Unlike software encoders that error out, NVENC writes a file that appears "
        "valid on disk but contains malformed NAL units that cannot be decoded. The GPU encoder's internal "
        "error handling does not propagate failures to the ffmpeg process."
    ))

    add_heading(doc, "Root Cause 2: Bash File Descriptor Interference", 3)
    add_para(doc, (
        "When the normalization function runs inside a while-read loop (iterating over video files), "
        "ffmpeg inherits the shell's file descriptors. ffmpeg reads from stdin by default, which "
        "interferes with the while-read loop's input stream. This causes the loop to skip files "
        "or terminate prematurely."
    ))
    add_code_block(doc, (
        "# BROKEN: ffmpeg inherits fd 0, interferes with while-read\n"
        'while IFS= read -r f; do\n'
        '  ffmpeg -i "$f" -c:v h264_nvenc ... "$out"   # eats stdin!\n'
        'done < <(find "$dir" -name "*.mp4")\n'
        "\n"
        "# FIXED: redirect ffmpeg stdin from /dev/null, use fd 3\n"
        'while IFS= read -r f <&3; do\n'
        '  ffmpeg -nostdin -i "$f" -c:v h264_nvenc ... "$out" < /dev/null\n'
        'done 3< <(find "$dir" -name "*.mp4")'
    ))

    add_heading(doc, "Root Cause 3: No Post-Encode Validation", 3)
    add_para(doc, (
        "The original code assumed that a successful ffmpeg exit code (0) meant the output was valid. "
        "This is not true for hardware encoders. The fix adds a mandatory validation step:"
    ))
    add_code_block(doc, (
        "# Post-encode validation — catches silent NVENC corruption\n"
        'if ! ffmpeg -v error -i "$out" -vframes 1 -f null - 2>/dev/null; then\n'
        '  log_warn "Corrupt output detected for ${bn}, falling back to raw copy"\n'
        '  rm -f "$out"\n'
        '  cp "$f" "${output_dir}/${bn}"  # raw copy as fallback\n'
        'fi'
    ))

    # 4.4 Segment-first
    add_heading(doc, "4.4 Segment-First Architecture", 2)
    add_para(doc, (
        "An important architectural decision was to segment videos BEFORE normalization, rather than "
        "normalizing the full video first. This approach has several advantages:"
    ))
    add_bullet(doc, "Speed: normalizing a 12-second segment is much faster than normalizing a 60-minute video")
    add_bullet(doc, "Parallelism: segments can be normalized independently")
    add_bullet(doc, "Fault isolation: a corrupt segment doesn't affect other segments from the same video")
    add_bullet(doc, "Efficiency: only segments that pass validation are normalized (skip excluded segments)")

    # 4.5 Failsafes
    add_heading(doc, "4.5 Failsafe Mechanisms", 2)
    add_para(doc, "The normalization module includes multiple layers of protection against failure:")
    add_bullet(doc, "Per-video timeout (default 600 seconds) with SIGKILL after 5-second grace period")
    add_bullet(doc, 'Input flags: -fflags +genpts+discardcorrupt -err_detect ignore_err')
    add_bullet(doc, "Increased probe: -probesize 10M -analyzeduration 10M for detecting unusual formats")
    add_bullet(doc, "Fallback to raw copy if encoding fails (ensures some output for every input)")
    add_bullet(doc, "Statistics tracking: norm_ok, norm_fail, norm_timeout, fallback_copy counters")

    doc.add_page_break()


def chapter_5(doc):
    add_heading(doc, "5. Video Segmentation", 1)

    add_para(doc, (
        "The VSP-LLM model processes video in fixed-length windows. Input segments that are too long "
        "degrade model performance (the attention mechanism's effectiveness drops with sequence length), "
        "while segments that are too short may not contain enough visual information for meaningful "
        "transcription. The segmentation system splits long videos into optimal-length segments with "
        "overlap to prevent information loss at segment boundaries."
    ))

    # 5.1 Why Segment
    add_heading(doc, "5.1 Why Segmentation Is Needed", 2)
    add_para(doc, (
        "The model was trained on LRS3 clips that are typically 3-12 seconds long. Processing a full "
        "60-second video as a single input leads to poor transcription quality because:"
    ))
    add_bullet(doc, "The AV-HuBERT encoder produces a very long feature sequence that exceeds the model's effective context window")
    add_bullet(doc, "Beam search becomes computationally expensive with long sequences (memory grows quadratically)")
    add_bullet(doc, "The model tends to hallucinate or repeat text on long inputs")
    add_bullet(doc, "Error in one part of the video contaminates the entire transcription")

    # 5.2 fast_segment.py
    add_heading(doc, "5.2 Fast Segmentation (fast_segment.py, 197 Lines)", 2)
    add_para(doc, (
        "The fast_segment.py script implements ultra-fast video segmentation using ffmpeg's codec-copy mode "
        "(-c copy). This means no re-encoding occurs — the script simply locates keyframe boundaries and "
        "splits the video into segments. This is orders of magnitude faster than re-encoding because "
        "it only manipulates the container format, not the video data."
    ))
    add_code_block(doc, (
        "# Core segmentation command (codec-copy, no re-encoding)\n"
        "ffmpeg -ss {start_time} -i {input_video} \\\n"
        "  -t {segment_duration} \\\n"
        "  -c copy \\\n"
        "  -avoid_negative_ts make_zero \\\n"
        "  {output_segment}"
    ))
    add_para(doc, (
        "The script handles edge cases such as videos shorter than the segment duration (passed through "
        "unchanged), videos with no audio track, and videos with variable frame rates."
    ))

    # 5.3 Overlapping Segmentation
    add_heading(doc, "5.3 Overlapping Segmentation (overlapping_segmentation.py, 390 Lines)", 2)
    add_para(doc, (
        "The overlapping segmentation module extends basic segmentation with time-based overlap windows. "
        "With 12-second segments and 2-second overlap, a 30-second video produces segments at:"
    ))
    add_code_block(doc, (
        "Segment 0:  0:00 - 0:12  (12 seconds)\n"
        "Segment 1:  0:10 - 0:22  (12 seconds, 2s overlap with seg 0)\n"
        "Segment 2:  0:20 - 0:30  (10 seconds, 2s overlap with seg 1)\n"
        "\n"
        "Overlap regions contain duplicate content that can be used for:\n"
        "  - Boundary word recovery (words split at segment edges)\n"
        "  - Confidence comparison (same content decoded twice)\n"
        "  - Conflict resolution (merge_overlapping_predictions.py)"
    ))

    # 5.4 Naming Convention
    add_heading(doc, "5.4 Segment Naming Convention", 2)
    add_para(doc, "Every segment follows a strict naming convention that encodes its source video and temporal location:")
    add_styled_table(doc,
        ["Component", "Format", "Example", "Meaning"],
        [
            ["Base name", "{original_video}", "Obama_speech", "Source video filename (without extension)"],
            ["Segment index", "_{idx:02d}_", "_00_", "Zero-indexed segment number"],
            ["Start frame", "{start:06d}", "000000", "Start position in frames"],
            ["End frame", "_{end:06d}", "_000300", "End position in frames"],
            ["Extension", ".mp4", ".mp4", "Always MP4"],
        ],
        col_widths=[1.2, 1.3, 1.5, 2.5]
    )
    add_para(doc, (
        'Full example: Obama_speech_00_000000_000300.mp4 — first segment of "Obama_speech.mp4", '
        "frames 0 to 300 (at 25 FPS = 12 seconds)."
    ))

    # 5.5 4s to 12s
    add_heading(doc, "5.5 Evolution: 4-Second to 12-Second Segments", 2)
    add_para(doc, (
        "The segmentation window was changed from 4 seconds to 12 seconds during development. The original "
        "4-second windows were too short for meaningful transcription — many segments contained only partial "
        "words or silence. The 12-second window provides substantially more context for the model."
    ))
    add_styled_table(doc,
        ["Aspect", "4-Second Segments", "12-Second Segments"],
        [
            ["Context per segment", "Very limited (~2-4 words)", "Adequate (~8-20 words)"],
            ["Segments per minute", "~15", "~5"],
            ["Boundary effects", "Many cut words", "Few cut words"],
            ["Model performance", "High error rate", "Better transcription quality"],
            ["Processing time", "More overhead per segment", "Fewer segments to process"],
        ],
        col_widths=[1.8, 2.3, 2.4]
    )

    # 5.6 Non-segmented bug
    add_heading(doc, "5.6 Non-Segmented Video Naming Bug", 2)
    add_para(doc, (
        "A bug was discovered where videos shorter than the segment duration (e.g., a 10-second clip with "
        "12-second windows) were passed through without segmentation but retained their original filename. "
        "This caused downstream stages to fail because they expected the segment naming convention. The fix "
        "ensures that even non-segmented videos are renamed to follow the convention "
        "(e.g., short_clip.mp4 becomes short_clip_00_000000_000250.mp4)."
    ))

    # 5.7 Why No Merging
    add_heading(doc, "5.7 Why Segments Are Not Merged Back", 2)
    add_para(doc, (
        "An early design choice was to keep segment-level outputs rather than merging predictions back "
        "into a single per-video transcription. This decision was made for four reasons:"
    ))
    add_bullet(doc, "Segment-level quality assessment: each segment gets its own WER score, making it easy to identify problematic segments")
    add_bullet(doc, "Overlap handling complexity: merging overlapping predictions requires conflict resolution (which words to keep at boundaries)")
    add_bullet(doc, "User value: burned videos with per-segment subtitles are more useful for manual review than a single long transcript")
    add_bullet(doc, "Debugging: segment-level output makes it easier to trace issues back to specific video moments")

    doc.add_page_break()


def chapter_6(doc):
    add_heading(doc, "6. Face Detection & Mouth Cropping", 1)

    add_para(doc, (
        "The VSP-LLM model requires tightly-cropped mouth region videos as input — specifically, 88x88 pixel "
        "grayscale crops centered on the speaker's mouth, resampled to 25 frames per second. Producing these "
        "crops from arbitrary video requires robust face detection, landmark estimation, spatial transformation, "
        "and temporal interpolation. This chapter documents the detection pipeline and the modifications made "
        "to improve its reliability in production conditions."
    ))

    # 6.1 Pipeline
    add_heading(doc, "6.1 Detection Pipeline Overview", 2)
    add_para(doc, "The face detection and cropping pipeline processes each video segment through four stages:")
    add_code_block(doc, (
        "1. Face Detection (detector.py)\n"
        "   └── Detect faces in every frame\n"
        "   └── Extract 4 keypoints per face (eyes + mouth corners)\n"
        "   └── Select largest face per frame\n"
        "\n"
        "2. Landmark Interpolation (video_process.py)\n"
        "   └── Fill in missing detections via temporal interpolation\n"
        "   └── Smooth landmark trajectories across frames\n"
        "\n"
        "3. Spatial Transformation (video_process.py)\n"
        "   └── Compute similarity transform from landmarks\n"
        "   └── Align and crop mouth region (88x88 pixels)\n"
        "\n"
        "4. Output Generation\n"
        "   └── Write grayscale mouth crop video at 25 FPS\n"
        "   └── Save as .mp4 in LRS3-compatible directory structure"
    ))

    # 6.2 MediaPipe
    add_heading(doc, "6.2 MediaPipe Detector (52 Lines)", 2)
    add_para(doc, (
        "The primary face detector uses Google's MediaPipe library with a two-tier detection strategy. "
        "This was an improvement over using a single detection model, which would fail on either "
        "distant or close-up faces depending on the model selection."
    ))
    add_code_block(doc, (
        "class LandmarksDetector:\n"
        "    def __init__(self):\n"
        "        # Two-tier strategy:\n"
        "        #   model_selection=1: full-range model (detects distant faces)\n"
        "        #   model_selection=0: short-range model (detects close-up faces)\n"
        "        self.min_detection_confidence = 0.5\n"
        "\n"
        "    def detect(self, frames):\n"
        "        # Try full-range model first\n"
        "        results = self._detect_with_model(frames, model_selection=1)\n"
        "        if all frames failed:\n"
        "            # Fallback to short-range model\n"
        "            results = self._detect_with_model(frames, model_selection=0)\n"
        "        if still all frames failed:\n"
        '            raise RuntimeError("Cannot detect any frames in the video")\n'
        "        return results  # numpy array [num_faces, 4 keypoints, (x, y)]"
    ))
    add_para(doc, (
        "The two-tier approach significantly improved detection rates in production. YouTube videos contain "
        "a wide range of camera distances — from news anchors at arm's length to speakers at podiums "
        "across a room. The full-range model handles distant faces, while the short-range model catches "
        "close-up selfie-style videos."
    ))

    # 6.3 Video Processing
    add_heading(doc, "6.3 Video Processing (video_process.py, 220 Lines)", 2)
    add_para(doc, (
        "The video processing module handles the transformation from detected face landmarks to "
        "cropped mouth region videos. Key operations include:"
    ))
    add_bullet_bold_value(doc, "Landmark Interpolation: ",
        "When face detection fails on individual frames (due to motion blur, occlusion, or camera "
        "transitions), the module interpolates missing landmarks from neighboring frames. This "
        "prevents jarring jumps in the crop position.")
    add_bullet_bold_value(doc, "Similarity Transform: ",
        "A 2D similarity transform (rotation, scaling, translation) is computed from the eye and "
        "mouth corner landmarks to align the face to a canonical orientation. This ensures the "
        "mouth region is consistently positioned regardless of head tilt or camera angle.")
    add_bullet_bold_value(doc, "cut_patch() Function: ",
        "Extracts the 88x88 pixel mouth region from the aligned frame. The crop is centered on "
        "the midpoint between the mouth corner landmarks, with a fixed-size window that captures "
        "the lips and immediate surrounding area.")
    add_bullet_bold_value(doc, "Temporal Resampling: ",
        "Input videos at various frame rates (24, 30, 60 FPS) are resampled to exactly 25 FPS "
        "to match the training data distribution.")

    # 6.4 RetinaFace
    add_heading(doc, "6.4 RetinaFace Alternative (260 Lines)", 2)
    add_para(doc, (
        "An alternative face detection backend using RetinaFace is also available. RetinaFace is "
        "a deep learning-based face detector that provides higher accuracy than MediaPipe, especially "
        "for difficult angles and partial occlusions, but at significantly slower inference speed."
    ))
    add_styled_table(doc,
        ["Aspect", "MediaPipe", "RetinaFace"],
        [
            ["detector.py", "52 lines", "43 lines"],
            ["video_process.py", "220 lines", "217 lines"],
            ["Speed", "Fast (CPU-based)", "Slower (GPU-based)"],
            ["Accuracy", "Good for frontal faces", "Better for difficult angles"],
            ["Dependencies", "mediapipe", "torch, torchvision"],
            ["Use Case", "Default (production)", "When MediaPipe fails"],
        ],
        col_widths=[1.5, 2.5, 2.5]
    )

    # 6.5 Transforms
    add_heading(doc, "6.5 Data Transforms (transforms.py, 171 Lines)", 2)
    add_para(doc, (
        "The transforms module provides data augmentation and preprocessing operations applied to "
        "mouth crop videos before they are fed to the model. These include spatial transforms "
        "(random crops, flips for training), normalization (mean subtraction, variance scaling), "
        "and format conversion (grayscale, tensor conversion)."
    ))

    # 6.6 Challenges
    add_heading(doc, "6.6 Production Challenges", 2)
    add_para(doc, (
        "Face detection in production encounters conditions rarely seen in research datasets:"
    ))
    add_bullet(doc, "Extreme camera angles: side profiles, overhead shots, low-angle interviews")
    add_bullet(doc, "Variable lighting: outdoor sunlight, dim indoor settings, backlighting")
    add_bullet(doc, "Partial occlusion: microphones, hands, other people crossing the frame")
    add_bullet(doc, "Multiple faces: group shots where the speaker must be identified")
    add_bullet(doc, "Camera motion: handheld footage with shake and rapid panning")
    add_bullet(doc, "Low resolution: old or heavily compressed YouTube videos")
    add_para(doc, (
        "The two-tier MediaPipe strategy with landmark interpolation handles most of these cases. "
        "For particularly difficult videos, the RetinaFace backend can be used as an alternative."
    ))

    doc.add_page_break()

def chapter_7(doc):
    add_heading(doc, "7. Reports & Burned Videos", 1)

    add_para(doc, (
        "A key differentiator of the Argos system is its comprehensive output generation. Rather than "
        "producing raw text files, the pipeline generates multi-format quality reports with semantic metrics "
        "and burned videos with color-coded subtitle overlays. These outputs enable non-technical users to "
        "assess transcription quality at a glance and identify which segments need manual correction."
    ))

    # 7.1 make_report.py
    add_heading(doc, "7.1 Report Generation (make_report.py, 809 Lines)", 2)
    add_para(doc, (
        "The make_report.py script is the largest custom Python script in the project. It generates "
        "comprehensive quality reports in five output formats from a single decode output JSON file:"
    ))
    add_styled_table(doc,
        ["Format", "File", "Audience", "Features"],
        [
            ["CSV", "report.csv", "Data analysis", "Per-segment metrics, machine-readable"],
            ["HTML", "report.html", "Visual review", "Color-coded table, alignment visualization"],
            ["ANSI", "report.ansi.txt", "Terminal users", "Colored terminal output for quick checks"],
            ["JSON", "metrics.json", "Automation", "Aggregate metrics, pipeline integration"],
            ["TXT", "report.txt", "Documentation", "Plain text summary"],
        ],
        col_widths=[0.7, 1.3, 1.2, 3.3]
    )

    # 7.2 Metrics
    add_heading(doc, "7.2 Metrics: WER, WWER, and NEA", 2)
    add_para(doc, (
        "The report system calculates three complementary metrics, each capturing a different aspect "
        "of transcription quality:"
    ))

    add_heading(doc, "Standard WER (Word Error Rate)", 3)
    add_para(doc, (
        "The standard metric from ASR evaluation. Calculated as the edit distance between the reference "
        "transcription and the hypothesis, divided by the number of reference words. A WER of 0% means "
        "perfect transcription; WER can exceed 100% when the hypothesis contains many inserted words."
    ))

    add_heading(doc, "WWER (Weighted Word Error Rate)", 3)
    add_para(doc, (
        "Standard WER treats all words equally, but in practice, content words (names, numbers, nouns) "
        "carry far more information than function words (the, a, is). WWER assigns weights based on "
        "part-of-speech and entity type:"
    ))
    add_styled_table(doc,
        ["Category", "Weight", "POS Tags / Entity Types", "Examples"],
        [
            ["HIGH VALUE", "2.0", "PROPN, NUM, Named Entities (PERSON, ORG, GPE, LOC, MONEY, DATE)", "Biden, 2024, Washington, $500"],
            ["MEDIUM VALUE", "1.0", "NOUN, VERB, ADJ, ADV", "speech, said, important, quickly"],
            ["LOW VALUE", "0.5", "DET, AUX, PRON, ADP, CONJ, PUNCT, stopwords", "the, is, he, of, and"],
        ],
        col_widths=[1.3, 0.7, 2.5, 2.0]
    )
    add_para(doc, (
        "With these weights, missing a person's name or a number counts twice as heavily as missing "
        "a determiner. This better reflects the real-world impact of transcription errors."
    ))

    add_heading(doc, "NEA (Named Entity Accuracy)", 3)
    add_para(doc, (
        "NEA measures how well the transcription captures semantically important tokens, using spaCy's "
        "NER (Named Entity Recognition) and POS (Part-of-Speech) tagging. It reports recall, precision, "
        "and F1 score for high-value tokens:"
    ))
    add_bullet_bold_value(doc, "NEA Recall: ", "What fraction of important reference tokens appear in the hypothesis?")
    add_bullet_bold_value(doc, "NEA Precision: ", "What fraction of important hypothesis tokens appear in the reference?")
    add_bullet_bold_value(doc, "NEA F1: ", "Harmonic mean of recall and precision")
    add_para(doc, (
        "When spaCy is unavailable (e.g., on HORIZON without internet), the report falls back to "
        "basic WER metrics. The spaCy installation in lib/outputs.sh tries local wheels first, "
        "then online pip, then gracefully degrades."
    ))

    # 7.3 Alignment
    add_heading(doc, "7.3 Color-Coded Alignment", 2)
    add_para(doc, (
        "Each word in the hypothesis is classified using an alignment algorithm that compares it "
        "against the reference transcription:"
    ))
    add_styled_table(doc,
        ["Tag", "Color", "Meaning", "Example"],
        [
            ["ok", "Green", "Word matches reference at correct position", 'ref: "the president spoke"  hyp: "the" → GREEN'],
            ["rep", "Yellow/Amber", "Word appears in reference but at wrong position", 'ref: "the president spoke"  hyp: "spoke the" → YELLOW'],
            ["ins", "Red", "Word does not appear in reference (hallucination)", 'ref: "the president spoke"  hyp: "the great president" → RED for "great"'],
        ],
        col_widths=[0.6, 1.0, 2.5, 2.4]
    )
    add_para(doc, (
        "In the HTML report, each word is rendered with its background color, creating an instant "
        "visual summary of transcription quality. A segment that is mostly green is well-transcribed; "
        "one that is mostly red contains significant hallucination."
    ))

    # 7.4 Part Naming
    add_heading(doc, "7.4 Segment Part Naming", 2)
    add_para(doc, (
        "When a source video is split into multiple segments, the reports label each segment with "
        'a human-readable part name: "Part 1", "Part 2", etc. This is derived from the segment '
        "naming convention — the segment index field determines the part number. Videos that were "
        "not segmented (shorter than the window size) are labeled without a part suffix."
    ))

    # 7.5 Decode Params
    add_heading(doc, "7.5 Decode Parameters in Reports", 2)
    add_para(doc, (
        "Each decode run saves its effective parameters (beam width, length penalty, repetition penalty, "
        "temperature, etc.) as a JSON file alongside the hypothesis output. The report generator reads "
        "this file and includes the parameters in the report header, ensuring full reproducibility. "
        "This was essential for the tuning experiments (Chapter 12), where seven different parameter "
        "configurations were compared."
    ))

    # 7.6 make_burn.py
    add_heading(doc, "7.6 Video Burning (make_burn.py, 510 Lines)", 2)
    add_para(doc, (
        "The make_burn.py script creates demonstration videos by overlaying transcription text as subtitles "
        "directly onto the original video. This produces a visual artifact where viewers can watch the "
        "speaker's lips while reading both the reference (ground truth) and hypothesis (model output) text."
    ))

    add_heading(doc, "Three-Strategy Video Source Lookup", 3)
    add_para(doc, (
        "Finding the correct source video for each segment requires a three-tier lookup strategy, "
        "because the video may be available in different forms:"
    ))
    add_bullet(doc, "1. Original video file (before segmentation) — best quality, full context")
    add_bullet(doc, "2. Normalized segment — post-normalization, correct format")
    add_bullet(doc, "3. Lip crop video — mouth region only, useful as fallback")

    add_heading(doc, "Subtitle Overlay", 3)
    add_para(doc, (
        "The script uses ffmpeg's drawtext filter to overlay two lines of text at the bottom of each video:"
    ))
    add_code_block(doc, (
        "# Reference text (white on semi-transparent black)\n"
        'ffmpeg -i input.mp4 \\\n'
        '  -vf "drawtext=text=\'REF: the president spoke today\':fontsize=16:\n'
        '       fontcolor=white:box=1:boxcolor=black@0.7:x=10:y=h-60,\n'
        '       drawtext=text=\'HYP: the great president spoke\':fontsize=16:\n'
        '       fontcolor=green:box=1:boxcolor=black@0.7:x=10:y=h-30" \\\n'
        '  -c:a copy output_burned.mp4'
    ))

    # 7.7 Per-video WER
    add_heading(doc, "7.7 Per-Video WER Calculation (300 Lines)", 2)
    add_para(doc, (
        "The calculate_per_video_wer.py script aggregates segment-level metrics into per-video scores. "
        "When overlapping segmentation is used, the same content appears in multiple segments. The script "
        "implements a deduplication algorithm that identifies overlapping regions and uses the best "
        "available transcription for each time window, avoiding double-counting words."
    ))

    # 7.8 Conflict Resolution
    add_heading(doc, "7.8 Conflict Resolution (289 + 149 Lines)", 2)
    add_para(doc, (
        "The merge_overlapping_predictions.py (289 lines) and generate_conflict_report.py (149 lines) "
        "scripts handle cases where overlapping segments produce different transcriptions for the same "
        "time window. The merge script uses a simple heuristic: prefer the transcription with lower WER "
        "against the reference. The conflict report documents all disagreements for manual review."
    ))

    doc.add_page_break()


def chapter_8(doc):
    add_heading(doc, "8. Modular Code Refactoring", 1)

    add_para(doc, (
        "By version 6, the pipeline script had grown to 823 lines of monolithic bash code. All stages, "
        "configuration, logging, error handling, and environment detection were interleaved in a single file. "
        "This made the code difficult to test, debug, and maintain — and critically, it prevented reuse "
        "between the EC2 development environment and the Linux container deployment. The modular refactoring "
        "(January 2026) was the most significant architectural improvement in the project."
    ))

    # 8.1 Problem
    add_heading(doc, "8.1 The Problem with Monolithic Code", 2)
    add_para(doc, "The 823-line monolithic script suffered from several concrete problems:")
    add_bullet(doc, "No unit testing: individual stages could not be tested in isolation")
    add_bullet(doc, "Path duplication: EC2 and container paths hardcoded throughout, requiring manual search-and-replace")
    add_bullet(doc, "Error cascading: a failure in one stage could leave the environment in an inconsistent state")
    add_bullet(doc, "No reuse: the fine-tuning pipeline needed the same preprocessing stages but couldn't import them")
    add_bullet(doc, "Debugging difficulty: with 823 lines of sequential logic, finding the source of a bug required reading the entire file")

    # 8.2 All Modules
    add_heading(doc, "8.2 The 12 Library Modules (1,562 Lines Total)", 2)
    add_para(doc, (
        "The refactoring extracted all reusable functionality into a lib/ directory with 12 modules. "
        "Each module is a self-contained bash script that exports specific functions and can be sourced "
        "independently:"
    ))
    add_styled_table(doc,
        ["Module", "Lines", "Key Functions", "Phase"],
        [
            ["common.sh", "81", "log_info, log_error, log_warn, log_stage, validate_directory", "Infrastructure"],
            ["config.sh", "87", "detect_environment(), get_base_path(), derived path functions", "Infrastructure"],
            ["venv_utils.sh", "36", "activate_venv(), deactivate_venv()", "Infrastructure"],
            ["archive.sh", "112", "archive_previous_run(), transcription preservation", "Preparation"],
            ["normalization.sh", "235", "run_normalization(), needs_tonemap(), GPU/CPU fallback", "Processing"],
            ["asr.sh", "229", "run_asr_transcription(), Steps 0.6/3/1.5", "Processing"],
            ["lrs3_prep.sh", "39", "run_lrs3_preparation()", "Processing"],
            ["manifests.sh", "80", "run_manifest_generation()", "Processing"],
            ["clustering.sh", "67", "run_clustering(), TRAIN_KMEANS toggle", "Processing"],
            ["decode.sh", "100", "run_vsp_decode(), Cython auto-build", "Decode"],
            ["outputs.sh", "165", "run_client_outputs(), spaCy install, reports + burns + lip crops", "Output"],
            ["test_all_modules.sh", "331", "37 automated tests across all modules", "Testing"],
        ],
        col_widths=[1.3, 0.5, 2.5, 0.8]
    )
    add_para(doc, (
        "The modules follow a consistent pattern: each file begins with a guard check to ensure "
        "it is being sourced (not executed directly), sources its dependencies (e.g., common.sh), "
        "and exports its public functions. This allows the main pipeline script to source only "
        "the modules it needs."
    ))

    # 8.3 Design Decisions
    add_heading(doc, "8.3 Key Design Decisions", 2)

    add_heading(doc, "Environment-Aware Configuration", 3)
    add_para(doc, (
        "The config.sh module automatically detects whether the pipeline is running on EC2 (/home/ubuntu/) "
        "or in a container (/workspace/) and configures all paths accordingly. This eliminates the need "
        "for manual path editing when moving between environments:"
    ))
    add_code_block(doc, (
        "detect_environment() {\n"
        '  if [[ -d "/workspace" ]]; then\n'
        '    ENV_TYPE="container"\n'
        '    BASE_PATH="/workspace"\n'
        "  else\n"
        '    ENV_TYPE="ec2"\n'
        '    BASE_PATH="/home/ubuntu"\n'
        "  fi\n"
        "  # All derived paths use BASE_PATH\n"
        '  VSP_DIR="${BASE_PATH}/VSP-LLM"\n'
        '  AUTO_AVSR_DIR="${BASE_PATH}/auto_avsr"\n'
        '  # ... etc\n'
        "}"
    ))

    add_heading(doc, "Virtual Environment Strategy", 3)
    add_para(doc, (
        "The ASR module (asr.sh) is self-contained: it activates the pre-process-venv internally "
        "and deactivates it when done. All other modules (stages 5-8) expect the caller to have "
        "already activated vsp-llm-yoad-venv. This minimizes venv activation overhead — the main "
        "pipeline activates the VSP venv once and uses it for multiple stages."
    ))

    add_heading(doc, "The stdout Contamination Bug", 3)
    add_para(doc, (
        "One of the most subtle bugs discovered during refactoring was that log_info() wrote to stdout "
        "instead of stderr. This caused logging output to be captured by command substitution "
        "operations, contaminating variables with log messages. For example:"
    ))
    add_code_block(doc, (
        "# BROKEN: log output mixed into the variable\n"
        'result=$(some_function)  # result = "[INFO] Processing...\\n/actual/path"\n'
        "\n"
        "# FIXED: log_info writes to stderr (fd 2)\n"
        'log_info() { echo "[$(date +%H:%M:%S)] INFO: $*" >&2; }\n'
        'result=$(some_function)  # result = "/actual/path"'
    ))

    # 8.4 Test Suite
    add_heading(doc, "8.4 Test Suite: 37 Automated Tests", 2)
    add_para(doc, (
        "The test suite (lib/test_all_modules.sh, 331 lines) validates every module's exports, "
        "functionality, and integration points. Tests use colored output for quick visual assessment:"
    ))
    add_styled_table(doc,
        ["Module", "Tests", "What Is Verified"],
        [
            ["common.sh", "5", "log_info/log_error formatting, validate_directory on existing and non-existent paths"],
            ["config.sh", "5", "ENV_TYPE detection, BASE_PATH setting, path function exports"],
            ["venv_utils.sh", "3", "activate_venv/deactivate_venv exports, real venv activation if available"],
            ["normalization.sh", "4", "Normalization function export, HDR detection, GPU/CPU fallback"],
            ["asr.sh", "4", "Whisper transcription, offline model loading, transcription reuse logic"],
            ["archive.sh", "3", "Previous run archiving, transcription preservation"],
            ["lrs3_prep.sh", "2", "LRS3 preparation function export and basic operation"],
            ["manifests.sh", "2", "Manifest generation function export"],
            ["clustering.sh", "2", "Clustering function export, TRAIN_KMEANS toggle"],
            ["decode.sh", "3", "Decode function export, Cython build check"],
            ["outputs.sh", "2", "Output function export, spaCy install detection"],
            ["Integration", "2", "Cross-module function calls, path consistency"],
        ],
        col_widths=[1.3, 0.5, 4.7]
    )

    # 8.5 Before/After
    add_heading(doc, "8.5 Before and After Metrics", 2)
    add_styled_table(doc,
        ["Metric", "Before (v6)", "After (v7)", "Improvement"],
        [
            ["Main script", "823 lines, 1 file", "501 lines, 1 file", "52% reduction"],
            ["Library code", "0 lines, 0 files", "1,562 lines, 12 files", "New reusable modules"],
            ["Total code", "823 lines", "2,063 lines", "More code but better organized"],
            ["Testability", "None", "37 automated tests", "Full coverage"],
            ["Environment support", "Manual path editing", "Auto-detection", "Zero configuration"],
            ["Fine-tuning reuse", "Not possible", "Shares all lib/ modules", "579-line script reuses entire lib/"],
        ],
        col_widths=[1.5, 1.5, 1.8, 1.7]
    )

    doc.add_page_break()

def chapter_9(doc):
    add_heading(doc, "9. Web UI & Server", 1)

    add_para(doc, (
        "The Argos web interface was built to make the pipeline accessible to non-technical users who need "
        "to process videos, review transcriptions, and download results without using the command line. "
        "The UI consists of a Python HTTP server backend and a vanilla JavaScript frontend — no external "
        "frameworks (no React, no Vue, no Flask) — to minimize dependencies and maximize portability. "
        "The complete UI spans 5,780 lines across 11 files."
    ))

    # 9.1 Architecture
    add_heading(doc, "9.1 Architecture: Zero-Dependency Web Stack", 2)
    add_para(doc, (
        "A deliberate architectural decision was made to avoid external web frameworks. The server uses "
        "Python's built-in http.server module (ThreadedHTTPServer), and the frontend uses vanilla JavaScript "
        "with no build tools, bundlers, or transpilers. This means the UI works on any machine with Python 3 "
        "installed — critical for the HORIZON server and standalone machines where installing npm, Node.js, "
        "or Python web frameworks is not possible."
    ))
    add_styled_table(doc,
        ["Component", "File", "Lines", "Technology"],
        [
            ["HTTP Server", "server.py", "1,124", "Python http.server (ThreadedHTTPServer)"],
            ["Frontend Logic", "app.js", "1,921", "Vanilla JavaScript (ES6+)"],
            ["HTML Structure", "index.html", "328", "Semantic HTML5"],
            ["Styling", "style.css", "1,191", "Mobile-first CSS3"],
            ["Pipeline Service", "pipeline_runner.py", "349", "subprocess + threading"],
            ["Progress Service", "progress_tracker.py", "206", "Regex-based log parsing"],
            ["Transcription Service", "transcription_manager.py", "280", "JSON + .wrd file management"],
            ["Validation Service", "validator.py", "374", "ffprobe-based video analysis"],
            ["Entry Point", "__main__.py", "5", "Python package entry"],
            ["Package Init", "__init__.py", "1", "Package marker"],
            ["Services Init", "services/__init__.py", "1", "Package marker"],
        ],
        col_widths=[1.5, 1.8, 0.5, 2.7]
    )

    # 9.2 State Machine
    add_heading(doc, "9.2 Six-Screen State Machine", 2)
    add_para(doc, (
        "The UI operates as a state machine with six mutually exclusive screens. Only one screen is "
        "visible at a time, and transitions occur based on user actions and pipeline events:"
    ))
    add_styled_table(doc,
        ["Screen", "Purpose", "Key Actions", "Next Screen"],
        [
            ["Welcome", "Upload videos, inspect input folder", "Drag-drop upload, remove videos, configure options", "Segment Review"],
            ["Segment Review", "Review segments after fast segmentation", "Transcribe segments, delete orphans, choose k-means mode", "Processing"],
            ["Processing", "Monitor pipeline execution", "View progress bar, read logs, cancel pipeline", "Complete or Error"],
            ["Transcription", "Per-segment review (legacy screen)", "Edit transcriptions, view segments", "Complete"],
            ["Complete", "Download results and review", "Download ZIP, open folder, start new batch", "Welcome"],
            ["Error", "Display error details", "View full logs, retry or reset", "Welcome"],
        ],
        col_widths=[1.0, 1.5, 2.3, 1.0]
    )

    # 9.3 server.py
    add_heading(doc, "9.3 Server (server.py, 1,124 Lines)", 2)
    add_para(doc, (
        "The server handles all HTTP requests through a single VSPRequestHandler class. It implements "
        "12 REST API endpoints (9 GET, 9 POST — some paths support both methods):"
    ))

    add_heading(doc, "GET Endpoints", 3)
    add_styled_table(doc,
        ["Endpoint", "Returns", "Purpose"],
        [
            ["/api/status", "JSON: video count, pipeline state, is_running", "Dashboard status polling"],
            ["/api/progress", "JSON: stage, percent, logs, errors", "Progress bar updates (500ms poll)"],
            ["/api/logs", "JSON: recent log lines", "Collapsible log viewer"],
            ["/api/golden-models", "JSON: available k-means models with metadata", "Golden model selector dropdown"],
            ["/api/download-output", "ZIP file stream", "Download all outputs as archive"],
            ["/api/transcription", "JSON: text, type, metadata", "Load transcription for editing"],
            ["/api/orphaned-transcriptions", "JSON: transcriptions without videos", "Orphan cleanup workflow"],
            ["/api/segments", "JSON: segment list with transcription status", "Segment review screen"],
            ["/api/segment-video/<id>", "Video file stream", "Lazy-loaded video preview in modal"],
        ],
        col_widths=[2.0, 2.0, 2.5]
    )

    add_heading(doc, "POST Endpoints", 3)
    add_styled_table(doc,
        ["Endpoint", "Input", "Action"],
        [
            ["/api/validate", "—", "Run ffprobe on all input videos, estimate segments"],
            ["/api/start", "JSON: train_kmeans, golden_model, segmentation, overlap, segment_only", "Launch pipeline subprocess"],
            ["/api/cancel", "—", "Send SIGTERM to pipeline process"],
            ["/api/reset", "—", "Clear pipeline state for fresh start"],
            ["/api/remove-video", "JSON: filename", "Move video to .excluded/ folder"],
            ["/api/transcription", "JSON: filename, text", "Save or delete transcription"],
            ["/api/save-segment-transcription", "JSON: filename, text", "Save with normalization"],
            ["/api/open-folder", "JSON: path", "Open in system file manager"],
            ["/api/upload", "Multipart form data", "Chunked file upload (1MB chunks)"],
        ],
        col_widths=[2.2, 2.5, 1.8]
    )

    add_heading(doc, "Manual Multipart Parser", 3)
    add_para(doc, (
        "The upload endpoint implements a custom multipart form data parser because Python's cgi.FieldStorage "
        "was deprecated in Python 3.11 and removed in Python 3.13. The custom parser reads the request body "
        "in 1MB chunks to prevent out-of-memory errors on large video files (the original implementation "
        "used .read(content_length) which loaded the entire file into RAM)."
    ))

    # 9.4 app.js
    add_heading(doc, "9.4 Frontend (app.js, 1,921 Lines)", 2)
    add_para(doc, "The frontend JavaScript implements several sophisticated UI patterns:")

    add_bullet_bold_value(doc, "Progress Polling (500ms): ",
        "The UI polls /api/progress every 500 milliseconds during pipeline execution. This interval "
        "provides responsive feedback without excessive server load. The progress bar uses smooth "
        "CSS transitions for visual continuity between updates.")
    add_bullet_bold_value(doc, "Drag-and-Drop Upload: ",
        "Files can be dragged onto the welcome screen. The upload uses XMLHttpRequest (not fetch) "
        "to support real-time progress events. For small files, simulated smooth progress fills "
        "the gap between actual server responses (8% increments for small, 2% for large files).")
    add_bullet_bold_value(doc, "Transcription Modal: ",
        "Clicking a segment opens a modal with a text area for editing transcriptions. The modal "
        "shows a live normalization preview (lowercase, alphanumeric + apostrophes only) and word count. "
        "A toggle button lazily loads the segment video for visual reference.")
    add_bullet_bold_value(doc, "Lazy Video Loading: ",
        "Segment videos in the review screen use IntersectionObserver with a 200px root margin to "
        "load only when scrolled into view. This prevents the browser from attempting to load hundreds "
        "of video files simultaneously.")
    add_bullet_bold_value(doc, "Orphan Transcription Cleanup: ",
        "When videos are removed but their transcriptions remain, the UI detects these orphans and "
        "offers the user a choice: delete the transcription or keep it for potential re-addition.")
    add_bullet_bold_value(doc, "K-Means Golden Model Selector: ",
        "A dropdown populated from /api/golden-models lets users choose a pre-trained k-means model "
        "instead of training a new one. Each model shows its dataset info and file size.")
    add_bullet_bold_value(doc, "Adaptive Upload Timeout: ",
        "Upload timeout scales with file size: 5 minutes base + 1 minute per 100MB. This prevents "
        "timeouts on large video files while keeping short timeouts for small files.")

    # 9.5 index.html
    add_heading(doc, "9.5 HTML Structure (index.html, 328 Lines)", 2)
    add_para(doc, (
        "The HTML file defines six screen sections (mutually exclusive via CSS display toggle), a drop "
        "zone overlay for drag-and-drop, and a transcription modal. The modal includes video preview, "
        "textarea, normalization preview, word count, and save/delete buttons. Status badges use CSS "
        "classes to indicate pipeline state (Ready, Processing, Complete, Error)."
    ))

    # 9.6 style.css
    add_heading(doc, "9.6 Styling (style.css, 1,191 Lines)", 2)
    add_para(doc, (
        "The stylesheet implements a complete design system with mobile-first responsive breakpoints. "
        "Key design elements include:"
    ))
    add_bullet(doc, "Color scheme: #2563eb (primary blue), #16a34a (success green), #dc2626 (error red)")
    add_bullet(doc, "Transcription badges: orange for auto-generated, green for manually edited")
    add_bullet(doc, "Progress bar: gradient fill with smooth transitions")
    add_bullet(doc, "Modal: fixed overlay with 70% black background, slide-in animation, max 650px width")
    add_bullet(doc, "Responsive: column layout below 600px, 95% modal width below 768px")
    add_bullet(doc, "Upload progress: per-file progress bars with checkmark or X status indicators")

    # 9.7 Upload Bug
    add_heading(doc, "9.7 Large File Upload Bug", 2)
    add_para(doc, (
        "One of the more impactful production bugs was the upload OOM (out of memory) crash. The original "
        "upload handler called request.rfile.read(content_length), which loaded the entire uploaded file "
        "into memory before writing it to disk. For a 2GB video file, this consumed 2GB of RAM instantly, "
        "crashing the server process."
    ))
    add_para(doc, "The fix involved two changes:")
    add_bullet(doc, "Chunked reading: read the request body in 1MB chunks, writing each chunk to disk immediately")
    add_bullet(doc, "Adaptive timeout: increase the socket timeout proportionally to file size (base 5 min + 1 min per 100MB)")

    doc.add_page_break()


def chapter_10(doc):
    add_heading(doc, "10. Backend Services", 1)

    add_para(doc, (
        "The web UI's backend is organized into four service modules, each handling a distinct aspect "
        "of pipeline management. These services are stateful Python classes instantiated once when the "
        "server starts and shared across all request handlers."
    ))

    # 10.1 PipelineRunner
    add_heading(doc, "10.1 PipelineRunner (349 Lines)", 2)
    add_para(doc, (
        "The PipelineRunner manages the lifecycle of the pipeline subprocess — starting, monitoring, "
        "and stopping the bash script execution. It operates as a singleton: only one pipeline run "
        "can be active at a time."
    ))

    add_heading(doc, "Subprocess Management", 3)
    add_para(doc, "The runner launches the pipeline with carefully configured environment variables:")
    add_code_block(doc, (
        "# Pipeline launch command\n"
        "bash /home/ubuntu/run_flat_english_pipeline.sh <INPUT_DIR>\n"
        "\n"
        "# Environment variables passed to subprocess:\n"
        "TRAIN_KMEANS=0|1           # Train new k-means or use existing\n"
        "GOLDEN_KMEANS=<path>       # Path to golden k-means model (optional)\n"
        "SEGMENTATION_ENABLED=0|1   # Enable video segmentation\n"
        "OVERLAP_ENABLED=0|1        # Enable overlap between segments\n"
        "SEGMENT_ONLY=0|1           # Stop after segmentation (for review)\n"
        "PYTHONUNBUFFERED=1         # Force unbuffered output for real-time logs"
    ))

    add_heading(doc, "Graceful Shutdown", 3)
    add_para(doc, (
        "When the user clicks Cancel, the runner follows a graceful shutdown sequence: send SIGTERM "
        "to the pipeline process, wait up to 10 seconds for it to finish the current stage, then "
        "send SIGKILL if it's still running. This prevents orphaned ffmpeg or Python processes."
    ))

    add_heading(doc, "Monitoring", 3)
    add_para(doc, (
        "The runner reads the pipeline's stdout line by line (line-buffered) in a background thread. "
        "Each line is passed to the ProgressTracker for stage detection. An inactivity warning is "
        "triggered if 10 minutes pass without any output from the pipeline."
    ))

    # 10.2 ProgressTracker
    add_heading(doc, "10.2 ProgressTracker (206 Lines)", 2)
    add_para(doc, (
        "The ProgressTracker parses pipeline log output in real time to determine which stage is currently "
        "executing and calculate an overall progress percentage."
    ))

    add_heading(doc, "ProgressState Dataclass", 3)
    add_code_block(doc, (
        "class ProgressState:\n"
        "    state: str          # IDLE, RUNNING, COMPLETED, FAILED, CANCELLED\n"
        "    current_stage_index: int\n"
        "    current_stage_id: str\n"
        "    current_stage_name: str\n"
        "    percent_complete: float   # 0-100\n"
        "    errors: list              # Last 10 errors\n"
        "    warnings: list            # Last 20 warnings\n"
        "    logs: list                # Last 1000 lines\n"
        "    segment_only: bool"
    ))

    add_heading(doc, "Stage Detection", 3)
    add_para(doc, (
        "Each pipeline stage prints a marker line (e.g., '[STAGE 3] Running Whisper ASR'). The tracker "
        "uses regex patterns (STAGE_MARKERS) to detect these markers and update the current stage. "
        "Progress is calculated by summing the weights of completed stages plus 50% of the current "
        "stage's weight, capped at 99% until the completion marker is detected."
    ))

    add_heading(doc, "Error Detection", 3)
    add_para(doc, (
        "The tracker watches for error patterns in the log output: ValueError, RuntimeError, CUDA errors, "
        "OOM messages, and other common failure modes. Detected errors are stored in the errors list and "
        "surfaced through the /api/progress endpoint."
    ))

    # 10.3 TranscriptionManager
    add_heading(doc, "10.3 TranscriptionManager (280 Lines)", 2)
    add_para(doc, (
        "The TranscriptionManager provides unified transcription storage with persistent metadata. "
        "All transcriptions (manual and auto-generated) are stored in a central .transcriptions/ "
        "directory within the input folder."
    ))

    add_heading(doc, "Storage Format", 3)
    add_code_block(doc, (
        "vsp_input/.transcriptions/\n"
        "├── video_name_00_000000_000300.wrd   # One word per line\n"
        "├── video_name_01_000300_000600.wrd\n"
        "└── metadata.json                      # Type, timestamps, word counts\n"
        "\n"
        "metadata.json structure:\n"
        "{\n"
        '  "transcriptions": {\n'
        '    "video_name_00_000000_000300.mp4": {\n'
        '      "type": "manual",          // or "auto"\n'
        '      "created_at": "2026-02-15T10:30:00Z",\n'
        '      "edited_at": "2026-02-16T14:20:00Z",\n'
        '      "word_count": 42\n'
        "    }\n"
        "  }\n"
        "}"
    ))

    add_heading(doc, "Text Normalization", 3)
    add_para(doc, (
        "When saving transcriptions, the manager normalizes text to match the model's expected format: "
        "lowercase, alphanumeric characters and apostrophes only. This ensures consistency between "
        "manual transcriptions and Whisper auto-transcriptions."
    ))

    add_heading(doc, "Orphan Detection", 3)
    add_para(doc, (
        "The get_orphaned_transcriptions() method identifies transcriptions that no longer have "
        "corresponding video files (e.g., because the video was excluded). The UI presents these "
        "orphans to the user for cleanup."
    ))

    # 10.4 Validator
    add_heading(doc, "10.4 Validator (374 Lines)", 2)
    add_para(doc, (
        "The Validator uses ffprobe to analyze input video files and determine their suitability "
        "for processing. It produces a detailed ValidationResult with per-video analysis."
    ))

    add_heading(doc, "Validation Checks", 3)
    add_bullet(doc, "Supported file extension (.mp4, .mkv, .webm, .mov, .m4v, .avi)")
    add_bullet(doc, "ffprobe can successfully read the file (not corrupt)")
    add_bullet(doc, "File contains at least one video stream")
    add_bullet(doc, "Duration is greater than 0 seconds")
    add_bullet(doc, "Resolution is at least 64x64 pixels (minimum for face detection)")

    add_heading(doc, "Segment Estimation", 3)
    add_para(doc, (
        "For each valid video, the validator estimates the number of segments that will be produced: "
        "ceil(duration / segment_duration). This is used by the UI to show the total expected segment "
        "count and to assess k-means viability (a warning is shown if fewer than 200 segments are expected)."
    ))

    add_heading(doc, "VideoInfo Dataclass", 3)
    add_para(doc, (
        "Each validated video is represented as a VideoInfo object containing: filename, path, "
        "duration_seconds, width, height, fps, has_audio, estimated_segments, valid (bool), error "
        "message, and transcription metadata (has_transcription, type, created_at, edited_at, word_count)."
    ))

    doc.add_page_break()


def chapter_11(doc):
    add_heading(doc, "11. Docker Container & Deployment", 1)

    add_para(doc, (
        "Deploying Argos to production machines required packaging the entire system — three Python repos, "
        "two virtual environments, model weights, system libraries, and the web UI — into a portable "
        "container that could run on any machine with an NVIDIA GPU. This chapter documents the Docker "
        "architecture, five container releases, 37 bugs fixed during deployment, and the testing "
        "infrastructure that ensures reliability."
    ))

    # 11.1 Architecture
    add_heading(doc, "11.1 Container Architecture", 2)
    add_para(doc, (
        "The Docker container is built on nvidia/cuda:12.8.0-base-ubuntu22.04, providing CUDA runtime "
        "libraries and GPU access through NVIDIA Container Toolkit."
    ))
    add_code_block(doc, (
        "# Dockerfile structure (262 lines)\n"
        "FROM nvidia/cuda:12.8.0-base-ubuntu22.04\n"
        "\n"
        "# System dependencies\n"
        "RUN apt-get update && apt-get install -y \\\n"
        "    python3 python3-pip python3-venv gcc make \\\n"
        "    ffmpeg git libglib2.0-0 libsm6 libxext6 \\\n"
        "    libxrender1 libsndfile1 libportaudio2\n"
        "\n"
        "# Virtual Environment 1: Preprocessing\n"
        "# PyTorch 2.8.0+cu128, triton 3.4.0, Whisper, MediaPipe\n"
        "# ~135 packages\n"
        "\n"
        "# Virtual Environment 2: Decoding\n"
        "# PyTorch 2.5.1+cu124, triton 3.1.0, fairseq, transformers\n"
        "# ~256 packages\n"
        "\n"
        "# Model weights, spaCy wheels, pipeline scripts, web UI\n"
        "COPY galaxy_export/ /workspace/"
    ))

    # 11.2 Galaxy Export
    add_heading(doc, "11.2 Galaxy Export Bundle (~65GB)", 2)
    add_para(doc, (
        "The galaxy export is the complete offline deployment package. At approximately 65GB, it includes "
        "everything needed to run the pipeline without internet access: both virtual environments with all "
        "packages, model weights (LLaMA-2-7B ~13GB, AV-HuBERT checkpoint ~1.2GB), the complete pipeline "
        "codebase, web UI, spaCy wheels for offline installation, and Whisper model cache."
    ))

    # 11.3 Container Versions
    add_heading(doc, "11.3 Five Container Releases", 2)
    add_styled_table(doc,
        ["Date", "Version Range", "Key Changes"],
        [
            ["Feb 3, 2026", "v1.0.0 — v1.0.10", "Initial release: basic pipeline, web UI, first deployment"],
            ["Feb 8, 2026", "v1.0.11 — v1.0.13", "Installation bug fixes: hardcoded paths, python3 compatibility, fairseq max_len, UI networking"],
            ["Feb 12, 2026", "v1.0.14 — v1.0.25", "Deployment fixes: CUDA OOM on 12GB GPU, desktop launcher, docker.conf, upload reliability"],
            ["Feb 15, 2026", "v1.0.26 — v1.0.35", "Final fixes: beam search OOM, inference tuning, report metrics, drag-drop upload, NVENC corruption"],
            ["Feb 17, 2026", "v1.0.36 — v1.0.37", "NVENC silent corruption fix, bash fd interference fix, post-encode validation"],
        ],
        col_widths=[1.2, 1.5, 3.8]
    )

    # 11.4 First Delivery
    add_heading(doc, "11.4 First Client Delivery", 2)
    add_para(doc, (
        "The first delivery package (FOR_CLIENT_OLD) included the OCI container image, a desktop "
        "launcher (.desktop file) for easy startup, and NVIDIA driver installers for Arch Linux "
        "compatibility. The container was loaded using docker load and launched with nvidia-docker "
        "or the NVIDIA Container Runtime."
    ))

    # 11.5 Bugs Table
    add_heading(doc, "11.5 Thirty-Seven Bugs Fixed", 2)
    add_para(doc, (
        "Over the course of five container releases, 37 bugs were identified and fixed. These bugs "
        "span installation, deployment, GPU compatibility, and user interface issues. The following "
        "table lists the most impactful bugs by category:"
    ))

    add_heading(doc, "Installation Bugs (1-13)", 3)
    add_styled_table(doc,
        ["Bug#", "Summary", "Impact"],
        [
            ["1-3", "Missing files during INSTALL.sh, hardcoded EC2 paths", "Container fails to start"],
            ["4-6", "python vs python3 incompatibility across 11 files", "Scripts crash on Ubuntu"],
            ["7-8", "Whisper offline model loading, spaCy offline install", "Pipeline fails on HORIZON"],
            ["9-10", "fairseq GenerationConfig schema (max_len field missing)", "Decode crashes"],
            ["11-13", "UI server binding 127.0.0.1 vs 0.0.0.0, host launcher", "UI inaccessible from host"],
        ],
        col_widths=[0.7, 3.5, 2.3]
    )

    add_heading(doc, "Deployment & GPU Bugs (14-25)", 3)
    add_styled_table(doc,
        ["Bug#", "Summary", "Impact"],
        [
            ["14", "CUDA OOM on 12GB GPU — memory accumulation across samples", "Decode crashes after ~50 segments"],
            ["15-17", "Desktop launcher: Terminal=true, execute permission, untrusted .desktop", "Users can't launch UI"],
            ["18-20", "Docker port conflicts, ENTRYPOINT double-bash, docker.conf parsing", "Container startup fails"],
            ["21-22", "fairseq PYTHONPATH override, wrong installation targeted", "Decode uses wrong fairseq"],
            ["23-25", "Open Results Folder, file upload reliability, cgi deprecation", "UI features broken"],
        ],
        col_widths=[0.7, 3.5, 2.3]
    )

    add_heading(doc, "Final Fixes (26-37)", 3)
    add_styled_table(doc,
        ["Bug#", "Summary", "Impact"],
        [
            ["26", "Beam search OOM on 12GB GPUs (no_repeat_ngram_size + dynamic max_len cap)", "Decode crashes on long segments"],
            ["27-28", "Inference tuning: max_len_a 3.0→2.0, max_len_b 300→200, repetition_penalty=1.2", "Reduced hallucination"],
            ["29-30", "Report part naming for segments, lip crop copy logic", "Reports show wrong segment labels"],
            ["31", "Log output stdout contamination (log_info to stderr)", "Pipeline paths corrupted"],
            ["32-33", "Drag-drop upload: manual multipart parser (Python 3.13 compatible)", "Upload crashes"],
            ["34-35", "Segment transcription persistence (Steps 0.6 and 1.5)", "Transcriptions lost between runs"],
            ["36-37", "NVENC silent corruption, bash fd interference, post-encode validation", "43% of segments corrupt"],
        ],
        col_widths=[0.7, 3.5, 2.3]
    )

    add_para(doc, (
        "The most impactful bugs were: CUDA OOM (#14/#26) which made decode impossible on 12GB GPUs, "
        "NVENC silent corruption (#36) which destroyed 43% of processed segments, log contamination (#31) "
        "which caused random pipeline failures, and the upload OOM (#25/#32) which crashed the server "
        "on large file uploads."
    ), bold=True)

    # 11.6 Build Automation
    add_heading(doc, "11.6 Build & Verification Automation", 2)
    add_styled_table(doc,
        ["Script", "Lines", "Purpose"],
        [
            ["build_container.sh", "92", "Docker image build with layer caching"],
            ["check_container_inventory.sh", "135", "Verify all required files present in container"],
            ["verify_container_sync.sh", "353", "Check EC2 changes are replicated to container"],
        ],
        col_widths=[2.2, 0.6, 3.7]
    )

    # 11.7 Installation Flow
    add_heading(doc, "11.7 Installation & Startup Flow", 2)
    add_code_block(doc, (
        "1. Load container image:\n"
        "   docker load < vsp_linux_container_FINAL_20260217.tar\n"
        "\n"
        "2. Run INSTALL.sh:\n"
        "   - Creates required directories\n"
        "   - Sets permissions\n"
        "   - Installs desktop launcher (optional)\n"
        "\n"
        "3. Start with vsp-start.sh:\n"
        "   - Launches Docker container with GPU access\n"
        "   - Mounts input/output directories\n"
        "   - Starts web UI server on port 8080\n"
        "\n"
        "4. Access UI:\n"
        "   - Open browser to http://localhost:8080\n"
        "   - Or click desktop icon (if installed)"
    ))

    # 11.8 Sync Items
    add_heading(doc, "11.8 Container Synchronization (28 Pending Items)", 2)
    add_para(doc, (
        "Every change made to the EC2 development environment must be explicitly replicated to the "
        "Linux container version. This synchronization requirement is documented in detail in "
        "docs/container-sync-changelog.md (1,329 lines), which lists 28 pending sync items with "
        "exact file paths, line numbers, and code diffs for each change."
    ))

    # 11.9 Testing
    add_heading(doc, "11.9 Testing Infrastructure", 2)
    add_para(doc, (
        "The project includes a comprehensive testing suite spanning five test scripts with a total "
        "of 1,731 lines of test code:"
    ))
    add_styled_table(doc,
        ["Script", "Lines", "Tests", "What It Verifies"],
        [
            ["test_all_modules.sh", "331", "37", "All lib/ module exports, functions, and integration"],
            ["test_pipeline_modules.sh", "350", "~20", "Module-level function tests in isolation"],
            ["test_ec2_environment.sh", "371", "~15", "EC2 environment: paths, venvs, GPU, dependencies"],
            ["test_pipeline_smoke.sh", "326", "~10", "End-to-end pipeline stages with test data"],
            ["verify_container_sync.sh", "353", "~15", "EC2 changes replicated to container version"],
        ],
        col_widths=[1.8, 0.5, 0.5, 3.7]
    )

    doc.add_page_break()

def chapter_12(doc):
    add_heading(doc, "12. Performance Evaluation & Tuning", 1)

    add_para(doc, (
        "Evaluating the Argos system's performance required processing a large dataset of real-world videos, "
        "establishing metrics that go beyond simple WER, running systematic parameter tuning experiments, "
        "and producing detailed analysis reports. This chapter documents the evaluation methodology, results, "
        "seven decode tuning experiments, six analysis reports, and the key insights that emerged."
    ))

    # 12.1 Setup
    add_heading(doc, "12.1 Evaluation Setup", 2)
    add_para(doc, (
        "The primary evaluation used approximately 1,000 YouTube videos from the AVSpeech dataset, "
        "which were processed through the pipeline to produce 860 evaluated segments. Whisper ASR "
        "(medium model) generated ground truth transcriptions for comparison against VSP-LLM's "
        "visual-only predictions."
    ))
    add_styled_table(doc,
        ["Parameter", "Value"],
        [
            ["Source Dataset", "AVSpeech (YouTube videos)"],
            ["Videos Processed", "~1,000"],
            ["Segments Evaluated", "860"],
            ["Ground Truth", "Whisper ASR (medium model, English)"],
            ["Decode Config", "beam=20, lenpen=0.0, rep_pen=1.2"],
            ["Evaluation Metrics", "WER, WWER, NEA (Recall, Precision, F1)"],
        ],
        col_widths=[2.0, 4.5]
    )

    # 12.2 Results
    add_heading(doc, "12.2 Results: Argos vs. Paper", 2)
    add_para(doc, (
        "The headline result is a significant performance gap between the paper's reported WER on "
        "LRS3 and Argos's WER on AVSpeech. This gap is expected and is primarily driven by domain "
        "mismatch rather than implementation error."
    ))
    add_styled_table(doc,
        ["Metric", "Paper (LRS3)", "Argos (AVSpeech)", "Gap Factor"],
        [
            ["Mean WER", "25.4%", "67.0%", "2.7x"],
            ["Dataset Type", "TED talks", "YouTube (wild)", "—"],
            ["Conditions", "Clean audio, frontal face", "Diverse angles, noise, lighting", "—"],
            ["Segments Usable (WER ≤ 20%)", "—", "11.4%", "—"],
            ["Segments Hallucinated (WER ≥ 100%)", "—", "20.6%", "—"],
            ["NEA F1", "—", "38.8%", "—"],
        ],
        col_widths=[2.0, 1.5, 1.7, 1.3]
    )

    # 12.3 WER Distribution
    add_heading(doc, "12.3 WER Distribution", 2)
    add_para(doc, (
        "Examining the distribution of WER scores across segments reveals a bimodal pattern: "
        "some segments are transcribed well while others are nearly or completely wrong."
    ))
    add_styled_table(doc,
        ["WER Range", "Category", "% of Segments", "Interpretation"],
        [
            ["0-20%", "Excellent / Usable", "11.4%", "Can be used with minor corrections"],
            ["21-40%", "Marginal", "17.4%", "Core meaning preserved but noisy"],
            ["41-60%", "Poor", "17.8%", "Some correct words but unreliable"],
            ["61-99%", "Unusable", "32.8%", "Too many errors for practical use"],
            ["≥100%", "Hallucinated", "20.6%", "Model generates unrelated text"],
        ],
        col_widths=[1.0, 1.3, 1.0, 3.2]
    )
    add_para(doc, (
        "The 20.6% hallucination rate is particularly noteworthy. These are cases where the LLaMA-2 "
        "decoder generates fluent, grammatically correct English text that has no relationship to "
        "what the speaker is saying. This happens when the visual features extracted by AV-HuBERT "
        "are ambiguous (poor face angle, motion blur, or the speaker's mouth is occluded)."
    ))

    # 12.4 Domain Gap
    add_heading(doc, "12.4 Domain Gap Analysis", 2)
    add_para(doc, (
        "The 2.7x performance gap between LRS3 and AVSpeech is caused by fundamental differences "
        "between the training and deployment domains:"
    ))
    add_styled_table(doc,
        ["Aspect", "LRS3 (Training Domain)", "AVSpeech (Deployment Domain)"],
        [
            ["Content", "TED talks", "YouTube videos (all genres)"],
            ["Camera", "Frontal, static, professional", "Variable angles, handheld, phone cameras"],
            ["Lighting", "Stage lighting, consistent", "Outdoor, indoor, backlighting, shadows"],
            ["Audio", "Clean studio recording", "Background noise, music, crosstalk"],
            ["Speakers", "Professional speakers", "Anyone on YouTube"],
            ["Face Position", "Centered, large in frame", "Variable size, sometimes partial"],
            ["Duration", "2-10 seconds", "Variable (requires segmentation)"],
        ],
        col_widths=[1.2, 2.6, 2.7]
    )
    add_para(doc, (
        "The model was fine-tuned by the paper's authors exclusively on LRS3 data (TED talks with "
        "clean conditions). When applied to YouTube videos with diverse real-world conditions, the "
        "visual features become less informative, and the LLaMA-2 decoder fills in the gaps with "
        "plausible but incorrect text (hallucination)."
    ))

    # 12.5 Tuning Experiments
    add_heading(doc, "12.5 Seven Decode Tuning Experiments", 2)
    add_para(doc, (
        "To explore whether decode-time parameters could improve transcription quality without "
        "retraining, seven experiments were run on a 100-video / 107-segment subset. Each experiment "
        "varied one or more generation parameters while keeping the rest at baseline values."
    ))
    add_styled_table(doc,
        ["Exp", "Name", "Parameters Changed", "Hypothesis"],
        [
            ["A", "Baseline", "beam=20, lenpen=0, rep_pen=1.2", "Current production settings"],
            ["B", "No Rep Penalty", "repetition_penalty=1.0", "Removing penalty may reduce error"],
            ["C", "LenPen +1.0", "lenpen=1.0", "Favor longer outputs (reduce truncation)"],
            ["D", "LenPen -0.5", "lenpen=-0.5", "Favor shorter, more confident outputs"],
            ["E", "Sampling Low Temp", "do_sample=true, temp=0.5, top_p=0.9", "Stochastic but conservative"],
            ["F", "Sampling Original", "do_sample=true, temp=1.0, top_p=0.9", "Match english_1k-style decode"],
            ["G", "Greedy Decode", "beam=1", "Simplest decode, no beam search overhead"],
        ],
        col_widths=[0.4, 1.3, 2.5, 2.3]
    )
    add_para(doc, (
        "The experiment framework (run_all_experiments.sh, 118 lines + run_experiment.sh, 101 lines) "
        "automates the full cycle for each configuration: set overrides, run decode, generate reports, "
        "and save config snapshots. A comparison table is built at the end showing WWER, NEA-Recall, "
        "and NEA-F1 across all experiments."
    ))

    # 12.6 Analysis Reports
    add_heading(doc, "12.6 Six Analysis Reports", 2)
    add_para(doc, (
        "In addition to the tuning experiments, six detailed analysis reports were written to explore "
        "different aspects of the system's performance and potential improvements:"
    ))
    add_styled_table(doc,
        ["#", "Title", "Key Finding / Recommendation"],
        [
            ["1", "Executive Assessment", "67% WER with 11.4% usable segments; domain gap is primary bottleneck"],
            ["2", "Hyperparameter Tuning", "Length penalty (lenpen) is highest-impact parameter; recommend 0.5-1.0"],
            ["3", "Prompt Engineering & Context Injection", "Analysis of context injection strategies for improved decode"],
            ["4", "Word-Level Confidence Scoring", "Methods for estimating per-word reliability from beam scores"],
            ["5", "Beam Search Aggregation & N-Best Fusion", "Using multiple beam hypotheses to improve accuracy"],
            ["6", "Fine-Tuning Analysis", "Domain adaptation on AVSpeech is essential; expected 15-25 WER improvement"],
        ],
        col_widths=[0.3, 2.2, 4.0]
    )

    # 12.7 Metrics Assessment
    add_heading(doc, "12.7 Metrics Assessment: Beyond WER", 2)
    add_para(doc, (
        "A key insight from the evaluation is that standard WER is insufficient for assessing "
        "transcription quality in a production context. Two additional metrics were developed:"
    ))
    add_bullet_bold_value(doc, "WWER (Weighted WER): ",
        "Weights entity tokens (names, numbers, locations) at 2x importance and function words "
        "(the, is, of) at 0.5x. A transcription that captures 'Biden met with Putin' but misses "
        "articles is more useful than one that gets articles right but hallucates names.")
    add_bullet_bold_value(doc, "NEA F1 (Named Entity Accuracy): ",
        "Measures how well high-value tokens are preserved. At 38.8% F1, the current system "
        "misses at least one important entity in 85% of segments — this is the primary quality "
        "gap for real-world usability.")

    # 12.8 Arabic Dead End
    add_heading(doc, "12.8 Arabic Language Attempt", 2)
    add_para(doc, (
        "An attempt was made to extend the system to Arabic language processing. However, this "
        "proved infeasible with the current model architecture: the LLaMA-2 tokenizer is designed "
        "for English text, and the fine-tuning was performed on English-only data (LRS3). The "
        "tokenizer cannot produce meaningful Arabic text tokens, and the visual features are "
        "language-agnostic but the decoder is not. Supporting Arabic would require either a "
        "multilingual LLM base or separate Arabic-trained model weights."
    ))

    # 12.9 POC Fine-Tune
    add_heading(doc, "12.9 POC Fine-Tune to Overfit", 2)
    add_para(doc, (
        "A preliminary fine-tuning experiment was conducted as a proof-of-concept to validate "
        "that the training pipeline and infrastructure work correctly. The model was trained on "
        "a small dataset with the goal of overfitting — achieving very low training loss regardless "
        "of generalization. This experiment confirmed that the fairseq-hydra-train integration, "
        "QLoRA adapter setup, and checkpoint saving all function correctly, paving the way for "
        "the full domain adaptation fine-tuning."
    ))

    # 12.10 Key Insights
    add_heading(doc, "12.10 Key Insights", 2)
    add_para(doc, "The evaluation and tuning work yielded several important conclusions:")
    add_bullet(doc, "The domain gap (LRS3 → AVSpeech) is the primary bottleneck, not the model architecture")
    add_bullet(doc, "Fine-tuning on in-domain AVSpeech data is expected to reduce WER by 15-25 percentage points")
    add_bullet(doc, "Decode-time parameter tuning (lenpen, beam width) can provide 5-10% relative improvement")
    add_bullet(doc, "WER alone is insufficient — WWER and NEA F1 better capture real-world transcription utility")
    add_bullet(doc, "The 20.6% hallucination rate is the most urgent quality issue to address")
    add_bullet(doc, "Expected post-fine-tuning WER: 40-50% (still above paper's 25.4% due to harder conditions)")

    doc.add_page_break()


def chapter_13(doc):
    add_heading(doc, "13. Fine-Tuning & Training Infrastructure", 1)

    add_para(doc, (
        "Based on the evaluation results (Chapter 12), fine-tuning the model on in-domain AVSpeech data "
        "is the most promising path to improving transcription quality. This chapter documents the training "
        "infrastructure that was built, the training configuration designed for domain adaptation, and the "
        "research that informed the fine-tuning strategy."
    ))

    # 13.1 Training Pipeline
    add_heading(doc, "13.1 Fine-Tuning Pipeline (579 Lines)", 2)
    add_para(doc, (
        "The run_avspeech_finetune_pipeline.sh script (579 lines) implements the complete end-to-end "
        "fine-tuning workflow. It reuses all lib/ modules from the inference pipeline for preprocessing "
        "(stages 0-6), then hands off to fairseq-hydra-train for the actual training."
    ))
    add_para(doc, "Key features of the fine-tuning pipeline:")
    add_bullet(doc, "Reuses all lib/ modules for preprocessing — same code path as inference pipeline")
    add_bullet(doc, "fairseq-hydra-train integration with custom task (vsp_llm_training)")
    add_bullet(doc, "Multi-GPU support via configurable NUM_GPUS environment variable (1-8 GPUs)")
    add_bullet(doc, "Checkpointing from pre-trained model (checkpoint_finetune.pt)")
    add_bullet(doc, "Dry-run mode for validating setup without launching training")
    add_bullet(doc, "Skip-preprocess option for re-training on already-preprocessed data")
    add_bullet(doc, "Training output directory management with automatic versioning")

    # 13.2 Training Config
    add_heading(doc, "13.2 Training Configuration", 2)
    add_para(doc, (
        "The training configuration (vsp-llm-avspeech-finetune.yaml, 136 lines) was designed specifically "
        "for domain adaptation from LRS3 to AVSpeech. It differs from the paper's original config in "
        "several important ways:"
    ))
    add_styled_table(doc,
        ["Parameter", "Paper Config", "Argos Config", "Rationale"],
        [
            ["freeze_finetune_updates", "18,000", "5,000", "Unfreeze encoder earlier for domain adaptation"],
            ["max_update", "30,000", "15,000", "Shorter schedule for smaller/diverse datasets"],
            ["warmup_steps", "10,000", "3,000", "Proportional to max_update"],
            ["decay_steps", "20,000", "12,000", "Tri-stage LR: warmup → hold → decay"],
            ["lr", "0.0005", "0.0005", "Same learning rate (Adam optimizer)"],
            ["batch_size", "1", "1", "Limited by GPU memory (7B parameter model)"],
            ["update_freq", "[1]", "[1] or [8]", "8 GPUs → freq=1; 1 GPU → freq=8 (grad accumulation)"],
            ["label_smoothing", "0.1", "0.1", "Prevent overconfident predictions"],
            ["seed", "1337", "1337", "Reproducibility"],
            ["save_interval_updates", "—", "2,500", "Checkpoint every 2,500 steps"],
            ["keep_interval_updates", "—", "3", "Keep last 3 checkpoints (disk space)"],
        ],
        col_widths=[1.6, 1.0, 1.0, 2.9]
    )

    # 13.3 QLoRA
    add_heading(doc, "13.3 QLoRA Architecture", 2)
    add_para(doc, (
        "The fine-tuning uses QLoRA (Quantized Low-Rank Adaptation) to make training feasible on "
        "consumer GPUs. Instead of updating all 7 billion parameters of LLaMA-2, QLoRA adds small "
        "trainable adapter matrices to each linear layer:"
    ))
    add_styled_table(doc,
        ["QLoRA Parameter", "Value", "Explanation"],
        [
            ["Rank (r)", "16", "Dimension of the low-rank decomposition"],
            ["Alpha (α)", "32", "Scaling factor (α/r = 2.0)"],
            ["Quantization", "4-bit NF4", "Base model quantized to 4 bits (NormalFloat4)"],
            ["Target Layers", "All linear layers", "LoRA adapters on q_proj, k_proj, v_proj, o_proj, etc."],
            ["Trainable Params", "~0.1% of total", "Only the LoRA adapter matrices are updated"],
            ["Memory Savings", "~75%", "4-bit quantization reduces memory from ~28GB to ~7GB"],
        ],
        col_widths=[1.5, 1.2, 3.8]
    )

    # 13.4 Model Architecture Detail
    add_heading(doc, "13.4 Full Model Architecture", 2)
    add_code_block(doc, (
        "┌─────────────────────────────────────────────────────────┐\n"
        "│                    VSP-LLM Architecture                 │\n"
        "├─────────────────────────────────────────────────────────┤\n"
        "│                                                         │\n"
        "│  Input: 88x88 grayscale mouth crops @ 25 FPS            │\n"
        "│                    │                                     │\n"
        "│                    ▼                                     │\n"
        "│  ┌─────────────────────────────────┐                    │\n"
        "│  │     AV-HuBERT Encoder           │                    │\n"
        "│  │     315M parameters             │                    │\n"
        "│  │     Output: 1024-dim features   │                    │\n"
        "│  └─────────────────────────────────┘                    │\n"
        "│                    │                                     │\n"
        "│                    ▼                                     │\n"
        "│  ┌─────────────────────────────────┐                    │\n"
        "│  │     Linear Projection           │                    │\n"
        "│  │     1024 → 4096 dimensions      │                    │\n"
        "│  └─────────────────────────────────┘                    │\n"
        "│                    │                                     │\n"
        "│                    ▼                                     │\n"
        "│  ┌─────────────────────────────────┐                    │\n"
        "│  │     LLaMA-2-7B + QLoRA          │                    │\n"
        "│  │     7B params (4-bit quantized)  │                    │\n"
        "│  │     + LoRA adapters (r=16, α=32) │                    │\n"
        "│  │     ~0.1% trainable              │                    │\n"
        "│  └─────────────────────────────────┘                    │\n"
        "│                    │                                     │\n"
        "│                    ▼                                     │\n"
        "│  Output: Text transcription (beam search or sampling)   │\n"
        "│                                                         │\n"
        "└─────────────────────────────────────────────────────────┘"
    ))

    # 13.5 AVSpeech Data Strategy
    add_heading(doc, "13.5 AVSpeech Data Strategy", 2)
    add_para(doc, (
        "The training data for domain adaptation was sourced from AVSpeech, a large-scale dataset "
        "of short video clips extracted from YouTube. Unlike LRS3 (which contains only TED talks), "
        "AVSpeech represents the diversity of real-world video: multiple camera angles, variable "
        "lighting, background noise, and speakers of all demographics."
    ))
    add_para(doc, (
        "The data was imported from an old server that had been previously mounted. The pipeline's "
        "preprocessing stages (segmentation, normalization, face detection, mouth cropping, ASR) "
        "prepare this raw data for training — the same stages used in the inference pipeline."
    ))

    # 13.6 Research Notes
    add_heading(doc, "13.6 Training Research Notes", 2)
    add_para(doc, (
        "Detailed research notes (docs/training-research-notes.md, 186 lines) document the analysis "
        "that informed the fine-tuning strategy:"
    ))
    add_bullet_bold_value(doc, "LoRA Rank Selection: ",
        "Rank 16 was chosen as optimal for the VSP-LLM model size. Lower ranks (4, 8) showed "
        "insufficient adaptation capacity in preliminary experiments, while higher ranks (32, 64) "
        "increased memory requirements without proportional quality gains.")
    add_bullet_bold_value(doc, "Angle Robustness: ",
        "Analysis of face detection success rates versus camera angle revealed that the model "
        "performs best with angles within ±30° of frontal. Fine-tuning on diverse-angle data "
        "is expected to improve robustness.")
    add_bullet_bold_value(doc, "Length Distribution: ",
        "Training segments should ideally be 5-12 seconds. Segments shorter than 3 seconds "
        "contain too few visual features, while segments longer than 15 seconds exceed the "
        "model's effective context window.")
    add_bullet_bold_value(doc, "Expected Improvement: ",
        "Based on domain adaptation literature and the evaluation results, fine-tuning on "
        "AVSpeech data is expected to reduce WER by 15-25 percentage points, bringing it "
        "to the 40-50% range on real-world YouTube content.")

    doc.add_page_break()


def chapter_14(doc):
    add_heading(doc, "14. Executive Summary", 1)

    add_para(doc, (
        "This chapter provides a high-level overview of the Argos project's scope, deliverables, "
        "and current state for management review."
    ))

    # 14.1 Project Scope
    add_heading(doc, "14.1 Project Scope & Timeline", 2)
    add_para(doc, (
        "The Argos project began on October 18, 2024, with the goal of building a production-ready "
        "visual speech processing (lip-reading) system based on the VSP-LLM research paper. All "
        "engineering work — pipeline development, web UI, Docker container, deployment, testing, and "
        "documentation — was carried out by Yoad Oxman as the sole developer."
    ))
    add_para(doc, (
        "The system takes raw video as input and produces text transcriptions by analyzing lip movements, "
        "with no audio input required. This enables transcription in scenarios where audio is unavailable, "
        "classified, or unreliable."
    ))

    # 14.2 Key Deliverables
    add_heading(doc, "14.2 Key Deliverables", 2)
    add_styled_table(doc,
        ["Deliverable", "Status", "Details"],
        [
            ["End-to-End Pipeline", "Complete (v7)", "8-stage pipeline, modular architecture (501 + 1,562 lines)"],
            ["Web UI", "Complete", "6 screens, 12 API endpoints, drag-drop upload, transcription management"],
            ["Docker Container", "5 versions deployed", "nvidia/cuda base, dual venvs, ~65GB galaxy export"],
            ["Performance Evaluation", "Complete", "860 segments, 6 analysis reports, 7 tuning experiments"],
            ["Report Generation", "Complete", "WER, WWER, NEA metrics in 5 formats (CSV/HTML/JSON/TXT/ANSI)"],
            ["Burned Videos", "Complete", "Subtitle overlays on original videos for visual review"],
            ["Fine-Tuning Infrastructure", "Ready", "579-line pipeline, training config, QLoRA setup"],
            ["Testing Suite", "Complete", "37 automated tests, 5 test scripts (1,731 lines)"],
            ["Documentation", "Comprehensive", "Architecture, dev guide, sync changelog, training notes, bug tracking"],
            ["4-Environment Support", "Operational", "EC2, 2 standalone Ubuntu, HORIZON (no internet)"],
        ],
        col_widths=[1.8, 1.2, 3.5]
    )

    # 14.3 Effort Estimate
    add_heading(doc, "14.3 Total Effort Estimate", 2)
    add_styled_table(doc,
        ["Phase", "Duration", "Key Activities"],
        [
            ["Literature Review", "~2-3 weeks", "VSP-LLM, AV-HuBERT, Auto-AVSR, fairseq, LRS3, LLaMA-2, QLoRA papers"],
            ["Failed Venv Builds", "~2 weeks", "Attempted venv setup on standalone machines (failed, led to container approach)"],
            ["Driver & CUDA Setup", "~1 month", "NVIDIA drivers, CUDA toolkit, GPU verification on standalone machines"],
            ["Pipeline Development", "Ongoing", "7 versions, from 240 to 501+1,562 lines, 13 custom Python scripts"],
            ["Web UI Development", "Ongoing", "5,780 lines across 11 files, 4 backend services"],
            ["Container & Deployment", "Ongoing", "5 releases, 37 bugs fixed, Dockerfile, galaxy export, installation scripts"],
            ["Evaluation & Tuning", "Ongoing", "860 segments, 7 experiments, 6 analysis reports"],
            ["Training Infrastructure", "Ongoing", "579-line pipeline, training config, research notes"],
        ],
        col_widths=[1.5, 1.0, 4.0]
    )

    # 14.4 Performance
    add_heading(doc, "14.4 Current Performance", 2)
    add_styled_table(doc,
        ["Metric", "Current Value", "Expected Post-Fine-Tuning"],
        [
            ["Mean WER", "67.0%", "40-50%"],
            ["Usable Segments (WER ≤ 20%)", "11.4%", "25-40%"],
            ["Hallucination Rate (WER ≥ 100%)", "20.6%", "<10%"],
            ["NEA F1 (Named Entity Accuracy)", "38.8%", "55-70%"],
            ["Paper WER (LRS3 benchmark)", "25.4%", "25.4% (same model on clean data)"],
        ],
        col_widths=[2.5, 1.5, 2.5]
    )

    # 14.5 Code Inventory
    add_heading(doc, "14.5 Complete Code Inventory", 2)
    add_para(doc, (
        "The following table summarizes the complete codebase written for the Argos project. "
        "All code was developed by Yoad Oxman:"
    ))
    add_styled_table(doc,
        ["Category", "Files", "Lines", "Key Components"],
        [
            ["Pipeline Scripts", "4", "1,299", "Main pipeline (501), fine-tune (579), tuning (219)"],
            ["Library Modules", "12", "1,562", "common, config, archive, normalization, asr, decode, outputs, etc."],
            ["Custom Python Scripts", "13", "3,342", "make_report, make_burn, fast_segment, transcribe_segments, etc."],
            ["Web UI + Server", "11", "5,780", "server.py (1,124), app.js (1,921), style.css (1,191), 4 services"],
            ["Face Detection", "4", "532", "MediaPipe detector (52+220), RetinaFace (43+217)"],
            ["Training Config", "1", "136", "vsp-llm-avspeech-finetune.yaml (Hydra config)"],
            ["Testing Scripts", "5", "1,731", "37 tests + smoke tests + environment tests + sync verification"],
            ["Deployment Scripts", "5", "~500", "build_container, check_inventory, monitor, resume"],
            ["Grand Total", "~55", "~14,882", "All custom production code"],
        ],
        col_widths=[1.5, 0.5, 0.5, 4.0]
    )

    # 14.6 Infrastructure
    add_heading(doc, "14.6 Infrastructure Summary", 2)
    add_styled_table(doc,
        ["Aspect", "Count", "Details"],
        [
            ["Deployment Environments", "4", "EC2, 2 standalone Ubuntu, HORIZON (no internet)"],
            ["Container Releases", "5", "Feb 3 → Feb 17, 2026 (v1.0.0 — v1.0.37)"],
            ["Bugs Fixed", "37", "Installation (13), deployment/GPU (12), final fixes (12)"],
            ["Container Sync Items", "28", "Pending EC2 → container replication"],
            ["Automated Tests", "37", "Covering all 12 lib/ modules"],
            ["Virtual Environments", "2", "pre-process-venv (135 pkgs) + vsp-llm-yoad-venv (256 pkgs)"],
            ["Upstream Repos Integrated", "3", "auto_avsr, VSP-LLM, av_hubert"],
            ["Git Commits", "40+", "From initial commit through v1.0.40+"],
        ],
        col_widths=[1.8, 0.6, 4.1]
    )

    # 14.7 Recommendations
    add_heading(doc, "14.7 Recommendations", 2)
    add_para(doc, "Based on the evaluation results and research analysis, the following next steps are recommended:")
    add_bullet_bold_value(doc, "1. Fine-Tune on AVSpeech: ",
        "Domain adaptation is the highest-impact improvement available. The training infrastructure "
        "is ready (Chapter 13). Expected improvement: 15-25 WER points.")
    add_bullet_bold_value(doc, "2. Evaluate RetinaFace: ",
        "Switching from MediaPipe to RetinaFace for face detection may improve mouth crop quality "
        "on difficult angles, potentially contributing 2-5 WER points improvement.")
    add_bullet_bold_value(doc, "3. Multi-Language Support: ",
        "Investigate multilingual LLM bases (e.g., LLaMA-3, Mistral) that could support Arabic "
        "and other languages. The current English-only tokenizer is the blocking factor.")
    add_bullet_bold_value(doc, "4. Segment Duration Optimization: ",
        "The hyperparameter tuning report identified segment duration as a high-impact data-level "
        "lever. Filtering segments to 5+ seconds and capping at 20-30 seconds may improve WER by 5-10%.")

    # 14.8 Milestones
    add_heading(doc, "14.8 Project Milestones", 2)
    add_styled_table(doc,
        ["Date", "Milestone"],
        [
            ["Oct 18, 2024", "Project start — literature review begins"],
            ["Oct-Nov 2024", "First pipeline version (~285 lines, galaxy_export)"],
            ["Nov-Dec 2024", "Pipeline iterations, face detection integration, AVSpeech processing"],
            ["Dec 29, 2025", 'Presentation: "Argos — A Working Pipeline" (19 slides)'],
            ["Jan 2026", "Modular refactoring (823 → 501+1,562 lines)"],
            ["Jan 2026", "Web UI development (5,780 lines)"],
            ["Feb 3, 2026", "First container deployment"],
            ["Feb 3-17, 2026", "Five container releases, 37 bugs fixed"],
            ["Feb 2026", "Performance evaluation (860 segments, 6 reports, 7 experiments)"],
            ["Feb 2026", "Fine-tuning infrastructure ready (579-line pipeline + config)"],
            ["Feb 2026", "HORIZON server operational (closed network, no internet)"],
        ],
        col_widths=[1.3, 5.2]
    )

    # Final note
    doc.add_paragraph()
    add_para(doc, (
        "This document represents a comprehensive record of all research and development work "
        "performed on the Argos visual speech processing system. It is intended as a living document "
        "that will be updated as the project progresses through fine-tuning, evaluation, and deployment."
    ), italic=True, color=C_GRAY)


# ═══════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════

def main():
    print("Generating Argos — The Orchard (Detailed Edition v2)...")
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Header/Footer
    add_header_footer(doc)

    # Build document
    create_cover_page(doc)
    create_toc(doc)

    chapter_1(doc)
    chapter_2(doc)
    chapter_3(doc)
    chapter_4(doc)
    chapter_5(doc)
    chapter_6(doc)
    chapter_7(doc)
    chapter_8(doc)
    chapter_9(doc)
    chapter_10(doc)
    chapter_11(doc)
    chapter_12(doc)
    chapter_13(doc)
    chapter_14(doc)

    # Save
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_FILE))
    print(f"Saved: {OUTPUT_FILE}")
    print(f"Size: {OUTPUT_FILE.stat().st_size / 1024:.0f} KB")


if __name__ == "__main__":
    main()
